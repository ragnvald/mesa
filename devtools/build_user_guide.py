"""Build MESA User Guide DOCX in English and Portuguese (Mozambique).

Recreates docs/MESA_User_Guide_en.docx and docs/MESA_User_Guide_pt.docx from
the canonical wiki content under ../mesa.wiki/. SVG figures are skipped on
purpose (no svg-to-raster dependency required); only PNG screenshots and
diagrams from the wiki are embedded. Run from the repo root:

    python devtools/build_user_guide.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
WIKI_IMG = REPO.parent / "mesa.wiki" / "images"
DOCS = REPO / "docs"

AUTHOR = "Ragnvald Larsen"
AFFILIATION = "Norwegian Environment Agency"
VERSION = "MESA 5.0.2"

PT_MONTHS = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
]


# ---------------------------------------------------------------------------
# Block helpers — content is language-keyed so a single tree feeds both files.
# ---------------------------------------------------------------------------

def H(level, en, pt):
    return ("H", level, {"en": en, "pt": pt})


def P(en, pt):
    return ("P", {"en": en, "pt": pt})


def UL(items_en, items_pt):
    return ("UL", {"en": items_en, "pt": items_pt})


def OL(items_en, items_pt):
    return ("OL", {"en": items_en, "pt": items_pt})


def IMG(filename, caption_en, caption_pt, width_cm=14.0, landscape=False):
    return ("IMG", filename, {"en": caption_en, "pt": caption_pt}, width_cm, landscape)


def TBL(headers_en, rows_en, headers_pt, rows_pt):
    return ("TBL", {"en": (headers_en, rows_en), "pt": (headers_pt, rows_pt)})


def PB():
    return ("PB",)


# ---------------------------------------------------------------------------
# Content tree — wiki-derived, both languages side-by-side.
# ---------------------------------------------------------------------------

def build_blocks():
    blocks = []
    add = blocks.append

    # =========================================================================
    # 1. Introduction (from Home.md)
    # =========================================================================
    add(H(1,
          "1. Introduction",
          "1. Introdução"))
    add(IMG("top_intro.png",
            "MESA — Mapping Environmentally Sensitive Assets.",
            "MESA — Mapeamento de Activos Ambientalmente Sensíveis.",
            width_cm=15.0))
    add(P(
        "The MESA method identifies and evaluates key environmental factors such as ecological significance, biodiversity, species abundance, habitat quality, and the presence of sensitive or protected areas. By combining these factors, the method supports a structured assessment of how pollution incidents may affect the environment.",
        "O método MESA identifica e avalia factores ambientais essenciais, tais como o significado ecológico, a biodiversidade, a abundância de espécies, a qualidade do habitat e a presença de áreas sensíveis ou protegidas. Ao combinar estes factores, o método permite uma avaliação estruturada de como os incidentes de poluição podem afectar o ambiente."))
    add(P(
        "Complementing the method, the MESA tool provides a Windows desktop workflow for conducting environmental sensitivity assessments and producing map- and report-ready outputs. The current MESA 5.0 build produces:",
        "Complementando o método, a ferramenta MESA disponibiliza um fluxo de trabalho de ambiente de trabalho Windows para realizar avaliações de sensibilidade ambiental e produzir resultados prontos para mapas e relatórios. A actual versão MESA 5.0 produz:"))
    add(UL(
        ["GeoParquet tables under output/geoparquet/ as the canonical analysis store",
         "MBTiles under output/mbtiles/ for fast map viewing",
         "A ready-to-open QGIS project under qgis/ for cartographic production",
         "Word reports under output/reports/"],
        ["Tabelas GeoParquet em output/geoparquet/ como repositório canónico da análise",
         "MBTiles em output/mbtiles/ para visualização rápida de mapas",
         "Um projecto QGIS pronto a abrir em qgis/ para produção cartográfica",
         "Relatórios Word em output/reports/"]))
    add(P(
        "If you use the packaged distribution, start with mesa.exe. If you are running from source, start with python mesa.py.",
        "Se utilizar a distribuição empacotada, inicie com mesa.exe. Se estiver a executar a partir do código-fonte, inicie com python mesa.py."))
    add(P(
        "MESA 5.0 is the current release line. The method continues to evolve, and ongoing method development will guide future tool development.",
        "MESA 5.0 é a linha de versão actual. O método continua a evoluir e o desenvolvimento contínuo do método orientará o desenvolvimento futuro da ferramenta."))

    add(H(2, "1.1 System requirements (minimum)", "1.1 Requisitos de sistema (mínimo)"))
    add(UL(
        ["Windows 11", "16 GB RAM", "8 CPU cores"],
        ["Windows 11", "16 GB de RAM", "8 núcleos de CPU"]))
    add(P(
        "Together, the MESA method and tool support better pollution preparedness, response planning, and environmental management. By systematically assessing and mapping environmental sensitivities, stakeholders can make better-informed decisions, prioritize mitigation, and improve coordination across teams.",
        "Em conjunto, o método e a ferramenta MESA apoiam uma melhor preparação para episódios de poluição, planeamento de resposta e gestão ambiental. Ao avaliar e mapear sistematicamente as sensibilidades ambientais, as partes interessadas podem tomar decisões mais informadas, priorizar a mitigação e melhorar a coordenação entre equipas."))
    add(IMG("Conflicting_interests.png",
            "Conflicting interests captured by sensitivity mapping.",
            "Interesses conflituantes captados pelo mapeamento de sensibilidades.",
            width_cm=12.0))
    add(PB())

    # =========================================================================
    # 2. What's new in MESA 5.0
    # =========================================================================
    add(H(1, "2. What's new in MESA 5.0", "2. Novidades no MESA 5.0"))
    add(P(
        "MESA 5.0 is a usability-, packaging-, and reliability-focused update of the desktop workflow. The headline gains are faster processing on real project data and a much lower crash rate on long runs, both driven by smarter memory and CPU balancing inside the pipeline. Around those two, the launcher has been simplified, screens have been deprecated or moved, and the report engine has had a substantial polish pass.",
        "O MESA 5.0 é uma actualização do fluxo de trabalho de ambiente de trabalho centrada na usabilidade, empacotamento e fiabilidade. Os principais ganhos são um processamento mais rápido em dados reais de projecto e uma taxa de falhas muito inferior em execuções longas, ambos resultado de um balanceamento mais inteligente de memória e CPU dentro da pipeline. Em torno destes dois eixos, o iniciador foi simplificado, alguns ecrãs foram removidos ou deslocados e o motor de relatórios foi substancialmente aperfeiçoado."))

    add(H(2, "2.1 Headline gains", "2.1 Principais ganhos"))
    add(H(3, "Substantially faster processing", "Processamento substancialmente mais rápido"))
    add(P(
        "Recent work in the data pipeline cuts avoidable temp-I/O, re-plans chunks more carefully, and tightens the flatten path. On larger projects this materially shortens processing time and reduces the feeling that runs are stuck in I/O-heavy phases. Exact improvement depends on dataset size, geometry complexity, disk speed, and hardware.",
        "Trabalho recente na pipeline de dados elimina I/O temporário evitável, replaneia blocos com maior cuidado e optimiza o caminho de achatamento. Em projectos maiores isto reduz materialmente o tempo de processamento e diminui a sensação de que as execuções ficam paradas em fases intensivas de I/O. A melhoria exacta depende do tamanho dos dados, da complexidade da geometria, da velocidade do disco e do hardware."))

    add(H(3, "Far fewer crashes on long runs", "Muito menos falhas em execuções longas"))
    add(P(
        "A 140+ GB peak-RSS swap stall on a 64 GB host showed where the previous memory budgets were too optimistic. Several layers now protect long-running jobs:",
        "Uma paragem por swap com pico de RSS superior a 140 GB num computador de 64 GB de RAM revelou onde os antigos limites de memória eram demasiado optimistas. Várias camadas protegem agora os trabalhos longos:"))
    add(UL(
        [
            "Runtime auto-tuning of worker counts derived from the live host fingerprint and the project's data fingerprint.",
            "Platform-aware memory budgets — Apple Silicon hosts use a tighter budget; Windows / Linux / Intel Mac hosts use a wider budget.",
            "Pre-flight RAM/swap check before the flatten stage, so the run aborts cleanly instead of entering swap-death.",
            "Per-pool memory panic watchdog for every multiprocessing pool.",
            "Process-lifetime sentinel watchdog as a last-resort backstop covering parent-side work between pools.",
            "Three-phase flatten split — huge / large / small partitions run with progressively wider parallelism.",
            "Sliver cleanup that drops zero-area precision artefacts from tbl_flat after polygonisation.",
        ],
        [
            "Auto-ajuste em tempo de execução do número de processos de trabalho com base na impressão digital do computador e dos dados do projecto.",
            "Limites de memória sensíveis à plataforma — máquinas Apple Silicon usam um limite mais apertado; máquinas Windows / Linux / Mac Intel usam um limite mais largo.",
            "Verificação prévia de RAM/swap antes da fase de achatamento, para que a execução aborte limpamente em vez de entrar em swap-death.",
            "Vigilante de pânico de memória por pool de multiprocessamento.",
            "Vigilante de tempo de vida do processo como rede de segurança final para trabalho do processo-pai entre pools.",
            "Achatamento em três fases — partições enormes / grandes / pequenas executam com paralelismo progressivamente mais amplo.",
            "Limpeza de fragmentos que remove artefactos de precisão de área zero de tbl_flat após a poligonização.",
        ]))
    add(P(
        "A single [auto-tune] block at the top of each run log summarises the platform, hardware, data fingerprint, and per-key decisions, so operators can audit at a glance instead of grepping scattered log lines.",
        "Um único bloco [auto-tune] no topo de cada registo de execução resume a plataforma, o hardware, a impressão digital dos dados e as decisões por chave, permitindo aos operadores auditar de relance em vez de procurar linhas de registo dispersas."))

    add(H(2, "2.2 Launcher and workflow", "2.2 Iniciador e fluxo de trabalho"))
    add(H(3, "Simpler tab structure", "Estrutura de separadores mais simples"))
    add(P(
        "The desktop launcher now has five tabs (down from the earlier six):",
        "O iniciador do ambiente de trabalho tem agora cinco separadores (anteriormente seis):"))
    add(UL(
        ["Workflows — launch the per-phase tools",
         "Status — project counters, recent activity, and headline metrics",
         "Manage data — backup, restore, clear output, and Publish to GeoNode",
         "Config — direct config.ini editor; also opens the Tune processing popup",
         "About — build banner and system profile"],
        ["Fluxos de trabalho — abrem as ferramentas de cada fase",
         "Estado — contadores do projecto, actividade recente e métricas principais",
         "Gerir dados — cópia de segurança, restauro, limpeza dos resultados e Publicar no GeoNode",
         "Config — editor directo de config.ini; abre também a janela Afinação do processamento",
         "Sobre — informação da versão e perfil do sistema"]))
    add(P(
        "Two former top-level tabs were extracted into popups so the launcher stays focused on day-to-day work: Tune processing is now a popup launched from the Config tab, and Publish to GeoNode is now a popup launched from the Manage data tab.",
        "Dois antigos separadores principais foram convertidos em janelas pop-up para manter o iniciador focado no trabalho do dia-a-dia: Afinação do processamento é agora uma janela aberta a partir do separador Config, e Publicar no GeoNode é agora uma janela aberta a partir do separador Gerir dados."))

    add(H(3, "Unified processing runner with numbered stages",
          "Executor de processamento unificado com fases numeradas"))
    add(UL(
        [
            "Numbered stages in both the run grid and the live status text, with rerun guidance when individual stages have already completed.",
            "Backfill is its own pipeline stage with its own checkbox alongside Prep / Intersect / Flatten.",
            "Normal / Advanced toggle hides per-stage selectors and the sliver-cleanup option behind a deliberate click.",
            "Directional headers in the Advanced grid show stage order at a glance.",
            "Progress map is launched from inside the runner window.",
            "CLI flags (--no-prep, --no-intersect, --no-flatten, --no-cleanup-slivers) mirror the GUI controls for headless runs.",
        ],
        [
            "Fases numeradas na grelha de execução e no texto de estado em tempo real, com orientação de re-execução quando as fases individuais já foram concluídas.",
            "Backfill é uma fase própria da pipeline com a sua própria caixa de selecção ao lado de Prep / Intersect / Flatten.",
            "Alternador Normal / Avançado oculta os selectores por fase e a opção de limpeza de fragmentos atrás de um clique deliberado.",
            "Cabeçalhos direccionais na grelha Avançada mostram a ordem das fases num só olhar.",
            "O Mapa de progresso é aberto a partir da própria janela do executor.",
            "Os argumentos de linha de comandos (--no-prep, --no-intersect, --no-flatten, --no-cleanup-slivers) espelham os controlos da interface gráfica para execuções sem ecrã.",
        ]))

    add(H(3, "Tune processing is safer", "Afinação do processamento mais segura"))
    add(P(
        "Tune processing follows an evaluate-before-commit pattern: Evaluate shows current versus advised values first, Commit changes is separate, and Restore previous tuning rolls back the last committed tune set.",
        "A Afinação do processamento segue um padrão avaliar-antes-de-confirmar: Avaliar mostra primeiro os valores actuais face aos valores aconselhados, Confirmar alterações é separado, e Restaurar afinação anterior reverte o último conjunto de afinação confirmado."))

    add(H(2, "2.3 Authoring and reporting", "2.3 Autoria e relatórios"))
    add(P(
        "Asset import has two convenience options on by default: Dissolve adjacent polygons (merges touching polygons that share identical attribute values) and automatic buffering of point and line inputs using default_point_buffer_m and default_line_buffer_m so zero-area features still contribute. The atlas helper now reads existing atlas state on open and includes a Delete action. The Word report engine has had a layout pass — index maps include legends, area maps are scaled down with explanatory text, atlas tile maps show the full grid with the current tile highlighted, and each atlas tile heading begins on a fresh page.",
        "A importação de Activos tem duas opções de conveniência activas por omissão: Dissolver polígonos adjacentes (combina polígonos contíguos que partilham valores de atributos idênticos) e amortecimento automático de entradas de pontos e linhas, utilizando default_point_buffer_m e default_line_buffer_m para que as feições de área zero continuem a contribuir. O auxiliar de atlas lê agora o estado existente do atlas ao abrir e inclui uma acção Eliminar. O motor de relatórios Word recebeu uma revisão de paginação — os mapas-índice incluem legendas, os mapas de área foram reduzidos com texto explicativo, os mapas dos azulejos do atlas mostram a grelha completa com o azulejo actual destacado, e cada cabeçalho de azulejo do atlas começa numa nova página."))
    add(PB())

    # =========================================================================
    # 3. Quickstart
    # =========================================================================
    add(H(1, "3. Quickstart", "3. Início rápido"))
    add(P(
        "This quick-start guide shows how to run the current MESA 5 workflow end-to-end: prepare inputs, configure parameters, run processing, and review outputs.",
        "Este guia de início rápido mostra como executar o fluxo de trabalho do MESA 5 de ponta a ponta: preparar entradas, configurar parâmetros, executar o processamento e rever os resultados."))
    add(IMG("process_overview.png",
            "MESA process overview — prepare data, configure, process, review.",
            "Visão geral do processo MESA — preparar dados, configurar, processar, rever.",
            width_cm=15.0))

    add(H(2, "3.1 The MESA tool", "3.1 A ferramenta MESA"))
    add(P(
        "MESA is operated from the desktop launcher. Use mesa.exe for the packaged distribution or python mesa.py when running from source. From the launcher you can import data, configure parameters, run processing, open interactive viewers, and export reports.",
        "O MESA é operado a partir do iniciador do ambiente de trabalho. Utilize mesa.exe para a distribuição empacotada ou python mesa.py quando executar a partir do código-fonte. A partir do iniciador pode importar dados, configurar parâmetros, executar o processamento, abrir visualizadores interactivos e exportar relatórios."))

    add(H(2, "3.2 Getting ready", "3.2 Preparação"))
    add(OL(
        ["System — Windows 11, minimum 16 GB RAM, minimum 8 CPU cores, and optional QGIS for atlas authoring.",
         "Workspace — unzip the latest MESA release (packaged) or clone the repository (source). Keep the folder structure intact.",
         "Disk space — plan for several GB free space for GeoParquet, MBTiles, logs, reports, and temporary processing files."],
        ["Sistema — Windows 11, mínimo 16 GB de RAM, mínimo 8 núcleos de CPU, e QGIS opcional para autoria do atlas.",
         "Espaço de trabalho — descomprima a versão mais recente do MESA (empacotada) ou clone o repositório (código-fonte). Mantenha a estrutura de pastas intacta.",
         "Espaço em disco — planeie vários GB de espaço livre para GeoParquet, MBTiles, registos, relatórios e ficheiros temporários de processamento."]))

    add(TBL(
        ["Data set", "Description", "Required"],
        [
            ["Environmental assets", "Natural or anthropogenic entities that provide ecological, socio-cultural, economic, or political value. Stored under input/asset/<collection>.", "Yes"],
            ["Geocode layer", "Spatial grids or polygons that aggregate sensitivity. H3 hexagons, quarter-degree grid cells, municipalities, or any non-overlapping polygons under input/geocode/.", "Yes"],
            ["Line data", "Roads, pipelines, coastlines, rivers, or similar linear assets stored under input/lines/.", "No"],
            ["Atlas tiles", "Optional atlas polygons for QGIS atlas exports and report rendering. Created inside MESA or imported from input/atlas/.", "No"],
            ["Analysis areas", "Optional study areas / groups used for side-by-side comparison and analysis outputs.", "No"],
        ],
        ["Conjunto de dados", "Descrição", "Obrigatório"],
        [
            ["Activos ambientais", "Entidades naturais ou antropogénicas que fornecem valor ecológico, sociocultural, económico ou político. Armazenadas em input/asset/<colecção>.", "Sim"],
            ["Camada de geocódigos", "Grelhas ou polígonos espaciais que agregam a sensibilidade. Hexágonos H3, células de grelha de quarto de grau, municípios ou polígonos não sobrepostos em input/geocode/.", "Sim"],
            ["Dados lineares", "Estradas, oleodutos, linhas de costa, rios ou activos lineares semelhantes em input/lines/.", "Não"],
            ["Azulejos do atlas", "Polígonos opcionais para exportações de atlas QGIS e relatórios. Criados dentro do MESA ou importados de input/atlas/.", "Não"],
            ["Áreas de análise", "Áreas de estudo / grupos opcionais utilizados para comparação lado a lado e resultados de análise.", "Não"],
        ]))

    add(H(2, "3.3 Importing data", "3.3 Importação de dados"))
    add(OL(
        ["Copy asset datasets into input/asset/.",
         "Copy geocode layers into input/geocode/.",
         "Copy line datasets into input/lines/ when you need segment-level sensitivity.",
         "Review config.ini if you need to adjust projections, tile zoom levels, buffers, or performance knobs."],
        ["Copie os conjuntos de dados de activos para input/asset/.",
         "Copie as camadas de geocódigos para input/geocode/.",
         "Copie os conjuntos de dados lineares para input/lines/ quando precisar de sensibilidade por segmento.",
         "Reveja config.ini se precisar de ajustar projecções, níveis de zoom dos azulejos, amortecimentos ou parâmetros de desempenho."]))
    add(P(
        "Then use Workflows → Prepare data (step 1):",
        "De seguida utilize Fluxos de trabalho → Preparar dados (passo 1):"))
    add(UL(
        ["Assets imports asset datasets and lets you maintain asset-group metadata.",
         "Geocodes imports geocode layers and can also generate or refresh supported grids.",
         "Lines imports and edits line data plus segmentation settings.",
         "Atlas imports or creates atlas polygons when you need atlas pages later."],
        ["Activos importa conjuntos de dados de activos e permite manter os metadados dos grupos de activos.",
         "Geocódigos importa camadas de geocódigos e pode também gerar ou actualizar grelhas suportadas.",
         "Linhas importa e edita dados lineares e definições de segmentação.",
         "Atlas importa ou cria polígonos de atlas quando precisar de páginas de atlas posteriormente."]))

    add(H(2, "3.4 Configuring the project", "3.4 Configurar o projecto"))
    add(P(
        "Use Workflows → Configure (step 2) before heavy runs. Parameters edits processing rules such as sensitivity weights, thresholds, and related settings. Analysis defines optional study-area polygons and named analysis groups.",
        "Utilize Fluxos de trabalho → Configurar (passo 2) antes de execuções pesadas. Parâmetros edita regras de processamento como pesos de sensibilidade, limiares e definições relacionadas. Análise define polígonos de área de estudo opcionais e grupos de análise nomeados."))

    add(H(2, "3.5 Running processing", "3.5 Execução do processamento"))
    add(OL(
        ["Open Workflows → Process (step 3) → Process.",
         "In the process runner, select the steps you want: Data processing (Prep / Intersect / Flatten), Tiles processing, Lines processing, Analysis processing.",
         "Leave Data processing enabled at minimum for a first full run.",
         "Enable Tiles processing if you want MBTiles in the same run.",
         "Enable Lines and/or Analysis processing when those datasets have been prepared.",
         "Use Progress map inside the process runner if you want a live minimap during longer runs.",
         "Use the Normal / Advanced toggle at the top of the runner to expose per-stage selectors and the sliver-cleanup option."],
        ["Abra Fluxos de trabalho → Processar (passo 3) → Processar.",
         "No executor de processamento, seleccione as fases pretendidas: Processamento de dados (Prep / Intersect / Flatten), Processamento de azulejos, Processamento de linhas, Processamento de análise.",
         "Mantenha pelo menos o Processamento de dados activado para uma primeira execução completa.",
         "Active o Processamento de azulejos se quiser MBTiles na mesma execução.",
         "Active o Processamento de linhas e/ou Análise quando esses dados estiverem preparados.",
         "Utilize o Mapa de progresso dentro do executor para ver um minimapa em tempo real em execuções mais longas.",
         "Utilize o alternador Normal / Avançado no topo do executor para expor os selectores por fase e a opção de limpeza de fragmentos."]))

    add(H(2, "3.6 Reviewing results", "3.6 Rever os resultados"))
    add(P(
        "After processing, use Workflows → Results (step 4): Asset map inspects asset layers and supporting data, Results map reviews processed outputs together with background layers, Compare study areas compares analysis groups side by side, and Report engine creates a Word report from the latest outputs.",
        "Após o processamento, utilize Fluxos de trabalho → Resultados (passo 4): Mapa de Activos inspecciona camadas de activos e dados de suporte, Mapa de Resultados revê resultados processados em conjunto com camadas de fundo, Comparar áreas de estudo compara grupos de análise lado a lado, e Motor de relatórios cria um relatório Word a partir dos resultados mais recentes."))

    add(H(2, "3.7 Backup and restore", "3.7 Cópia de segurança e restauro"))
    add(UL(
        ["Create backup saves input/, output/, and config.ini into one ZIP archive.",
         "Restore backup replaces the current input/, output/, and config.ini with the selected archive.",
         "Clear output removes generated outputs while keeping inputs and configuration."],
        ["Criar cópia de segurança guarda input/, output/ e config.ini num único arquivo ZIP.",
         "Restaurar cópia de segurança substitui input/, output/ e config.ini pelo arquivo seleccionado.",
         "Limpar resultados remove os resultados gerados, mantendo intactos os dados de entrada e a configuração."]))
    add(P(
        "Restore is destructive. Make sure no MESA windows are actively using the project folder before restoring.",
        "O restauro é destrutivo. Assegure-se de que nenhuma janela do MESA está a utilizar a pasta do projecto antes de restaurar."))
    add(PB())

    # =========================================================================
    # 4. User interface
    # =========================================================================
    add(H(1, "4. Desktop user interface", "4. Interface gráfica"))
    add(P(
        "The MESA 5 desktop launcher opens on Workflows and organizes the project into four practical phases: prepare data, configure the project, run processing, and review results.",
        "O iniciador MESA 5 abre no separador Fluxos de trabalho e organiza o projecto em quatro fases práticas: preparar dados, configurar o projecto, executar o processamento e rever os resultados."))
    add(IMG("ui_workflows.png",
            "Desktop launcher — Workflows tab.",
            "Iniciador do ambiente de trabalho — separador Fluxos de trabalho.",
            width_cm=15.0))
    add(IMG("ui_status.png",
            "Status tab — counters, recent activity, headline metrics.",
            "Separador Estado — contadores, actividade recente, métricas principais.",
            width_cm=15.0))
    add(IMG("ui_manage.png",
            "Manage data tab — backup, restore, clear output, GeoNode publish.",
            "Separador Gerir dados — cópia de segurança, restauro, limpeza, publicação GeoNode.",
            width_cm=15.0))
    add(IMG("ui_config.png",
            "Config tab — direct editor for config.ini, with Tune processing popup.",
            "Separador Config — editor directo de config.ini, com janela Afinação do processamento.",
            width_cm=15.0))
    add(IMG("ui_about.png",
            "About tab — build banner and system profile.",
            "Separador Sobre — informação da versão e perfil do sistema.",
            width_cm=15.0))

    add(H(2, "4.1 Prepare data (step 1)", "4.1 Preparar dados (passo 1)"))
    add(IMG("overview_prepare_data.png",
            "Prepare data — assets, geocodes, atlas.",
            "Preparar dados — activos, geocódigos, atlas.",
            width_cm=15.0))
    add(UL(
        ["Assets — import area assets from input/asset/ and maintain asset-group metadata. Two convenience options are on by default: Dissolve adjacent polygons (removes source-grid moiré) and automatic buffering of point and line inputs.",
         "Geocodes — import geocode layers or generate supported grids used by analysis and reporting.",
         "Atlas — import or create atlas polygons and maintain atlas page metadata. The helper now surfaces existing atlas state on open and includes a Delete action."],
        ["Activos — importar activos de área a partir de input/asset/ e manter metadados dos grupos. Duas opções de conveniência estão activas por omissão: Dissolver polígonos adjacentes (remove o efeito moiré da grelha de origem) e amortecimento automático de entradas de pontos e linhas.",
         "Geocódigos — importar camadas ou gerar grelhas suportadas usadas pela análise e pelos relatórios.",
         "Atlas — importar ou criar polígonos de atlas e manter os metadados das páginas. O auxiliar mostra agora o estado existente ao abrir e inclui uma acção Eliminar."]))
    add(IMG("ui_asset_manage.png",
            "Asset manager — group metadata and per-layer controls.",
            "Gestor de activos — metadados de grupo e controlos por camada.",
            width_cm=15.0))
    add(IMG("ui_geocode_create.png",
            "Geocode creation — generate or import geocode layers.",
            "Criação de geocódigos — gerar ou importar camadas.",
            width_cm=15.0))
    add(IMG("ui_atlas_manage.png",
            "Atlas manager — atlas pages and metadata.",
            "Gestor de atlas — páginas e metadados.",
            width_cm=15.0))

    add(H(2, "4.2 Configure (step 2)", "4.2 Configurar (passo 2)"))
    add(IMG("overview_configure_analysis.png",
            "Configure — parameters, lines, analysis.",
            "Configurar — parâmetros, linhas, análise.",
            width_cm=15.0))
    add(UL(
        ["Parameters — adjust processing rules such as weights, thresholds, and related setup values.",
         "Lines — import and edit line datasets such as transport, rivers, utilities, or coastlines.",
         "Analysis — define study-area polygons and named analysis groups."],
        ["Parâmetros — ajustar regras de processamento como pesos, limiares e valores relacionados.",
         "Linhas — importar e editar conjuntos lineares como transportes, rios, infra-estruturas ou linhas costeiras.",
         "Análise — definir polígonos de áreas de estudo e grupos de análise nomeados."]))
    add(IMG("ui_processing_setup.png",
            "Parameters helper — sensitivity and index weights.",
            "Auxiliar de parâmetros — pesos de sensibilidade e índices.",
            width_cm=15.0))
    add(IMG("ui_line_manage.png",
            "Line manager — line layers and segmentation settings.",
            "Gestor de linhas — camadas lineares e definições de segmentação.",
            width_cm=15.0))
    add(IMG("ui_analysis_setup.png",
            "Analysis setup — study areas and analysis groups.",
            "Configuração da análise — áreas de estudo e grupos.",
            width_cm=15.0))

    add(H(2, "4.3 Process (step 3)", "4.3 Processar (passo 3)"))
    add(IMG("overview_run_processing.png",
            "Process — unified processing runner.",
            "Processar — executor de processamento unificado.",
            width_cm=15.0))
    add(P(
        "The runner can execute, in order: Data processing (prep, intersect, flatten, backfill); Tiles processing; Lines processing; Analysis processing. Each materially-costly stage has its own checkbox in the run grid, so any one of them can be re-run without forcing the others to repeat. A Normal / Advanced toggle at the top keeps the default view simple. Worker counts for each stage are decided at runtime by the auto-tune step unless explicit values are set in config.ini.",
        "O executor pode executar, por ordem: Processamento de dados (prep, intersect, flatten, backfill); Processamento de azulejos; Processamento de linhas; Processamento de análise. Cada fase materialmente custosa tem a sua própria caixa na grelha de execução, pelo que qualquer uma pode ser executada novamente sem forçar a repetição das restantes. O alternador Normal / Avançado no topo mantém a vista simples. As contagens de processos de trabalho de cada fase são decididas em tempo de execução pelo passo de auto-ajuste, salvo se forem definidos valores explícitos em config.ini."))
    add(IMG("ui_processing_pipeline_run.png",
            "Processing pipeline runner — numbered stages and per-stage controls.",
            "Executor da pipeline de processamento — fases numeradas e controlos por fase.",
            width_cm=15.0,
            landscape=False))

    add(H(2, "4.4 Results (step 4)", "4.4 Resultados (passo 4)"))
    add(IMG("overview_review_and_report.png",
            "Results — asset map, results map, compare, report engine.",
            "Resultados — mapa de activos, mapa de resultados, comparação, motor de relatórios.",
            width_cm=15.0))
    add(UL(
        ["Asset map — inspect asset layers and supporting data.",
         "Results map — review processed outputs together with basemaps / background layers.",
         "Compare study areas — compare analysis groups side by side.",
         "Report engine — create a Word report from the latest results."],
        ["Mapa de activos — inspeccionar camadas de activos e dados de suporte.",
         "Mapa de resultados — rever resultados processados junto com mapas-base.",
         "Comparar áreas de estudo — comparar grupos de análise lado a lado.",
         "Motor de relatórios — criar um relatório Word a partir dos resultados mais recentes."]))
    add(IMG("ui_asset_map_view.png",
            "Asset map view.",
            "Vista do mapa de activos.",
            width_cm=15.0))
    add(IMG("ui_map_overview.png",
            "Results map view.",
            "Vista do mapa de resultados.",
            width_cm=15.0))
    add(IMG("ui_analysis_present.png",
            "Compare study areas.",
            "Comparar áreas de estudo.",
            width_cm=15.0))
    add(IMG("ui_report_generate.png",
            "Report engine.",
            "Motor de relatórios.",
            width_cm=15.0))

    add(H(2, "4.5 Tune processing popup", "4.5 Janela Afinação do processamento"))
    add(IMG("ui_tune_processing_popup.png",
            "Tune processing popup — evaluate, commit, restore.",
            "Janela Afinação do processamento — avaliar, confirmar, restaurar.",
            width_cm=15.0))
    add(P(
        "Evaluate compares current values with advised values based on detected CPU and RAM. Commit changes writes the advised values into config.ini. Restore previous tuning restores the last saved pre-tuning values.",
        "Avaliar compara valores actuais com valores aconselhados com base no CPU e na RAM detectados. Confirmar alterações escreve os valores aconselhados em config.ini. Restaurar afinação anterior repõe os valores guardados antes da última afinação."))
    add(PB())

    # =========================================================================
    # 5. Data and data quality
    # =========================================================================
    add(H(1, "5. Data and data quality", "5. Dados e qualidade dos dados"))
    add(P(
        "Most MESA results trace back to the quality of the inputs. Two layers with the same geographic extent and the same sensitivity values can produce very different mosaics depending on how their geometry was prepared and what attributes survive into MESA's processing tables.",
        "A maior parte dos resultados do MESA decorre da qualidade das entradas. Duas camadas com a mesma extensão geográfica e os mesmos valores de sensibilidade podem produzir mosaicos muito diferentes consoante a forma como a sua geometria foi preparada e os atributos que sobrevivem nas tabelas de processamento do MESA."))
    add(IMG("data_folders.png",
            "MESA input data and folder layout.",
            "Dados de entrada e estrutura de pastas do MESA.",
            width_cm=15.0))

    add(H(2, "5.1 Input data folders", "5.1 Pastas de dados de entrada"))
    add(TBL(
        ["Folder", "Contents", "Required"],
        [
            ["input/asset/", "Area-asset datasets, organised into one subfolder per asset group.", "Yes"],
            ["input/geocode/", "Geocode layers (analysis grids / polygons).", "Yes (or generate inside MESA)"],
            ["input/lines/", "Line datasets — pipelines, rivers, roads, transmission lines.", "Optional"],
            ["input/atlas/", "Atlas-tile polygons used by QGIS atlas layouts.", "Optional"],
            ["input/images/", "Project imagery referenced by reports.", "Optional"],
            ["input/evaluate_landuse/", "Land-use classification inputs (advanced).", "Optional"],
        ],
        ["Pasta", "Conteúdo", "Obrigatório"],
        [
            ["input/asset/", "Conjuntos de activos de área, organizados numa subpasta por grupo de activos.", "Sim"],
            ["input/geocode/", "Camadas de geocódigos (grelhas / polígonos de análise).", "Sim (ou gerar dentro do MESA)"],
            ["input/lines/", "Conjuntos lineares — oleodutos, rios, estradas, linhas eléctricas.", "Opcional"],
            ["input/atlas/", "Polígonos de azulejo de atlas usados em layouts QGIS.", "Opcional"],
            ["input/images/", "Imagens de projecto referenciadas nos relatórios.", "Opcional"],
            ["input/evaluate_landuse/", "Entradas de classificação de uso do solo (avançado).", "Opcional"],
        ]))

    add(H(2, "5.2 Supported file formats", "5.2 Formatos de ficheiro suportados"))
    add(P(
        "MESA reads inputs through GeoPandas and Fiona/PyOGRIO, so common vector formats work: GeoPackage (.gpkg), Shapefile (.shp + sidecars), GeoParquet (.parquet), and a few others. GeoPackage is the safest input format. Shapefile works but truncates: column names are silently shortened to 10 characters, string values longer than 254 characters are clipped, and field types may be downgraded.",
        "O MESA lê entradas através de GeoPandas e Fiona/PyOGRIO, suportando formatos vectoriais comuns: GeoPackage (.gpkg), Shapefile (.shp + ficheiros associados), GeoParquet (.parquet), entre outros. O GeoPackage é o formato de entrada mais seguro. O Shapefile funciona mas trunca: os nomes de coluna são silenciosamente reduzidos a 10 caracteres, os valores de cadeias de caracteres acima de 254 caracteres são cortados e alguns tipos de campos podem ser rebaixados."))

    add(H(2, "5.3 Coordinate systems and projections", "5.3 Sistemas de coordenadas e projecções"))
    add(UL(
        ["working_projection_epsg (default 4326) — the projection MESA uses for fast spatial work in lat/lon.",
         "area_projection_epsg (default 3035 for Europe) — equal-area projection used to compute areas in m². Pick one appropriate to your study region.",
         "output_projection_epsg (default 4326) — the projection MESA writes outputs in."],
        ["working_projection_epsg (omissão 4326) — projecção utilizada pelo MESA para trabalho espacial rápido em lat/lon.",
         "area_projection_epsg (omissão 3035 para a Europa) — projecção de igual área utilizada para calcular áreas em m². Escolha uma adequada à região de estudo.",
         "output_projection_epsg (omissão 4326) — projecção em que o MESA escreve os resultados."]))

    add(H(2, "5.4 Preparing your data", "5.4 Preparação dos dados"))
    add(P(
        "Ecological / organic asset polygons work best — wetland outlines, settlement footprints, river polygons, protected-area boundaries. The mosaic stays compact and the result map reads as discrete sensitivity zones. Gridded / raster-derived asset layers are the failure mode: per-cell visual moiré, and mosaic-face collapse if the layer has no honest classifier column. The takeaway: raster-derived asset layers without a classifying attribute are a known hazard. MESA handles them conservatively (no dissolve → moiré), but the right fix is upstream.",
        "Os polígonos de activos ecológicos / orgânicos funcionam melhor — contornos de zonas húmidas, áreas de povoamento, polígonos de rios, limites de áreas protegidas. O mosaico mantém-se compacto e o mapa final lê-se como zonas discretas de sensibilidade. As camadas de activos derivadas de grelha / raster são o modo de falha: efeito moiré célula a célula e colapso de faces do mosaico quando a camada não tem uma coluna classificadora honesta. Conclusão: camadas de activos derivadas de raster sem atributo classificador são um perigo conhecido. O MESA trata-as de forma conservadora (sem dissolução → moiré), mas a correcção correcta é a montante."))
    add(IMG("artefacts.png",
            "How gridded vector data creates map artefacts.",
            "Como os dados vectoriais em grelha geram artefactos cartográficos.",
            width_cm=15.0))

    add(H(2, "5.5 What MESA does at import", "5.5 O que o MESA faz na importação"))
    add(OL(
        ["Geometry validation — every geometry passes through make_valid() with a buffer(0) fallback.",
         "Smart-key dissolve — classifies attribute columns as uniform or diverging; merges touching polygons that share uniform-attribute values.",
         "Buffer points and lines (always on for non-polygon inputs) — Points and LineStrings are buffered using default_point_buffer_m and default_line_buffer_m before intersection."],
        ["Validação geométrica — toda a geometria passa por make_valid() com um buffer(0) como recurso.",
         "Dissolução por chave inteligente — classifica colunas de atributos como uniformes ou divergentes; combina polígonos contíguos que partilham valores uniformes.",
         "Amortecimento de pontos e linhas (sempre activo para entradas não-poligonais) — Pontos e LineStrings recebem um amortecimento usando default_point_buffer_m e default_line_buffer_m antes da intersecção."]))

    add(H(2, "5.6 Output data tables", "5.6 Tabelas de resultados"))
    add(TBL(
        ["Table", "Contents"],
        [
            ["tbl_asset_group", "One row per asset group (subfolder under input/asset/)."],
            ["tbl_asset_object", "One row per imported asset polygon, with attributes and serialised attributes string."],
            ["tbl_geocode_group", "One row per geocode group (basic_mosaic, H3 levels, ...)."],
            ["tbl_geocode_object", "One row per geocode cell, with parent group reference."],
            ["tbl_stacked", "Where assets touch geocodes — one row per (asset, geocode) pair."],
            ["tbl_flat", "Flattened analysis result — one row per cell with summary sensitivity (max, mean, count, codes, area_m²)."],
            ["tbl_system_capabilities", "Host fingerprint snapshot used by About and auto_tune."],
        ],
        ["Tabela", "Conteúdo"],
        [
            ["tbl_asset_group", "Uma linha por grupo de activos (subpasta em input/asset/)."],
            ["tbl_asset_object", "Uma linha por polígono de activo importado, com atributos e cadeia serializada de atributos."],
            ["tbl_geocode_group", "Uma linha por grupo de geocódigos (basic_mosaic, níveis H3, ...)."],
            ["tbl_geocode_object", "Uma linha por célula de geocódigo, com referência ao grupo-pai."],
            ["tbl_stacked", "Onde os activos tocam os geocódigos — uma linha por par (activo, geocódigo)."],
            ["tbl_flat", "Resultado de análise achatado — uma linha por célula com sensibilidade resumida (máx., média, contagem, códigos, área_m²)."],
            ["tbl_system_capabilities", "Fotografia da impressão digital do computador usada por Sobre e auto_tune."],
        ]))
    add(PB())

    # =========================================================================
    # 6. Method
    # =========================================================================
    add(H(1, "6. Method", "6. Método"))
    add(IMG("sensitivity_table.png",
            "MESA — Mapping Environmentally Sensitive Assets.",
            "MESA — Mapeamento de Activos Ambientalmente Sensíveis.",
            width_cm=12.0))
    add(P(
        "MESA (Mapping Environmentally Sensitive Assets) is a method intended for a broad audience of individuals working on the development and use of environmental sensitivity maps. The method provides a step-by-step protocol for developing an environmental sensitivity atlas based on a standardized methodology that was developed following a review of several existing methods.",
        "O MESA (Mapeamento de Activos Ambientalmente Sensíveis) é um método dirigido a um público alargado de pessoas que trabalham no desenvolvimento e utilização de mapas de sensibilidade ambiental. O método fornece um protocolo passo a passo para desenvolver um atlas de sensibilidade ambiental baseado numa metodologia padronizada, desenvolvida após a revisão de vários métodos existentes."))

    add(H(2, "6.1 Introduction", "6.1 Introdução"))
    add(P(
        "Environmental Sensitivity Atlases display the relative sensitivity of places to one or more clearly defined pressures. Historically, many sensitivity atlases were developed for coastal emergency response and acute pollution scenarios. In MESA, the same underlying idea is used more broadly for early-stage screening of an area of interest by relating it to mapped assets.",
        "Os Atlas de Sensibilidade Ambiental mostram a sensibilidade relativa dos lugares a uma ou mais pressões claramente definidas. Historicamente, muitos atlas foram desenvolvidos para resposta a emergências costeiras e cenários de poluição aguda. No MESA, a mesma ideia é utilizada mais amplamente para a triagem inicial de uma área de interesse, relacionando-a com os activos mapeados."))
    add(P(
        "In this broader planning context it is important to be explicit about the chain from activity → pressure → pathway → asset. Sensitivity is not an inherent property of a location: it is conditional on what pressure is being considered, how and when it is expressed, and which assets are valued.",
        "Neste contexto mais amplo de planeamento é importante ser explícito quanto à cadeia actividade → pressão → via → activo. A sensibilidade não é uma propriedade inerente a um local: depende da pressão considerada, da forma e momento em que é expressa, e dos activos valorizados."))
    add(IMG("process_overview.png",
            "Process overview — methodological steps.",
            "Visão geral do processo — passos metodológicos.",
            width_cm=15.0))

    add(H(2, "6.2 Importance, susceptibility, sensitivity", "6.2 Importância, susceptibilidade, sensibilidade"))
    add(IMG("calculation_mesa.png",
            "MESA calculation — importance × susceptibility = sensitivity.",
            "Cálculo MESA — importância × susceptibilidade = sensibilidade.",
            width_cm=12.0))
    add(UL(
        ["Asset importance — an asset's value based on its rarity, significance, functional and intrinsic value. Quantified with a value from 1 to 5.",
         "Asset susceptibility — the degree to which an asset will be affected by a pressure, based on the predicted severity of the impact and the asset's ability to recover. Quantified from 1 to 5.",
         "Sensitivity — the product of importance and susceptibility, also expressed on a 1–5 scale."],
        ["Importância do activo — o seu valor baseado na raridade, significado, valor funcional e intrínseco. Quantificada de 1 a 5.",
         "Susceptibilidade do activo — o grau em que um activo será afectado por uma pressão, com base na severidade prevista do impacte e na capacidade de recuperação. Quantificada de 1 a 5.",
         "Sensibilidade — o produto de importância e susceptibilidade, também expressa numa escala 1–5."]))

    add(H(2, "6.3 Intended use", "6.3 Utilização prevista"))
    add(P(
        "Knowing whether the atlas is intended to be used at a strategic (national planning), operational (project management) or tactical (emergency response) level helps identify the appropriate resolution and assets to include, as well as providing a filter for identifying relevant pressures.",
        "Saber se o atlas se destina a uso estratégico (planeamento nacional), operacional (gestão de projecto) ou táctico (resposta a emergências) ajuda a identificar a resolução e os activos adequados a incluir, e fornece um filtro para identificar as pressões relevantes."))

    add(H(2, "6.4 Pressures", "6.4 Pressões"))
    add(P(
        "The pressures considered will determine which assets are most sensitive. Sources of impact must be clearly identified by users prior to mapping. Table 6.1 lists potential pressures from oil & gas operations across terrestrial, coastal, and marine realms.",
        "As pressões consideradas determinam que activos são mais sensíveis. As fontes de impacte devem ser claramente identificadas antes do mapeamento. A Tabela 6.1 lista pressões potenciais das operações de petróleo e gás nos domínios terrestre, costeiro e marinho."))
    add(TBL(
        ["Terrestrial", "Coastal", "Marine"],
        [
            ["Oil Spill", "Oil Spill", "Oil Spill"],
            ["Habitat Loss", "Habitat Loss", "Habitat Loss"],
            ["Habitat Fragmentation", "Habitat Fragmentation", "Habitat Fragmentation"],
            ["Disturbance", "Disturbance", "Disturbance"],
            ["Atmospheric emissions", "Atmospheric emissions", "Atmospheric emissions"],
            ["Aquatic pollution (excl. oil)", "Aquatic pollution (excl. oil)", "Aquatic pollution (excl. oil)"],
            ["Soil erosion and degradation", "Soil erosion", "Increased biological resource use"],
            ["Airborne particulates (e.g. dust)", "Airborne particulates (e.g. dust)", ""],
            ["Increased biological resource use", "Increased biological resource use", ""],
        ],
        ["Terrestre", "Costeiro", "Marinho"],
        [
            ["Derrame de petróleo", "Derrame de petróleo", "Derrame de petróleo"],
            ["Perda de habitat", "Perda de habitat", "Perda de habitat"],
            ["Fragmentação de habitat", "Fragmentação de habitat", "Fragmentação de habitat"],
            ["Perturbação", "Perturbação", "Perturbação"],
            ["Emissões atmosféricas", "Emissões atmosféricas", "Emissões atmosféricas"],
            ["Poluição aquática (excl. petróleo)", "Poluição aquática (excl. petróleo)", "Poluição aquática (excl. petróleo)"],
            ["Erosão e degradação do solo", "Erosão do solo", "Maior uso de recursos biológicos"],
            ["Partículas em suspensão (p. ex. poeira)", "Partículas em suspensão (p. ex. poeira)", ""],
            ["Maior uso de recursos biológicos", "Maior uso de recursos biológicos", ""],
        ]))

    add(H(2, "6.5 Assets", "6.5 Activos"))
    add(P(
        "While identifying all features that meet the broad definition of ecological assets is ideal, there is a need to find a pragmatic starting point that captures key biodiversity values. Recommended ecological assets include areas with protected status, biodiversity designations (KBA, AZE), Critical Habitats per IFC PS6, areas supporting Threatened Species, habitats of high biodiversity (mangroves, forests, coral reefs), and other ecologically important assets.",
        "Embora seja ideal identificar todas as feições que cabem na definição ampla de activos ecológicos, é necessário um ponto de partida pragmático que capte os principais valores de biodiversidade. Os activos ecológicos recomendados incluem áreas com estatuto de protecção, designações de biodiversidade (KBA, AZE), Habitats Críticos segundo a IFC PS6, áreas que suportam Espécies Ameaçadas, habitats de elevada biodiversidade (mangais, florestas, recifes de coral) e outros activos ecologicamente importantes."))

    add(H(2, "6.6 Geocoding and segmentation", "6.6 Geocodificação e segmentação"))
    add(P(
        "Geocode layers define how MESA aggregates results spatially. Drop polygon datasets (Quarter Degree Grid Cells, municipal boundaries, H3 hexagons, or any custom tessellation) into input/geocode/ and reference them in config.ini. Linear features follow a similar pattern: place line datasets under input/lines/. MESA buffers and optionally segments these features and stores results in tbl_segment_flat.parquet.",
        "As camadas de geocódigos definem como o MESA agrega os resultados espacialmente. Coloque conjuntos poligonais (células de grelha de quarto de grau, limites municipais, hexágonos H3 ou qualquer outra tesselação) em input/geocode/ e referencie-os em config.ini. As feições lineares seguem um padrão semelhante: coloque conjuntos lineares em input/lines/. O MESA aplica amortecimento e, opcionalmente, segmenta estas feições, armazenando os resultados em tbl_segment_flat.parquet."))
    add(PB())

    # =========================================================================
    # 7. Indexes
    # =========================================================================
    add(H(1, "7. Indexes", "7. Índices"))
    add(P(
        "An index is a function that maps one or more measured or derived variables to a single summary value. Index construction typically involves selecting variables, normalising them, applying weights, and aggregating them. Indices reduce complexity but introduce modelling choices: weighting, scaling, and aggregation rules can substantially affect interpretation and ranking.",
        "Um índice é uma função que mapeia uma ou mais variáveis medidas ou derivadas para um único valor sumário. A construção de um índice envolve normalmente a selecção de variáveis, a sua normalização, a aplicação de pesos e a sua agregação. Os índices reduzem a complexidade mas introduzem escolhas de modelação: regras de ponderação, escala e agregação podem afectar substancialmente a interpretação e a ordenação."))
    add(P(
        "MESA produces a set of indices and indicators that summarise the intensity and composition of overlapping assets within each spatial reporting unit (a geocode polygon, and in some workflows along buffered/segmented lines). Three normalised indices serve as primary overlays — Importance index (index_importance), Sensitivity index (index_sensitivity), and OWA index (index_owa) for precautionary addition — alongside four supplementary per-cell indicators that provide raw maxima and counts: Sensitivity (sensitivity_max), Importance (importance_max), Asset groups (asset_groups_total), and Asset objects (assets_overlap_total). All seven are stored as fields in tbl_flat.parquet and rendered as overlays in the desktop viewer.",
        "O MESA produz um conjunto de índices e indicadores que resumem a intensidade e a composição de activos sobrepostos em cada unidade espacial (um polígono de geocódigo, e em alguns fluxos de trabalho ao longo de linhas amortecidas/segmentadas). Três índices normalizados funcionam como sobreposições principais — Índice de Importância (index_importance), Índice de Sensibilidade (index_sensitivity) e Índice OWA (index_owa) para adição precaucional — juntamente com quatro indicadores complementares por célula que fornecem máximos e contagens em estado bruto: Sensibilidade (sensitivity_max), Importância (importance_max), Grupos de activos (asset_groups_total) e Objectos de activo (assets_overlap_total). Todos os sete são guardados como campos em tbl_flat.parquet e renderizados como sobreposições no visualizador desktop."))
    add(P(
        "For map presentation, an index value of 0 indicates the absence of relevant overlaps and is not symbolised in the index overlays — those units therefore remain blank in map outputs.",
        "Para a apresentação cartográfica, um valor de índice igual a 0 indica a ausência de sobreposições relevantes e não é simbolizado nas sobreposições de índice — essas unidades aparecem em branco nos mapas."))

    add(H(2, "7.1 Importance index (index_importance)", "7.1 Índice de Importância (index_importance)"))
    add(P(
        "The importance index is an aggregate indicator of how important the overlapping assets are within a spatial unit. Each overlapping asset contributes an importance class (typically 1–5). For each spatial unit, MESA counts overlaps per importance class, applies a user-configurable weight per class, and normalises the raw score to 0–100 within the geocode group.",
        "O Índice de Importância é um indicador agregado de quão importantes são os activos sobrepostos numa unidade espacial. Cada activo sobreposto contribui com uma classe de importância (tipicamente 1–5). Para cada unidade, o MESA conta as sobreposições por classe, aplica um peso configurável por classe e normaliza a pontuação bruta para o intervalo 0–100 dentro do grupo de geocódigos."))

    add(H(2, "7.2 Sensitivity index (index_sensitivity)", "7.2 Índice de Sensibilidade (index_sensitivity)"))
    add(P(
        "The sensitivity index summarises the degree to which overlapping assets are sensitive to a specified pressure, taking both importance and susceptibility into account (commonly an importance × susceptibility product). Counts across product values are weighted, summed, and normalised to 0–100 within the geocode group, using the same max-based scaling as importance.",
        "O Índice de Sensibilidade resume o grau de sensibilidade dos activos sobrepostos a uma pressão especificada, tomando em conta tanto a importância como a susceptibilidade (tipicamente um produto importância × susceptibilidade). As contagens por valor de produto são ponderadas, somadas e normalizadas para o intervalo 0–100 dentro do grupo de geocódigos, usando a mesma escala baseada no máximo da importância."))

    add(H(2, "7.3 OWA index (index_owa)", "7.3 Índice OWA (index_owa)"))
    add(P(
        "OWA is used as a precautionary ranking device: locations that contain any very high sensitivity overlaps are prioritised ahead of locations that merely contain many moderate overlaps. This corresponds to a lexicographic comparison of the count-vector of sensitivity classes from highest to lowest. The OWA index is intentionally non-compensatory: a small number of high-sensitivity overlaps can dominate the ranking, so it should be paired with contextual reporting before management decisions.",
        "O OWA é utilizado como mecanismo de classificação precaucional: locais que contenham quaisquer sobreposições de sensibilidade muito alta são priorizados face a locais com muitas sobreposições moderadas. Isto corresponde a uma comparação lexicográfica do vector de contagens das classes de sensibilidade, da mais alta para a mais baixa. O Índice OWA é intencionalmente não-compensatório: um pequeno número de sobreposições de sensibilidade alta pode dominar a classificação, pelo que deve ser acompanhado de relatórios contextuais antes de tomar decisões de gestão."))

    add(H(2, "7.4 Sensitivity (sensitivity_max)", "7.4 Sensibilidade (sensitivity_max)"))
    add(P(
        "The single highest sensitivity value present in the overlap stack of a spatial unit. Sensitivity per overlap is the product of importance and susceptibility (typically 1–25); this layer reports the worst-case overlap. It complements the Sensitivity index — the max highlights the single highest class present, while the index highlights accumulated weighted overlap. A cell can show Sensitivity = 25 with a low Sensitivity index (one extreme overlap in an otherwise sparse cell), or Sensitivity = 6 with a high Sensitivity index (many mid-class overlaps stacked).",
        "O valor de sensibilidade mais alto presente na sobreposição de uma unidade espacial. A sensibilidade por sobreposição é o produto da importância e da susceptibilidade (tipicamente 1–25); esta camada reporta a pior sobreposição. Complementa o Índice de Sensibilidade — o máximo destaca a classe mais alta presente, enquanto o índice destaca a sobreposição acumulada e ponderada. Uma célula pode apresentar Sensibilidade = 25 com Índice de Sensibilidade baixo (uma sobreposição extrema numa célula de outra forma esparsa), ou Sensibilidade = 6 com Índice de Sensibilidade alto (muitas sobreposições de classe média empilhadas)."))

    add(H(2, "7.5 Importance (importance_max)", "7.5 Importância (importance_max)"))
    add(P(
        "The maximum importance class among the overlapping assets in a spatial unit. Importance classes are typically integers in 1..5. This layer answers: what is the highest-importance asset class present in this cell? Pair with the Importance index to distinguish a single standout asset (high max, low index) from broad accumulation (mid max, high index).",
        "A classe de importância máxima entre os activos sobrepostos numa unidade espacial. As classes de importância são tipicamente inteiros em 1..5. Esta camada responde: qual é a classe de importância mais alta presente nesta célula? Combine com o Índice de Importância para distinguir um único activo destacado (máximo alto, índice baixo) de uma acumulação mais ampla (máximo médio, índice alto)."))

    add(H(2, "7.6 Asset groups (asset_groups_total)", "7.6 Grupos de activos (asset_groups_total)"))
    add(P(
        "The number of distinct asset groups whose footprints overlap a spatial unit. An asset group is a logical category of features (e.g., wetlands, industrial sites); multiple objects from the same group count as one. This layer is a diversity indicator: high values mark cells where many kinds of features coincide (a multi-themed hotspot), independent of how many individual objects are involved. The names of the contributing groups are stored in the asset_group_names field.",
        "O número de grupos de activos distintos cujas pegadas se sobrepõem a uma unidade espacial. Um grupo de activos é uma categoria lógica de feições (por exemplo, zonas húmidas, locais industriais); múltiplos objectos do mesmo grupo contam como um. Esta camada é um indicador de diversidade: valores altos marcam células onde coincidem muitos tipos de feições (um ponto crítico multi-temático), independentemente do número de objectos individuais envolvidos. Os nomes dos grupos contribuintes são guardados no campo asset_group_names."))

    add(H(2, "7.7 Asset objects (assets_overlap_total)", "7.7 Objectos de activo (assets_overlap_total)"))
    add(P(
        "The total number of individual asset objects (polygons, lines, points) overlapping a spatial unit. Two objects from the same asset group both count. This layer is a density indicator: high values mark cells where many features are stacked, regardless of group diversity. Pair with Asset groups to distinguish 'lots of one thing' (1 group, 50 objects) from 'diverse mix, but thin' (5 groups, 5 objects).",
        "O número total de objectos de activo individuais (polígonos, linhas, pontos) que se sobrepõem a uma unidade espacial. Dois objectos do mesmo grupo de activos contam ambos. Esta camada é um indicador de densidade: valores altos marcam células onde muitas feições estão empilhadas, independentemente da diversidade de grupos. Combine com Grupos de activos para distinguir 'muito de uma coisa' (1 grupo, 50 objectos) de 'mistura diversa, mas fina' (5 grupos, 5 objectos)."))
    add(PB())

    # =========================================================================
    # 8. QGIS
    # =========================================================================
    add(H(1, "8. QGIS", "8. QGIS"))
    add(P(
        "MESA ships with a QGIS project template (qgis/mesa.qgz). The project loads GeoParquet tables from output/geoparquet/ (analysis outputs) and MBTiles from output/mbtiles/ (fast rendering layers). Paths are relative to the qgis/ folder, so keep the standard MESA folder structure.",
        "O MESA inclui um modelo de projecto QGIS (qgis/mesa.qgz). O projecto carrega tabelas GeoParquet de output/geoparquet/ (resultados de análise) e MBTiles de output/mbtiles/ (camadas de renderização rápida). Os caminhos são relativos à pasta qgis/, pelo que deve manter a estrutura padrão do MESA."))
    add(P(
        "Prerequisites: run MESA processing at least once so the output/ folders exist and contain data; install QGIS 3.36 or newer (the project is authored with QGIS 3.44). Open QGIS by double-clicking qgis/mesa.qgz.",
        "Pré-requisitos: executar o processamento MESA pelo menos uma vez para que as pastas output/ existam com dados; instalar o QGIS 3.36 ou superior (o projecto foi criado com QGIS 3.44). Abra o QGIS fazendo duplo clique em qgis/mesa.qgz."))

    add(H(2, "8.1 QGIS opened", "8.1 QGIS aberto"))
    add(P(
        "When QGIS opens, the project can optionally pan/zoom to your data. This depends on QGIS macros: if you see a prompt about macros, enable them to allow the startup logic to run. If you prefer not to enable macros, you can still use the project — just zoom to your layers manually. If layers appear broken or empty, confirm that output/geoparquet/ and output/mbtiles/ were created by a successful MESA run.",
        "Quando o QGIS abre, o projecto pode opcionalmente fazer pan/zoom até aos seus dados. Isto depende das macros do QGIS: se aparecer um aviso sobre macros, active-as para permitir a execução da lógica de arranque. Se preferir não activar as macros, pode continuar a usar o projecto — basta fazer zoom manualmente. Se as camadas aparecerem partidas ou vazias, confirme que output/geoparquet/ e output/mbtiles/ foram criadas por uma execução bem sucedida do MESA."))
    add(IMG("qgis_default.png",
            "Default QGIS project view.",
            "Vista padrão do projecto QGIS.",
            width_cm=15.0,
            landscape=False))

    add(H(2, "8.2 Segments", "8.2 Segmentos"))
    add(P(
        "The Segments group visualizes results for processed line assets (roads, rivers, coastlines, pipelines, etc.). These layers are backed by GeoParquet tables under output/geoparquet/.",
        "O grupo Segmentos visualiza os resultados de activos lineares processados (estradas, rios, linhas costeiras, oleodutos, etc.). Estas camadas são suportadas pelas tabelas GeoParquet em output/geoparquet/."))
    add(IMG("qgis_segments.png",
            "QGIS — segments group.",
            "QGIS — grupo de segmentos.",
            width_cm=15.0,
            landscape=False))

    add(H(2, "8.3 Importance", "8.3 Importância"))
    add(P(
        "The Importance layers show the calculated importance index and/or categorical outputs. They are typically driven by tbl_flat.parquet (polygons) and related tables.",
        "As camadas de Importância mostram o índice de importância calculado e/ou os resultados categóricos. São tipicamente impulsionadas por tbl_flat.parquet (polígonos) e tabelas relacionadas."))
    add(IMG("qgis_importance.png",
            "QGIS — importance overlay.",
            "QGIS — sobreposição de importância.",
            width_cm=15.0,
            landscape=False))

    add(H(2, "8.4 Assets", "8.4 Activos"))
    add(P(
        "The Assets group shows imported assets (often with multiple geometry types) loaded from the GeoParquet store. Use these layers to QA imports and to understand which datasets contribute to hotspots.",
        "O grupo de Activos mostra os activos importados (frequentemente com múltiplos tipos de geometria) carregados a partir do repositório GeoParquet. Utilize estas camadas para validar importações e para perceber quais os conjuntos que contribuem para os pontos críticos."))
    add(IMG("qgis_assets.png",
            "QGIS — assets group.",
            "QGIS — grupo de activos.",
            width_cm=15.0,
            landscape=False))
    add(PB())

    # =========================================================================
    # 9. Definitions
    # =========================================================================
    add(H(1, "9. Definitions", "9. Definições"))
    add(H(2, "9.1 Sensitivity mapping", "9.1 Mapeamento de sensibilidade"))
    add(P(
        "Sensitivity mapping is a structured, anticipatory method for integrating environmental, social, and cultural considerations into spatial decision-making. It identifies assets that society values — habitats, species, ecosystem functions, cultural heritage, livelihoods — and evaluates their sensitivity to specific pressures or interventions. Sensitivity mapping is explicitly designed for early-stage planning and screening: it does not replace detailed impact assessments, but supports strategic choices by making trade-offs visible.",
        "O mapeamento de sensibilidade é um método estruturado e antecipatório para integrar considerações ambientais, sociais e culturais na tomada de decisão espacial. Identifica os activos que a sociedade valoriza — habitats, espécies, funções de ecossistema, património cultural, meios de subsistência — e avalia a sua sensibilidade a pressões ou intervenções específicas. Está explicitamente desenhado para planeamento e triagem em fase inicial: não substitui as avaliações de impacte detalhadas, mas apoia escolhas estratégicas tornando visíveis as compensações."))

    add(H(2, "9.2 Geocodes", "9.2 Geocódigos"))
    add(P(
        "In MESA, geocodes are polygon layers used as the common reference units for aggregating results (indices, counts, summaries). The built-in geocode basic_mosaic is a non-overlapping polygon mosaic produced from the combined footprint of all imported asset objects. Other geocodes you can use include Quarter Degree Grid Cells (QDGC), administrative boundaries, H3 hexagons, and any custom polygon set intended as a set of analysis zones (ideally non-overlapping).",
        "No MESA, os geocódigos são camadas poligonais utilizadas como unidades de referência comum para agregar resultados (índices, contagens, resumos). O geocódigo nativo basic_mosaic é um mosaico de polígonos não sobrepostos produzido a partir da pegada combinada de todos os activos importados. Pode ainda usar outros geocódigos: Células de Grelha de Quarto de Grau (QDGC), limites administrativos, hexágonos H3 e qualquer conjunto de polígonos personalizado destinado a servir como zonas de análise (idealmente não sobrepostas)."))

    add(H(2, "9.3 Lines", "9.3 Linhas"))
    add(P(
        "Lines in MESA are used to analyze sensitivity along linear features (rivers, roads, coastlines, power lines, pipelines, etc.). To represent their footprint or risk area, MESA buffers the input lines before analysis. Buffer width is configurable in config.ini. Lines can also be split into segments so results vary along the feature; MESA stores per-segment metrics in output/geoparquet/tbl_segment_flat.parquet.",
        "As linhas no MESA são utilizadas para analisar a sensibilidade ao longo de feições lineares (rios, estradas, linhas costeiras, linhas eléctricas, oleodutos, etc.). Para representar a sua pegada ou área de risco, o MESA aplica um amortecimento às linhas antes da análise. A largura do amortecimento é configurável em config.ini. As linhas podem também ser divididas em segmentos para que os resultados variem ao longo da feição; as métricas por segmento são guardadas em output/geoparquet/tbl_segment_flat.parquet."))
    add(IMG("segments.png",
            "Line segments — sensitivity along linear features.",
            "Segmentos de linha — sensibilidade ao longo de feições lineares.",
            width_cm=14.0))

    add(H(2, "9.4 Atlas", "9.4 Atlas"))
    add(P(
        "The QGIS atlas functionality automates the creation of map series based on a given layout and data, ideal for presenting comprehensive spatial analyses. By integrating the QGIS atlas with MESA, users can efficiently produce a series of maps that align with the MESA method, facilitating better visualization and decision-making in conservation planning and resource management. ESRI's ArcGIS Pro offers comparable functionality through its Map Series tool.",
        "A funcionalidade de atlas do QGIS automatiza a criação de séries de mapas baseadas num determinado layout e em determinados dados, ideal para apresentar análises espaciais abrangentes. Ao integrar o atlas QGIS com o MESA, os utilizadores podem produzir eficientemente uma série de mapas alinhados com o método MESA, facilitando uma melhor visualização e tomada de decisão no planeamento de conservação e gestão de recursos. O ArcGIS Pro da ESRI oferece funcionalidade comparável através da ferramenta Map Series."))

    add(H(2, "9.5 Other definitions", "9.5 Outras definições"))
    add(UL(
        ["Environmental Sensitivity Atlas — a collection of maps and supporting narrative text presenting spatial data on the sensitivity of ecological and/or socio-economic assets to a specific pressure.",
         "Environmental Sensitivity — combination of susceptibility and importance of an affected asset that signifies the potential impact of a given pressure.",
         "Pressure (stressor) — the source of potential impact from an activity (habitat loss, disturbance, pollution, altered water levels, ...).",
         "Ecological Asset — naturally occurring entities that provide ecological functions or services, including those without monetary value.",
         "Socio-Economic Asset — natural or anthropogenic entities providing social, cultural, economic or political value.",
         "Importance — an asset's value at global, national, or local scale in relation to rarity, significance, functional and intrinsic value.",
         "Susceptibility — the degree to which an asset will be affected by a pressure, based on severity of impact and recovery ability.",
         "Sensitivity — overall rating of consequences of allowing an impact to occur. Combines importance and susceptibility for a given pressure.",
         "Importance index — normalised (1–100) score summarising weighted importance across overlapping assets.",
         "Sensitivity index — normalised (1–100) score combining importance and susceptibility per geocode group."],
        ["Atlas de Sensibilidade Ambiental — colectânea de mapas e texto narrativo de suporte que apresenta dados espaciais sobre a sensibilidade de activos ecológicos e/ou socioeconómicos a uma pressão específica.",
         "Sensibilidade Ambiental — combinação da susceptibilidade e da importância de um activo afectado, indicando o impacte potencial de uma dada pressão.",
         "Pressão (factor de stress) — fonte de impacte potencial de uma actividade (perda de habitat, perturbação, poluição, alteração de níveis de água, …).",
         "Activo Ecológico — entidades naturais que prestam funções ou serviços ecológicos, incluindo as que não têm valor monetário.",
         "Activo Socioeconómico — entidades naturais ou antropogénicas com valor social, cultural, económico ou político.",
         "Importância — valor de um activo à escala global, nacional ou local, em relação à raridade, ao significado, ao valor funcional e intrínseco.",
         "Susceptibilidade — grau em que um activo é afectado por uma pressão, com base na severidade do impacte e na capacidade de recuperação.",
         "Sensibilidade — classificação global das consequências de permitir um impacte. Combina importância e susceptibilidade para uma dada pressão.",
         "Índice de Importância — pontuação normalizada (1–100) que resume a importância ponderada nos activos sobrepostos.",
         "Índice de Sensibilidade — pontuação normalizada (1–100) que combina importância e susceptibilidade por grupo de geocódigos."]))
    add(PB())

    # =========================================================================
    # 10. Advanced
    # =========================================================================
    add(H(1, "10. Advanced", "10. Avançado"))
    add(H(2, "10.1 Workspace location overrides (MESA_BASE_DIR)",
          "10.1 Sobreposições da localização do espaço de trabalho (MESA_BASE_DIR)"))
    add(P(
        "By default, MESA assumes the workspace folders (input/, output/, qgis/, config.ini) live next to mesa.py (source) or next to mesa.exe (packaged). If you want to keep the application code in one folder but your data workspace somewhere else, set the MESA_BASE_DIR environment variable to the workspace root before launching. When set, MESA reads config.ini and writes output/ under that folder.",
        "Por omissão, o MESA assume que as pastas do espaço de trabalho (input/, output/, qgis/, config.ini) estão junto a mesa.py (código-fonte) ou a mesa.exe (empacotado). Se quiser manter o código da aplicação numa pasta e os dados noutra, defina a variável de ambiente MESA_BASE_DIR para a raiz do espaço de trabalho antes de iniciar. Quando definida, o MESA lê config.ini e escreve output/ dentro dessa pasta."))

    add(H(2, "10.2 Performance knobs (config.ini)", "10.2 Parâmetros de desempenho (config.ini)"))
    add(UL(
        ["chunk_size — processing chunk size for heavy spatial operations.",
         "max_workers, flatten_max_workers, flatten_small_max_workers, tiles_max_workers, backfill_max_workers — per-stage worker caps. Leave at 0 to let auto_tune decide; set a positive integer to override exactly.",
         "h3_max_cells — safeguard for extremely fine H3 resolutions.",
         "tiles_minzoom / tiles_maxzoom — output MBTiles zoom range.",
         "mem_target_frac / approx_gb_per_worker — memory targeting for the auto-tuner."],
        ["chunk_size — tamanho do bloco de processamento para operações espaciais pesadas.",
         "max_workers, flatten_max_workers, flatten_small_max_workers, tiles_max_workers, backfill_max_workers — limites de processos de trabalho por fase. Deixe a 0 para que o auto_tune decida; defina um inteiro positivo para forçar.",
         "h3_max_cells — protecção para resoluções H3 extremamente finas.",
         "tiles_minzoom / tiles_maxzoom — intervalo de zoom dos MBTiles produzidos.",
         "mem_target_frac / approx_gb_per_worker — alvo de memória para o auto-ajuste."]))

    add(H(2, "10.3 Memory safety nets", "10.3 Redes de segurança de memória"))
    add(UL(
        ["Pre-flight check before flatten — evaluates two signals and aborts only when both fail: system-wide RAM use above flatten_preflight_max_vm_percent AND available RAM below flatten_preflight_avail_safety_factor × estimated peak (estimated from on-disk tbl_stacked size and the auto-tuned worker counts). Pagefile/swap above flatten_preflight_max_swap_gb is still its own hard signal. Small datasets on busy desktops now go through; big datasets on starved hosts still abort cleanly.",
         "Per-pool panic watchdog (mem_panic_percent / mem_panic_grace_secs) — if RAM crosses the threshold past the grace period, the pool is terminated cleanly.",
         "Process-lifetime sentinel (mem_lifetime_panic_percent / mem_lifetime_panic_grace_secs) — last-resort backstop covering parent-side work between pools."],
        ["Verificação prévia antes do flatten — avalia dois sinais e aborta apenas quando ambos falham: utilização global de RAM acima de flatten_preflight_max_vm_percent E memória disponível abaixo de flatten_preflight_avail_safety_factor × pico estimado (calculado a partir do tamanho em disco de tbl_stacked e do número de trabalhadores auto-ajustado). O swap acima de flatten_preflight_max_swap_gb continua a ser um sinal forte por si só. Conjuntos de dados pequenos em computadores ocupados passam agora a verificação; conjuntos grandes em computadores com pouca memória continuam a abortar limpamente.",
         "Vigilante de pânico por pool (mem_panic_percent / mem_panic_grace_secs) — se a RAM ultrapassar o limite além do período de tolerância, o pool é encerrado limpamente.",
         "Sentinela do tempo de vida do processo (mem_lifetime_panic_percent / mem_lifetime_panic_grace_secs) — rede de segurança final que cobre o trabalho do processo-pai entre pools."]))

    add(H(2, "10.4 Mosaic union batching", "10.4 Lotes de união do mosaico"))
    add(UL(
        ["mosaic_coverage_union_batch — bigger batches (~2000–4000 on 64+ GB hosts) produce fewer, larger intermediate unions.",
         "mosaic_line_union_max_partials — higher cap (32–64) means fewer reduction rounds.",
         "mosaic_coverage_union = false — escape hatch that skips the coverage-reduction stage entirely and falls back to STRtree filtering."],
        ["mosaic_coverage_union_batch — lotes maiores (~2000–4000 em máquinas com 64+ GB) produzem uniões intermédias menores em número e maiores em tamanho.",
         "mosaic_line_union_max_partials — limite mais alto (32–64) reduz o número de rondas de redução.",
         "mosaic_coverage_union = false — saída de emergência que ignora completamente a fase de redução de cobertura e usa filtragem STRtree."]))

    add(H(2, "10.5 Basemap and tile caching", "10.5 Mapas-base e cache de azulejos"))
    add(UL(
        ["xyz — use the built-in downloader with persistent cache under output/tile_cache/.",
         "contextily — use contextily if installed.",
         "auto — prefer contextily, fall back to xyz."],
        ["xyz — usa o transferidor incorporado com cache persistente em output/tile_cache/.",
         "contextily — usa contextily se estiver instalado.",
         "auto — prefere contextily, com xyz como recurso."]))
    add(P(
        "First-run cost: the first report run on a new AOI (or after the cache has expired) must download every basemap tile the report touches from tile.openstreetmap.org. With each map covering several hundred tiles across multiple zoom levels, this can add minutes to tens of minutes to the run depending on report scope, network speed, and OSM tile-server response. Subsequent runs that touch the same area read from output/tile_cache/ and complete much faster. Tiles are cached for 30 days (TILE_CACHE_MAX_AGE_DAYS in code/report_generate.py); older tiles are re-fetched the next time they are needed. The cache is keyed by zoom/x/y, so entries are reusable across atlases and reports that overlap geographically.",
        "Custo da primeira execução: a primeira geração de relatório numa nova AOI (ou após a cache ter expirado) tem de descarregar todos os azulejos do mapa-base que o relatório precisa, a partir de tile.openstreetmap.org. Como cada mapa cobre várias centenas de azulejos em vários níveis de zoom, isto pode acrescentar minutos a dezenas de minutos ao tempo de execução, conforme o âmbito do relatório, a velocidade da rede e a resposta do servidor de azulejos do OSM. Execuções subsequentes que toquem a mesma área leem de output/tile_cache/ e são muito mais rápidas. Os azulejos são guardados durante 30 dias (TILE_CACHE_MAX_AGE_DAYS em code/report_generate.py); azulejos mais antigos são novamente transferidos quando voltarem a ser necessários. A cache é indexada por zoom/x/y, pelo que as entradas são reutilizáveis entre atlas e relatórios que cobrem áreas sobrepostas."))
    add(UL(
        ["Run the report once on a representative AOI before a workshop or live demo so the cache is warm.",
         "The cache is portable: copying output/tile_cache/ between machines or projects covering the same area skips the download phase entirely.",
         "If you regenerate reports frequently and tile-server response is slow, prefer a single wide first run that covers the eventual report extent rather than running narrow then expanding.",
         "Be respectful of the OpenStreetMap tile-server usage policy (https://operations.osmfoundation.org/policies/tiles/); the local cache is partly there for that reason."],
        ["Execute o relatório uma vez numa AOI representativa antes de um workshop ou demonstração ao vivo para que a cache esteja preparada.",
         "A cache é portátil: copiar output/tile_cache/ entre máquinas ou projectos que cobrem a mesma área evita totalmente a fase de descarga.",
         "Se gerar relatórios com frequência e a resposta do servidor de azulejos for lenta, prefira uma primeira execução ampla que cubra a extensão final do relatório, em vez de começar com uma área pequena e ir alargando.",
         "Respeite a política de utilização do servidor de azulejos do OpenStreetMap (https://operations.osmfoundation.org/policies/tiles/); a cache local existe parcialmente por essa razão."]))
    add(PB())

    # =========================================================================
    # 11. Troubleshooting
    # =========================================================================
    add(H(1, "11. Troubleshooting", "11. Resolução de problemas"))
    add(H(2, "11.1 What 'slow' is normal", "11.1 O que conta como lentidão normal"))
    add(P(
        "MESA works with large geospatial datasets and uses a Python stack (GeoPandas/Shapely/PyProj/Arrow) that can take time to import and warm up. Startup and the first opening of heavy tools may feel slow (tens of seconds). Once loaded, the UI and map interaction are usually responsive and stable. Prefer an SSD and keep the MESA workspace on a local disk; avoid network drives or highly synced folders such as OneDrive if performance is poor.",
        "O MESA trabalha com grandes conjuntos de dados geoespaciais e utiliza uma pilha Python (GeoPandas/Shapely/PyProj/Arrow) que pode demorar a carregar. O arranque e a primeira abertura de ferramentas pesadas podem parecer lentos (dezenas de segundos). Após o carregamento, a interface e a interacção com mapas costumam ser ágeis e estáveis. Prefira um SSD e mantenha o espaço de trabalho do MESA num disco local; evite unidades de rede ou pastas muito sincronizadas como OneDrive se o desempenho for fraco."))

    add(H(2, "11.2 Maps: window does not open / opens blank",
          "11.2 Mapas: a janela não abre ou abre em branco"))
    add(P(
        "MESA's built-in map viewers use pywebview with WebView2 (Edge Chromium). WebView2 is not Google Chrome and does not depend on Chrome. Check Windows Settings → Apps → Installed apps for Microsoft Edge WebView2 Runtime. If missing, run Windows Update fully and reboot, or install the Evergreen WebView2 Runtime (or via 'winget install Microsoft.WebView2Runtime').",
        "Os visualizadores de mapas embutidos do MESA usam pywebview com WebView2 (Edge Chromium). O WebView2 não é o Google Chrome e não depende dele. Verifique em Definições do Windows → Aplicações → Aplicações instaladas se Microsoft Edge WebView2 Runtime está presente. Se faltar, execute totalmente o Windows Update e reinicie, ou instale o WebView2 Runtime Evergreen (ou via 'winget install Microsoft.WebView2Runtime')."))
    add(P(
        "If the map opens but shows no layers, the analysis map renders overlays from MBTiles under output/mbtiles/. If that folder is empty, run Workflows → Process with Area enabled and verify .mbtiles files exist.",
        "Se o mapa abrir mas não mostrar camadas, o mapa de análise usa MBTiles em output/mbtiles/. Se essa pasta estiver vazia, execute Fluxos de trabalho → Processar com Área activada e confirme que existem ficheiros .mbtiles."))
    add(P(
        "Offline behaviour (MESA 5.0.2 and later): map windows ship with the Leaflet library bundled locally, so they render even on hosts with no internet access or with corporate proxies that block public CDNs. When the WebView reports the host is offline, a thin orange banner appears at the top of the map with the text 'Offline – showing cached map tiles only. Pan to areas you have viewed online before.' OpenStreetMap basemap tiles continue to be served from MESA's on-disk tile cache (output/cache/osm_tiles), so areas previously viewed online remain available. Areas never visited online will appear as a grey grid until network access returns; the editor controls themselves work either way.",
        "Comportamento sem ligação (MESA 5.0.2 e posterior): as janelas de mapa incluem a biblioteca Leaflet localmente, pelo que funcionam mesmo em computadores sem acesso à Internet ou com proxies empresariais que bloqueiem CDN públicos. Quando o WebView indica que o computador está offline, surge uma faixa laranja fina no topo do mapa com o texto 'Offline – showing cached map tiles only. Pan to areas you have viewed online before.' Os azulejos do mapa-base OpenStreetMap continuam a ser servidos a partir da cache local em output/cache/osm_tiles, pelo que áreas previamente vistas com ligação permanecem disponíveis. Áreas nunca visitadas online aparecerão como uma grelha cinzenta até voltar a haver rede; os controlos de edição em si funcionam em ambos os casos."))

    add(H(2, "11.3 'Nothing happens' when clicking a map button",
          "11.3 'Nada acontece' ao clicar num botão de mapa"))
    add(P(
        "In the packaged app, some helper tools are launched in-process and others are started as separate processes from tools/. Open PowerShell in the MESA folder and run the standalone helper directly to see errors: tools\\map_overview.exe, tools\\asset_map_view.exe, tools\\analysis_setup.exe, tools\\line_manage.exe, tools\\tiles_create_raster.exe.",
        "Na aplicação empacotada, alguns auxiliares são lançados no mesmo processo e outros como processos separados a partir de tools/. Abra a PowerShell na pasta do MESA e execute o auxiliar autónomo directamente para ver os erros: tools\\map_overview.exe, tools\\asset_map_view.exe, tools\\analysis_setup.exe, tools\\line_manage.exe, tools\\tiles_create_raster.exe."))

    add(H(2, "11.4 Common Windows setup pitfalls",
          "11.4 Armadilhas comuns na instalação Windows"))
    add(UL(
        ["Running from a protected folder (e.g. C:\\Program Files) can prevent MESA from writing output/ and logs. Prefer a user-writable workspace folder.",
         "Unzipping 'partial' releases — keep the shipped folder structure intact (input/, output/, tools/, system_resources/, qgis/, config.ini).",
         "On fresh Windows installs, maps may fail until all Windows updates and runtimes (including WebView2) are in place."],
        ["Executar a partir de uma pasta protegida (p. ex. C:\\Program Files) pode impedir o MESA de escrever em output/ e nos registos. Prefira uma pasta de espaço de trabalho com permissão de escrita do utilizador.",
         "Descompactar versões 'parciais' — mantenha a estrutura de pastas entregue (input/, output/, tools/, system_resources/, qgis/, config.ini).",
         "Em instalações Windows novas, os mapas podem falhar até que todas as actualizações Windows e runtimes (incluindo WebView2) estejam instalados."]))

    add(H(2, "11.5 Processing aborts at flatten with PRE-FLIGHT ABORT",
          "11.5 O processamento aborta no flatten com PRE-FLIGHT ABORT"))
    add(P(
        "The flatten stage refuses to start if the host already looks too memory-pressured to finish safely. As of MESA 5.0.2 the check evaluates two signals together — system-wide RAM use AND a dataset-aware headroom estimate — and aborts only when both fail. The swap-residue check is still a separate hard signal. The log line names the gate that tripped it, for example 'vm.percent 78.4% > 60% AND avail 6.2 GB < 1.25 × est_peak 8.0 GB = 10.0 GB' or 'swap_used 7.3 GB > 5.0 GB'. Close other heavy applications, wait a minute for swap to drain, restart the MESA launcher, and rerun Process. If pre-flight keeps tripping on a host with plenty of free RAM, raise flatten_preflight_avail_safety_factor (default 1.25) or flatten_preflight_max_vm_percent (default 60) in config.ini; the abort message itself names both knobs.",
        "A fase flatten recusa-se a iniciar se o computador já estiver sob demasiada pressão de memória para terminar com segurança. A partir do MESA 5.0.2, a verificação avalia dois sinais em conjunto — utilização global de RAM E uma estimativa de margem que considera o tamanho do conjunto de dados — e só aborta quando ambos falham. O resíduo de swap continua a ser um sinal independente. A linha de registo identifica o critério que disparou, por exemplo 'vm.percent 78.4% > 60% AND avail 6.2 GB < 1.25 × est_peak 8.0 GB = 10.0 GB' ou 'swap_used 7.3 GB > 5.0 GB'. Feche outras aplicações pesadas, aguarde um minuto para o swap baixar, reinicie o iniciador MESA e volte a executar Processar. Se o pre-flight continuar a disparar num computador com bastante RAM livre, aumente flatten_preflight_avail_safety_factor (predefinição 1.25) ou flatten_preflight_max_vm_percent (predefinição 60) em config.ini; a própria mensagem de aborto refere os dois parâmetros."))

    add(H(2, "11.6 'Reducing coverage' looks stuck",
          "11.6 'Reducing coverage' parece encravado"))
    add(P(
        "Basic_mosaic's coverage-reduction stage runs single-threaded GEOS unions. On larger projects this can take hours, and the heartbeat goes silent during a single very-large union call. Confirm the process is alive in Task Manager (CPU at ~85% on a single core means it is grinding, not deadlocked). For long-term improvement, run Tune processing → Evaluate → Commit changes — advised values for mosaic_coverage_union_batch and mosaic_line_union_max_partials scale with the host's RAM tier and dramatically reduce reduction rounds. For an extreme dataset, set mosaic_coverage_union = false in config.ini to skip the coverage-reduction stage entirely.",
        "A fase de redução de cobertura do basic_mosaic executa uniões GEOS num único thread. Em projectos maiores pode demorar horas, e o heartbeat fica em silêncio durante uma única chamada de união muito grande. Confirme no Gestor de Tarefas que o processo está vivo (CPU em ~85% num único núcleo significa que está a processar, não bloqueado). Para melhorias de longo prazo, execute Afinação do processamento → Avaliar → Confirmar alterações — os valores aconselhados para mosaic_coverage_union_batch e mosaic_line_union_max_partials escalam com o nível de RAM do computador e reduzem drasticamente o número de rondas. Para um conjunto de dados extremo, defina mosaic_coverage_union = false em config.ini para ignorar completamente a fase de redução de cobertura."))

    add(H(2, "11.7 Still stuck?", "11.7 Ainda com problemas?"))
    add(P(
        "When reporting a problem, include: Windows version (Home/Pro + build); whether you run packaged or from source; whether Microsoft Edge WebView2 Runtime is installed; whether output/mbtiles contains .mbtiles files; the newest lines from log.txt in the MESA folder.",
        "Ao reportar um problema, inclua: versão do Windows (Home/Pro + build); se executa empacotado ou a partir do código-fonte; se o Microsoft Edge WebView2 Runtime está instalado; se output/mbtiles contém ficheiros .mbtiles; as linhas mais recentes de log.txt na pasta MESA."))

    return blocks


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_cover(doc, lang):
    if lang == "en":
        title = "MESA User Guide"
        subtitle = "Mapping Environmentally Sensitive Assets"
        version_line = f"{VERSION} — desktop tool"
        author_label = "Author"
        affil_label = "Affiliation"
        date_label = "Issued"
        today = date.today()
        date_text = today.strftime("%B %Y")
    else:
        title = "Manual do Utilizador MESA"
        subtitle = "Mapeamento de Activos Ambientalmente Sensíveis"
        version_line = f"{VERSION} — ferramenta de ambiente de trabalho"
        author_label = "Autor"
        affil_label = "Afiliação"
        date_label = "Edição"
        today = date.today()
        date_text = f"{PT_MONTHS[today.month - 1]} de {today.year}"

    # Cover page logo
    logo = WIKI_IMG / "mesa_logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo), width=Cm(8))

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(version_line)
    run.font.size = Pt(14)

    # Spacer paragraphs
    for _ in range(6):
        doc.add_paragraph()

    # Author block
    for label, value in [
        (author_label, AUTHOR),
        (affil_label, AFFILIATION),
        (date_label, date_text),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(13)
        run.bold = True
        run = p.add_run(value)
        run.font.size = Pt(13)

    # Page break to first content page
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_toc(doc, lang):
    title = "Contents" if lang == "en" else "Índice"
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(fld_char_end)

    placeholder = doc.add_paragraph()
    placeholder.add_run(
        "(Right-click → Update Field in Microsoft Word to populate the table of contents.)"
        if lang == "en" else
        "(Clique com o botão direito → Actualizar Campo no Microsoft Word para preencher o índice.)"
    ).italic = True

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_orientation(section, landscape):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def render(blocks, lang, out_path):
    doc = Document()

    # Default styles
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    _add_cover(doc, lang)
    _add_toc(doc, lang)

    current_landscape = False

    for block in blocks:
        kind = block[0]

        if kind == "H":
            _, level, content = block
            doc.add_heading(content[lang], level=level)

        elif kind == "P":
            _, content = block
            doc.add_paragraph(content[lang])

        elif kind == "UL":
            _, content = block
            for item in content[lang]:
                doc.add_paragraph(item, style="List Bullet")

        elif kind == "OL":
            _, content = block
            for item in content[lang]:
                doc.add_paragraph(item, style="List Number")

        elif kind == "IMG":
            _, filename, caption, width_cm, want_landscape = block
            img_path = WIKI_IMG / filename
            if want_landscape != current_landscape:
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                _set_orientation(new_section, want_landscape)
                current_landscape = want_landscape
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(str(img_path), width=Cm(width_cm))
                except Exception as exc:
                    p.add_run(f"[image load failed: {filename}: {exc}]")
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(caption[lang])
                cap_run.italic = True
                cap_run.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                p.add_run(f"[Figure missing: {filename}]").italic = True

        elif kind == "TBL":
            _, content = block
            headers, rows = content[lang]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(header)
                run.bold = True
                _set_cell_shading(cell, "1F3A5F")
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for row_idx, row in enumerate(rows, start=1):
                for col_idx, value in enumerate(row):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = value
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            doc.add_paragraph()

        elif kind == "PB":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

    # Restore portrait at the very end if we ended on a landscape figure
    if current_landscape:
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_orientation(new_section, False)

    doc.save(str(out_path))
    print(f"  wrote {out_path}")


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    blocks = build_blocks()
    print("Building MESA User Guide...")
    render(blocks, "en", DOCS / "MESA_User_Guide_en.docx")
    render(blocks, "pt", DOCS / "MESA_User_Guide_pt.docx")
    print("Done.")


if __name__ == "__main__":
    main()
