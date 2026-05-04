import locale
try:
    locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')
except Exception:
    # Fall back silently if locale isn't available on this system
    pass

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QGroupBox, QLabel, QPushButton, QPlainTextEdit, QCheckBox, QRadioButton,
    QComboBox, QProgressBar, QFrame, QSizePolicy, QButtonGroup, QMessageBox,
)
from PySide6.QtGui import QIcon, QFont
from PySide6.QtCore import Qt, QTimer, Signal, QObject
from asset_manage import apply_shared_stylesheet

import geopandas as gpd
import pandas as pd
import configparser
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
from matplotlib.path import Path as MplPath
try:
    from mpl_toolkits.axes_grid1.inset_locator import inset_axes
except Exception:
    inset_axes = None
import numpy as np
import argparse
import os
import sys
try:
    import contextily as ctx  # optional; heavy deps (rasterio) may be absent in EXE
except Exception:
    ctx = None
from matplotlib.patches import Rectangle, PathPatch
from matplotlib.ticker import MaxNLocator

from PIL import Image as PILImage
import re
import threading
from pathlib import Path
import subprocess
import io, math, time, urllib.request
import sqlite3
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ModuleNotFoundError as exc:
    msg = (
        "Report engine dependency missing: python-docx\n\n"
        "Install it in the same environment that runs MESA, for example:\n"
        "  pip install python-docx\n\n"
        "If you use the bundled MESA requirements, add/install it from:\n"
        "  requirements_compile_win311.txt (or requirements_all_win311.txt)\n\n"
        f"Original error: {exc}"
    )
    print(msg, file=sys.stderr, flush=True)
    try:
        _app = QApplication.instance() or QApplication([])
        QMessageBox.critical(None, "MESA report engine", msg)
    except Exception:
        pass
    raise

# ---------------- UI / sizing constants ----------------
MAX_MAP_PX_HEIGHT = 2000           # hard cap for saved map PNG height (px)
RIBBON_CM_HEIGHT   = 0.6           # ribbon display height inside PDF (cm)
ATLAS_FIGURE_INCHES = (7.2, 7.2)   # atlas tiles render smaller to fit more per page
ATLAS_DOC_WIDTH_SCALE = 0.75       # reduce displayed width of atlas images inside documents

TILE_CACHE_MAX_AGE_DAYS = 30       # discard cached OSM tiles older than this (<=0 keeps forever)

OSM_ATTRIBUTION_TEXT = "© OpenStreetMap contributors"

_OSM_TILE_CACHE_LOGGED = False

# Basemap mode (set from config):
# - xyz: force built-in XYZ downloader with output/tile_cache
# - contextily: force contextily when available (fallback to xyz if not)
# - auto: prefer contextily, fallback to xyz
_REPORT_BASEMAP_MODE = "xyz"

# Inset styling (for overview/context insets on maps). Can be overridden via config.ini.
_REPORT_INSET_BORDER_COLOR = "#1f1f1f"
_REPORT_INSET_BORDER_LW = 1.2
_REPORT_INSET_SHADOW_ALPHA = 0.18
_REPORT_INSET_SHADOW_DX = 0.006
_REPORT_INSET_SHADOW_DY = 0.008

# ---------------- GUI / globals ----------------
log_widget = None        # QPlainTextEdit (set when GUI is up)
progress_var = None      # QProgressBar (set when GUI is up)
last_report_path = None
link_var = None          # QLabel used as hyperlink

_gui_window = None       # ReportGeneratorWindow instance (global ref for thread-safe access)


class _ReportSignals(QObject):
    """Thread-safe signals for updating the GUI from background threads."""
    log_message = Signal(str)
    progress_update = Signal(float)
    link_update = Signal(str)


_signals: _ReportSignals | None = None

# When running long jobs in a background thread, set_progress() messages were
# only written to the GUI widget (not to log.txt). Keep a best-effort base_dir
# so progress is also persisted to disk.
_LOG_BASE_DIR: str | None = None

SENSITIVITY_ORDER = ['A', 'B', 'C', 'D', 'E']
SENSITIVITY_UNKNOWN_COLOR = "#FF00F2"
_SENSITIVITY_NUMERIC_RANGES: list[tuple[str, float, float]] = []


# ---------------- Analysis presentation helpers (graphs per analysis area) ----------------
_ANALYSIS_GROUP_TABLE = "tbl_analysis_group.parquet"
_ANALYSIS_POLYGON_TABLE = "tbl_analysis_polygons.parquet"
_ANALYSIS_FLAT_TABLE = "tbl_analysis_flat.parquet"


def _analysis_read_table(gpq_dir: str, filename: str, *, geo: bool = False):
    path = os.path.join(gpq_dir, filename)
    if not os.path.exists(path):
        return None
    try:
        if geo:
            return gpd.read_parquet(path)
        return pd.read_parquet(path)
    except Exception:
        try:
            gdf = gpd.read_parquet(path)
            return pd.DataFrame(gdf.drop(columns=[c for c in ("geometry",) if c in gdf.columns]))
        except Exception:
            return None


def _analysis_group_choices(gpq_dir: str) -> list[tuple[str, str]]:
    df = _analysis_read_table(gpq_dir, _ANALYSIS_GROUP_TABLE, geo=False)
    if df is None or df.empty or "id" not in df.columns:
        return []
    choices: list[tuple[str, str]] = []
    for _, row in df.iterrows():
        gid = str(row.get("id", "") or "").strip()
        if not gid:
            continue
        name = str(row.get("name", "") or "").strip()
        display = name if name else gid
        choices.append((gid, f"{display} ({gid})"))
    return choices


def _analysis_group_title(gpq_dir: str, group_id: str) -> str:
    df = _analysis_read_table(gpq_dir, _ANALYSIS_GROUP_TABLE, geo=False)
    if df is None or df.empty:
        return str(group_id)
    try:
        match = df[df["id"].astype(str) == str(group_id)]
        if match.empty:
            return str(group_id)
        name = str(match.iloc[0].get("name", "") or "").strip()
        return name or str(group_id)
    except Exception:
        return str(group_id)


def _analysis_latest_run_filter(df: pd.DataFrame, group_id: str) -> pd.DataFrame:
    """Filter an analysis table to the latest run for the given group.

    tbl_analysis_flat.parquet can contain multiple analysis runs; without filtering,
    maps/statistics may mix runs.
    """
    if df is None or getattr(df, "empty", True):
        return df
    if "analysis_group_id" not in df.columns:
        return df
    try:
        mask = df["analysis_group_id"].astype(str) == str(group_id)
        sub = df.loc[mask].copy()
    except Exception:
        return df
    if sub.empty:
        return sub

    # Prefer run_id when present; otherwise use timestamp.
    if "analysis_run_id" in sub.columns:
        try:
            # Pick the most recent timestamp within each run_id, then choose the newest run.
            if "analysis_timestamp" in sub.columns:
                ts = pd.to_datetime(sub["analysis_timestamp"], errors="coerce")
                sub["__ts"] = ts
                run_ts = sub.groupby("analysis_run_id", dropna=False)["__ts"].max()
                chosen = run_ts.sort_values().index[-1]
                sub = sub[sub["analysis_run_id"] == chosen].copy()
                sub.drop(columns=["__ts"], inplace=True, errors="ignore")
                return sub
            # No timestamp: fall back to last run_id in sorted order
            chosen = sorted(sub["analysis_run_id"].astype(str).unique())[-1]
            return sub[sub["analysis_run_id"].astype(str) == chosen].copy()
        except Exception:
            sub.drop(columns=["__ts"], inplace=True, errors="ignore")
            return sub

    if "analysis_timestamp" in sub.columns:
        try:
            ts = pd.to_datetime(sub["analysis_timestamp"], errors="coerce")
            if ts.notna().any():
                latest = ts.max()
                return sub.loc[ts == latest].copy()
        except Exception:
            pass
    return sub


def _analysis_polygons_for_group(gpq_dir: str, group_id: str) -> gpd.GeoDataFrame:
    # tbl_analysis_flat.parquet contains *many* geometries per polygon (typically per geocode cell)
    # and is not suitable for rendering analysis area boundaries. We instead use it only to
    # determine which analysis_polygon_id are part of the selected group/run, then fetch the
    # actual polygon boundaries from tbl_analysis_polygons.parquet.
    polygon_ids: list[str] = []
    gdf_flat = _analysis_read_table(gpq_dir, _ANALYSIS_FLAT_TABLE, geo=True)
    if gdf_flat is not None and not getattr(gdf_flat, "empty", True):
        try:
            flat_df = pd.DataFrame(gdf_flat)
            flat_df = _analysis_latest_run_filter(flat_df, group_id)
            if "analysis_geocode" in flat_df.columns:
                # Match the rest of the report (basic_mosaic totals)
                mask = flat_df["analysis_geocode"].astype(str).str.lower() == "basic_mosaic"
                flat_df = flat_df.loc[mask].copy()
            if "analysis_polygon_id" in flat_df.columns:
                polygon_ids = [
                    str(v).strip()
                    for v in flat_df["analysis_polygon_id"].dropna().astype(str).unique().tolist()
                    if str(v).strip()
                ]
        except Exception:
            polygon_ids = []

    gdf = _analysis_read_table(gpq_dir, _ANALYSIS_POLYGON_TABLE, geo=True)
    if gdf is None or getattr(gdf, "empty", True):
        return gpd.GeoDataFrame()

    # First try: explicit polygon ids from the flat table.
    if polygon_ids and "id" in gdf.columns:
        try:
            sub = gdf[gdf["id"].astype(str).isin(set(polygon_ids))].copy()
            if not sub.empty:
                return gpd.GeoDataFrame(sub)
        except Exception:
            pass

    # Fallback: use group_id link in the polygons table.
    if "group_id" not in gdf.columns:
        return gpd.GeoDataFrame(gdf)
    try:
        subset = gdf[gdf["group_id"].astype(str) == str(group_id)].copy()
        return gpd.GeoDataFrame(subset)
    except Exception:
        return gpd.GeoDataFrame()


def _analysis_flat_for_group(gpq_dir: str, group_id: str) -> pd.DataFrame:
    gdf = _analysis_read_table(gpq_dir, _ANALYSIS_FLAT_TABLE, geo=True)
    if gdf is None or getattr(gdf, "empty", True):
        return pd.DataFrame()
    df = pd.DataFrame(gdf.drop(columns=[c for c in ("geometry",) if c in gdf.columns]))
    if "analysis_group_id" not in df.columns:
        return pd.DataFrame()
    try:
        subset = _analysis_latest_run_filter(df, group_id)
    except Exception:
        return pd.DataFrame()
    if "analysis_geocode" in subset.columns:
        mask = subset["analysis_geocode"].astype(str).str.lower() == "basic_mosaic"
        subset = subset.loc[mask]
    return subset


def _analysis_flat_geo_for_group(gpq_dir: str, group_id: str) -> gpd.GeoDataFrame:
    """GeoDataFrame subset of tbl_analysis_flat for mapping.

    Note: tbl_analysis_flat geometry is typically per geocode-cell polygon (many per analysis polygon).
    This is ideal for visualising sensitivity overlays within the analysis area boundary.
    """
    gdf = _analysis_read_table(gpq_dir, _ANALYSIS_FLAT_TABLE, geo=True)
    if gdf is None or getattr(gdf, "empty", True) or "geometry" not in gdf.columns:
        return gpd.GeoDataFrame()
    try:
        df = pd.DataFrame(gdf)
        df = _analysis_latest_run_filter(df, group_id)
        if "analysis_geocode" in df.columns:
            mask = df["analysis_geocode"].astype(str).str.lower() == "basic_mosaic"
            df = df.loc[mask].copy()
        out = gpd.GeoDataFrame(df, geometry="geometry", crs=getattr(gdf, "crs", None))
        out = out.dropna(subset=["geometry"]).copy()
        return out
    except Exception:
        return gpd.GeoDataFrame()


def _analysis_sensitivity_totals_km2(flat_df: pd.DataFrame) -> pd.DataFrame:
    if flat_df is None or flat_df.empty:
        return pd.DataFrame(columns=["Code", "Description", "Area (km²)"])
    code_col = "sensitivity_code_max" if "sensitivity_code_max" in flat_df.columns else "sensitivity_code"
    desc_col = "sensitivity_description_max" if "sensitivity_description_max" in flat_df.columns else (
        "sensitivity_description" if "sensitivity_description" in flat_df.columns else None
    )
    area_col = "analysis_area_m2" if "analysis_area_m2" in flat_df.columns else None
    if code_col not in flat_df.columns or area_col is None:
        return pd.DataFrame(columns=["Code", "Description", "Area (km²)"])

    tmp = flat_df.copy()
    tmp[area_col] = pd.to_numeric(tmp[area_col], errors="coerce").fillna(0.0)
    tmp[code_col] = tmp[code_col].astype(str).str.strip().str.upper().replace({"NAN": "", "NONE": ""})

    grouped = tmp.groupby(code_col, dropna=False)[area_col].sum()
    rows: list[dict[str, object]] = []
    for code, area_m2 in grouped.items():
        code_u = str(code or "").strip().upper()
        if not code_u:
            continue
        if float(area_m2 or 0.0) <= 0:
            continue
        desc = ""
        if desc_col and desc_col in tmp.columns:
            try:
                desc_vals = tmp.loc[tmp[code_col].astype(str).str.upper() == code_u, desc_col].dropna()
                if not desc_vals.empty:
                    desc = str(desc_vals.iloc[-1]).strip()
            except Exception:
                desc = ""
        rows.append({
            "Code": code_u,
            "Description": desc,
            "Area (km²)": float(area_m2) / 1_000_000.0,
        })
    out = pd.DataFrame(rows, columns=["Code", "Description", "Area (km²)"])
    if out.empty:
        return out
    # Sort codes consistently (A–E), then any other codes, then UNKNOWN/blank.
    order_map = {c: i for i, c in enumerate(SENSITIVITY_ORDER)}

    def _code_key(val: str) -> tuple[int, str]:
        s = str(val or "").strip().upper()
        if s in order_map:
            return (0, f"{order_map[s]:02d}")
        if s in ("", "NONE", "NAN"):
            return (3, "")
        if s == "UNKNOWN":
            return (2, s)
        return (1, s)

    out["__code_key"] = out["Code"].map(_code_key)
    out.sort_values(["__code_key"], inplace=True, ignore_index=True)
    out.drop(columns=["__code_key"], inplace=True)
    return out


def _analysis_code_color(code: str, palette_A2E: dict) -> str:
    fallback = None
    try:
        fallback = (palette_A2E or {}).get('UNKNOWN')
    except Exception:
        fallback = None
    if not fallback:
        fallback = SENSITIVITY_UNKNOWN_COLOR if 'SENSITIVITY_UNKNOWN_COLOR' in globals() else "#6c757d"

    try:
        entry = (palette_A2E or {}).get(str(code).strip().upper())
        if isinstance(entry, str) and entry.strip():
            return entry.strip()
        if isinstance(entry, dict) and entry.get("color"):
            return str(entry["color"]).strip() or fallback
    except Exception:
        pass
    return fallback


def _analysis_write_area_map_png(
    polygons_gdf: gpd.GeoDataFrame,
    out_path: str,
    *,
    base_dir: str | None,
    config_path: str | None,
    palette_A2E: dict,
    desc_A2E: dict,
    overlay_alpha: float = 0.65,
    flat_cells_gdf: gpd.GeoDataFrame | None = None,
) -> bool:
    """Render analysis area map with basemap + sensitivity overlay (no MBTiles).

    - Border is drawn from tbl_analysis_polygons.parquet (true analysis polygons).
    - Sensitivity overlay is derived from tbl_analysis_flat.parquet (geocode-cell polygons).
    """
    try:
        if polygons_gdf is None or polygons_gdf.empty or "geometry" not in polygons_gdf.columns:
            return False
        gdf = polygons_gdf.dropna(subset=["geometry"]).copy()
        if gdf.empty:
            return False

        g3857 = _safe_to_3857(gdf)
        if g3857.empty:
            return False

        bounds_3857 = _expand_bounds(g3857.total_bounds, pad_ratio=0.08)

        # Union geometry for clipping/masking
        clip_geom = None
        try:
            if hasattr(g3857, "union_all"):
                clip_geom = g3857.union_all()
            else:
                clip_geom = g3857.unary_union
        except Exception:
            clip_geom = None

        fig_h_in = 4.2
        fig_w_in = 5.8
        dpi = 180
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        minx, miny, maxx, maxy = bounds_3857
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)
        try:
            ax.set_aspect('equal', adjustable='box')
        except Exception:
            pass

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        # Sensitivity overlay from flat geocode-cell polygons (dissolved by sensitivity_code_max)
        try:
            if flat_cells_gdf is not None and not getattr(flat_cells_gdf, 'empty', True) and 'geometry' in flat_cells_gdf.columns:
                cells = flat_cells_gdf.dropna(subset=['geometry']).copy()
                if not cells.empty:
                    cells_3857 = _safe_to_3857(cells)
                    if not cells_3857.empty:
                        code_col = 'sensitivity_code_max' if 'sensitivity_code_max' in cells_3857.columns else (
                            'sensitivity_code' if 'sensitivity_code' in cells_3857.columns else None
                        )
                        if code_col:
                            cells_3857['__sens_code'] = (
                                cells_3857[code_col].astype(str).str.strip().str.upper().replace({'NAN': '', 'NONE': ''})
                            )
                            cells_3857 = cells_3857[cells_3857['__sens_code'] != ''].copy()
                            if not cells_3857.empty:
                                # Dissolve to reduce draw cost and make output stable.
                                dissolved = cells_3857.dissolve(by='__sens_code', as_index=False)
                                dissolved['__order'] = dissolved['__sens_code'].map({c: i for i, c in enumerate(SENSITIVITY_ORDER)}).fillna(99)
                                dissolved = dissolved.sort_values(['__order', '__sens_code'])
                                for _, row in dissolved.iterrows():
                                    code = str(row.get('__sens_code', '')).strip().upper()
                                    geom = row.get('geometry')
                                    if not code or geom is None:
                                        continue
                                    col = _analysis_code_color(code, palette_A2E)
                                    try:
                                        gpd.GeoSeries([geom], crs=dissolved.crs).plot(
                                            ax=ax,
                                            color=col,
                                            alpha=float(max(0.05, min(0.95, overlay_alpha))),
                                            edgecolor='none',
                                            linewidth=0.0,
                                            zorder=10,
                                        )
                                    except Exception:
                                        pass
        except Exception:
            pass

        # Border outline from analysis polygons (single union outline)
        try:
            gborder = g3857.copy()
            union_geom = None
            try:
                if hasattr(gborder, "union_all"):
                    union_geom = gborder.union_all()
                else:
                    union_geom = gborder.unary_union
            except Exception:
                union_geom = None
            if union_geom is not None and not getattr(union_geom, 'is_empty', True):
                gpd.GeoSeries([union_geom], crs=gborder.crs).boundary.plot(
                    ax=ax,
                    color="#2f3e46",
                    linewidth=1.3,
                    alpha=0.9,
                    zorder=30,
                )
        except Exception:
            pass

        _add_map_decorations(ax, bounds_3857, base_dir=base_dir, add_inset=False)
        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        return False


def _geom_to_mpl_path(geom) -> MplPath | None:
    """Convert a (Multi)Polygon-like geometry to a Matplotlib Path."""
    try:
        if geom is None:
            return None
        if getattr(geom, "is_empty", True):
            return None
        gtype = getattr(geom, "geom_type", "")
        if gtype == "Polygon":
            polygons = [geom]
        elif gtype == "MultiPolygon":
            polygons = list(getattr(geom, "geoms", []) or [])
        else:
            return None

        vertices: list[tuple[float, float]] = []
        codes: list[int] = []

        def _add_ring(coords):
            pts = list(coords)
            if len(pts) < 3:
                return
            # Ensure closed ring
            if pts[0] != pts[-1]:
                pts.append(pts[0])
            vertices.extend([(float(x), float(y)) for x, y in pts])
            codes.extend([MplPath.MOVETO] + [MplPath.LINETO] * (len(pts) - 2) + [MplPath.CLOSEPOLY])

        for poly in polygons:
            try:
                _add_ring(poly.exterior.coords)
                for interior in list(getattr(poly, "interiors", []) or []):
                    _add_ring(interior.coords)
            except Exception:
                continue

        if not vertices:
            return None
        return MplPath(vertices, codes)
    except Exception:
        return None


def _analysis_write_sensitivity_bar_png(totals_df: pd.DataFrame, palette_A2E: dict, out_path: str, title: str) -> bool:
    try:
        if totals_df is None or totals_df.empty:
            return False
        raw = totals_df[["Code", "Area (km²)"]].copy()
        raw["Code"] = raw["Code"].astype(str).str.strip().str.upper()

        order_map = {c: i for i, c in enumerate(SENSITIVITY_ORDER)}

        def _code_key(val: str) -> tuple[int, str]:
            s = str(val or "").strip().upper()
            if s in order_map:
                return (0, f"{order_map[s]:02d}")
            if s in ("", "NONE", "NAN"):
                return (3, "")
            if s == "UNKNOWN":
                return (2, s)
            return (1, s)

        raw["__code_key"] = raw["Code"].map(_code_key)
        raw.sort_values(["__code_key"], inplace=True)

        codes = raw["Code"].astype(str).tolist()
        areas = raw["Area (km²)"] .astype(float).tolist()
        colors = [_analysis_code_color(c, palette_A2E) for c in codes]

        fig, ax = plt.subplots(figsize=(5.8, 2.8), dpi=170)
        ax.bar(codes, areas, color=colors, edgecolor="#333333", linewidth=0.4)
        ax.set_title(title, fontsize=11)
        ax.set_ylabel("Area (km²)")
        ax.yaxis.set_major_locator(MaxNLocator(nbins=5))
        ax.grid(axis="y", alpha=0.25)
        fig.tight_layout(pad=0.6)
        fig.savefig(out_path, bbox_inches="tight")
        plt.close(fig)
        return True
    except Exception:
        return False


def _analysis_write_relative_share_png(
    left_totals_df: pd.DataFrame,
    right_totals_df: pd.DataFrame,
    palette_A2E: dict,
    out_path: str,
    left_label: str,
    right_label: str,
) -> bool:
    """Render relative (percentage) composition charts for two areas."""
    try:
        def _to_map(df: pd.DataFrame) -> dict[str, float]:
            if df is None or df.empty:
                return {}
            m: dict[str, float] = {}
            for _, r in df.iterrows():
                code = str(r.get('Code', '')).strip().upper()
                if not code or code in {'NONE', 'NAN'}:
                    continue
                try:
                    km2 = float(r.get('Area (km²)', 0.0) or 0.0)
                except Exception:
                    km2 = 0.0
                if km2 > 0:
                    m[code] = m.get(code, 0.0) + km2
            return m

        left_map = _to_map(left_totals_df)
        right_map = _to_map(right_totals_df)
        codes = [c for c in SENSITIVITY_ORDER if left_map.get(c, 0) > 0 or right_map.get(c, 0) > 0]
        # Add any extra codes deterministically
        extras = sorted({*left_map.keys(), *right_map.keys()} - set(SENSITIVITY_ORDER))
        codes = codes + extras
        if not codes:
            return False

        left_total = sum(left_map.get(c, 0.0) for c in codes)
        right_total = sum(right_map.get(c, 0.0) for c in codes)
        if left_total <= 0 and right_total <= 0:
            return False

        def _shares(area_map: dict[str, float], total: float) -> list[float]:
            if total <= 0:
                return [0.0 for _ in codes]
            return [max(area_map.get(c, 0.0), 0.0) / total for c in codes]

        left_sh = _shares(left_map, left_total)
        right_sh = _shares(right_map, right_total)

        fig, ax = plt.subplots(figsize=(6.4, 2.2), dpi=170)
        ax.set_axis_off()

        y_positions = [1.0, 0.35]
        labels = [left_label, right_label]
        bar_h = 0.18
        for y, shares, label in zip(y_positions, [left_sh, right_sh], labels):
            x = 0.0
            for code, share in zip(codes, shares):
                if share <= 0:
                    continue
                col = _analysis_code_color(code, palette_A2E)
                ax.add_patch(Rectangle((x, y), share, bar_h, facecolor=col, edgecolor='white', lw=0.6, alpha=0.95))
                if share >= 0.08:
                    ax.text(x + share / 2, y + bar_h / 2, f"{code} {share*100:.0f}%", ha='center', va='center', fontsize=8, color='white')
                x += share
            ax.text(-0.02, y + bar_h / 2, label, ha='right', va='center', fontsize=9)

        ax.set_xlim(-0.12, 1.02)
        ax.set_ylim(0, 1.35)
        ax.text(0.0, 1.28, "Relative sensitivity composition (100%)", fontsize=10, fontweight='bold')
        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        try:
            plt.close('all')
        except Exception:
            pass
        return False


def _analysis_write_sankey_difference_png(
    left_totals_df: pd.DataFrame,
    right_totals_df: pd.DataFrame,
    palette_A2E: dict,
    out_path: str,
    left_label: str,
    right_label: str,
    ribbon_alpha: float = 0.45,
) -> bool:
    """Sankey-style comparison: same-code ribbons between left/right blocks (like analysis_present.py)."""
    try:
        def _to_map(df: pd.DataFrame) -> dict[str, float]:
            if df is None or df.empty:
                return {}
            m: dict[str, float] = {}
            for _, r in df.iterrows():
                code = str(r.get('Code', '')).strip().upper()
                if not code or code in {'NONE', 'NAN'}:
                    continue
                try:
                    km2 = float(r.get('Area (km²)', 0.0) or 0.0)
                except Exception:
                    km2 = 0.0
                if km2 > 0:
                    m[code] = m.get(code, 0.0) + km2
            return m

        left_map = _to_map(left_totals_df)
        right_map = _to_map(right_totals_df)
        codes_in_data = [c for c in SENSITIVITY_ORDER if left_map.get(c, 0) > 0 or right_map.get(c, 0) > 0]
        extras = sorted({*left_map.keys(), *right_map.keys()} - set(SENSITIVITY_ORDER))
        codes = codes_in_data + extras
        if not codes:
            return False

        left_total = sum(left_map.get(c, 0.0) for c in codes)
        right_total = sum(right_map.get(c, 0.0) for c in codes)
        scale_total = max(left_total, right_total)
        if scale_total <= 0:
            return False

        def _build_blocks(area_map: dict[str, float]) -> dict[str, dict[str, float]]:
            # Build stacked blocks top-down so sensitivity ordering reads A (top) -> E (bottom).
            values = [max(area_map.get(code, 0.0), 0.0) for code in codes]
            min_height = 0.010
            gap = 0.012
            heights = [max((v / scale_total), min_height) if v > 0 else 0.0 for v in values]

            non_zero = [h for h in heights if h > 0]
            span = sum(non_zero) + gap * (len(non_zero) - 1 if len(non_zero) > 1 else 0)

            y_bottom = 0.06
            y_top = 0.88
            available = max(0.05, y_top - y_bottom)
            scale = 1.0 if span <= available else available / span

            blocks: dict[str, dict[str, float]] = {}
            cursor_top = y_top
            for code, height, value in zip(codes, heights, values):
                if height <= 0:
                    continue
                h = height * scale
                blocks[code] = {'start': cursor_top - h, 'end': cursor_top, 'value': value, 'height': h}
                cursor_top -= h + gap * scale
            return blocks

        left_blocks = _build_blocks(left_map)
        right_blocks = _build_blocks(right_map)

        # NOTE: In the PDF renderer we never scale images *up* (only down).
        # So to make this diagram larger on the page we must render a reasonably
        # large PNG and with a taller aspect ratio (more height relative to width).
        fig, ax = plt.subplots(figsize=(10.0, 4.0), dpi=170)
        ax.axis('off')

        # Push the two stacks closer to the edges to avoid large left/right buffers in the exported PNG.
        x_left, x_right = 0.12, 0.88
        bar_w = 0.14

        def _draw_blocks(blocks: dict[str, dict[str, float]], x: float, align: str):
            for code, info in blocks.items():
                col = _analysis_code_color(code, palette_A2E)
                rect = PathPatch(
                    MplPath(
                        [
                            (x - bar_w / 2, info['start']),
                            (x + bar_w / 2, info['start']),
                            (x + bar_w / 2, info['end']),
                            (x - bar_w / 2, info['end']),
                            (x - bar_w / 2, info['start']),
                        ],
                        [MplPath.MOVETO, MplPath.LINETO, MplPath.LINETO, MplPath.LINETO, MplPath.CLOSEPOLY],
                    ),
                    facecolor=col,
                    edgecolor='#333333',
                    lw=0.5,
                    alpha=0.90,
                )
                ax.add_patch(rect)
                label_y = (info['start'] + info['end']) / 2
                # Keep labels inside the bars; labels outside the axes make bbox_inches='tight' add margins.
                if align == 'left':
                    label_x = x - bar_w / 2 + 0.012
                    ha = 'left'
                else:
                    label_x = x + bar_w / 2 - 0.012
                    ha = 'right'
                ax.text(label_x, label_y, f"{code}", fontsize=9, ha=ha, va='center', color='black')

        def _draw_ribbon(lb: dict[str, float], rb: dict[str, float], color: str, delta_km2: float):
            verts = [
                (x_left + bar_w / 2, lb['start']),
                (x_right - bar_w / 2, rb['start']),
                (x_right - bar_w / 2, rb['end']),
                (x_left + bar_w / 2, lb['end']),
                (x_left + bar_w / 2, lb['start']),
            ]
            patch = PathPatch(
                MplPath(verts, [MplPath.MOVETO, MplPath.LINETO, MplPath.LINETO, MplPath.LINETO, MplPath.CLOSEPOLY]),
                facecolor=color,
                alpha=float(max(0.05, min(0.90, ribbon_alpha))),
                edgecolor='none',
            )
            ax.add_patch(patch)
            mid_x = (x_left + x_right) / 2
            mid_y = (lb['start'] + lb['end'] + rb['start'] + rb['end']) / 4
            # Keep labels compact; show 0.00 for tiny deltas
            try:
                dv = float(delta_km2)
            except Exception:
                dv = 0.0
            if not math.isfinite(dv) or abs(dv) < 0.005:
                txt = '0.00 km²'
            else:
                txt = f"{dv:+.2f} km²"
            ax.text(mid_x, mid_y, txt, fontsize=8, ha='center', va='center', color='black')

        _draw_blocks(left_blocks, x_left, 'left')
        _draw_blocks(right_blocks, x_right, 'right')

        for code in codes:
            lb = left_blocks.get(code)
            rb = right_blocks.get(code)
            if not lb or not rb:
                continue
            col = _analysis_code_color(code, palette_A2E)
            delta = left_map.get(code, 0.0) - right_map.get(code, 0.0)
            _draw_ribbon(lb, rb, col, delta)

        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.text(x_left, 0.96, left_label, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(x_right, 0.96, right_label, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(0.5, 0.99, "Area difference Sankey", ha='center', va='top', fontsize=10)

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        try:
            plt.close('all')
        except Exception:
            pass
        return False

class ReportEngine:
    """
    Helper responsible for rendering all cartographic artefacts and tracking temporary files.
    """

    def __init__(self,
                 base_dir: str,
                 tmp_dir: str,
                 palette: dict,
                 desc: dict,
                 config_path: str):
        self.base_dir = base_dir
        self.tmp_dir = Path(tmp_dir)
        self.tmp_dir.mkdir(parents=True, exist_ok=True)
        self.palette = palette
        self.desc = desc
        self.config_path = config_path
        self.generated_paths: list[Path] = []

    def register(self, path: str | Path):
        self.generated_paths.append(Path(path))

    def make_path(self, *parts: str, suffix: str = ".png") -> str:
        name = "_".join(parts) + suffix
        full = self.tmp_dir / name
        self.register(full)
        return str(full)

    def cleanup(self):
        for path in self.generated_paths:
            try:
                if path.exists():
                    path.unlink()
            except Exception:
                pass
        self.generated_paths.clear()

    def render_geocode_maps(self,
                            flat_df: gpd.GeoDataFrame,
                            set_progress_callback) -> tuple[list, list, list]:
        pages: list = []
        intro_table = None
        groups = []
        if flat_df.empty or 'name_gis_geocodegroup' not in flat_df.columns:
            return pages, intro_table, groups

        # Report maps: only basic_mosaic (or configured basic_group_name)
        try:
            cfg_local = read_config(self.config_path)
            basic_name = (cfg_local['DEFAULT'].get('basic_group_name', 'basic_mosaic') or 'basic_mosaic').strip()
        except Exception:
            basic_name = 'basic_mosaic'

        all_groups = (flat_df['name_gis_geocodegroup']
                      .astype('string')
                      .dropna().unique().tolist())
        all_groups = sorted(all_groups)

        # Prefer configured basic group if present, else fall back to a literal basic_mosaic, else first group.
        chosen = None
        for cand in (basic_name, 'basic_mosaic'):
            if any(str(g).strip().lower() == str(cand).strip().lower() for g in all_groups):
                chosen = cand
                break
        if chosen is None and all_groups:
            chosen = str(all_groups[0])
        groups = [chosen] if chosen else []

        counts = (flat_df.groupby('name_gis_geocodegroup')
                  .agg(total=('name_gis_geocodegroup', 'size'),
                       populated=('geometry', lambda s: int(s.notna().sum())))
                  .reset_index())
        # Only show the chosen (basic) group in the intro table.
        try:
            if chosen:
                mask = counts['name_gis_geocodegroup'].astype('string').str.strip().str.lower() == str(chosen).strip().lower()
                counts = counts[mask].copy()
        except Exception:
            pass

        intro_table = [["Geocode category", "Total objects", "With geometry"]]
        for _, row in counts.sort_values('name_gis_geocodegroup').iterrows():
            intro_table.append([
                str(row['name_gis_geocodegroup']),
                int(row['total']),
                int(row['populated']),
            ])

        fixed_bounds_3857 = compute_fixed_bounds_3857(flat_df, base_dir=self.base_dir)

        done = 0
        for gname in groups:
            # Use MBTiles rasters when available (preferred)
            safe = _safe_name(gname)
            mb_root = os.path.join(self.base_dir, "output", "mbtiles")

            # Other maps (non-index): these should come before the index sections in the report.
            # Each tuple: (mbtiles suffix, heading title, output slug, per-map intro text).
            # The intro text explains what each per-cell aggregation means and how to read it,
            # since the maps share a basemap and look superficially similar otherwise.
            layers = [
                (
                    "sensitivity_max", "Sensitive areas (A–E)", "sensitivity",
                    "For each geocode cell, this map shows the <b>highest sensitivity class</b> (A–E, derived from "
                    "the <code>sensitivity_max</code> field) found among all assets overlapping the cell. It is a "
                    "worst-case view: a single A-rated asset is enough to colour the cell as A, even if many "
                    "lower-rated assets also overlap. Pair this with the <b>Sensitivity index</b> later in the "
                    "report — the max highlights the single highest class present, while the index reflects "
                    "accumulated weighted overlap. The polygon-fallback variant (used when raster tiles are "
                    "missing) shows the same data drawn as analysis polygons rather than rasterised tiles.",
                ),
                (
                    "importance_max", "Importance (max)", "importance_max",
                    "Per-cell <b>highest importance class</b> (typically 1–5, from <code>importance_max</code>) "
                    "among overlapping assets. Use this to spot where the most important features are, regardless "
                    "of how many lower-importance assets share the cell. Pair with the <b>Importance index</b> "
                    "later in the report to distinguish a single standout asset (high max, low index) from broad "
                    "accumulation (mid max, high index).",
                ),
                (
                    "groupstotal", "# asset groups", "groupstotal",
                    "Per-cell <b>count of distinct asset groups</b> overlapping the cell (from "
                    "<code>asset_groups_total</code>). This is a <b>diversity</b> indicator: high values indicate "
                    "cells where many different kinds of features coincide (a multi-themed hotspot), independent "
                    "of how many individual asset objects are involved. The names of the contributing groups are "
                    "kept in <code>asset_group_names</code>.",
                ),
                (
                    "assetstotal", "# asset objects", "assetstotal",
                    "Per-cell <b>count of asset objects</b> overlapping the cell (from "
                    "<code>assets_overlap_total</code>). This is a <b>density</b> indicator: high values mean many "
                    "individual features are stacked here, regardless of group diversity. Pair with "
                    "<b>&ldquo;# asset groups&rdquo;</b> to distinguish &ldquo;lots of one thing&rdquo; (1 group, 50 objects) from "
                    "&ldquo;diverse mix, but thin&rdquo; (5 groups, 5 objects).",
                ),
            ]

            pages.append(('heading(2)', f"Other maps: {gname}"))
            pages.append(('text',
                "Per-cell summary maps for the <b>basic_mosaic</b> grouping, rendered from MBTiles with a "
                "basemap underlay. Each map shows a different aggregation of the assets that overlap each "
                "geocode cell — the <b>worst-case</b> sensitivity class, the highest <b>importance</b> "
                "value, and counts of overlapping <b>asset groups</b> and <b>asset objects</b>. They are "
                "complementary views of the same input data; read them together rather than in isolation."))
            pages.append(('rule', None))

            for suffix, title, slug, intro in layers:
                mb_path = os.path.join(mb_root, f"{safe}_{suffix}.mbtiles")
                out_png = self.make_path("geocode", safe, slug)
                ok = False
                note = ""
                if os.path.exists(mb_path):
                    ok, note = render_mbtiles_to_png_best_fit(mb_path, out_png, base_dir=self.base_dir)
                else:
                    note = f"Missing MBTiles: {os.path.basename(mb_path)}"

                if ok and _file_ok(out_png):
                    pages.append(('heading(3)', title))
                    pages.append(('text', intro))
                    pages.append(('image_map', (title, out_png)))
                    pages.append(('new_page', None))
                else:
                    # Fallback (only for sensitivity): render polygons (no borders) if MBTiles missing
                    if suffix in ("sensitivity_max", "sensitivity"):
                        sub = flat_df[flat_df['name_gis_geocodegroup'].astype('string').str.strip().str.lower() == str(gname).strip().lower()].copy()
                        ok_poly = draw_group_map_sensitivity(sub, gname, self.palette, self.desc, out_png, fixed_bounds_3857, base_dir=self.base_dir)
                        if ok_poly and _file_ok(out_png):
                            pages.append(('heading(3)', title))
                            pages.append(('text',
                                f"{intro} <i>Raster tiles for this layer were unavailable, so the same "
                                "data has been redrawn from the analysis polygons.</i>"))
                            pages.append(('image_map', (title, out_png)))
                            pages.append(('new_page', None))
                            continue
                    pages.append(('heading(3)', title))
                    pages.append(('text', f"Could not render map: {note or 'unknown error'}"))
                    pages.append(('new_page', None))

            done += 1
            if set_progress_callback:
                set_progress_callback(done, max(1, len(groups)))

        return pages, intro_table, groups

    def render_data_extent_overview(self) -> list:
        """Load the dissolved data extent written by Stage 1 and render an
        overview map. Returns an order_list fragment ready to splice into the
        report; empty list if the parquet is missing or unreadable."""
        extent_path = os.path.join(self.base_dir, "output", "geoparquet",
                                   "tbl_data_extent.parquet")
        if not os.path.exists(extent_path):
            write_to_log(
                "tbl_data_extent.parquet not found — re-run Stage 1 (Prep) to "
                "regenerate. Skipping area overview map.",
                self.base_dir,
            )
            return []
        try:
            extent_gdf = gpd.read_parquet(extent_path)
        except Exception as exc:
            write_to_log(f"Cannot read tbl_data_extent.parquet: {exc}", self.base_dir)
            return []
        png = self.make_path("data_extent", "overview")
        if not draw_data_extent_overview_map(extent_gdf, png, base_dir=self.base_dir):
            return []
        if not _file_ok(png):
            return []
        try:
            n_parts = sum(1 for _ in extent_gdf.geometry.iloc[0].geoms) \
                if extent_gdf.geometry.iloc[0].geom_type == 'MultiPolygon' else 1
        except Exception:
            n_parts = 1
        intro = (
            "Outline of the project's data area, computed at processing-time as "
            "the dissolved union of all input asset and geocode geometries. "
            f"Shown as {n_parts} polygon{'s' if n_parts != 1 else ''}; disjoint "
            "regions are preserved rather than collapsed into a single hull."
        )
        return [
            ('heading(2)', "Area overview"),
            ('text', intro),
            ('image', ("Project data extent", png)),
            ('rule', None),
        ]

    def render_atlas_maps(self,
                          flat_df: gpd.GeoDataFrame,
                          atlas_df: gpd.GeoDataFrame,
                          include_atlas_maps: bool,
                          set_progress_callback) -> tuple[list, str | None]:
        pages = []
        atlas_geocode_selected = 'basic_mosaic'
        if not include_atlas_maps or atlas_df is None or atlas_df.empty:
            return pages, atlas_geocode_selected

        flat_geos = flat_df if flat_df is not None else gpd.GeoDataFrame()
        flat_polys_3857 = gpd.GeoDataFrame()
        if not flat_geos.empty and 'geometry' in flat_geos.columns:
            flat_polys = flat_geos[flat_geos.geometry.type.isin(['Polygon','MultiPolygon'])].copy()
            if 'name_gis_geocodegroup' in flat_polys.columns:
                flat_polys['name_gis_geocodegroup'] = (
                    flat_polys['name_gis_geocodegroup']
                    .astype('string')
                    .str.strip()
                )
            flat_polys_3857 = _safe_to_3857(flat_polys)

        atlas_crs = atlas_df.crs
        atlas_df = atlas_df[atlas_df['geometry'].notna()].copy()

        polys_for_atlas = flat_polys_3857
        if atlas_geocode_selected and not flat_polys_3857.empty:
            if 'name_gis_geocodegroup' in flat_polys_3857.columns:
                atlas_mask = flat_polys_3857['name_gis_geocodegroup'].astype('string').str.lower() == atlas_geocode_selected.lower()
                filtered = flat_polys_3857[atlas_mask].copy()
                if filtered.empty:
                    write_to_log("No atlas polygons found for basic_mosaic. Using all polygons.", self.base_dir)
                else:
                    polys_for_atlas = filtered
            else:
                write_to_log("Atlas polygons missing 'name_gis_geocodegroup'; using all polygons.", self.base_dir)

        bounds = compute_fixed_bounds_3857(flat_df, base_dir=self.base_dir)
        overview_png = self.make_path("atlas", "overview")
        ok_overview = draw_atlas_overview_map(atlas_df, atlas_crs, polys_for_atlas, overview_png, bounds, base_dir=self.base_dir)
        if ok_overview and _file_ok(overview_png):
            text = "Overview of all atlas tiles within the study area."
            pages += [
                ('heading(3)', "Atlas overview"),
                ('text', text),
                ('image', ("Atlas tiles overview", overview_png)),
                ('new_page', None),
            ]

        # Reproject every atlas tile once so each per-tile inset can show
        # the rest of the grid as faint context lines around the highlighted
        # current tile.
        all_tiles_3857: gpd.GeoDataFrame | None = None
        try:
            tiles_gdf = gpd.GeoDataFrame(
                {"geometry": atlas_df["geometry"].values},
                geometry="geometry",
                crs=atlas_crs,
            )
            all_tiles_3857 = _safe_to_3857(tiles_gdf)
            if all_tiles_3857 is None or all_tiles_3857.empty:
                all_tiles_3857 = None
        except Exception as exc:
            write_to_log(f"Could not prepare atlas context tiles for inset: {exc}", self.base_dir)
            all_tiles_3857 = None

        atlas_total = len(atlas_df)
        for idx, tile_row in atlas_df.iterrows():
            safe_tile = _safe_name(tile_row.get('name_gis') or f"atlas_{idx+1}")
            sens_png = self.make_path("atlas", safe_tile, "sens")

            ok_sens = draw_atlas_map(
                tile_row,
                atlas_crs,
                polys_for_atlas,
                self.palette,
                self.desc,
                sens_png,
                bounds,
                base_dir=self.base_dir,
                all_tiles_3857=all_tiles_3857,
            )

            has_entries = False
            title_raw = tile_row.get('title_user') or tile_row.get('name_gis') or safe_tile
            tile_id = tile_row.get('name_gis') or safe_tile
            heading = f"Atlas tile: {title_raw}" if str(title_raw) == str(tile_id) else f"Atlas tile: {title_raw} ({tile_id})"
            # Force each atlas tile sub-chapter onto its own page so the
            # heading, description, and map stay together. Atlas tiles use
            # the 'image' kind (not 'image_map'), so the automatic
            # page-break-before-heading logic in compile_docx does not fire
            # for them — emit the page break explicitly here.
            pages.append(('new_page', None))
            pages.append(('heading(3)', heading))
            info_parts = []
            if isinstance(tile_row.get('description'), str) and tile_row.get('description').strip():
                info_parts.append(tile_row.get('description').strip())
            info_parts.append("Inset highlights tile within the study area.")
            pages.append(('text', " ".join(info_parts)))

            if ok_sens and _file_ok(sens_png):
                pages.append(('text',
                    "Cells coloured by sensitivity class (A–E palette), clipped to this tile. The "
                    "small inset shows where the tile sits inside the study area; the main image is "
                    "the same data the &ldquo;Other maps&rdquo; section shows zoomed out for the full extent."))
                pages.append(('image', ("Sensitivity atlas map", sens_png)))
                has_entries = True
            if has_entries:
                pages.append(('spacer', 1))
            else:
                pages.pop()  # remove text
                pages.pop()  # remove heading
                pages.pop()  # remove new_page

            if set_progress_callback:
                set_progress_callback(idx+1, max(1, atlas_total))

        return pages, atlas_geocode_selected

    def render_index_statistics(self,
                               flat_df: gpd.GeoDataFrame,
                               cfg: configparser.ConfigParser,
                               set_progress_callback=None) -> list:
        """Create one page per index showing basic_mosaic distribution."""
        pages: list = []
        if flat_df is None or flat_df.empty:
            return pages

        basic_name = (cfg["DEFAULT"].get("basic_group_name", "basic_mosaic") or "basic_mosaic").strip()

        candidates = [
            ("index_importance", "Importance index", "index_importance"),
            ("index_sensitivity", "Sensitivity index", "index_sensitivity"),
            ("index_owa", "OWA index", "index_owa"),
        ]

        available = [(col, title, mb_suffix) for (col, title, mb_suffix) in candidates if col in flat_df.columns]
        if not available:
            return pages

        # Per-index narrative shown above each map. Each entry contains:
        #   stats_lead -> one-sentence "what this index is", prepended to the stats page.
        #   map_intro  -> paragraph(s) shown above the map, explaining meaning, scale,
        #                 and how to read the colours. Grounded in the computation in
        #                 code/processing_internal.py:_compute_index_scores_from_stacked
        #                 and :_compute_index_owa_from_counts.
        index_descriptions = {
            "index_importance": {
                "stats_lead":
                    "The <b>Importance index</b> aggregates the <i>importance</i> attribute of every asset "
                    "that overlaps each geocode cell into a single 0–100 score per cell.",
                "map_intro": (
                    "The <b>Importance index</b> highlights where assets that have been rated as <b>important</b> "
                    "concentrate on the map. For each geocode cell, MESA performs three steps: "
                    "<b>(1) count</b> overlapping assets per importance class, "
                    "<b>(2) weight</b> those counts using <code>index_importance_weights</code> from <b>Parameters</b> "
                    "(in <code>config.ini</code>) and sum to a raw score, and "
                    "<b>(3) rank</b> the cell against all others in the current <b>{basic}</b> grouping by "
                    "rescaling to <b>0–100</b> (the most-loaded cell becomes 100; cells with no important "
                    "assets stay at 0).",
                    "<i>Worked example.</i> Suppose a cell overlaps 3 assets rated importance class 5 and the "
                    "configured weight for class 5 is 50. The cell&rsquo;s raw score is 3 &times; 50 = 150. If "
                    "150 happens to be the largest raw score anywhere in the geocode group, the cell shows "
                    "<b>100</b>; every other cell scales down in proportion to its own raw score. Real weights "
                    "and class counts come from your configuration, so the index is a <b>transparent expression "
                    "of stakeholder choices</b> rather than an objective measurement.",
                    "Read the colour ramp as a <b>relative</b> measure: darker / higher cells are the hotspots of "
                    "important features inside this study area, not absolute importance scores. Use this map to "
                    "spot clusters of high-value features (e.g. conservation targets, high-value infrastructure) "
                    "and to compare the spatial reach of importance against sensitivity and OWA on the next pages.",
                ),
            },
            "index_sensitivity": {
                "stats_lead":
                    "The <b>Sensitivity index</b> aggregates the <i>sensitivity</i> attribute of every asset "
                    "that overlaps each geocode cell into a single 0–100 score per cell.",
                "map_intro": (
                    "The <b>Sensitivity index</b> shows where features that are <b>vulnerable to pressure or change</b> "
                    "concentrate on the map. For each geocode cell, MESA performs three steps: "
                    "<b>(1) count</b> overlaps per sensitivity-product value (importance × susceptibility, in "
                    "{{1, 2, 3, …, 25}}), "
                    "<b>(2) weight</b> those counts using <code>index_sensitivity_weights</code> from <b>Parameters</b> "
                    "and sum to a raw score, and "
                    "<b>(3) rank</b> the cell against all others in the current <b>{basic}</b> grouping by rescaling "
                    "to <b>0–100</b>.",
                    "Defaults for the sensitivity weights are <b>flat</b> because the sensitivity-product values themselves "
                    "already encode magnitude — a product of 25 is intrinsically 25× a product of 1. With flat defaults "
                    "the index reduces to a weighted count of sensitive overlaps. Raise individual weights to "
                    "over-emphasise particular sensitivity levels (e.g. push the weight on 25 to make cells with "
                    "extreme overlaps rise sharply).",
                    "Read the colour ramp as a <b>relative</b> measure: a 100 means &ldquo;the most sensitive cell here&rdquo;, "
                    "not an absolute sensitivity rating. Use this map to target survey and mitigation effort and to "
                    "screen plans against the most vulnerable cells. Pair it with <b>Sensitive areas (A–E)</b> earlier "
                    "in the report — that map shows the single highest sensitivity class present in each cell, while "
                    "this index reflects accumulated weighted overlap.",
                ),
            },
            "index_owa": {
                "stats_lead":
                    "The <b>OWA index</b> (Ordered Weighted Average, precautionary form) ranks each cell on a 0–100 "
                    "scale by giving more weight to its highest-sensitivity overlaps than to its many low ones.",
                "map_intro": (
                    "The <b>OWA index</b> is a <b>precautionary</b> companion to the Sensitivity index. Instead of "
                    "summing weighted counts, it ranks cells <b>lexicographically</b> on their per-class count "
                    "vector, examined from the highest sensitivity class downwards. The ranks are rescaled to "
                    "<b>0–100</b> within the current <b>{basic}</b> grouping; cells with no overlapping sensitive "
                    "assets remain at 0.",
                    "<i>Worked example.</i> A cell containing even <b>one</b> overlap at sensitivity 25 outranks "
                    "every cell that has zero overlaps at 25, regardless of how many lower-class overlaps the "
                    "second cell has. Among cells that all have one overlap at 25, the next deciding criterion is "
                    "the count at 24, then 23, and so on. The rule is fixed — there are no tunable weights for "
                    "OWA — and uses only the sensitivity counts produced earlier in the pipeline.",
                    "Read this map as the answer to &ldquo;<i>where would the worst case be worst?</i>&rdquo; Compared with "
                    "the Sensitivity index, OWA pushes isolated high-sensitivity hits up the ranking and is the "
                    "more conservative choice when any extreme-sensitivity overlap should dominate the result. Use "
                    "it for precautionary screening and red-flag mapping.",
                ),
            },
        }

        mb_root = os.path.join(self.base_dir, "output", "mbtiles")
        safe_basic = _safe_name(basic_name)

        total = len(available)
        for i, (col, title, mb_suffix) in enumerate(available, start=1):
            descr = index_descriptions.get(col, {})
            stats_lead = descr.get("stats_lead", "")
            map_intro_paragraphs = descr.get("map_intro", ())
            if isinstance(map_intro_paragraphs, str):
                map_intro_paragraphs = (map_intro_paragraphs,)

            out_png = self.make_path("index", _safe_name(col), "distribution")
            ok, note = create_index_area_distribution_chart(
                flat_df,
                index_col=col,
                output_path=out_png,
                basic_group_name=basic_name,
                base_dir=self.base_dir,
            )
            if ok and _file_ok(out_png):
                stats_text = (
                    (f"{stats_lead} " if stats_lead else "") +
                    f"The chart below is computed from <b>tbl_flat</b> for geocode group <b>{basic_name}</b> only. "
                    "Bars show total polygon area (km²) per index value. "
                    "The line shows the number of categories per index value (A–E if available; otherwise the number of cells)."
                )
                pages += [
                    ('heading(2)', f"{title} – statistics"),
                    ('text', stats_text),
                    ('image', (f"{title} – area distribution", out_png)),
                    ('new_page', None),
                ]
            else:
                msg = note or "(chart not available)"
                pages += [
                    ('heading(2)', f"{title} – statistics"),
                    ('text', f"Could not generate chart for <b>{col}</b>: {msg}"),
                    ('new_page', None),
                ]

            # Index map page (placed immediately after its index statistics page)
            mb_path = os.path.join(mb_root, f"{safe_basic}_{mb_suffix}.mbtiles")
            map_png = self.make_path("index", safe_basic, _safe_name(mb_suffix), "map")
            okm = False
            note_m = ""
            if os.path.exists(mb_path):
                okm, note_m = render_mbtiles_to_png_best_fit(mb_path, map_png, base_dir=self.base_dir)
            else:
                note_m = f"Missing MBTiles: {os.path.basename(mb_path)}"

            pages.append(('heading(2)', f"{title} – map"))
            if map_intro_paragraphs:
                for paragraph in map_intro_paragraphs:
                    pages.append(('text', paragraph.format(basic=basic_name)))
            else:
                pages.append(('text', f"Map of <b>{title}</b> rendered from the basic_mosaic MBTiles overlay."))
            if okm and _file_ok(map_png):
                pages.append(('image_map', (f"{title} – map", map_png)))
                # Legend strip below the map. Filename includes "_legend" so
                # compile_docx routes it through the legend-width branch.
                legend_png = self.make_path("index", _safe_name(col), "legend")
                legend_kind = "importance" if col == "index_importance" else "sensitivity"
                legend_caption = f"Index value 0–100 (relative within {basic_name})"
                if _build_index_legend_png(
                    legend_png,
                    kind=legend_kind,
                    cfg_path=self.config_path,
                    caption=legend_caption,
                ):
                    pages.append(('image', (f"{title} – legend", legend_png)))
            else:
                pages.append(('text', f"Could not render map: {note_m or 'unknown error'}"))
            pages.append(('new_page', None))

            if set_progress_callback:
                try:
                    set_progress_callback(i, total)
                except Exception:
                    pass

        return pages

    def render_lines_overview(self,
                              lines_df: gpd.GeoDataFrame,
                              segments_df: gpd.GeoDataFrame,
                              palette: dict,
                              base_dir: str) -> list:
        """Create a short overview section for all lines (map + distribution + summary table)."""
        pages: list = []
        if lines_df is None or segments_df is None:
            return pages
        if lines_df.empty or segments_df.empty:
            return pages
        if 'name_gis' not in lines_df.columns or 'geometry' not in lines_df.columns:
            return pages
        if 'name_gis' not in segments_df.columns:
            return pages

        segs = segments_df.copy()
        if 'geometry' in segs.columns:
            segs = segs[segs['geometry'].notna()].copy()

        if segs.empty:
            return pages

        # Normalize max sensitivity code (robust against missing code vs numeric).
        segs['__sens_code_max'] = segs.apply(
            lambda row: _normalize_sensitivity_code(
                row.get('sensitivity_code_max'),
                row.get('sensitivity_max')
            ),
            axis=1
        )

        # Overview map + distribution image.
        overview_map = self.make_path("lines", "overview", "map")
        dist_png = self.make_path("lines", "overview", "distribution")
        ok_map = draw_lines_overview_map(segs, palette, overview_map, mode='max', pad_ratio=0.08, base_dir=base_dir)
        create_sensitivity_summary(segs['__sens_code_max'], palette, dist_png)

        if ok_map and _file_ok(overview_map):
            pages.append(('heading(3)', "All lines – overview map"))
            pages.append(('text',
                "All line segments shown together, coloured by the <b>highest sensitivity class</b> "
                "(A–E) found at any point along the segment. This is a worst-case overview: a segment "
                "is drawn red (A) if it touches even one A-rated cell, regardless of how short that "
                "stretch is. Use this map to scan the network for hot stretches before drilling into "
                "individual lines on the following pages."))
            pages.append(('image_map', ("Lines overview (max sensitivity)", overview_map)))

        if _file_ok(dist_png):
            pages.append(('heading(3)', "All lines – sensitivity distribution"))
            pages.append(('image', ("Maximum sensitivity – distribution", dist_png)))

        # Summary table per line.
        pages.append(('heading(3)', "Per-line summary"))

        counts = (
            segs.groupby(['name_gis', '__sens_code_max'])
            .size()
            .unstack(fill_value=0)
        )
        for code in ['A', 'B', 'C', 'D', 'E']:
            if code not in counts.columns:
                counts[code] = 0

        seg_count = segs.groupby('name_gis').size().rename('__segments').to_frame()
        per_line = lines_df[['name_gis', 'geometry']].copy()
        per_line['name_gis'] = per_line['name_gis'].astype('string')
        try:
            g_len = gpd.GeoDataFrame(per_line[['geometry']].copy(), geometry='geometry', crs=lines_df.crs)
            g_len = g_len[g_len.geometry.notna()].copy()
            g_len_3857 = _safe_to_3857(g_len)
            per_line['length_m'] = g_len_3857.geometry.length.reindex(per_line.index).fillna(0.0)
        except Exception:
            per_line['length_m'] = 0.0
        per_line = per_line.merge(seg_count, left_on='name_gis', right_index=True, how='left')
        per_line = per_line.merge(counts[['A', 'B', 'C', 'D', 'E']], left_on='name_gis', right_index=True, how='left')
        per_line = per_line.fillna(0)

        def _km(v) -> str:
            try:
                return f"{float(v)/1000.0:,.2f}"
            except Exception:
                return "0.00"

        table = [["Line", "Length (km)", "Segments", "A", "B", "C", "D", "E"]]
        try:
            per_line = per_line.sort_values('name_gis')
        except Exception:
            pass

        for _, r in per_line.iterrows():
            table.append([
                str(r.get('name_gis', '')),
                _km(r.get('length_m', 0.0)),
                int(r.get('__segments', 0) or 0),
                int(r.get('A', 0) or 0),
                int(r.get('B', 0) or 0),
                int(r.get('C', 0) or 0),
                int(r.get('D', 0) or 0),
                int(r.get('E', 0) or 0),
            ])

        pages.append(('table_data', ("Line overview", table)))
        pages.append(('new_page', None))
        return pages

    def render_segments(self,
                        lines_df: gpd.GeoDataFrame,
                        segments_df: gpd.GeoDataFrame,
                        palette: dict,
                        base_dir: str,
                        set_progress_callback):
        pages_lines = []
        log_data = []
        if (lines_df.empty or segments_df.empty or
            {'name_gis','segment_id','sensitivity_code_max','sensitivity_code_min'}.issubset(segments_df.columns) is False or
            'name_gis' not in lines_df.columns or 'geometry' not in lines_df.columns):
            return pages_lines, log_data

        # tbl_lines may not have a precomputed length_m column; derive it from geometry.
        lines_work = lines_df.copy()
        if 'length_m' not in lines_work.columns:
            try:
                g_len = gpd.GeoDataFrame(lines_work[['geometry']].copy(), geometry='geometry', crs=lines_df.crs)
                g_len = g_len[g_len.geometry.notna()].copy()
                g_len_3857 = _safe_to_3857(g_len)
                lines_work['length_m'] = g_len_3857.geometry.length.reindex(lines_work.index).fillna(0.0)
            except Exception:
                lines_work['length_m'] = 0.0

        total = len(lines_work)
        for n, (_, line) in enumerate(lines_work.iterrows(), start=1):
            ln_visible = line['name_gis']
            ln_safe = _safe_name(ln_visible)
            length_m = float(line.get('length_m', 0) or 0)
            length_km = length_m / 1000.0
            try:
                write_to_log(f"Rendering line segments ({n}/{total}): {ln_visible}", base_dir)
            except Exception:
                pass

            segment_records = sort_segments_numerically(segments_df[segments_df['name_gis'] == ln_visible]).copy()

            if not segment_records.empty:
                segment_records['__sens_code_max'] = segment_records.apply(
                    lambda row: _normalize_sensitivity_code(
                        row.get('sensitivity_code_max'),
                        row.get('sensitivity_max')
                    ),
                    axis=1
                )
                segment_records['__sens_code_min'] = segment_records.apply(
                    lambda row: _normalize_sensitivity_code(
                        row.get('sensitivity_code_min'),
                        row.get('sensitivity_min')
                    ),
                    axis=1
                )

            seg_map_max = self.make_path("line", ln_safe, "segments_max_inset")
            seg_map_min = self.make_path("line", ln_safe, "segments_min_inset")

            ok_max = draw_line_segments_map_with_context_inset(
                segments_df,
                ln_visible,
                palette,
                seg_map_max,
                mode='max',
                pad_ratio=0.20,
                context_pad_ratio=1.00,
                rect_buffer_ratio=0.03,
                base_dir=base_dir,
            )
            ok_min = draw_line_segments_map_with_context_inset(
                segments_df,
                ln_visible,
                palette,
                seg_map_min,
                mode='min',
                pad_ratio=0.20,
                context_pad_ratio=1.00,
                rect_buffer_ratio=0.03,
                base_dir=base_dir,
            )

            max_stats_img = self.make_path("line", ln_safe, "max_dist")
            min_stats_img = self.make_path("line", ln_safe, "min_dist")
            max_codes = segment_records.get('__sens_code_max', segment_records.get('sensitivity_code_max'))
            min_codes = segment_records.get('__sens_code_min', segment_records.get('sensitivity_code_min'))
            if max_codes is None:
                max_codes = pd.Series(dtype=object)
            if min_codes is None:
                min_codes = pd.Series(dtype=object)
            create_sensitivity_summary(max_codes, palette, max_stats_img)
            create_sensitivity_summary(min_codes, palette, min_stats_img)

            max_img = self.make_path("line", ln_safe, "max_ribbon")
            min_img = self.make_path("line", ln_safe, "min_ribbon")
            create_line_statistic_image(ln_visible, max_codes, palette, length_m, max_img)
            create_line_statistic_image(ln_visible, min_codes, palette, length_m, min_img)

            first_page = [
                ('heading(2)', f"Line: {ln_visible}"),
                ('text',
                    f"This section summarises sensitivity along <b>{ln_visible}</b> "
                    f"(total length <b>{length_km:.2f} km</b>, split into <b>{len(segment_records)}</b> segments). "
                    "Each segment is rated by the sensitivity classes of the cells it crosses; the "
                    "<b>maximum</b> class is the worst class encountered along the segment, and is the "
                    "right view when even a brief touch on a sensitive cell matters. The context inset "
                    "shows where this line sits inside the wider study area, and the ribbon below the "
                    "distribution lays the line out flat from start to end so peaks are easy to locate."),
            ]
            if ok_max and _file_ok(seg_map_max):
                first_page.append(('image', ("Segments colored by maximum sensitivity (with context inset)", seg_map_max)))
            first_page.append(('image', ("Maximum sensitivity – distribution", max_stats_img)))
            first_page.append(('text', f"Distance (km): 0 – {length_km/2:.1f} – {length_km:.1f}"))
            first_page.append(('image_ribbon', ("Maximum sensitivity – along line", max_img)))
            first_page.append(('new_page', None))

            second_page = [
                ('heading(2)', f"Line: {ln_visible} (continued)"),
                ('text',
                    "Same line, this time rated by the <b>minimum sensitivity class</b> per segment — the "
                    "best (least sensitive) class encountered along the segment. The minimum view is the "
                    "right one when you need to know whether a segment ever passes through robust ground "
                    "at all; comparing it side-by-side with the maximum view from the previous page tells "
                    "you whether high sensitivity is everywhere on the line or only in pockets."),
            ]
            if ok_min and _file_ok(seg_map_min):
                second_page.append(('image', ("Segments colored by minimum sensitivity (with context inset)", seg_map_min)))
            second_page.append(('image', ("Minimum sensitivity – distribution", min_stats_img)))
            second_page.append(('text', f"Distance (km): 0 – {length_km/2:.1f} – {length_km:.1f}"))
            second_page.append(('image_ribbon', ("Minimum sensitivity – along line", min_img)))
            second_page.append(('new_page', None))

            pages_lines += first_page + second_page

            for _, seg in segment_records.iterrows():
                log_data.append({
                    'line_name': ln_visible,
                    'segment_id': seg['segment_id'],
                    'sensitivity_code_max': seg['sensitivity_code_max'],
                    'sensitivity_code_min': seg['sensitivity_code_min']
                })

            if set_progress_callback:
                try:
                    set_progress_callback(n, total)
                except Exception:
                    pass

        return pages_lines, log_data

# Primary UI color (Steel Blue by default)
PRIMARY_HEX = "#4682B4"        # Steel blue
LIGHT_PRIMARY_HEX = "#6fa6cf"  # lighter steel-blue

# ---------------- Config + path helpers ----------------
def read_config(file_name: str) -> configparser.ConfigParser:
    cfg = configparser.ConfigParser()
    try:
        cfg.read(file_name, encoding="utf-8")
    except Exception as e:
        print(f"[mesa] warn: config could not be read from {file_name}: {e}", flush=True)
    if "DEFAULT" not in cfg:
        cfg["DEFAULT"] = {}
    return cfg

def config_path(base_dir: str) -> str:
    """Flat config: <base>/config.ini"""
    return os.path.join(base_dir, "config.ini")

def parquet_dir_from_cfg(base_dir: str, cfg: configparser.ConfigParser) -> str:
    sub = cfg["DEFAULT"].get("parquet_folder", "output/geoparquet")
    required = ("tbl_asset_object.parquet", "tbl_asset_group.parquet", "tbl_flat.parquet")

    base_path = Path(base_dir)
    sub_path = Path(sub)

    candidates: list[Path] = []
    seen: set[str] = set()

    def _register(path: Path):
        try:
            resolved = path.resolve()
        except Exception:
            resolved = path
        key = str(resolved)
        if key not in seen:
            seen.add(key)
            candidates.append(resolved)

    if sub_path.is_absolute():
        _register(sub_path)
    else:
        ancestors = [base_path]
        ancestors.extend(list(base_path.parents)[:4])
        for ancestor in ancestors:
            _register(ancestor / sub_path)
            _register(ancestor / "code" / sub_path)

    primary = candidates[0] if candidates else (base_path / sub_path)
    for cand in candidates:
        if all((cand / req).exists() for req in required):
            if cand != primary:
                write_to_log(f"Using parquet folder at {cand}", base_dir)
            return str(cand)

    try:
        primary.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    if all((primary / req).exists() for req in required):
        return str(primary)

    write_to_log(f"Expected parquet tables not found; using {primary}", base_dir)
    return str(primary)

# ---------------- Base dir normalization ----------------
def normalize_base_dir(path_str: str) -> str:
    """Return the logical app root for both .py and packaged .exe runs.
    - If started inside tools/ or system/ or code/, climb to parent.
    - Then climb up a few levels to find a folder containing 'output' & 'input'
      or containing 'tools' and 'config.ini'.
    """
    try:
        p = Path(path_str).resolve()
    except Exception:
        return path_str

    if p.name.lower() in ("tools", "system", "code"):
        p = p.parent

    q = p
    for _ in range(4):
        try:
            if (q / "output").exists() and (q / "input").exists():
                return str(q)
            if (q / "tools").exists() and (q / "config.ini").exists():
                return str(q)
        except Exception:
            pass
        q = q.parent
    return str(p)

# ---------------- Logging ----------------
def write_to_log(message: str, base_dir: str | None = None):
    timestamp = datetime.datetime.now().strftime("%Y.%m.%d %H:%M:%S")
    formatted = f"{timestamp} - {message}"
    try:
        if base_dir:
            with open(os.path.join(base_dir, "log.txt"), "a", encoding="utf-8") as f:
                f.write(formatted + "\n")
    except Exception:
        print(formatted, flush=True)
    # GUI log update (thread-safe via signal)
    try:
        if _signals is not None:
            _signals.log_message.emit(formatted)
    except Exception:
        # Never let logging break report generation.
        pass

_OUTPUT_ROOT: Path | None = None

def _output_candidates(base_dir: str | None) -> list[Path]:
    base = Path(base_dir) if base_dir else Path.cwd()
    try:
        base = base.resolve()
    except Exception:
        pass
    candidates: list[Path] = []
    # Prefer <base>/output (workspace-root output) whenever possible.
    # Only fall back to <base>/code/output if needed (e.g. when running with base_dir pointing at repo root
    # but output exists only under code/).
    candidates.append((base / "output").resolve())
    if base.name.lower() != "code":
        code_dir = (base / "code")
        if code_dir.exists():
            candidates.append((code_dir / "output").resolve())
    return candidates

def _resolve_output_root(base_dir: str | None) -> Path:
    global _OUTPUT_ROOT
    if _OUTPUT_ROOT is not None:
        return _OUTPUT_ROOT
    candidates = _output_candidates(base_dir)
    for idx, cand in enumerate(candidates):
        if cand.exists():
            _OUTPUT_ROOT = cand
            if idx > 0 and base_dir:
                write_to_log(f"Using fallback output folder: {cand}", base_dir)
            return _OUTPUT_ROOT
    target = candidates[0]
    try:
        target.mkdir(parents=True, exist_ok=True)
        if base_dir:
            write_to_log(f"Created output folder at {target}", base_dir)
    except Exception:
        pass
    _OUTPUT_ROOT = target
    return _OUTPUT_ROOT

def output_subpath(base_dir: str | None, *parts: str) -> Path:
    root = _resolve_output_root(base_dir)
    path = root.joinpath(*parts)
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    return path

# ---------------- Small utilities ----------------
def _file_ok(path: str) -> bool:
    try:
        return os.path.isfile(path) and os.path.getsize(path) > 0
    except Exception:
        return False

def _safe_name(name: str) -> str:
    # Windows-safe file name (keep visible name unchanged in titles)
    return re.sub(r'[^A-Za-z0-9_.-]+', '_', str(name))

def _cfg_getboolean(cfg: configparser.ConfigParser,
                    section: str,
                    option: str,
                    default: bool = False) -> bool:
    """
    Robust boolean reader for configparser with permissive parsing.
    Accepts truthy values like '1', 'true', 'yes', 'on'.
    """
    try:
        raw = cfg.get(section, option, fallback=None)
    except Exception:
        raw = None
    if raw is None:
        return default
    val = str(raw).strip().lower()
    if val in {"1", "true", "yes", "on"}:
        return True
    if val in {"0", "false", "no", "off"}:
        return False
    return default

def _cfg_getfloat(cfg: configparser.ConfigParser,
                  section: str,
                  option: str,
                  default: float | None = None) -> float | None:
    """Robust float reader for configparser values."""
    try:
        raw = cfg.get(section, option, fallback=None)
    except Exception:
        raw = None
    if raw is None:
        return default
    try:
        return float(str(raw).strip())
    except Exception:
        return default

def _dpi_for_fig_height(fig_height_in: float, px_cap: int = MAX_MAP_PX_HEIGHT, min_dpi: int = 110, max_dpi: int = 300) -> int:
    """
    Choose a DPI so that fig_height_in * dpi <= px_cap, bounded by min/max dpi.
    """
    try:
        dpi_by_cap = int(px_cap / max(fig_height_in, 1e-6))
        return max(min_dpi, min(max_dpi, dpi_by_cap))
    except Exception:
        return 150

# ---------------- Palette / descriptions ----------------
def read_sensitivity_palette_and_desc(file_name: str):
    """
    Reads color palette and descriptions for A-E from config.ini.
    Returns: (colors: dict{A..E->hex, 'UNKNOWN'->hex}, desc: dict{A..E->str})
    """
    global SENSITIVITY_UNKNOWN_COLOR, _SENSITIVITY_NUMERIC_RANGES
    cfg = configparser.ConfigParser()
    try:
        cfg.read(file_name, encoding="utf-8")
    except Exception as e:
        print(f"[mesa] warn: sensitivity palette config could not be read from {file_name}: {e}", flush=True)

    unknown_col = '#BDBDBD'
    if cfg.has_section('VALID_VALUES'):
        try:
            unknown_col = cfg['VALID_VALUES'].get('category_colour_unknown', unknown_col).strip() or unknown_col
        except Exception:
            pass
    SENSITIVITY_UNKNOWN_COLOR = unknown_col

    ranges: list[tuple[str, float, float]] = []

    def _parse_range(range_str: str):
        try:
            parts = re.split(r'\s*[-]\s*', range_str.strip())
            if len(parts) == 2:
                lo = float(parts[0])
                hi = float(parts[1])
                if lo > hi:
                    lo, hi = hi, lo
                return lo, hi
        except Exception:
            return None
        return None

    colors_map, desc_map = {}, {}
    for code in ['A','B','C','D','E']:
        if cfg.has_section(code):
            col = cfg[code].get('category_colour', '').strip() or unknown_col
            colors_map[code] = col
            desc_map[code] = cfg[code].get('description', '').strip()
            range_str = cfg[code].get('range', '').strip()
            parsed = _parse_range(range_str) if range_str else None
            if parsed:
                ranges.append((code, parsed[0], parsed[1]))
        else:
            colors_map[code] = unknown_col
            desc_map[code] = ''
    colors_map['UNKNOWN'] = unknown_col
    _SENSITIVITY_NUMERIC_RANGES = ranges
    return colors_map, desc_map

def _normalize_sensitivity_code(code_val, numeric_val):
    if code_val is not None and not pd.isna(code_val):
        try:
            code_str = str(code_val).strip().upper()
        except Exception:
            code_str = str(code_val).upper()
        if code_str in SENSITIVITY_ORDER:
            return code_str
    numeric_code = _sensitivity_code_from_numeric(numeric_val)
    if numeric_code in SENSITIVITY_ORDER:
        return numeric_code
    return 'UNKNOWN'

def _prepare_sensitivity_annotations(df: gpd.GeoDataFrame,
                                     code_column: str,
                                     numeric_column: str) -> gpd.GeoDataFrame:
    if df.empty:
        return df.copy()
    codes = df.apply(
        lambda row: _normalize_sensitivity_code(
            row.get(code_column),
            row.get(numeric_column)
        ),
        axis=1
    )
    out = df.copy()
    out['__sens_code'] = codes.fillna('UNKNOWN')
    return out


def _colors_from_annotations(gdf: gpd.GeoDataFrame,
                             palette: dict[str, str]) -> pd.Series:
    """
    Map normalized sensitivity codes stored in '__sens_code' to palette colors.
    Falls back to the configured UNKNOWN color when missing.
    """
    fallback = palette.get('UNKNOWN', SENSITIVITY_UNKNOWN_COLOR)

    def _lookup(code):
        if pd.isna(code):
            return fallback
        try:
            key = str(code).strip().upper()
        except Exception:
            key = str(code).upper()
        return palette.get(key, fallback)

    return gdf['__sens_code'].apply(_lookup)

def _sensitivity_code_from_numeric(value: float | int | None) -> str | None:
    try:
        if value is None or (isinstance(value, float) and not math.isfinite(value)):
            return None
        val = float(value)
    except Exception:
        return None
    for code, lo, hi in _SENSITIVITY_NUMERIC_RANGES:
        if lo <= val <= hi:
            return code
    return None

def _resolve_sensitivity_color(code_value,
                               numeric_value,
                               palette: dict[str, str]) -> str:
    if code_value is not None:
        try:
            code_str = str(code_value).strip().upper()
        except Exception:
            code_str = str(code_value).upper()
        if code_str:
            color = palette.get(code_str)
            if color:
                return color
    numeric_code = _sensitivity_code_from_numeric(numeric_value)
    if numeric_code:
        color = palette.get(numeric_code)
        if color:
            return color
    return palette.get('UNKNOWN', SENSITIVITY_UNKNOWN_COLOR)



# ---------------- Generic plotting utils ----------------
def _safe_to_3857(gdf: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    if gdf.empty:
        return gdf
    g = gdf.copy()
    try:
        if g.crs is None:
            g = g.set_crs(4326)
        try:
            epsg = g.crs.to_epsg()
        except Exception:
            epsg = None
        if epsg != 3857:
            g = g.to_crs(3857)
    except Exception:
        try:
            g = g.set_crs(3857, allow_override=True)
        except Exception:
            pass
    return g

def _tile_cache_root(base_dir: str | None) -> Path:
    """
    Resolve the cache directory used for XYZ tiles. Defaults to <base>/output/tile_cache.
    """
    cache = output_subpath(base_dir, "tile_cache")
    try:
        cache.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass

    # Help users locate the cache on disk (log once per process).
    global _OSM_TILE_CACHE_LOGGED
    if not _OSM_TILE_CACHE_LOGGED and base_dir:
        _OSM_TILE_CACHE_LOGGED = True
        try:
            write_to_log(f"OSM tile cache folder: {cache}", base_dir)
        except Exception:
            pass
    return cache

def _load_cached_tile(cache_path: Path):
    if not cache_path.exists():
        return None
    if TILE_CACHE_MAX_AGE_DAYS > 0:
        try:
            age = time.time() - cache_path.stat().st_mtime
            if age > TILE_CACHE_MAX_AGE_DAYS * 86400:
                try:
                    cache_path.unlink()
                except Exception:
                    pass
                return None
        except Exception:
            return None
    try:
        with PILImage.open(cache_path) as pil_img:
            return pil_img.convert("RGB")
    except Exception:
        try:
            cache_path.unlink()
        except Exception:
            pass
        return None

def _save_tile_to_cache(cache_path: Path, data: bytes) -> None:
    try:
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        with open(cache_path, "wb") as f:
            f.write(data)
    except Exception:
        pass

def _plot_basemap(ax, crs_epsg=3857, base_dir: str | None = None):
    """Add a basemap under current axes.
    Prefers contextily; if unavailable, fetches OSM Web Mercator tiles and composites them.
    """
    mode = str(_REPORT_BASEMAP_MODE or "xyz").strip().lower()

    # Preferred path: contextily (unless forced xyz)
    if mode != "xyz" and ctx is not None:
        try:
            # We render attribution as plain text under the map in the DOCX output.
            # Try to suppress attribution embedded in the map image.
            try:
                ctx.add_basemap(
                    ax,
                    crs=f"EPSG:{crs_epsg}",
                    source=ctx.providers.OpenStreetMap.Mapnik,
                    attribution=None,
                )
            except TypeError:
                ctx.add_basemap(ax, crs=f"EPSG:{crs_epsg}", source=ctx.providers.OpenStreetMap.Mapnik)

            if base_dir:
                try:
                    write_to_log("Basemap source: contextily (no local XYZ tile cache used)", base_dir)
                except Exception:
                    pass

            # Defensive cleanup: remove any attribution text artists if present
            try:
                for t in list(getattr(ax, "texts", []) or []):
                    s = (getattr(t, "get_text", lambda: "")() or "")
                    if ("openstreetmap" in s.lower()) or ("©" in s) or ("copyright" in s.lower()):
                        try:
                            t.remove()
                        except Exception:
                            pass
            except Exception:
                pass
            return
        except Exception as e:
            write_to_log(f"Basemap via contextily failed, falling back to tiles: {e}", base_dir)

    # If contextily is forced, do not proceed with XYZ fallback.
    if mode == "contextily":
        return

    # Fallback: simple OSM tile fetch/composite in EPSG:3857 only
    if int(crs_epsg) != 3857:
        write_to_log("Basemap fallback only supports EPSG:3857; skipping.", base_dir)
        return

    try:
        # Current view in meters (WebMercator)
        minx, maxx = ax.get_xlim()
        miny, maxy = ax.get_ylim()

        def merc_to_lonlat(x, y):
            R = 6378137.0
            lon = (x / R) * 180.0 / math.pi
            lat = (2.0 * math.atan(math.exp(y / R)) - math.pi/2.0) * 180.0 / math.pi
            return lon, lat

        def lonlat_to_tile(lon, lat, z):
            n = 2 ** z
            x = int((lon + 180.0) / 360.0 * n)
            lat = max(-85.05112878, min(85.05112878, float(lat)))
            lat_rad = math.radians(lat)
            y = int((1.0 - math.log(math.tan(lat_rad) + 1.0 / math.cos(lat_rad)) / math.pi) / 2.0 * n)
            return x, y

        def tile_bounds_lonlat(z, x, y):
            n = 2.0 ** z
            minlon = x / n * 360.0 - 180.0
            maxlon = (x + 1) / n * 360.0 - 180.0
            def tiley_to_lat(t):
                Y = t / n
                lat_rad = math.atan(math.sinh(math.pi * (1 - 2 * Y)))
                return math.degrees(lat_rad)
            maxlat = tiley_to_lat(y)
            minlat = tiley_to_lat(y + 1)
            return (minlon, minlat, maxlon, maxlat)

        def lonlat_to_merc(lon, lat):
            R = 6378137.0
            x = lon * math.pi / 180.0 * R
            lat = max(-85.05112878, min(85.05112878, float(lat)))
            y = math.log(math.tan((90.0 + lat) * math.pi / 360.0)) * R
            return x, y

        # Decide a zoom level based on axis pixel size to keep labels crisp.
        # Aim for mosaic pixel dims >= 1.5x axes pixels (oversample to reduce blur).
        bbox = ax.get_window_extent()
        try:
            ax_pix_w = int(bbox.width)
            ax_pix_h = int(bbox.height)
        except Exception:
            # Fallback: estimate from figure size
            ax_pix_w = int(ax.figure.dpi * ax.figure.get_size_inches()[0])
            ax_pix_h = int(ax.figure.dpi * ax.figure.get_size_inches()[1])

        lon0, lat1 = merc_to_lonlat(minx, maxy)
        lon1, lat0 = merc_to_lonlat(maxx, miny)

        target_w = max(512, int(ax_pix_w * 1.5))
        target_h = max(512, int(ax_pix_h * 1.5))

        z = 3
        for test in range(3, 20):
            x0, y0 = lonlat_to_tile(lon0, lat0, test)
            x1, y1 = lonlat_to_tile(lon1, lat1, test)
            w = (abs(x1 - x0) + 1) * 256
            h = (abs(y1 - y0) + 1) * 256
            if w >= target_w and h >= target_h:
                z = test
                break

        tile_limit = 400 if getattr(sys, "frozen", False) else 900
        def _tile_range(z_level):
            tx0, ty0 = lonlat_to_tile(lon0, lat0, z_level)
            tx1, ty1 = lonlat_to_tile(lon1, lat1, z_level)
            return (min(tx0, tx1), max(tx0, tx1), min(ty0, ty1), max(ty0, ty1))

        xmin, xmax, ymin, ymax = _tile_range(z)
        tile_span = (xmax - xmin + 1) * (ymax - ymin + 1)
        while tile_span > tile_limit and z > 3:
            z -= 1
            xmin, xmax, ymin, ymax = _tile_range(z)
            tile_span = (xmax - xmin + 1) * (ymax - ymin + 1)

        if tile_span > tile_limit:
            try:
                ax.set_facecolor("#f5f5f5")
            except Exception:
                pass
            write_to_log(f"Basemap fallback skipped; tile span too large ({tile_span} tiles @ z{z}).", base_dir)
            return

        if tile_span > tile_limit * 0.6:
            write_to_log(f"Basemap fallback using reduced zoom z{z} ({tile_span} tiles).", base_dir)

        TILE = 256
        W, H = (xmax - xmin + 1) * TILE, (ymax - ymin + 1) * TILE

        mosaic = PILImage.new("RGB", (W, H), (240, 240, 240))

        ua = ("Mozilla/5.0 (compatible; MESA-Report/1.0)")
        opener = urllib.request.build_opener()
        opener.addheaders = [("User-Agent", ua)]
        urllib.request.install_opener(opener)

        def tile_url(z, x, y):
            return f"https://tile.openstreetmap.org/{z}/{x}/{y}.png"

        if base_dir:
            try:
                write_to_log("Basemap source: XYZ tiles (OpenStreetMap) with local cache", base_dir)
            except Exception:
                pass
        cache_root = _tile_cache_root(base_dir)
        cache_hits = 0
        fetched = 0

        for xi, x in enumerate(range(xmin, xmax+1)):
            for yi, y in enumerate(range(ymin, ymax+1)):
                cache_path = cache_root / str(z) / str(x) / f"{y}.png"
                img = _load_cached_tile(cache_path)
                if img is not None:
                    cache_hits += 1
                else:
                    try:
                        url = tile_url(z, x, y)
                        with urllib.request.urlopen(url, timeout=5) as resp:
                            data = resp.read()
                        img = PILImage.open(io.BytesIO(data)).convert("RGB")
                        _save_tile_to_cache(cache_path, data)
                        fetched += 1
                    except Exception:
                        img = None
                if img is not None:
                    mosaic.paste(img, (xi*TILE, yi*TILE))

        if (cache_hits or fetched) and base_dir:
            try:
                write_to_log(f"Basemap tiles z{z}: {cache_hits} cached, {fetched} fetched ({tile_span} total).", base_dir)
            except Exception:
                pass

        # Extent of mosaic in Mercator meters (XYZ scheme)
        west  = tile_bounds_lonlat(z, xmin, ymin)[0]  # minlon of top-left tile
        north = tile_bounds_lonlat(z, xmin, ymin)[3]  # maxlat of top-left tile
        east  = tile_bounds_lonlat(z, xmax, ymax)[2]  # maxlon of bottom-right tile
        south = tile_bounds_lonlat(z, xmax, ymax)[1]  # minlat of bottom-right tile

        lx, north_y = lonlat_to_merc(west,  north)
        rx, south_y = lonlat_to_merc(east,  south)

        # Show with correct extent; origin='upper' so row 0 is at 'north'
        ax.imshow(mosaic, extent=[lx, rx, south_y, north_y], origin="upper",
                  interpolation='bilinear', resample=True)
    except Exception as e:
        write_to_log(f"Basemap fallback failed: {e}", base_dir)

def _expand_bounds(bounds, pad_ratio=0.08):
    minx, miny, maxx, maxy = bounds
    dx, dy = maxx - minx, maxy - miny
    if dx <= 0 or dy <= 0:
        return bounds
    pad_x, pad_y = dx * pad_ratio, dy * pad_ratio
    return (minx - pad_x, miny - pad_y, maxx + pad_x, maxy + pad_y)

# ---------------- MBTiles rendering (raster) ----------------
WEBMERCATOR_MAX_LAT = 85.05112878
MAX_MAP_PX_WIDTH = 2000

def _clip_latlon_bounds(bounds: tuple[float, float, float, float]) -> tuple[float, float, float, float]:
    minlon, minlat, maxlon, maxlat = bounds
    minlon = max(-180.0, min(180.0, float(minlon)))
    maxlon = max(-180.0, min(180.0, float(maxlon)))
    minlat = max(-WEBMERCATOR_MAX_LAT, min(WEBMERCATOR_MAX_LAT, float(minlat)))
    maxlat = max(-WEBMERCATOR_MAX_LAT, min(WEBMERCATOR_MAX_LAT, float(maxlat)))
    if maxlon < minlon:
        minlon, maxlon = maxlon, minlon
    if maxlat < minlat:
        minlat, maxlat = maxlat, minlat
    # Avoid exact-180/85 edges causing off-by-one tile selection
    eps = 1e-9
    minlon = max(-180.0 + eps, minlon)
    maxlon = min(180.0 - eps, maxlon)
    minlat = max(-WEBMERCATOR_MAX_LAT + eps, minlat)
    maxlat = min(WEBMERCATOR_MAX_LAT - eps, maxlat)
    return minlon, minlat, maxlon, maxlat

def _lonlat_to_tile_fraction(lon: float, lat: float, z: int) -> tuple[float, float]:
    lat = max(-WEBMERCATOR_MAX_LAT, min(WEBMERCATOR_MAX_LAT, float(lat)))
    lon = max(-180.0, min(180.0, float(lon)))
    n = 2 ** int(z)
    x = (lon + 180.0) / 360.0 * n
    lat_rad = math.radians(lat)
    y = (1.0 - math.log(math.tan(lat_rad) + (1.0 / math.cos(lat_rad))) / math.pi) / 2.0 * n
    return x, y

def _mbtiles_metadata(conn: sqlite3.Connection) -> dict[str, str]:
    meta: dict[str, str] = {}
    try:
        cur = conn.cursor()
        cur.execute("SELECT name, value FROM metadata")
        for name, value in cur.fetchall():
            if name is None:
                continue
            meta[str(name)] = "" if value is None else str(value)
    except Exception:
        pass
    return meta

def _parse_mbtiles_bounds(meta: dict[str, str]) -> tuple[float, float, float, float] | None:
    raw = (meta.get("bounds") or "").strip()
    if not raw:
        return None
    try:
        parts = [float(x.strip()) for x in raw.split(",")]
        if len(parts) != 4:
            return None
        return _clip_latlon_bounds((parts[0], parts[1], parts[2], parts[3]))
    except Exception:
        return None

def _parse_mbtiles_zoom(meta: dict[str, str], key: str, default: int) -> int:
    try:
        return int(float((meta.get(key) or str(default)).strip()))
    except Exception:
        return int(default)

def _get_mbtiles_tile_bytes(conn: sqlite3.Connection, z: int, x: int, y_xyz: int) -> bytes | None:
    """Fetch tile bytes, trying TMS row convention first, then XYZ."""
    try:
        cur = conn.cursor()
        n = 2 ** int(z)
        y_tms = (n - 1) - int(y_xyz)
        cur.execute(
            "SELECT tile_data FROM tiles WHERE zoom_level=? AND tile_column=? AND tile_row=?",
            (int(z), int(x), int(y_tms)),
        )
        row = cur.fetchone()
        if row and row[0] is not None:
            return bytes(row[0])
        cur.execute(
            "SELECT tile_data FROM tiles WHERE zoom_level=? AND tile_column=? AND tile_row=?",
            (int(z), int(x), int(y_xyz)),
        )
        row = cur.fetchone()
        if row and row[0] is not None:
            return bytes(row[0])
    except Exception:
        return None
    return None

def lonlat_to_merc(lon: float, lat: float) -> tuple[float, float]:
    """Convert lon/lat (EPSG:4326) to WebMercator meters (EPSG:3857)."""
    R = 6378137.0
    lon_f = float(lon)
    lat_f = max(-85.05112878, min(85.05112878, float(lat)))
    x = lon_f * math.pi / 180.0 * R
    y = math.log(math.tan((90.0 + lat_f) * math.pi / 360.0)) * R
    return x, y


def merc_to_lonlat(x: float, y: float) -> tuple[float, float]:
    """Convert WebMercator meters (EPSG:3857) to lon/lat (EPSG:4326)."""
    R = 6378137.0
    lon = (float(x) / R) * 180.0 / math.pi
    lat = (2.0 * math.atan(math.exp(float(y) / R)) - math.pi / 2.0) * 180.0 / math.pi
    return lon, lat

def _nice_scale_length_m(target_m: float) -> float:
    """Pick a human-friendly scale bar length in meters (1/2/5 * 10^n) <= target."""
    try:
        t = float(target_m)
    except Exception:
        return 0.0
    if t <= 0:
        return 0.0
    exp = 10 ** math.floor(math.log10(t))
    for m in (5, 2, 1):
        val = m * exp
        if val <= t:
            return float(val)
    return float(exp)


def _style_inset_box(fig, inset_ax,
                     *,
                     border_color: str | None = None,
                     border_lw: float | None = None,
                     shadow_color: str = '#000000',
                     shadow_alpha: float | None = None,
                     shadow_dx: float | None = None,
                     shadow_dy: float | None = None):
    """Apply a consistent dark border + subtle drop shadow to an inset axes.

    The shadow is drawn as figure-level patches on purpose.
    Many report maps are saved with bbox_inches='tight'. If we draw the shadow outside
    the inset using path-effects, the shadow can be cropped away because tight-bbox
    calculation does not account for path-effects extents.

    To avoid the occasional "faulty"/misaligned shadow, we compute the inset position
    from the renderer after a canvas draw (in display coords), then transform to figure
    coords.
    """
    try:
        if fig is None or inset_ax is None:
            return
        # Resolve defaults from config-driven globals.
        if border_color is None:
            border_color = _REPORT_INSET_BORDER_COLOR
        if border_lw is None:
            border_lw = _REPORT_INSET_BORDER_LW
        if shadow_alpha is None:
            shadow_alpha = _REPORT_INSET_SHADOW_ALPHA
        if shadow_dx is None:
            shadow_dx = _REPORT_INSET_SHADOW_DX
        if shadow_dy is None:
            shadow_dy = _REPORT_INSET_SHADOW_DY

        # Compute bbox in figure coordinates.
        bbox = None
        try:
            # Ensure layout is finalized before reading window extents.
            canvas = getattr(fig, "canvas", None)
            if canvas is not None:
                try:
                    canvas.draw()
                except Exception:
                    pass
                try:
                    renderer = canvas.get_renderer()
                    bb_disp = inset_ax.get_window_extent(renderer=renderer)
                    bbox = bb_disp.transformed(fig.transFigure.inverted())
                except Exception:
                    bbox = None
        except Exception:
            bbox = None
        if bbox is None:
            bbox = inset_ax.get_position()

        # Shadow behind (slightly down/right)
        shadow_patch = Rectangle(
            (bbox.x0 + float(shadow_dx), bbox.y0 - float(shadow_dy)),
            bbox.width,
            bbox.height,
            transform=fig.transFigure,
            facecolor=shadow_color,
            edgecolor='none',
            alpha=float(shadow_alpha),
            zorder=15,
        )
        try:
            shadow_patch.set_clip_on(False)
            shadow_patch.set_in_layout(False)
        except Exception:
            pass
        fig.patches.append(shadow_patch)

        # Dark border in front
        border_patch = Rectangle(
            (bbox.x0, bbox.y0),
            bbox.width,
            bbox.height,
            transform=fig.transFigure,
            facecolor='none',
            edgecolor=border_color,
            linewidth=float(border_lw),
            alpha=0.95,
            zorder=25,
        )
        try:
            border_patch.set_clip_on(False)
            border_patch.set_in_layout(False)
        except Exception:
            pass
        fig.patches.append(border_patch)
    except Exception:
        return

def _add_map_decorations(ax,
                         extent_3857: tuple[float, float, float, float],
                         base_dir: str | None = None,
                         add_inset: bool = True,
                         inset_context_tiles_3857: gpd.GeoDataFrame | None = None):
    """Add north arrow (top-right), scale bar (bottom-left) and optional inset overview.

    inset_context_tiles_3857: optional collection of polygons (already in
    EPSG:3857) to draw inside the inset as weak context lines, plotted
    behind the highlighted current-extent rectangle. Used by the atlas
    tile maps so each tile inset shows where the current tile sits in
    relation to the rest of the atlas grid.
    """
    try:
        west_x, east_x, south_y, north_y = extent_3857
        width_m = float(east_x - west_x)
        height_m = float(north_y - south_y)
        if width_m <= 0 or height_m <= 0:
            return

        # North arrow (axes fraction coordinates)
        try:
            ax.annotate(
                "",
                xy=(0.95, 0.94),
                xytext=(0.95, 0.82),
                xycoords=ax.transAxes,
                textcoords=ax.transAxes,
                arrowprops=dict(arrowstyle='-|>', lw=1.6, color='black'),
                zorder=40,
            )
            ax.text(
                0.95, 0.95, "N",
                transform=ax.transAxes,
                ha='center', va='bottom',
                fontsize=12, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.15', fc='white', alpha=0.65, ec='none'),
                zorder=41,
            )
        except Exception:
            pass

        # Scale bar (data coordinates, bottom-left)
        try:
            pad_x = 0.04 * width_m
            pad_y = 0.04 * height_m
            target = 0.22 * width_m
            bar_len = _nice_scale_length_m(target)
            if bar_len > 0:
                x0 = west_x + pad_x
                y0 = south_y + pad_y
                x1 = x0 + bar_len
                tick_h = 0.012 * height_m
                ax.plot([x0, x1], [y0, y0], color='black', lw=3.0, zorder=40, solid_capstyle='butt')
                ax.plot([x0, x0], [y0, y0 + tick_h], color='black', lw=2.0, zorder=40)
                ax.plot([x1, x1], [y0, y0 + tick_h], color='black', lw=2.0, zorder=40)
                if bar_len >= 1000:
                    label = f"{bar_len/1000.0:g} km"
                else:
                    label = f"{bar_len:g} m"
                ax.text(
                    x0, y0 + tick_h * 1.4, label,
                    ha='left', va='bottom', fontsize=10,
                    bbox=dict(boxstyle='round,pad=0.15', fc='white', alpha=0.65, ec='none'),
                    zorder=41,
                )
        except Exception:
            pass

        # Inset overview (if available)
        if add_inset and inset_axes is not None:
            try:
                # Atlas tile maps (identified by the context-tiles parameter) use
                # a smaller inset so it doesn't cover too much of the tile content;
                # other map types keep the larger overview at 28%.
                inset_size = "18%" if inset_context_tiles_3857 is not None else "28%"
                iax = inset_axes(ax, width=inset_size, height=inset_size, loc='upper left', borderpad=0.8)
                try:
                    iax.set_zorder(20)
                except Exception:
                    pass
                iax.set_axis_off()
                # Consistent inset styling (dark frame + subtle shadow)
                _style_inset_box(ax.figure, iax)

                # Ensure the inset has a full-axes background so bbox tightening does not collapse.
                try:
                    iax.add_patch(Rectangle((0, 0), 1, 1, transform=iax.transAxes,
                                            facecolor='white', edgecolor='none', zorder=0))
                except Exception:
                    pass
                # Expand view to provide context
                pad_ix = 2.5 * width_m
                pad_iy = 2.5 * height_m
                iax.set_xlim(west_x - pad_ix, east_x + pad_ix)
                iax.set_ylim(south_y - pad_iy, north_y + pad_iy)
                _plot_basemap(iax, crs_epsg=3857, base_dir=base_dir)
                # Draw the rest of the atlas grid as weak context lines so the
                # reader can see how the highlighted tile sits relative to the
                # other tiles, not just the basemap.
                if inset_context_tiles_3857 is not None and not inset_context_tiles_3857.empty:
                    try:
                        inset_context_tiles_3857.boundary.plot(
                            ax=iax,
                            color="#555555",
                            linewidth=0.4,
                            alpha=0.55,
                            zorder=40,
                        )
                    except Exception:
                        pass
                rect = Rectangle((west_x, south_y), width_m, height_m,
                                 fill=False, edgecolor=PRIMARY_HEX, linewidth=1.6, alpha=0.95, zorder=50)
                iax.add_patch(rect)
            except Exception:
                pass
    except Exception:
        return

def render_mbtiles_to_png_best_fit(mbtiles_path: str,
                                  out_png: str,
                                  base_dir: str | None = None) -> tuple[bool, str]:
    """Stitch MBTiles into a single PNG for the MBTiles bounds.

    Chooses the highest zoom that fits within MAX_MAP_PX_WIDTH/MAX_MAP_PX_HEIGHT.
    """
    try:
        if not os.path.exists(mbtiles_path):
            return False, "mbtiles not found"
        conn = sqlite3.connect(mbtiles_path)
        try:
            meta = _mbtiles_metadata(conn)
            bounds = _parse_mbtiles_bounds(meta)
            if bounds is None:
                return False, "MBTiles missing bounds"
            minz = _parse_mbtiles_zoom(meta, "minzoom", 0)
            maxz = _parse_mbtiles_zoom(meta, "maxzoom", 19)
            tile_size = 256

            minlon, minlat, maxlon, maxlat = bounds

            # Pick best fit zoom
            chosen_z = maxz
            while chosen_z > minz:
                x_w, y_n = _lonlat_to_tile_fraction(minlon, maxlat, chosen_z)
                x_e, y_s = _lonlat_to_tile_fraction(maxlon, minlat, chosen_z)
                x0 = int(math.floor(min(x_w, x_e)))
                x1 = int(math.floor(max(x_w, x_e)))
                y0 = int(math.floor(min(y_n, y_s)))
                y1 = int(math.floor(max(y_n, y_s)))
                w_px = (x1 - x0 + 1) * tile_size
                h_px = (y1 - y0 + 1) * tile_size
                tiles = (x1 - x0 + 1) * (y1 - y0 + 1)
                if w_px <= MAX_MAP_PX_WIDTH and h_px <= MAX_MAP_PX_HEIGHT and tiles <= 400:
                    break
                chosen_z -= 1

            z = max(minz, chosen_z)

            x_w, y_n = _lonlat_to_tile_fraction(minlon, maxlat, z)
            x_e, y_s = _lonlat_to_tile_fraction(maxlon, minlat, z)
            x0 = int(math.floor(min(x_w, x_e)))
            x1 = int(math.floor(max(x_w, x_e)))
            y0 = int(math.floor(min(y_n, y_s)))
            y1 = int(math.floor(max(y_n, y_s)))

            width = (x1 - x0 + 1) * tile_size
            height = (y1 - y0 + 1) * tile_size
            if width <= 0 or height <= 0:
                return False, "Invalid tile bounds"

            canvas = PILImage.new("RGBA", (width, height), (0, 0, 0, 0))
            blank = PILImage.new("RGBA", (tile_size, tile_size), (0, 0, 0, 0))

            for xt in range(x0, x1 + 1):
                for yt in range(y0, y1 + 1):
                    tile_bytes = _get_mbtiles_tile_bytes(conn, z, xt, yt)
                    if tile_bytes:
                        try:
                            tile_img = PILImage.open(io.BytesIO(tile_bytes)).convert("RGBA")
                        except Exception:
                            tile_img = blank
                    else:
                        tile_img = blank
                    px = (xt - x0) * tile_size
                    py = (yt - y0) * tile_size
                    canvas.paste(tile_img, (px, py))

            # Crop to exact bounds
            left = int(round((min(x_w, x_e) - x0) * tile_size))
            right = int(round((max(x_w, x_e) - x0) * tile_size))
            top = int(round((min(y_n, y_s) - y0) * tile_size))
            bottom = int(round((max(y_n, y_s) - y0) * tile_size))
            left = max(0, min(width - 1, left))
            right = max(left + 1, min(width, right))
            top = max(0, min(height - 1, top))
            bottom = max(top + 1, min(height, bottom))
            canvas = canvas.crop((left, top, right, bottom))

            # Final safety resize
            if canvas.size[0] > MAX_MAP_PX_WIDTH or canvas.size[1] > MAX_MAP_PX_HEIGHT:
                ratio = min(MAX_MAP_PX_WIDTH / canvas.size[0], MAX_MAP_PX_HEIGHT / canvas.size[1])
                new_w = max(1, int(canvas.size[0] * ratio))
                new_h = max(1, int(canvas.size[1] * ratio))
                canvas = canvas.resize((new_w, new_h), PILImage.LANCZOS)

            # Composite over a basemap for context
            try:
                fig_h_in = 8.2
                fig_w_in = 6.2
                dpi = 180
                fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
                ax.set_axis_off()

                # Set extent first (basemap fetch depends on ax limits)
                west_x, north_y = lonlat_to_merc(minlon, maxlat)
                east_x, south_y = lonlat_to_merc(maxlon, minlat)
                ax.set_xlim(west_x, east_x)
                ax.set_ylim(south_y, north_y)
                try:
                    ax.set_aspect('equal', adjustable='box')
                except Exception:
                    pass

                # Background basemap in WebMercator
                _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

                # Overlay MBTiles image in correct extent
                ax.imshow(
                    canvas,
                    extent=[west_x, east_x, south_y, north_y],
                    origin="upper",
                    interpolation='nearest',
                    zorder=10,
                )
                _add_map_decorations(ax, (west_x, east_x, south_y, north_y), base_dir=base_dir, add_inset=True)
                # tight_layout frequently warns with inset axes; bbox_inches='tight' is sufficient here.
                plt.savefig(out_png, bbox_inches='tight')
                plt.close(fig)
                return True, ""
            except Exception:
                try:
                    plt.close('all')
                except Exception:
                    pass
                # Fallback: save the stitched overlay only
                canvas.save(out_png)
                return True, ""
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        write_to_log(f"MBTiles render failed ({mbtiles_path}): {e}", base_dir)
        return False, str(e)

# ---------------- tbl_flat (per-geocode maps) ----------------
def load_tbl_flat(parquet_dir: str, base_dir: str | None = None) -> gpd.GeoDataFrame:
    fp = os.path.join(parquet_dir, "tbl_flat.parquet")
    if not os.path.exists(fp):
        write_to_log("tbl_flat.parquet not found.", base_dir)
        return gpd.GeoDataFrame()
    try:
        gdf = gpd.read_parquet(fp)
        if 'geometry' not in gdf.columns or gdf.geometry.is_empty.all():
            write_to_log("tbl_flat has no valid geometry.", base_dir)
            return gpd.GeoDataFrame()
        return gdf
    except Exception as e:
        write_to_log(f"Failed reading tbl_flat: {e}", base_dir)
        return gpd.GeoDataFrame()

def compute_fixed_bounds_3857(flat_df: gpd.GeoDataFrame, base_dir: str | None = None):
    try:
        poly = flat_df[flat_df.geometry.type.isin(['Polygon','MultiPolygon'])].copy()
        if poly.empty:
            write_to_log("No polygons in tbl_flat to compute fixed bounds; using all geometry.", base_dir)
            poly = flat_df.copy()
        g3857 = _safe_to_3857(poly)
        if g3857.empty:
            return None
        b = g3857.total_bounds
        b = _expand_bounds(b, pad_ratio=0.08)
        return b
    except Exception as e:
        write_to_log(f"Failed computing fixed bounds: {e}", base_dir)
        return None

def draw_group_map_sensitivity(gdf_group: gpd.GeoDataFrame,
                               group_name: str,
                               palette: dict,
                               desc: dict,
                               out_path: str,
                               fixed_bounds_3857=None,
                               base_dir: str | None = None,
                               alpha: float = 0.95):
    try:
        g = gdf_group.copy()
        g = g[g.geometry.type.isin(['Polygon','MultiPolygon'])]
        if g.empty:
            write_to_log(f"[{group_name}] No polygon geometries for sensitivity map.", base_dir)
            return False

        g = _safe_to_3857(g)
        fig_h_in = 10.0  # square to preserve detail
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        if fixed_bounds_3857:
            minx, miny, maxx, maxy = fixed_bounds_3857
            ax.set_xlim(minx, maxx)
            ax.set_ylim(miny, maxy)

        # Draw basemap first, then polygons on top for clarity
        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        annotated = _prepare_sensitivity_annotations(g, 'sensitivity_code_max', 'sensitivity_max')
        if annotated.empty:
            write_to_log(f"[{group_name}] No sensitivity records to plot.", base_dir)
            plt.close(fig)
            return False

        colors = _colors_from_annotations(annotated, palette)
        annotated.plot(ax=ax,
                       color=colors,
                       edgecolor='none',
                       linewidth=0.0,
                       alpha=float(max(0.0, min(1.0, alpha))),
                       zorder=10)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=True)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Sensitivity map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Sensitivity map failed for {group_name}: {e}", base_dir)
        plt.close('all')
        return False

def draw_atlas_map(tile_row: pd.Series,
                   atlas_crs,
                   atlas_polys_3857: gpd.GeoDataFrame,
                   palette: dict,
                   desc: dict,
                   out_path: str,
                   global_bounds_3857=None,
                   base_dir: str | None = None,
                   all_tiles_3857: gpd.GeoDataFrame | None = None) -> bool:
    """
    Render a single atlas tile sensitivity map using shared cartography.
    """
    try:
        geom = tile_row.get('geometry', None)
        tile_name = tile_row.get('name_gis', '?')
        if geom is None or getattr(geom, "is_empty", True):
            write_to_log(f"[Atlas {tile_name}] Missing geometry; skipping map.", base_dir)
            return False

        tile_gdf = gpd.GeoDataFrame([{'geometry': geom}], geometry='geometry', crs=atlas_crs)
        if tile_gdf.crs is None:
            tile_gdf.set_crs(4326, inplace=True)
        tile_3857 = _safe_to_3857(tile_gdf)
        if tile_3857.empty or tile_3857.geometry.is_empty.all():
            write_to_log(f"[Atlas {tile_name}] Reprojection failed; skipping map.", base_dir)
            return False

        tile_bounds = _expand_bounds(tile_3857.total_bounds, pad_ratio=0.05)

        subset = gpd.GeoDataFrame()
        if atlas_polys_3857 is not None and not atlas_polys_3857.empty:
            try:
                mask = atlas_polys_3857.geometry.intersects(tile_3857.iloc[0].geometry)
                subset = atlas_polys_3857[mask].copy()
            except Exception as e:
                write_to_log(f"[Atlas {tile_name}] Intersection failed: {e}", base_dir)
                subset = gpd.GeoDataFrame()

        fig_w_in, fig_h_in = ATLAS_FIGURE_INCHES
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()
        ax.set_xlim(tile_bounds[0], tile_bounds[2])
        ax.set_ylim(tile_bounds[1], tile_bounds[3])

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        drew_data = False
        if not subset.empty and "sensitivity_code_max" in subset.columns:
            subset_ann = _prepare_sensitivity_annotations(subset, "sensitivity_code_max", "sensitivity_max")
            if subset_ann.empty:
                write_to_log(f"[Atlas {tile_name}] Sensitivity annotations empty.", base_dir)
            else:
                colors = _colors_from_annotations(subset_ann, palette)
                subset_ann.plot(ax=ax,
                                color=colors,
                                edgecolor="none",
                                linewidth=0.0,
                                alpha=0.75,
                                zorder=10)
                drew_data = True
        elif not subset.empty:
            subset.plot(ax=ax,
                        color=SENSITIVITY_UNKNOWN_COLOR,
                        edgecolor="none",
                        linewidth=0.0,
                        alpha=0.75,
                        zorder=10)
            drew_data = True
        if not drew_data:
            plt.close(fig)
            write_to_log(f"[Atlas {tile_name}] No drawable data for sensitivity map.", base_dir)
            return False

        tile_3857.boundary.plot(ax=ax, edgecolor=PRIMARY_HEX, linewidth=1.6, zorder=12)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(
                ax,
                (minx, maxx, miny, maxy),
                base_dir=base_dir,
                add_inset=True,
                inset_context_tiles_3857=all_tiles_3857,
            )
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Atlas map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Atlas map failed for {tile_row.get('name_gis', '?')}: {e}", base_dir)
        plt.close('all')
        return False

def draw_atlas_overview_map(atlas_df: gpd.GeoDataFrame,
                            atlas_crs,
                            atlas_polys_3857: gpd.GeoDataFrame,
                            out_path: str,
                            global_bounds_3857=None,
                            base_dir: str | None = None) -> bool:
    """
    Draw an overview map showing every atlas tile in relation to the full dataset extent.
    """
    try:
        if atlas_df.empty or 'geometry' not in atlas_df.columns:
            write_to_log("Atlas overview skipped (no geometries).", base_dir)
            return False

        atlas_gdf = atlas_df.copy()
        atlas_gdf = atlas_gdf[atlas_gdf['geometry'].notna()]
        if atlas_gdf.empty:
            write_to_log("Atlas overview skipped (all geometries empty).", base_dir)
            return False

        if atlas_crs is not None:
            try:
                atlas_gdf = atlas_gdf.set_crs(atlas_crs, allow_override=True)
            except Exception:
                pass
        atlas_3857 = _safe_to_3857(atlas_gdf)
        if atlas_3857.empty:
            write_to_log("Atlas overview reprojection failed; skipping.", base_dir)
            return False

        if global_bounds_3857:
            bounds = global_bounds_3857
        else:
            bounds = _expand_bounds(atlas_3857.total_bounds, pad_ratio=0.08)

        fig_h_in = 10.0
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()
        ax.set_xlim(bounds[0], bounds[2])
        ax.set_ylim(bounds[1], bounds[3])

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        if atlas_polys_3857 is not None and not atlas_polys_3857.empty:
            atlas_polys_3857.plot(ax=ax, facecolor="#d9d9d9", edgecolor='white',
                                 linewidth=0.15, alpha=0.5, zorder=8)

        atlas_3857.boundary.plot(ax=ax, edgecolor=PRIMARY_HEX, linewidth=1.0, alpha=0.9, zorder=12)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=False)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Atlas overview map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Atlas overview map failed: {e}", base_dir)
        plt.close('all')
        return False


def draw_data_extent_overview_map(extent_gdf: gpd.GeoDataFrame,
                                  out_path: str,
                                  base_dir: str | None = None) -> bool:
    """Render the project's overall data extent (precomputed in Stage 1 as
    tbl_data_extent.parquet) on a basemap. Used at the top of the report so
    a reader can see at a glance how big the AOI is and whether it has
    disjoint regions."""
    try:
        if extent_gdf is None or extent_gdf.empty or 'geometry' not in extent_gdf.columns:
            write_to_log("Data extent overview skipped (no geometry).", base_dir)
            return False

        extent_3857 = _safe_to_3857(extent_gdf[extent_gdf.geometry.notna()].copy())
        if extent_3857.empty:
            write_to_log("Data extent reprojection failed; skipping overview.", base_dir)
            return False

        bounds = _expand_bounds(extent_3857.total_bounds, pad_ratio=0.10)

        fig_h_in = 10.0
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()
        ax.set_xlim(bounds[0], bounds[2])
        ax.set_ylim(bounds[1], bounds[3])

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        # Translucent fill so the basemap stays legible inside the AOI.
        extent_3857.plot(ax=ax, facecolor=PRIMARY_HEX, edgecolor=PRIMARY_HEX,
                         linewidth=1.5, alpha=0.22, zorder=10)
        extent_3857.boundary.plot(ax=ax, edgecolor=PRIMARY_HEX,
                                  linewidth=1.5, alpha=0.95, zorder=12)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=False)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Data extent overview map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Data extent overview map failed: {e}", base_dir)
        plt.close('all')
        return False

# ---------------- Lines: context + segments map ----------------
def _normalize_bounds_aspect(bounds: tuple[float, float, float, float],
                             *,
                             min_aspect: float = 0.70,
                             max_aspect: float = 1.30) -> tuple[float, float, float, float]:
    """Expand bounds so that dy/dx stays within [min_aspect, max_aspect].

    This prevents very tall/narrow (or wide/flat) extents from creating maps that
    end up tiny in DOCX when clamped by max height.
    """
    minx, miny, maxx, maxy = bounds
    dx = float(maxx - minx)
    dy = float(maxy - miny)
    if dx <= 0 or dy <= 0:
        return bounds
    cx = (minx + maxx) / 2.0
    cy = (miny + maxy) / 2.0
    aspect = dy / dx

    # Too tall: widen X.
    if aspect > float(max_aspect):
        target_dx = dy / float(max_aspect)
        dx2 = target_dx / 2.0
        return (cx - dx2, miny, cx + dx2, maxy)

    # Too wide: widen Y.
    if aspect < float(min_aspect):
        target_dy = dx * float(min_aspect)
        dy2 = target_dy / 2.0
        return (minx, cy - dy2, maxx, cy + dy2)

    return bounds


def draw_line_segments_map_with_context_inset(segments_df: gpd.GeoDataFrame,
                                              line_name: str,
                                              palette: dict,
                                              out_path: str,
                                              mode: str = 'max',
                                              pad_ratio: float = 0.20,
                                              context_pad_ratio: float = 1.00,
                                              rect_buffer_ratio: float = 0.03,
                                              base_dir: str | None = None) -> bool:
    """Draw the segments map as the main map, with a context inset showing its location."""
    try:
        if segments_df.empty or 'geometry' not in segments_df.columns:
            write_to_log(f"[{line_name}] No geometry in segments; skipping segments map.", base_dir)
            return False

        col = 'sensitivity_code_max' if mode == 'max' else 'sensitivity_code_min'
        segs = segments_df[(segments_df['name_gis'] == line_name) & (segments_df['geometry'].notna())].copy()
        if segs.empty:
            write_to_log(f"[{line_name}] No segments found for segments map.", base_dir)
            return False

        segs = _safe_to_3857(segs)
        if segs.empty:
            write_to_log(f"[{line_name}] Segments reprojection failed.", base_dir)
            return False

        b = tuple(segs.total_bounds)
        # Main map bounds: pad, then normalize aspect so it doesn't become ultra-narrow.
        main_bounds = _expand_bounds(b, pad_ratio)
        main_bounds = _normalize_bounds_aspect(main_bounds, min_aspect=0.70, max_aspect=1.30)
        minx, miny, maxx, maxy = main_bounds

        fig_h_in = 10.0
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)

        # Ensure there's always a full-axes artist so bbox_inches='tight' doesn't
        # collapse into a thin strip when basemap fetching is unavailable.
        ax.add_patch(Rectangle((0, 0), 1, 1, transform=ax.transAxes,
                               facecolor='white', edgecolor='none', zorder=0))

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        polys = segs[segs.geometry.type.isin(['Polygon', 'MultiPolygon'])]
        lines = segs[segs.geometry.type.isin(['LineString', 'MultiLineString'])]
        value_col = 'sensitivity_max' if mode == 'max' else 'sensitivity_min'

        if not polys.empty:
            polys_ann = _prepare_sensitivity_annotations(polys, col, value_col)
            colors_polys = _colors_from_annotations(polys_ann, palette)
            polys_ann.plot(
                ax=ax,
                color=colors_polys,
                edgecolor='white',
                linewidth=0.4,
                alpha=0.95,
                zorder=12,
            )

        if not lines.empty:
            lines_ann = _prepare_sensitivity_annotations(lines, col, value_col)
            line_colors = _colors_from_annotations(lines_ann, palette)
            line_colors_rgba = [mcolors.to_rgba(c, alpha=1.0) for c in line_colors]
            try:
                lines.plot(ax=ax, color='white', linewidth=4.2, alpha=0.9, zorder=13)
            except Exception:
                pass
            lines.plot(ax=ax, color=line_colors_rgba, linewidth=2.4, alpha=1.0, zorder=14)

        try:
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=False)
        except Exception:
            pass

        # ---- Context inset (shows where the main map sits) ----
        try:
            inset_w = "34%"
            inset_h = "34%"
            if inset_axes is not None:
                axins = inset_axes(ax, width=inset_w, height=inset_h, loc="upper left", borderpad=0.8)
            else:
                # Fallback: manual inset in axes fraction
                axins = fig.add_axes([0.10, 0.62, 0.30, 0.30])
            # Ensure the inset sits above the main map layers.
            try:
                axins.set_zorder(20)
            except Exception:
                pass
            axins.set_axis_off()
            axins.add_patch(Rectangle((0, 0), 1, 1, transform=axins.transAxes,
                                      facecolor='white', edgecolor='none', zorder=0))

            ctx_bounds = _expand_bounds(b, context_pad_ratio)
            ctx_bounds = _normalize_bounds_aspect(ctx_bounds, min_aspect=0.70, max_aspect=1.30)
            cminx, cminy, cmaxx, cmaxy = ctx_bounds
            axins.set_xlim(cminx, cmaxx)
            axins.set_ylim(cminy, cmaxy)

            _plot_basemap(axins, crs_epsg=3857, base_dir=base_dir)

            try:
                segs.boundary.plot(ax=axins, edgecolor=PRIMARY_HEX, linewidth=1.0, alpha=0.95, zorder=8)
            except Exception:
                pass

            # Rectangle showing the main-map extent (slightly expanded)
            bx0, by0, bx1, by1 = main_bounds
            bw, bh = (bx1 - bx0), (by1 - by0)
            rx0 = bx0 - bw * rect_buffer_ratio
            ry0 = by0 - bh * rect_buffer_ratio
            rw = bw + 2 * bw * rect_buffer_ratio
            rh = bh + 2 * bh * rect_buffer_ratio
            rect = Rectangle((rx0, ry0), rw, rh, fill=False, edgecolor='red', linewidth=1.6, alpha=0.95, zorder=20)
            axins.add_patch(rect)

            _style_inset_box(fig, axins)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight', pad_inches=0.02)
        plt.close(fig)
        write_to_log(f"Segments map ({mode}) with inset saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Segments map with inset failed for {line_name}: {e}", base_dir)
        plt.close('all')
        return False

# ---------------- Lines: overview across all lines ----------------
def draw_lines_overview_map(segments_df: gpd.GeoDataFrame,
                            palette: dict,
                            out_path: str,
                            mode: str = 'max',
                            pad_ratio: float = 0.08,
                            base_dir: str | None = None) -> bool:
    """Draw a single overview map for all line segments, colored by sensitivity (max or min)."""
    try:
        if segments_df is None or segments_df.empty or 'geometry' not in segments_df.columns:
            write_to_log("Lines overview map skipped (no segment geometries).", base_dir)
            return False

        col = 'sensitivity_code_max' if mode == 'max' else 'sensitivity_code_min'
        value_col = 'sensitivity_max' if mode == 'max' else 'sensitivity_min'
        segs = segments_df[segments_df['geometry'].notna()].copy()
        if segs.empty:
            write_to_log("Lines overview map skipped (all segment geometries empty).", base_dir)
            return False

        segs = _safe_to_3857(segs)
        if segs.empty:
            write_to_log("Lines overview map reprojection failed; skipping.", base_dir)
            return False

        fig_h_in = 10.0
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        b = segs.total_bounds
        minx, miny, maxx, maxy = _expand_bounds(b, pad_ratio)
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        polys = segs[segs.geometry.type.isin(['Polygon','MultiPolygon'])]
        lines = segs[segs.geometry.type.isin(['LineString','MultiLineString'])]

        if not polys.empty:
            polys_ann = _prepare_sensitivity_annotations(polys, col, value_col)
            colors_polys = _colors_from_annotations(polys_ann, palette)
            polys_ann.plot(ax=ax,
                           color=colors_polys,
                           edgecolor='white',
                           linewidth=0.35,
                           alpha=0.90,
                           zorder=12)

        if not lines.empty:
            lines_ann = _prepare_sensitivity_annotations(lines, col, value_col)
            line_colors = _colors_from_annotations(lines_ann, palette)
            line_colors_rgba = [mcolors.to_rgba(c, alpha=1.0) for c in line_colors]
            try:
                lines.plot(ax=ax, color='white', linewidth=4.0, alpha=0.90, zorder=13)
            except Exception:
                pass
            lines.plot(ax=ax, color=line_colors_rgba, linewidth=2.2, alpha=1.0, zorder=14)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=True)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Lines overview map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Lines overview map failed: {e}", base_dir)
        plt.close('all')
        return False

# ---------------- Existing (kept / refined) ----------------
def get_color_from_code(code, color_codes):
    return _resolve_sensitivity_color(code, None, color_codes)

def sort_segments_numerically(segments):
    def extract_number(segment_id):
        m = re.search(r'_(\d+)$', str(segment_id))
        return int(m.group(1)) if m else float('inf')
    s = segments.copy()
    s['sort_key'] = s['segment_id'].apply(extract_number)
    return s.sort_values(by='sort_key').drop(columns=['sort_key'])

def create_line_statistic_image(line_name, sensitivity_series, color_codes, length_m, output_path):
    """
    Generate ribbon PNG with NO axes/ticks/labels.
    Axes fill the full figure area; saved with bbox_inches=None and pad_inches=0
    so every PNG has an identical pixel canvas regardless of content.
    """
    segment_count = max(1, len(sensitivity_series))
    fig_h_in = 0.236  # ~6 mm
    dpi = 300
    fig_w_in = 12
    fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)

    # Fill figure with axes; no margins
    ax.set_position([0, 0, 1, 1])
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.axis('off')

    for i, code in enumerate(sensitivity_series):
        color = get_color_from_code(code, color_codes)
        ax.add_patch(plt.Rectangle((i/segment_count, 0), 1/segment_count, 1, color=color))

    # Save with constant canvas (no tight bbox, no padding)
    plt.savefig(output_path, dpi=dpi, bbox_inches=None, pad_inches=0)
    plt.close(fig)

def resize_image(image_path, max_width_px, max_height_px):
    """
    Resample the stored image file to be within (max_width_px, max_height_px) in *pixels*.
    """
    try:
        with PILImage.open(image_path) as img:
            width, height = img.size
            ratio = min(float(max_width_px) / width, float(max_height_px) / height)
            if ratio >= 1.0:
                return  # already within limits
            new_w, new_h = max(1, int(width * ratio)), max(1, int(height * ratio))
            img = img.resize((new_w, new_h), PILImage.LANCZOS)
            img.save(image_path)
    except Exception as e:
        write_to_log(f"Image resize failed for {image_path}: {e}")

def create_sensitivity_summary(sensitivity_series, color_codes, output_path):
    counts = sensitivity_series.value_counts().reindex(['A','B','C','D','E']).fillna(0)
    total = max(1, len(sensitivity_series))
    fallback_color = color_codes.get('UNKNOWN', SENSITIVITY_UNKNOWN_COLOR)

    fig, (ax_text, ax_bar) = plt.subplots(2, 1, figsize=(10, 1.7), height_ratios=[0.45, 0.55])
    plt.subplots_adjust(hspace=0.15)

    parts = [f"{c}: {int(counts[c])} ({counts[c]/total*100:.1f}%)" for c in ['A','B','C','D','E']]
    summary = "Distribution: " + " | ".join(parts)
    ax_text.text(0.5, 0.5, summary, ha='center', va='center', fontsize=10)
    ax_text.axis('off')

    left = 0
    for c in ['A','B','C','D','E']:
        w = counts[c] / total
        ax_bar.barh(0, w, left=left, color=color_codes.get(c, fallback_color), edgecolor='white')
        left += w
    ax_bar.set_xlim(0, 1); ax_bar.set_ylim(-0.5, 0.5); ax_bar.axis('off')

    plt.savefig(output_path, bbox_inches='tight', dpi=110)
    plt.close(fig)

def _build_index_legend_png(
    output_path: str,
    *,
    kind: str,
    cfg_path: str | Path,
    caption: str = "",
) -> bool:
    """Render a horizontal colour-bar legend matching the MBTiles index ramp.

    kind: "importance" -> green ramp (matches importance_max).
          "sensitivity" / "owa" -> A->E warm ramp (red..yellow), low to high.

    The gradient is rebuilt from the same source as
    code/tiles_create_raster.py so the legend stays in sync with the actual
    map colours when the config palette is edited.
    """
    try:
        from tiles_create_raster import (
            INDEX_GRADIENT_STEPS,
            IMPORTANCE_MAX_HEX,
            build_importance_max_palette,
            build_index_gradient_from_palette,
            hex_to_rgba,
            read_sensitivity_palette_from_config,
        )
    except Exception as exc:
        try:
            print(f"[report] legend builder: cannot import gradient helpers: {exc}", flush=True)
        except Exception:
            pass
        return False

    try:
        if str(kind).lower() == "importance":
            gradient = build_index_gradient_from_palette(
                build_importance_max_palette(alpha=1.0),
                steps=INDEX_GRADIENT_STEPS,
                order=[1, 2, 3, 4, 5],
                fallback=hex_to_rgba(IMPORTANCE_MAX_HEX[1], 1.0),
            )
        else:
            gradient = build_index_gradient_from_palette(
                read_sensitivity_palette_from_config(Path(cfg_path), alpha=1.0),
                steps=INDEX_GRADIENT_STEPS,
            )

        rgb = np.array(
            [[c[0] / 255.0, c[1] / 255.0, c[2] / 255.0] for c in gradient],
            dtype=float,
        )
        bar = rgb.reshape(1, -1, 3)

        # Render at 13 cm wide x ~1.4 cm tall so it sits as a thin strip below
        # the 13 cm-wide map without dominating the page.
        fig_w_in = 13.0 / 2.54
        fig_h_in = 1.4 / 2.54
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=200)
        ax.imshow(bar, aspect="auto", extent=(0.0, 100.0, 0.0, 1.0))
        ax.set_yticks([])
        ax.set_xticks([0, 25, 50, 75, 100])
        ax.set_xticklabels([
            "0\n(no overlap)",
            "25",
            "50",
            "75",
            "100\n(highest)",
        ])
        ax.tick_params(axis="x", labelsize=6, length=2, pad=1)
        for spine in ax.spines.values():
            spine.set_visible(False)
        if caption:
            ax.set_xlabel(caption, fontsize=7)

        fig.subplots_adjust(left=0.02, right=0.98, top=0.92, bottom=0.32)
        fig.savefig(output_path, bbox_inches="tight", dpi=200)
        plt.close(fig)
        return os.path.exists(output_path) and os.path.getsize(output_path) > 0
    except Exception as exc:
        try:
            print(f"[report] legend builder failed: {exc}", flush=True)
        except Exception:
            pass
        try:
            plt.close("all")
        except Exception:
            pass
        return False


def create_index_area_distribution_chart(flat_df: gpd.GeoDataFrame,
                                        index_col: str,
                                        output_path: str,
                                        basic_group_name: str = "basic_mosaic",
                                        base_dir: str | None = None) -> tuple[bool, str]:
    """Render a chart of total area per index value for basic_mosaic only.

    Bars: total polygon area (km²) per index value (0..100)
    Line (2nd axis): number of "categories" within each index value.
      - If sensitivity_code_max exists: distinct A–E count per index value.
      - Else: number of geocode cells per index value.
    Returns (ok, note).
    """
    try:
        if flat_df is None or flat_df.empty:
            return False, "tbl_flat is empty."
        if index_col not in flat_df.columns:
            return False, f"Missing column: {index_col}."
        if "area_m2" not in flat_df.columns:
            return False, "Missing column: area_m2."

        df = flat_df.copy()
        if "name_gis_geocodegroup" in df.columns and basic_group_name:
            mask = df["name_gis_geocodegroup"].astype("string").str.strip().str.lower() == str(basic_group_name).strip().lower()
            df = df[mask].copy()

        if df.empty:
            return False, f"No rows for geocode group '{basic_group_name}'."

        df["__idx__"] = pd.to_numeric(df[index_col], errors="coerce")
        df["__idx__"] = df["__idx__"].round().clip(0, 100)
        df["__area_m2__"] = pd.to_numeric(df["area_m2"], errors="coerce")
        df = df.dropna(subset=["__idx__", "__area_m2__"])
        if df.empty:
            return False, "No numeric index/area rows to chart."

        df["__idx__"] = df["__idx__"].astype(int)

        # Area (km²) per index value (0..100)
        area_by = df.groupby("__idx__")["__area_m2__"].sum()
        x = np.arange(0, 101)
        area_km2 = np.array([float(area_by.get(i, 0.0)) / 1e6 for i in x], dtype=float)

        # "Categories" per index value
        cat_label = "# categories"
        if "sensitivity_code_max" in df.columns:
            codes = df["sensitivity_code_max"].astype("string").str.strip().str.upper()
            tmp = pd.DataFrame({"idx": df["__idx__"], "code": codes})
            # count distinct A–E per idx (ignore blanks)
            tmp = tmp[tmp["code"].isin(["A", "B", "C", "D", "E"])].copy()
            cats_by = tmp.groupby("idx")["code"].nunique()
            cats = np.array([int(cats_by.get(i, 0)) for i in x], dtype=int)
        else:
            cat_label = "# cells"
            cnt_by = df.groupby("__idx__").size()
            cats = np.array([int(cnt_by.get(i, 0)) for i in x], dtype=int)

        total_area_km2 = float(area_km2.sum())
        nonzero_values = int(np.count_nonzero(area_km2 > 0))

        fig, ax1 = plt.subplots(figsize=(11.0, 6.2), dpi=140)
        ax2 = ax1.twinx()

        ax1.bar(x, area_km2, color=PRIMARY_HEX, alpha=0.85, width=1.0, linewidth=0)
        ax2.plot(x, cats, color="#222222", linewidth=1.5)

        ax1.set_xlim(0, 100)
        ax1.set_xlabel("Index value")
        ax1.set_ylabel("Total area (km²)")
        ax2.set_ylabel(cat_label)

        ax1.set_title(f"{index_col} distribution (basic_mosaic)\nTotal area: {total_area_km2:.1f} km² | Values present: {nonzero_values}")

        ax1.grid(True, axis="y", linestyle=":", alpha=0.35)
        ax1.set_xticks(np.arange(0, 101, 10))

        # Tighten y2 for category scale
        try:
            ax2.set_ylim(0, max(1, int(cats.max()) + 1))
        except Exception:
            pass

        fig.tight_layout()
        plt.savefig(output_path, bbox_inches="tight")
        plt.close(fig)
        return True, ""
    except Exception as e:
        try:
            plt.close('all')
        except Exception:
            pass
        write_to_log(f"Index distribution chart failed for {index_col}: {e}", base_dir)
        return False, str(e)

def debug_atlas_sample(base_dir: str,
                       cfg: configparser.ConfigParser,
                       palette: dict[str, str],
                       desc: dict[str, str],
                       tile_name: str,
                       sample_size: int | None = None) -> str | None:
    """
    Render a quick atlas sensitivity map for a single tile using an optional
    subset of polygons. Helpful when debugging palette application.
    """
    gpq_dir = parquet_dir_from_cfg(base_dir, cfg)
    flat_df = load_tbl_flat(gpq_dir, base_dir=base_dir)
    if flat_df.empty:
        write_to_log("debug_atlas_sample: tbl_flat missing or empty.", base_dir)
        return None

    atlas_path = os.path.join(gpq_dir, "tbl_atlas.parquet")
    if not os.path.exists(atlas_path):
        write_to_log("debug_atlas_sample: tbl_atlas.parquet not found.", base_dir)
        return None

    try:
        atlas_df = gpd.read_parquet(atlas_path)
    except Exception as e:
        write_to_log(f"debug_atlas_sample: failed to read tbl_atlas ({e})", base_dir)
        return None

    if atlas_df.empty or 'geometry' not in atlas_df.columns:
        write_to_log("debug_atlas_sample: atlas table empty or missing geometry.", base_dir)
        return None

    tile_row = atlas_df[atlas_df['name_gis'] == tile_name]
    if tile_row.empty:
        write_to_log(f"debug_atlas_sample: tile '{tile_name}' not found in atlas.", base_dir)
        return None
    tile_row = tile_row.iloc[0]

    flat_polys = flat_df[flat_df.geometry.type.isin(['Polygon','MultiPolygon'])].copy()
    if flat_polys.empty:
        write_to_log("debug_atlas_sample: no polygon geometries in tbl_flat.", base_dir)
        return None

    if 'name_gis_geocodegroup' in flat_polys.columns:
        flat_polys['name_gis_geocodegroup'] = (
            flat_polys['name_gis_geocodegroup']
            .astype('string')
            .str.strip()
        )
        mask_lvl = flat_polys['name_gis_geocodegroup'].astype('string').str.lower() == 'basic_mosaic'
        filtered = flat_polys[mask_lvl].copy()
        if filtered.empty:
            write_to_log("debug_atlas_sample: no basic_mosaic polygons found; using full set.", base_dir)
        else:
            flat_polys = filtered

    flat_polys_3857 = _safe_to_3857(flat_polys)
    if flat_polys_3857.empty:
        write_to_log("debug_atlas_sample: reprojection yielded empty polygon set.", base_dir)
        return None

    tile_gdf = gpd.GeoDataFrame([tile_row], geometry='geometry', crs=atlas_df.crs)
    tile_3857 = _safe_to_3857(tile_gdf)
    if tile_3857.empty or tile_3857.geometry.is_empty.all():
        write_to_log(f"debug_atlas_sample: tile '{tile_name}' reprojection failed.", base_dir)
        return None

    try:
        intersects_mask = flat_polys_3857.geometry.intersects(tile_3857.iloc[0].geometry)
    except Exception as e:
        write_to_log(f"debug_atlas_sample: intersection failed ({e}).", base_dir)
        return None

    subset = flat_polys_3857[intersects_mask].copy()
    if subset.empty:
        write_to_log(f"debug_atlas_sample: no polygons intersect tile '{tile_name}'.", base_dir)
        return None

    if sample_size and sample_size > 0 and len(subset) > sample_size:
        subset = subset.sample(sample_size, random_state=0)
        write_to_log(f"debug_atlas_sample: down-sampled to {len(subset)} polygons.", base_dir)
    else:
        write_to_log(f"debug_atlas_sample: using {len(subset)} polygons.", base_dir)

    counts = _prepare_sensitivity_annotations(subset, 'sensitivity_code_max', 'sensitivity_max')['__sens_code'].value_counts(dropna=False)
    write_to_log(f"debug_atlas_sample sensitivity distribution: {counts.to_dict()}", base_dir)

    bounds = compute_fixed_bounds_3857(flat_df, base_dir=base_dir)
    out_path = output_subpath(base_dir, "tmp", f"debug_atlas_{_safe_name(tile_name)}.png")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    ok = draw_atlas_map(tile_row, atlas_df.crs, subset, palette, desc, out_path, bounds, base_dir=base_dir)
    if ok:
        write_to_log(f"debug_atlas_sample: saved {out_path}", base_dir)
        return out_path
    return None

def fetch_asset_group_statistics(asset_group_df: gpd.GeoDataFrame, asset_object_df: gpd.GeoDataFrame):
    """
    Returns per sensitivity code:
      - Number of Asset Objects
      - Active asset groups (distinct groups having ≥1 object)
    """
    if asset_group_df.empty:
        return pd.DataFrame(columns=['Sensitivity Code','Sensitivity Description','Active asset groups','Number of Asset Objects'])

    grp = asset_group_df[['id','sensitivity_code','sensitivity_description']].copy()

    if not asset_object_df.empty and 'ref_asset_group' in asset_object_df.columns:
        cnt = asset_object_df.groupby('ref_asset_group').size().rename('asset_objects_nr')
        merged = grp.merge(cnt, left_on='id', right_index=True, how='left')
    else:
        merged = grp.copy()
        merged['asset_objects_nr'] = 0

    merged['asset_objects_nr'] = merged['asset_objects_nr'].fillna(0).astype(int)

    active_groups = (merged[merged['asset_objects_nr'] > 0]
                     .groupby(['sensitivity_code','sensitivity_description'])
                     .agg(active_groups=('id','nunique'))
                     .reset_index())

    objects_per_code = (merged.groupby(['sensitivity_code','sensitivity_description'])
                        .agg(asset_objects=('asset_objects_nr','sum'))
                        .reset_index())

    out = objects_per_code.merge(active_groups,
                                 on=['sensitivity_code','sensitivity_description'],
                                 how='left').fillna({'active_groups':0})

    out = out.rename(columns={
        'sensitivity_code':'Sensitivity Code',
        'sensitivity_description':'Sensitivity Description',
        'active_groups':'Active asset groups',
        'asset_objects':'Number of Asset Objects'
    }).sort_values('Sensitivity Code')

    # Column order is read by the docx renderer straight from the DataFrame, so
    # match the empty-stub schema (Code, Description, Active asset groups,
    # Number of Asset Objects).
    out = out[['Sensitivity Code', 'Sensitivity Description',
               'Active asset groups', 'Number of Asset Objects']]
    return out

def format_area(row):
    if row['geometry_type'] in ['Polygon', 'MultiPolygon']:
        try:
            area = float(row['total_area'])
        except Exception:
            return "-"
        return _format_area_m2(area)
    return "-"


def _format_area_m2(area_m2: float) -> str:
    """Format an area in m² as either m² (<0.1 km²) or km² (max 4 decimals)."""
    try:
        a = float(area_m2)
    except Exception:
        return "-"
    if not math.isfinite(a) or a < 0:
        return "-"
    # Use m² for small areas (<1 km² == 1,000,000 m²)
    if a < 1_000_000:
        return f"{a:.0f} m²"
    km2 = a / 1_000_000.0
    s = f"{km2:.4f}"  # max 4 decimals
    # Trim trailing zeros while keeping <= 4 decimals.
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return f"{s} km²"


def _format_area_km2(area_km2: float) -> str:
    """Format an area in km² as either m² (<1 km²) or km² (max 4 decimals)."""
    try:
        km2 = float(area_km2)
    except Exception:
        return "-"
    if not math.isfinite(km2) or km2 < 0:
        return "-"
    if km2 < 1.0:
        return _format_area_m2(km2 * 1_000_000.0)
    s = f"{km2:.4f}"  # max 4 decimals
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return f"{s} km²"


def _analysis_assets_within_area_table(
    asset_objects_df: gpd.GeoDataFrame,
    asset_groups_df: gpd.GeoDataFrame,
    area_polys_gdf: gpd.GeoDataFrame,
) -> list[list[str]] | None:
    """Return a table of polygon asset area inside the given analysis area polygon(s)."""
    try:
        if asset_objects_df is None or asset_objects_df.empty:
            return None
        if asset_groups_df is None or asset_groups_df.empty:
            return None
        if area_polys_gdf is None or area_polys_gdf.empty or 'geometry' not in area_polys_gdf.columns:
            return None
        if 'ref_asset_group' not in asset_objects_df.columns:
            return None

        base = asset_groups_df.drop(columns=[c for c in ['geometry'] if c in asset_groups_df.columns])
        ao = asset_objects_df.merge(base, left_on='ref_asset_group', right_on='id', how='left')
        if 'geometry' not in ao.columns:
            return None

        ao = ao[ao.geometry.notna()].copy()
        if ao.empty:
            return None
        ao['geometry_type'] = ao.geometry.type
        ao_poly = ao[ao['geometry_type'].isin(['Polygon', 'MultiPolygon'])].copy()
        if ao_poly.empty:
            return None

        # Ensure CRS, then compute intersection areas in an equal-area projection.
        if ao_poly.crs is None:
            ao_poly = ao_poly.set_crs(4326)
        ap = area_polys_gdf.dropna(subset=['geometry']).copy()
        if ap.empty:
            return None
        if ap.crs is None:
            ap = ap.set_crs(ao_poly.crs)

        try:
            ao_poly = ao_poly.to_crs('ESRI:54009')
            ap = ap.to_crs('ESRI:54009')
        except Exception:
            # If reprojection fails, proceed in current CRS (areas may be less accurate)
            pass

        try:
            if hasattr(ap, "union_all"):
                area_union = ap.union_all()
            else:
                area_union = ap.unary_union
        except Exception:
            area_union = None
        if area_union is None or getattr(area_union, 'is_empty', True):
            return None

        # Fast pre-filter by intersection
        try:
            ao_poly = ao_poly[ao_poly.intersects(area_union)].copy()
        except Exception:
            pass
        if ao_poly.empty:
            return None

        area_gdf = gpd.GeoDataFrame({'geometry': [area_union]}, geometry='geometry', crs=ao_poly.crs)
        try:
            inter = gpd.overlay(ao_poly, area_gdf, how='intersection')
        except Exception:
            # Fallback: compute per-geometry intersection without overlay
            try:
                ao_poly['__geom_i'] = ao_poly.geometry.apply(lambda g: g.intersection(area_union))
                inter = ao_poly.drop(columns=['geometry']).copy()
                inter = gpd.GeoDataFrame(inter, geometry='__geom_i', crs=ao_poly.crs)
                inter = inter.rename_geometry('geometry')
            except Exception:
                return None

        if inter is None or inter.empty:
            return None

        try:
            inter['area_m2'] = inter.geometry.area
        except Exception:
            return None

        title_col = 'title_fromuser' if 'title_fromuser' in inter.columns else None
        if title_col is None:
            return None

        agg = (inter.groupby([title_col], dropna=False)
               .agg(area_m2=('area_m2', 'sum'), objects=('geometry', 'size'))
               .reset_index())
        if agg.empty:
            return None

        agg = agg.sort_values('area_m2', ascending=False)

        table: list[list[str]] = [["Asset", "Area within area", "# objects (polygons)"]]
        for _, r in agg.iterrows():
            title = str(r[title_col]) if not pd.isna(r[title_col]) else "(Unknown)"
            area_str = _format_area_m2(float(r['area_m2']))
            try:
                obj_n = int(r['objects'])
            except Exception:
                obj_n = 0
            table.append([title, area_str, str(obj_n)])
        return table
    except Exception:
        return None

def calculate_group_statistics(asset_object_df: gpd.GeoDataFrame, asset_group_df: gpd.GeoDataFrame):
    if asset_object_df.empty or asset_group_df.empty:
        return pd.DataFrame(columns=['Title','Code','Description','Type','Total area','# objects'])
    base = asset_group_df.drop(columns=[c for c in ['geometry'] if c in asset_group_df.columns])
    gdf = asset_object_df.merge(base, left_on='ref_asset_group', right_on='id', how='left')
    gdf['geometry_type'] = gdf.geometry.type
    try:
        gdf = gdf.to_crs('ESRI:54009')
    except Exception:
        pass

    gdf = gdf.copy()
    gdf['area'] = np.nan
    poly_mask = gdf['geometry_type'].isin(['Polygon', 'MultiPolygon'])
    if poly_mask.any():
        try:
            gdf.loc[poly_mask, 'area'] = gdf.loc[poly_mask, 'geometry'].area
        except Exception:
            gdf.loc[poly_mask, 'area'] = gdf.loc[poly_mask, 'geometry'].area
    gdf2 = gdf

    stats = gdf2.groupby(
        ['title_fromuser','sensitivity_code','sensitivity_description','geometry_type'],
        dropna=False
    ).agg(total_area=('area','sum'), object_count=('geometry','size')).reset_index()

    stats['total_area'] = stats['total_area'].fillna(0)
    stats['total_area'] = stats.apply(format_area, axis=1)
    stats = stats.rename(columns={
        'title_fromuser':'Title',
        'sensitivity_code':'Code',
        'sensitivity_description':'Description',
        'geometry_type':'Type',
        'total_area':'Total area',
        'object_count':'# objects'
    }).sort_values('Code')
    return stats

def export_to_excel(df, fp):
    if not fp.lower().endswith('.xlsx'):
        fp += '.xlsx'
    df.to_excel(fp, index=False)


def build_geocode_overview_table(parquet_dir: str) -> list | None:
    """Return a list-of-rows (header + body) summarising every geocode group.

    One row per geocode group (basic_mosaic, H3_R6, ..., custom grids) with
    object count drawn from tbl_geocode_object.parquet. Friendly title and
    description come from tbl_geocode_group.parquet when available. Returns
    None when neither source file is available, so callers can skip the
    section cleanly.
    """
    group_pq  = os.path.join(parquet_dir, "tbl_geocode_group.parquet")
    object_pq = os.path.join(parquet_dir, "tbl_geocode_object.parquet")

    counts = {}
    try:
        if os.path.exists(object_pq):
            try:
                gdf = gpd.read_parquet(object_pq)
            except Exception:
                gdf = pd.read_parquet(object_pq)
            if gdf is not None and not gdf.empty and 'name_gis_geocodegroup' in gdf.columns:
                series = gdf.groupby('name_gis_geocodegroup').size()
                counts = {str(k): int(v) for k, v in series.items()}
    except Exception:
        counts = {}

    meta_rows = []
    try:
        if os.path.exists(group_pq):
            try:
                meta = pd.read_parquet(group_pq)
            except Exception:
                meta = None
            if meta is not None and not meta.empty:
                cols = set(meta.columns)
                for _, r in meta.iterrows():
                    name = str(r.get('name_gis_geocodegroup', '')).strip()
                    if not name:
                        continue
                    title = str(r.get('title_user', '')).strip() if 'title_user' in cols else ''
                    description = str(r.get('description', '')).strip() if 'description' in cols else ''
                    meta_rows.append((name, title, description))
    except Exception:
        meta_rows = []

    seen = {name for name, _t, _d in meta_rows}
    for name in counts.keys():
        if name not in seen:
            meta_rows.append((name, '', ''))

    if not meta_rows:
        return None

    meta_rows.sort(key=lambda t: (0 if t[0].lower() == 'basic_mosaic' else 1, t[0]))
    rows = [["Geocode group", "Title", "Description", "# objects"]]
    for name, title, description in meta_rows:
        rows.append([name, title, description, counts.get(name, 0)])
    return rows

def fetch_lines_and_segments(parquet_dir: str):
    lines_pq    = os.path.join(parquet_dir, "tbl_lines.parquet")
    segments_pq = os.path.join(parquet_dir, "tbl_segment_flat.parquet")
    try:
        lines_df = gpd.read_parquet(lines_pq) if os.path.exists(lines_pq) else gpd.GeoDataFrame()
    except Exception:
        lines_df = gpd.GeoDataFrame()
    try:
        segments_df = gpd.read_parquet(segments_pq) if os.path.exists(segments_pq) else gpd.GeoDataFrame()
    except Exception:
        segments_df = gpd.GeoDataFrame()
    return lines_df, segments_df

def _clean_docx_text(txt: str) -> str:
    if txt is None:
        return ""
    s = str(txt)
    s = s.replace("<br/>", "\n").replace("<br>", "\n")
    s = re.sub(r"</?b>", "", s)
    s = re.sub(r"</?i>", "", s)
    s = re.sub(r"</?code>", "", s)
    # HTML entities like &ldquo; were leaking through into the rendered DOCX
    # because we strip tags but never decode entities. Decode after tag removal
    # so any &amp;, &ldquo;, &rdquo;, &hellip;, etc. become their Unicode form.
    import html as _html
    s = _html.unescape(s)
    return s

def compile_docx(output_docx: str, order_list: list):
    """Export the report to a Word document (.docx).

    Notes:
    - The table of contents (TOC) is a Word field and updates when opened/updated in Word.
    - Only real section headings (from 'heading(n)') use Word Heading styles.
      Image/table titles are bold paragraphs so they do NOT appear in the TOC.
    """
    doc = Document()
    doc.core_properties.title = "MESA Report"
    body_style = doc.styles["Normal"]
    body_style.font.name = "Calibri"
    body_style.font.size = Pt(10)
    try:
        body_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    except Exception:
        pass

    MAX_IMAGE_WIDTH_CM = 16.0
    # Atlas tiles previously rendered up to 12 x 18 cm (full A4 page minus
    # margins) and forced their heading + description onto a separate page.
    # Cap the height so heading + intro text + tile fit on one page; width
    # is still scaled by ATLAS_DOC_WIDTH_SCALE * MAX_IMAGE_WIDTH_CM.
    MAX_ATLAS_HEIGHT_CM = 12.0
    # Line maps are inserted together with distribution + ribbon.
    # Allow a taller cap now that we use a single combined map (segments + context inset).
    MAX_LINE_MAP_HEIGHT_CM = 14.0
    # Overview / index / "other maps" area maps used to fill the full text width
    # at 10 cm tall, which forced their explanatory text onto a separate page.
    # Shrink them to a panel size so heading + intro paragraphs + map fit on one
    # page together. Width is independent of MAX_IMAGE_WIDTH_CM for this kind.
    MAX_OVERVIEW_MAP_WIDTH_CM = 13.0
    MAX_OVERVIEW_MAP_HEIGHT_CM = 9.5
    # Index legends are emitted right after the map and should align visually
    # with it, so they share the overview width.
    MAX_LEGEND_WIDTH_CM = MAX_OVERVIEW_MAP_WIDTH_CM

    def _is_line_map_image(title: str | None, path: str | None) -> bool:
        if not path:
            return False
        base = os.path.basename(str(path)).lower()
        if "_context" in base or base.endswith("context.png") or "context" in base:
            return True
        if "segments_" in base or "_segments" in base or "segments" in base:
            return True
        if title:
            t = str(title).lower()
            if "geographical context" in t:
                return True
            if "segments" in t and "sensitivity" in t:
                return True
        return False

    def _docx_add_picture_with_height_clamp(img_path: str, *, target_width_cm: float, max_height_cm: float):
        """Add a picture at up to target_width_cm, but reduce width if needed to keep height <= max_height_cm."""
        out_w_cm = float(target_width_cm)
        try:
            with PILImage.open(img_path) as im:
                w_px, h_px = im.size
            if w_px and h_px and w_px > 0 and h_px > 0:
                aspect = float(h_px) / float(w_px)
                out_h_cm = out_w_cm * aspect
                if out_h_cm > float(max_height_cm):
                    out_h_cm = float(max_height_cm)
                    out_w_cm = out_h_cm / aspect
        except Exception:
            # If we can't read image metadata, fall back to the target width.
            out_w_cm = float(target_width_cm)
        doc.add_picture(img_path, width=Cm(out_w_cm))

    def add_heading(level: int, text: str):
        clean = _clean_docx_text(text)
        lvl = max(1, min(4, int(level)))
        doc.add_heading(clean, level=lvl)

    def add_text(text: str):
        clean = _clean_docx_text(text)
        para = doc.add_paragraph(clean)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_bold_title(text: str):
        clean = _clean_docx_text(text)
        p = doc.add_paragraph()
        r = p.add_run(clean)
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_rule():
        # Earlier iterations of the template used a row of "―" characters as a
        # visual separator between sections. The current template is laid out
        # with proper headings, page breaks, and Word's natural paragraph
        # spacing, so the rule reads as a stray artefact. Existing
        # ('rule', None) entries in order_list are kept for backwards
        # compatibility but render as nothing.
        return

    def add_table(data):
        if not data:
            return
        rows = len(data)
        cols = len(data[0]) if rows else 0
        if rows <= 0 or cols <= 0:
            return
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for r in range(rows):
            for c in range(cols):
                try:
                    table.cell(r, c).text = _clean_docx_text(data[r][c])
                except Exception:
                    table.cell(r, c).text = ""
        # Bold header row
        try:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for rr in p.runs:
                        rr.font.bold = True
        except Exception:
            pass

    def _add_field(paragraph, instr: str, placeholder: str = ""):
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), instr)
        if placeholder:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = str(placeholder)
            r.append(t)
            fld.append(r)
        paragraph._p.append(fld)

    def _is_map_image(path: str) -> bool:
        name = os.path.basename(path).lower()
        return any(tok in name for tok in ['map_', '_segments_', '_context', 'mbtiles'])

    def _is_atlas_image(path: str) -> bool:
        return 'atlas' in os.path.basename(path).lower()

    def _is_legend_image(path: str) -> bool:
        return '_legend' in os.path.basename(path).lower()

    def _needs_osm_attribution(path: str) -> bool:
        return _is_map_image(path) or _is_atlas_image(path)

    def _add_osm_attribution_paragraph():
        p = doc.add_paragraph()
        r = p.add_run(OSM_ATTRIBUTION_TEXT)
        r.italic = True
        r.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _heading_starts_map_block(idx: int, max_ahead: int = 10) -> bool:
        try:
            if idx < 0 or idx >= len(order_list):
                return False
            k0, _ = order_list[idx]
            if not str(k0).startswith("heading"):
                return False
            for j in range(idx + 1, min(len(order_list), idx + 1 + max_ahead)):
                kj, _pj = order_list[j]
                if kj == "new_page":
                    return False
                if kj == "image_map":
                    return True
                if str(kj).startswith("heading"):
                    return False
            return False
        except Exception:
            return False

    has_any_content = False
    last_was_page_break = False

    for idx, (kind, payload) in enumerate(order_list):
        if kind == 'toc':
            # TOC should start on a new page.
            if has_any_content and not last_was_page_break:
                doc.add_page_break()
                last_was_page_break = True
            add_bold_title("Contents")
            p = doc.add_paragraph()
            _add_field(p, 'TOC \\o "2-3" \\h \\z \\u', placeholder="(Put your marker here and press F9 to update the table of contents.)")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            has_any_content = True
            last_was_page_break = False

        elif isinstance(kind, str) and kind.startswith("heading"):
            # Keep map sections from starting at end of previous page.
            if _heading_starts_map_block(idx) and has_any_content and not last_was_page_break:
                doc.add_page_break()
                last_was_page_break = True
            try:
                level = int(kind[kind.find('(') + 1:kind.find(')')])
            except Exception:
                level = 2
            add_heading(level, payload)
            has_any_content = True
            last_was_page_break = False

        elif kind == "text":
            add_text(payload)
            has_any_content = True
            last_was_page_break = False

        elif kind in ("table", "table_data", "table_data_small"):
            title = None
            data = None
            if kind == "table":
                if isinstance(payload, str) and os.path.exists(payload):
                    try:
                        df = pd.read_excel(payload)
                        data = [df.columns.tolist()] + df.fillna("").astype(str).values.tolist()
                    except Exception:
                        data = None
                else:
                    data = payload
            else:
                try:
                    title, data = payload
                except Exception:
                    title, data = None, None
            if title:
                add_bold_title(title)
            if data:
                add_table(data)
            else:
                add_text("(table data unavailable)")
            has_any_content = True
            last_was_page_break = False

        elif kind in ("image", "image_ribbon", "image_map"):
            # payload: (title, path)
            try:
                title, path = payload
            except Exception:
                title, path = None, None
            if title:
                add_bold_title(title)
            if path and isinstance(path, str) and os.path.exists(path):
                try:
                    width_cm = float(MAX_IMAGE_WIDTH_CM)
                    base_name_lower = os.path.basename(path).lower()

                    # Atlas images slightly narrower
                    if _is_atlas_image(path):
                        width_cm *= float(ATLAS_DOC_WIDTH_SCALE)

                    # Ribbons: fixed height (keep consistent visual size)
                    if kind == "image_ribbon":
                        doc.add_picture(path, width=Cm(width_cm), height=Cm(float(RIBBON_CM_HEIGHT)))
                    # Index-map legend strip: align width with the map above it.
                    elif _is_legend_image(path):
                        doc.add_picture(path, width=Cm(float(MAX_LEGEND_WIDTH_CM)))
                    # Atlas: clamp height to keep pages readable
                    elif _is_atlas_image(path):
                        with PILImage.open(path) as im:
                            w_px, h_px = im.size
                        if w_px > 0 and h_px > 0:
                            aspect = float(h_px) / float(w_px)
                            out_w_cm = float(width_cm)
                            out_h_cm = out_w_cm * aspect
                            if out_h_cm > float(MAX_ATLAS_HEIGHT_CM):
                                out_h_cm = float(MAX_ATLAS_HEIGHT_CM)
                                out_w_cm = out_h_cm / aspect
                            doc.add_picture(path, width=Cm(out_w_cm))
                        else:
                            doc.add_picture(path, width=Cm(width_cm))
                    # Line context/segment maps: clamp height (portrait maps can overflow page)
                    elif _is_line_map_image(title, path):
                        _docx_add_picture_with_height_clamp(
                            path,
                            target_width_cm=width_cm,
                            max_height_cm=float(MAX_LINE_MAP_HEIGHT_CM),
                        )
                    # Overview / area maps: rendered as a panel so heading + intro
                    # paragraphs share the page with the map rather than getting
                    # pushed onto a separate page by a full-width image.
                    elif kind == "image_map":
                        _docx_add_picture_with_height_clamp(
                            path,
                            target_width_cm=float(MAX_OVERVIEW_MAP_WIDTH_CM),
                            max_height_cm=float(MAX_OVERVIEW_MAP_HEIGHT_CM),
                        )
                    else:
                        doc.add_picture(path, width=Cm(width_cm))

                    try:
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                    if _needs_osm_attribution(path):
                        _add_osm_attribution_paragraph()
                except Exception:
                    add_text(f"(image unavailable: {os.path.basename(path)})")
            else:
                if path:
                    add_text(f"(image missing: {os.path.basename(str(path))})")
                else:
                    add_text("(image missing)")
            has_any_content = True
            last_was_page_break = False

        elif kind == "rule":
            add_rule()
            has_any_content = True
            last_was_page_break = False

        elif kind == "spacer":
            try:
                n = int(payload) if payload is not None else 1
            except Exception:
                n = 1
            n = max(1, min(20, n))
            for _ in range(n):
                doc.add_paragraph("")
            has_any_content = True
            last_was_page_break = False

        elif kind == "new_page":
            # Avoid consecutive page breaks which can create empty pages.
            if last_was_page_break:
                continue
            doc.add_page_break()
            has_any_content = True
            last_was_page_break = True

        else:
            add_text(str(payload))
            has_any_content = True
            last_was_page_break = False

    doc.save(output_docx)


def _read_text_file(path: str) -> str:
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception:
        try:
            with open(path, 'r', encoding='cp1252') as f:
                return f.read()
        except Exception:
            return ""


_URL_RE = re.compile(r"\bhttps?://[^\s<>]+", re.IGNORECASE)


def _urls_to_footnotes(text: str) -> str:
    """Replace inline URLs with numeric footnotes appended at the end."""
    if not text:
        return ""
    urls: list[str] = []

    def _repl(m: re.Match) -> str:
        url = str(m.group(0))
        urls.append(url)
        return f"[{len(urls)}]"

    out = _URL_RE.sub(_repl, str(text))
    if not urls:
        return out
    foot = "\n\nFootnotes:\n" + "\n".join([f"[{i}] {u}" for i, u in enumerate(urls, start=1)])
    return out + foot

def set_progress(pct: float, message: str | None = None):
    try:
        pct = max(0.0, min(100.0, float(pct)))
        if message:
            write_to_log(message, _LOG_BASE_DIR)
        # Thread-safe progress update via signal
        try:
            if _signals is not None:
                _signals.progress_update.emit(pct)
        except Exception:
            pass
    except Exception:
        pass

# ---------------- Core: generate report ----------------
def generate_report(base_dir: str,
                    config_file: str,
                    palette_A2E: dict,
                    desc_A2E: dict,
                    report_mode: str | None = None,
                    include_assets: bool = True,
                    include_other_maps: bool = True,
                    include_index_statistics: bool = True,
                    include_lines_and_segments: bool = True,
                    include_atlas_maps: bool | None = None,
                    include_analysis_presentation: bool = False,
                    analysis_mode: str = "single",
                    analysis_area_left: str | None = None,
                    analysis_area_right: str | None = None):
    engine: ReportEngine | None = None
    try:
        global _LOG_BASE_DIR
        _LOG_BASE_DIR = base_dir
        set_progress(3, "Initializing report generation …")
        cfg       = read_config(config_file)

        # Basemap mode: xyz (default), contextily, auto
        global _REPORT_BASEMAP_MODE
        try:
            _REPORT_BASEMAP_MODE = (cfg['DEFAULT'].get('report_basemap_mode', 'xyz') or 'xyz').strip().lower()
        except Exception:
            _REPORT_BASEMAP_MODE = 'xyz'
        if _REPORT_BASEMAP_MODE not in {'xyz', 'contextily', 'auto'}:
            _REPORT_BASEMAP_MODE = 'xyz'
        write_to_log(f"Report basemap mode: {_REPORT_BASEMAP_MODE}", base_dir)

        # Inset styling (shadow + border). Configurable via config.ini.
        global _REPORT_INSET_BORDER_COLOR, _REPORT_INSET_BORDER_LW
        global _REPORT_INSET_SHADOW_ALPHA, _REPORT_INSET_SHADOW_DX, _REPORT_INSET_SHADOW_DY
        try:
            v = (cfg['DEFAULT'].get('report_inset_border_color', '') or '').strip()
            if v:
                _REPORT_INSET_BORDER_COLOR = v
        except Exception:
            pass
        try:
            v = _cfg_getfloat(cfg, 'DEFAULT', 'report_inset_border_lw', default=_REPORT_INSET_BORDER_LW)
            if v is not None and float(v) > 0:
                _REPORT_INSET_BORDER_LW = float(v)
        except Exception:
            pass
        try:
            v = _cfg_getfloat(cfg, 'DEFAULT', 'report_inset_shadow_alpha', default=_REPORT_INSET_SHADOW_ALPHA)
            if v is not None:
                _REPORT_INSET_SHADOW_ALPHA = max(0.0, min(1.0, float(v)))
        except Exception:
            pass
        try:
            v = _cfg_getfloat(cfg, 'DEFAULT', 'report_inset_shadow_dx', default=_REPORT_INSET_SHADOW_DX)
            if v is not None:
                _REPORT_INSET_SHADOW_DX = abs(float(v))
        except Exception:
            pass
        try:
            v = _cfg_getfloat(cfg, 'DEFAULT', 'report_inset_shadow_dy', default=_REPORT_INSET_SHADOW_DY)
            if v is not None:
                _REPORT_INSET_SHADOW_DY = abs(float(v))
        except Exception:
            pass

        include_atlas_cfg = _cfg_getboolean(cfg, 'DEFAULT', 'report_include_atlas_maps', default=False)
        if include_atlas_maps is None:
            if report_mode is None:
                include_atlas_maps = include_atlas_cfg
            else:
                include_atlas_maps = (str(report_mode).lower() == "detailed")
        else:
            include_atlas_maps = bool(include_atlas_maps)
        atlas_geocode_selected: str | None = None
        write_to_log(f"Report mode selected: {'Detailed (atlas included)' if include_atlas_maps else 'General maps only'}", base_dir)
        gpq_dir   = parquet_dir_from_cfg(base_dir, cfg)
        tmp_dir_path = output_subpath(base_dir, 'tmp')
        try:
            tmp_dir_path.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        tmp_dir   = str(tmp_dir_path)
        set_progress(6, "Paths prepared.")

        # Ensure the OSM tile cache folder exists in the expected place, even when basemap is served
        # via contextily (which doesn't use our XYZ-tile cache).
        try:
            _tile_cache_root(base_dir)
        except Exception:
            pass

        # Load key tables (conditionally)
        asset_objects_df = gpd.GeoDataFrame()
        asset_groups_df = gpd.GeoDataFrame()
        need_assets_tables = bool(include_assets or include_analysis_presentation)
        if need_assets_tables:
            asset_object_pq = os.path.join(gpq_dir, "tbl_asset_object.parquet")
            asset_group_pq  = os.path.join(gpq_dir, "tbl_asset_group.parquet")
            try:
                asset_objects_df = gpd.read_parquet(asset_object_pq) if os.path.exists(asset_object_pq) else gpd.GeoDataFrame()
            except Exception:
                asset_objects_df = gpd.GeoDataFrame()
            try:
                # tbl_asset_group is metadata-only and does not carry geo metadata,
                # so geopandas's read_parquet raises "Missing geo metadata".
                # Fall back to plain pandas; downstream code only consumes columns,
                # never geometry, on this table.
                if os.path.exists(asset_group_pq):
                    try:
                        asset_groups_df = gpd.read_parquet(asset_group_pq)
                    except Exception:
                        asset_groups_df = pd.read_parquet(asset_group_pq)
                else:
                    asset_groups_df = gpd.GeoDataFrame()
            except Exception:
                asset_groups_df = gpd.GeoDataFrame()
            set_progress(10, "Loaded asset tables.")
        else:
            set_progress(8, "Skipping asset tables (not selected).")

        # ---- Assets statistics ----
        object_stats_xlsx = os.path.join(tmp_dir, 'asset_object_statistics.xlsx')
        ag_stats_xlsx = os.path.join(tmp_dir, 'asset_group_statistics.xlsx')
        if include_assets:
            write_to_log("Computing asset object statistics …", base_dir)
            group_stats_df = calculate_group_statistics(asset_objects_df, asset_groups_df)
            export_to_excel(group_stats_df, object_stats_xlsx)

            ag_stats_df = fetch_asset_group_statistics(asset_groups_df, asset_objects_df)
            export_to_excel(ag_stats_df, ag_stats_xlsx)
            set_progress(22, "Asset stats ready.")
        else:
            write_to_log("Skipping asset statistics (not selected).", base_dir)
            set_progress(12)

        need_flat = bool(include_other_maps or include_index_statistics or include_lines_and_segments or include_atlas_maps)
        flat_df = gpd.GeoDataFrame()
        if need_flat:
            flat_df = load_tbl_flat(gpq_dir, base_dir=base_dir)

        need_engine = bool(include_other_maps or include_index_statistics or include_lines_and_segments or include_atlas_maps)
        if need_engine:
            engine = ReportEngine(base_dir, tmp_dir, palette_A2E, desc_A2E, config_file)

        # ---- Lines & segments (grouped per line with context & segments maps) ----
        pages_lines = []
        pages_lines_overview = []
        if include_lines_and_segments and engine is not None:
            lines_pq = os.path.join(gpq_dir, "tbl_lines.parquet")
            segments_pq = os.path.join(gpq_dir, "tbl_segment_flat.parquet")
            if not os.path.exists(lines_pq):
                write_to_log(f"Parquet table tbl_lines not found at {lines_pq}", base_dir)
            if not os.path.exists(segments_pq):
                write_to_log(f"Parquet table tbl_segment_flat not found at {segments_pq}", base_dir)

            lines_df, segments_df = fetch_lines_and_segments(gpq_dir)
            try:
                write_to_log(f"Lines table: {len(lines_df)} rows; Segments table: {len(segments_df)} rows", base_dir)
            except Exception:
                pass

            try:
                set_progress(23, "Rendering line overview …")
                pages_lines_overview = engine.render_lines_overview(lines_df, segments_df, palette_A2E, base_dir)
            except Exception as exc:
                write_to_log(f"Line overview failed: {exc}", base_dir)
                pages_lines_overview = []

            def _progress_segments(done: int, total: int):
                set_progress(25 + int(10 * done / max(1, total)), f"Rendering line segment maps ({done}/{total})")

            try:
                # Ensure we always move past 23% when starting the segments stage.
                _progress_segments(0, len(lines_df) if lines_df is not None else 0)
            except Exception:
                pass

            pages_lines, log_data = engine.render_segments(lines_df, segments_df, palette_A2E, base_dir, _progress_segments)

            if pages_lines:
                if log_data:
                    log_df = pd.DataFrame(log_data)
                    log_xlsx = os.path.join(tmp_dir, 'line_segment_log.xlsx')
                    export_to_excel(log_df, log_xlsx)
                    write_to_log(f"Segments log exported to {log_xlsx}", base_dir)
                else:
                    write_to_log("Line/segment pages created, but segment log is empty (no segments matched lines).", base_dir)
            else:
                # Provide a more actionable reason summary.
                reason_bits = []
                if lines_df is None or getattr(lines_df, 'empty', True):
                    reason_bits.append("tbl_lines is empty")
                if segments_df is None or getattr(segments_df, 'empty', True):
                    reason_bits.append("tbl_segment_flat is empty")
                if segments_df is not None and not getattr(segments_df, 'empty', True):
                    missing = [c for c in ['name_gis','segment_id','sensitivity_code_max','sensitivity_code_min'] if c not in segments_df.columns]
                    if missing:
                        reason_bits.append(f"tbl_segment_flat missing columns: {', '.join(missing)}")
                if lines_df is not None and not getattr(lines_df, 'empty', True):
                    missing_l = [c for c in ['name_gis','length_m','geometry'] if c not in lines_df.columns]
                    if missing_l:
                        reason_bits.append(f"tbl_lines missing columns: {', '.join(missing_l)}")
                msg = "; ".join(reason_bits) if reason_bits else "missing/invalid inputs"
                write_to_log(f"Skipping line/segment pages ({msg}).", base_dir)
            set_progress(35, "Lines/segments processed.")
        else:
            write_to_log("Skipping lines/segments section (not selected).", base_dir)
            set_progress(25)

        # ---- Atlas maps ----
        atlas_df = gpd.GeoDataFrame()
        atlas_geocode_selected = None
        if include_atlas_maps:
            atlas_pq = os.path.join(gpq_dir, "tbl_atlas.parquet")
            if os.path.exists(atlas_pq):
                try:
                    atlas_df = gpd.read_parquet(atlas_pq)
                except Exception as e:
                    write_to_log(f"Failed reading tbl_atlas: {e}", base_dir)
                    atlas_df = gpd.GeoDataFrame()
            else:
                write_to_log("Parquet table tbl_atlas not found; skipping atlas maps.", base_dir)

        atlas_geocode_pref = 'basic_mosaic' if include_atlas_maps else None

        def _progress_atlas(done: int, total: int):
            if include_atlas_maps:
                set_progress(37 + int(8 * done / max(1, total)), f"Rendering atlas tiles ({done}/{total})")

        atlas_pages = []
        if engine is not None:
            atlas_pages, atlas_geocode_selected = engine.render_atlas_maps(flat_df, atlas_df, include_atlas_maps, _progress_atlas)
        if atlas_pages:
            write_to_log("Per-atlas maps created.", base_dir)
        elif include_atlas_maps:
            write_to_log("Atlas maps requested but none were rendered.", base_dir)
        if include_atlas_maps:
            set_progress(45, "Atlas maps processed.")
        else:
            # Atlas is intentionally excluded in Basic/General report mode.
            set_progress(38)

        # ---- Other maps (basic_mosaic only) ----
        def _progress_geocode(done: int, total: int):
            start = 45 if include_atlas_maps else 38
            set_progress(start + int(30 * done / max(1, total)), f"Rendered maps for group {done}/{total}")

        geocode_pages = []
        geocode_intro_table = None
        geocode_groups = []
        if include_other_maps and engine is not None:
            geocode_pages, geocode_intro_table, geocode_groups = engine.render_geocode_maps(flat_df, _progress_geocode)
            if geocode_pages:
                write_to_log("Per-geocode maps created.", base_dir)
            else:
                write_to_log("tbl_flat missing or no 'name_gis_geocodegroup'; skipping per-geocode maps.", base_dir)
            set_progress(75 if include_atlas_maps else 68, "Per-geocode maps completed.")
        else:
            write_to_log("Skipping other maps section (not selected).", base_dir)
            set_progress(55 if include_atlas_maps else 45)

        # ---- Index statistics (basic_mosaic only; each index includes its map page) ----
        # Pick up where the per-geocode block actually ended so the bar moves
        # monotonically forward. The previous block ends at one of four marks
        # depending on whether atlas and other-maps ran:
        #   atlas + other_maps : 75    atlas + no other_maps : 55
        #   no atlas + other_maps : 68    no atlas + no other_maps : 45
        if include_atlas_maps:
            _index_start = 75 if include_other_maps else 55
        else:
            _index_start = 68 if include_other_maps else 45

        def _progress_indexes(done: int, total: int):
            # keep this light: just a small bump in the progress bar
            set_progress(_index_start + int(5 * done / max(1, total)),
                         f"Rendering index charts ({done}/{total})")

        index_pages = []
        if include_index_statistics and engine is not None:
            index_pages = engine.render_index_statistics(flat_df, cfg, _progress_indexes)
            # Settle on a clear "indexes done" mark so the next leap to the
            # compose phase starts from a known baseline.
            set_progress(_index_start + 5, "Index charts completed.")
        else:
            write_to_log("Skipping index statistics section (not selected).", base_dir)
        # ---- Compose PDF ----
        timestamp = datetime.datetime.now().strftime("%Y.%m.%d %H:%M:%S")

        assets_area_note = (
            "Note on areas: areas are computed only for polygon geometries. "
            "Points and lines have no surface area, therefore their 'Total area' is shown as “–”. "
            "Line and point assets are still counted in '# objects' and can be visualized on maps."
        )

        contents_lines = []
        if include_assets:
            contents_lines.append("- Assets overview & statistics")
        if include_analysis_presentation:
            mode = str(analysis_mode or "single").strip().lower()
            if mode == "compare":
                contents_lines.append("- Analysis presentation (compare two areas)")
            elif mode == "all":
                contents_lines.append("- Analysis presentation (all areas)")
            else:
                contents_lines.append("- Analysis presentation (single area)")
        if include_other_maps:
            contents_lines.append("- Other maps (basic_mosaic)")
        if include_index_statistics:
            contents_lines.append("- Index statistics + maps (basic_mosaic)")
        if include_lines_and_segments:
            contents_lines.append("- Line overview (summary across all lines)")
        if include_atlas_maps and atlas_pages:
            contents_lines.append("- Atlas tile maps (per atlas object with inset)")
        if include_lines_and_segments:
            contents_lines.append("- Lines & segments (grouped by line; context & segments maps, distributions, ribbons)")
        contents_text = "<br/>".join(contents_lines)

        # Templates live under docs/templates/ from MESA 5.0.2+; fall back to
        # docs/ for older project layouts where the file sat at the docs root.
        about_path = os.path.join(base_dir, "docs", "templates", "report_about.md")
        if not os.path.exists(about_path):
            about_path = os.path.join(base_dir, "docs", "report_about.md")
        about_text = _urls_to_footnotes(_read_text_file(about_path)).strip() if os.path.exists(about_path) else ""

        # Build initial order_list up to About section
        order_list = [
            ('heading(1)', "MESA report"),
            ('text', f"Timestamp: {timestamp}"),
            ('spacer', 2),
            ('heading(2)', "About this report"),
            ('text', about_text or "(About text missing: docs/templates/report_about.md)"),
            ('rule', None),
            # Insert a Word TOC field after About section
            ('toc', None),
            ('new_page', None),
        ]

        # After all sections are added to order_list, insert the dynamic Contents page
        # (This requires a second pass after order_list is fully built, so patch the PDF builder to do this)

        # Area overview: dissolved outline of all input data, computed in
        # Stage 1 (Prep) and stored at output/geoparquet/tbl_data_extent.parquet.
        # Slotted in here so readers see "where does this study cover" before
        # the Assets / Geocodes inventory tables.
        try:
            extent_pages = engine.render_data_extent_overview()
        except Exception as exc:
            write_to_log(f"Area overview render failed: {exc}", base_dir)
            extent_pages = []
        if extent_pages:
            order_list.extend(extent_pages)
            order_list.append(('new_page', None))

        if include_assets:
            order_list.extend([
                ('heading(2)', "Assets – overview"),
                ('text',
                    "This section inventories the <b>input</b> data that drives the rest of the report: "
                    "the asset objects loaded into MESA, broken down by their asset group. The first "
                    "table lists the groups with object counts and (where geometry is available) total "
                    "polygon area; the second table summarises how those objects are distributed across "
                    "the sensitivity classes A–E. These two tables together describe what is in the "
                    "stack before any of the per-cell aggregations on the following pages are computed."),
                ('text', assets_area_note),
                ('table', object_stats_xlsx),
                ('spacer', 1),
                ('text',
                    "<b>Asset groups – sensitivity distribution.</b> For each group, the count of "
                    "objects per sensitivity class (A–E) plus the number of <i>active</i> groups "
                    "(groups containing at least one object). Use this to confirm that the inputs "
                    "match expectations before reading the maps."),
                ('table', ag_stats_xlsx),
                ('new_page', None),
            ])

            # Geocodes – overview: one row per geocode group with object count.
            # Sourced from tbl_geocode_group (metadata) and tbl_geocode_object
            # (geometry rows), so it is independent of whether tbl_flat has
            # been loaded for this particular report run.
            geocode_overview_rows = build_geocode_overview_table(gpq_dir)
            if geocode_overview_rows is not None:
                order_list.extend([
                    ('heading(2)', "Geocodes – overview"),
                    ('text',
                        "<b>What is a geocode in MESA?</b> A <i>geocode group</i> is a set of "
                        "non-overlapping polygons that MESA uses as a common spatial reporting unit. "
                        "Every analytical output in this report — sensitivity maxima, importance maxima, "
                        "the three normalised indices, atlas tiles — is computed by overlaying assets "
                        "onto these polygons and aggregating per cell. Different geocode groups give "
                        "different views of the same input data: an asset-derived mosaic is shaped "
                        "exactly to where features actually occur, while a regular hex grid lets you "
                        "compare cell-for-cell across study areas at a fixed cell size."),
                    ('text',
                        "<b>basic_mosaic — the default reporting unit.</b> <code>basic_mosaic</code> is "
                        "produced by MESA itself from the imported assets: their footprints are "
                        "buffered, unioned, and then polygonised into the smallest set of "
                        "non-overlapping atomic faces that still distinguish every overlap pattern. "
                        "Every face represents an area where a specific combination of one or more "
                        "asset objects is present, and areas with no assets are not part of the mosaic "
                        "at all. The shape of basic_mosaic is therefore <i>data-driven</i>: a face is "
                        "as large or small as the underlying asset boundaries demand. This makes it "
                        "the most faithful representation of where assets actually occur, which is why "
                        "it is the default basis for the &ldquo;Other maps&rdquo;, Index, and Atlas "
                        "sections of this report."),
                    ('text',
                        "<b>H3 hexagon grids (H3_R6 … H3_R9).</b> These are standardised hierarchical "
                        "hex tessellations from Uber's open H3 spec. R6 is the coarsest (cells "
                        "≈36 km², about the size of a small district); each level halves the cell "
                        "size, so R9 cells are ≈0.1 km². H3 grids are useful when you need consistent "
                        "cell sizes across runs or when comparing study areas — a value at H3_R8 in "
                        "one project is directly comparable to H3_R8 in another, regardless of the "
                        "asset footprint. The fields used for ranking and weighting are identical to "
                        "those on basic_mosaic, only aggregated to a different polygon set."),
                    ('text',
                        "<b>Custom geocode sets.</b> Any polygon dataset placed under "
                        "<code>input/geocode/</code> and registered with MESA becomes a geocode group "
                        "with the same machinery — administrative boundaries, statistical units, "
                        "Quarter Degree Grid Cells, marine spatial planning blocks, and so on. "
                        "Provided the polygons are non-overlapping and cover the assets, the rest of "
                        "the report works without modification."),
                    ('text',
                        "The table below lists every group MESA knows about for this run, together "
                        "with the number of geocode objects (cells) it contains. <code># objects</code> "
                        "is the row count in <code>tbl_geocode_object.parquet</code> for that group "
                        "— a higher value means more cells and therefore a finer spatial resolution."),
                    ('table_data', ("Geocode groups and object counts", geocode_overview_rows)),
                    ('new_page', None),
                ])

        if include_analysis_presentation:
            write_to_log("Preparing analysis presentation section …", base_dir)
            try:
                gpq_dir_analysis = gpq_dir
                groups = _analysis_group_choices(gpq_dir_analysis)
                id_list = [gid for gid, _disp in groups]

                def _add_single(group_id: str):
                    group_title = _analysis_group_title(gpq_dir_analysis, group_id)
                    order_list.append(('heading(2)', f"Analysis – {group_title}"))

                    polys = _analysis_polygons_for_group(gpq_dir_analysis, group_id)
                    flat_cells = _analysis_flat_geo_for_group(gpq_dir_analysis, group_id)
                    minimap_path = os.path.join(tmp_dir, f"analysis_minimap_{group_id}.png")
                    if _analysis_write_area_map_png(
                        polys,
                        minimap_path,
                        base_dir=base_dir,
                        config_path=config_file,
                        palette_A2E=palette_A2E,
                        desc_A2E=desc_A2E,
                        overlay_alpha=0.65,
                        flat_cells_gdf=flat_cells,
                    ):
                        order_list.append(('image_map', ("Area overview (basemap + sensitivity)", minimap_path)))
                    else:
                        order_list.append(('text', "(Area map unavailable: missing analysis polygons and/or basemap/mbtiles.)"))

                    assets_tbl = _analysis_assets_within_area_table(asset_objects_df, asset_groups_df, polys)
                    if assets_tbl:
                        order_list.append(('table_data', ("Assets within area (polygon area)", assets_tbl)))

                    flat_sub = _analysis_flat_for_group(gpq_dir_analysis, group_id)
                    totals = _analysis_sensitivity_totals_km2(flat_sub)
                    chart_path = os.path.join(tmp_dir, f"analysis_sensitivity_max_{group_id}.png")
                    if _analysis_write_sensitivity_bar_png(totals, palette_A2E, chart_path, "Sensitivity (max) totals"):
                        order_list.append(('image', ("Sensitivity (max) totals", chart_path)))
                    else:
                        order_list.append(('text', "(Sensitivity chart unavailable: no analysis results.)"))

                    if totals is not None and not totals.empty:
                        cols = ["Code", "Description", "Area"]
                        rows = []
                        for _, r in totals.iterrows():
                            code = str(r.get("Code", ""))
                            desc = str(r.get("Description", ""))
                            area = _format_area_km2(r.get("Area (km²)", 0.0))
                            rows.append([code, desc, area])
                        table = [cols] + rows
                        order_list.append(('table_data', ("Sensitivity totals (max)", table)))
                    order_list.append(('new_page', None))

                mode = str(analysis_mode or "single").strip().lower()
                if mode == "all":
                    if not id_list:
                        order_list.append(('heading(2)', "Analysis"))
                        order_list.append(('text', "(No analysis groups found. Run analysis_setup.py / processing_pipeline_run.py first.)"))
                        order_list.append(('new_page', None))
                    else:
                        for gid in id_list:
                            _add_single(gid)
                elif mode == "compare":
                    left = (analysis_area_left or "").strip()
                    right = (analysis_area_right or "").strip()
                    if not left or not right:
                        order_list.append(('heading(2)', "Analysis"))
                        order_list.append(('text', "(Compare mode selected, but one or both areas were not chosen.)"))
                        order_list.append(('new_page', None))
                    else:
                        left_title = _analysis_group_title(gpq_dir_analysis, left)
                        right_title = _analysis_group_title(gpq_dir_analysis, right)
                        order_list.append(('heading(2)', f"Analysis – Compare: {left_title} vs {right_title}"))

                        left_polys = _analysis_polygons_for_group(gpq_dir_analysis, left)
                        right_polys = _analysis_polygons_for_group(gpq_dir_analysis, right)
                        left_cells = _analysis_flat_geo_for_group(gpq_dir_analysis, left)
                        right_cells = _analysis_flat_geo_for_group(gpq_dir_analysis, right)
                        left_map = os.path.join(tmp_dir, f"analysis_minimap_{left}.png")
                        right_map = os.path.join(tmp_dir, f"analysis_minimap_{right}.png")
                        if _analysis_write_area_map_png(
                            left_polys,
                            left_map,
                            base_dir=base_dir,
                            config_path=config_file,
                            palette_A2E=palette_A2E,
                            desc_A2E=desc_A2E,
                            overlay_alpha=0.65,
                            flat_cells_gdf=left_cells,
                        ):
                            order_list.append(('image_map', (f"Area map – {left_title}", left_map)))

                        left_assets_tbl = _analysis_assets_within_area_table(asset_objects_df, asset_groups_df, left_polys)
                        if left_assets_tbl:
                            order_list.append(('table_data', (f"Assets within area – {left_title}", left_assets_tbl)))

                        if _analysis_write_area_map_png(
                            right_polys,
                            right_map,
                            base_dir=base_dir,
                            config_path=config_file,
                            palette_A2E=palette_A2E,
                            desc_A2E=desc_A2E,
                            overlay_alpha=0.65,
                            flat_cells_gdf=right_cells,
                        ):
                            order_list.append(('image_map', (f"Area map – {right_title}", right_map)))

                        right_assets_tbl = _analysis_assets_within_area_table(asset_objects_df, asset_groups_df, right_polys)
                        if right_assets_tbl:
                            order_list.append(('table_data', (f"Assets within area – {right_title}", right_assets_tbl)))

                        left_totals = _analysis_sensitivity_totals_km2(_analysis_flat_for_group(gpq_dir_analysis, left))
                        right_totals = _analysis_sensitivity_totals_km2(_analysis_flat_for_group(gpq_dir_analysis, right))
                        left_chart = os.path.join(tmp_dir, f"analysis_sensitivity_max_{left}.png")
                        right_chart = os.path.join(tmp_dir, f"analysis_sensitivity_max_{right}.png")
                        _analysis_write_sensitivity_bar_png(left_totals, palette_A2E, left_chart, f"Sensitivity (max) – {left_title}")
                        _analysis_write_sensitivity_bar_png(right_totals, palette_A2E, right_chart, f"Sensitivity (max) – {right_title}")
                        if os.path.exists(left_chart):
                            order_list.append(('image', (f"Sensitivity (max) – {left_title}", left_chart)))
                        if os.path.exists(right_chart):
                            order_list.append(('image', (f"Sensitivity (max) – {right_title}", right_chart)))

                        # Relative-size (100%) composition chart
                        rel_chart = os.path.join(tmp_dir, f"analysis_sensitivity_relative_{left}_vs_{right}.png")
                        if _analysis_write_relative_share_png(left_totals, right_totals, palette_A2E, rel_chart, left_title, right_title):
                            order_list.append(('image', ("Sensitivity composition (relative)", rel_chart)))

                        # Sankey-style difference diagram (ribbons between left/right blocks)
                        sankey_chart = os.path.join(tmp_dir, f"analysis_sensitivity_sankey_{left}_vs_{right}.png")
                        if _analysis_write_sankey_difference_png(left_totals, right_totals, palette_A2E, sankey_chart, left_title, right_title):
                            order_list.append(('image', ("Area difference Sankey", sankey_chart)))

                        lm = {r["Code"]: float(r["Area (km²)"]) for _, r in left_totals.iterrows()} if left_totals is not None and not left_totals.empty else {}
                        rm = {r["Code"]: float(r["Area (km²)"]) for _, r in right_totals.iterrows()} if right_totals is not None and not right_totals.empty else {}
                        order_map = {c: i for i, c in enumerate(SENSITIVITY_ORDER)}

                        def _code_key(val: str) -> tuple[int, str]:
                            s = str(val or "").strip().upper()
                            if s in order_map:
                                return (0, f"{order_map[s]:02d}")
                            if s in ("", "NONE", "NAN"):
                                return (3, "")
                            if s == "UNKNOWN":
                                return (2, s)
                            return (1, s)

                        all_codes = sorted(set(lm.keys()) | set(rm.keys()), key=_code_key)
                        comp_rows = [["Code", f"{left_title} (km²)", f"{right_title} (km²)", "Difference (L - R)"]]
                        for code in all_codes:
                            lv = lm.get(code, 0.0)
                            rv = rm.get(code, 0.0)
                            comp_rows.append([code, f"{lv:,.2f}", f"{rv:,.2f}", f"{(lv-rv):+,.2f}"])
                        order_list.append(('table_data', ("Sensitivity comparison (max)", comp_rows)))
                        order_list.append(('new_page', None))
                else:
                    gid = (analysis_area_left or "").strip() or (id_list[0] if id_list else "")
                    if not gid:
                        order_list.append(('heading(2)', "Analysis"))
                        order_list.append(('text', "(No analysis groups found. Run analysis_setup.py / processing_pipeline_run.py first.)"))
                        order_list.append(('new_page', None))
                    else:
                        _add_single(gid)
            except Exception as exc:
                write_to_log(f"Analysis presentation section failed: {exc}", base_dir)
                order_list.append(('heading(2)', "Analysis"))
                order_list.append(('text', "(Analysis presentation could not be generated; check log for details.)"))
                order_list.append(('new_page', None))

        if include_other_maps:
            order_list.extend([
                ('heading(2)', "Other maps"),
                ('text',
                 "These maps describe the asset stack itself, before any composite index is computed. "
                 "For each <b>basic_mosaic</b> cell they show the worst-case sensitivity class, the "
                 "highest importance value, and counts of overlapping asset groups and asset objects. "
                 "Other geocode groupings are deliberately omitted to keep the basemap consistent. "
                 "Treat this section as the &ldquo;raw inputs&rdquo; view; the index sections later in the report "
                 "build on the same data."),
            ])

            if geocode_intro_table is not None:
                order_list.append(('table_data', ("Available geocode categories and object counts", geocode_intro_table)))
                order_list.append(('rule', None))

            order_list.extend(geocode_pages)

        if include_index_statistics:
            order_list.extend([
                ('heading(2)', "Index statistics"),
                ('text',
                    "MESA produces three <b>normalised composite indices</b> that condense the asset stack into a "
                    "single 0–100 score, sitting alongside the four <b>per-cell summary indicators</b> "
                    "(Sensitive areas (A–E), Importance (max), # asset groups, # asset objects) shown on the "
                    "preceding &ldquo;Other maps&rdquo; pages. The <b>Importance index</b> answers &ldquo;where are the most "
                    "valuable features?&rdquo;, the <b>Sensitivity index</b> answers &ldquo;where is the most that "
                    "could be harmed?&rdquo;, and the <b>OWA index</b> is a precautionary variant that lets a "
                    "single very-sensitive overlap dominate the result. All three are scaled relative to "
                    "the current study area (most-loaded cell = 100), so the colour ramps are comparable "
                    "<i>within</i> the report but not across different runs. The indices reflect ranking; the "
                    "supplementary indicators reflect raw maxima and counts — read them together to ground-truth "
                    "what each index summarises."),
                ('text',
                    "Each index gets a statistics page (area distribution chart) followed by its map. "
                    "Read the three together: an area where Importance and Sensitivity are both high is "
                    "a high-stakes hotspot; an area where only OWA is elevated tells you a single rare "
                    "but extreme overlap is driving the signal."),
            ])

        if index_pages:
            order_list.extend(index_pages)
        else:
            order_list.append(('text', "No index statistics could be generated (missing data/columns)."))
            order_list.append(('new_page', None))

        # Insert line overview immediately before the Atlas section.
        if include_lines_and_segments and pages_lines_overview:
            order_list.append(('heading(2)', "Line overview"))
            order_list.append(('text',
                "A bird's-eye view of every line analysed in this report. The overview map colours all "
                "segments by their <b>maximum</b> sensitivity class, the distribution chart shows how "
                "much segment length sits in each class, and the per-line table breaks the same totals "
                "down by individual line. Use this section to compare lines against each other before "
                "drilling into any single line in the next section."))
            order_list.append(('rule', None))
            order_list.extend(pages_lines_overview)

        # Place line details immediately after the line overview.
        if include_lines_and_segments and pages_lines:
            order_list.append(('heading(2)', "Lines and segments"))
            order_list.append(('text',
                "One pair of pages per line: a maximum-sensitivity view followed by a minimum-sensitivity "
                "view. Each page shows a segments map (with a context inset placing the line in the wider "
                "study area), a sensitivity distribution chart, and a horizontal <b>ribbon</b> that lays "
                "the line out flat from start to end so peaks are easy to locate against distance markers. "
                "Read maximum and minimum together: if both are high, sensitivity is uniformly elevated; "
                "if they diverge, the line passes through both robust and sensitive ground."))
            order_list.append(('rule', None))
            order_list.append(('new_page', None))
            order_list.extend(pages_lines)

        if atlas_pages:
            order_list.append(('heading(2)', "Atlas maps"))
            atlas_intro = (
                "Atlas tiles split the study area into named sub-areas — typically prepared in advance so "
                "that each tile covers a manageable extent at usable detail. Every tile gets its own page "
                "with a sensitivity map (A–E palette) and a small inset locating the tile inside the full "
                "study area. Use this section when you need to see local detail that the study-area-wide "
                "maps necessarily smooth over.")
            order_list.append(('text', atlas_intro))
            order_list.extend(atlas_pages)

        # Avoid creating a trailing blank page in DOCX/PDF when the report ends with a page break.
        while order_list and isinstance(order_list[-1], tuple) and order_list[-1][0] == 'new_page':
            order_list.pop()

        set_progress(86, "Composing Word report …")
        ts_docx = datetime.datetime.now().strftime("%Y-%m-%d_%H%M")
        output_docx_path = output_subpath(base_dir, "reports", f"MESA-report_{ts_docx}.docx")
        output_docx = str(output_docx_path)
        compile_docx(output_docx, order_list)
        engine.cleanup()
        engine = None
        set_progress(100, "Report completed.")

        global last_report_path
        last_report_path = output_docx
        write_to_log(f"Word report created: {output_docx}", base_dir)

        try:
            if _signals is not None:
                _signals.link_update.emit("Open report folder")
        except Exception:
            pass

    except Exception as e:
        write_to_log(f"ERROR during report generation: {e}", base_dir)
        try:
            import traceback

            tb = traceback.format_exc()
            for line in (tb or "").splitlines():
                if line.strip():
                    write_to_log(line, base_dir)
        except Exception:
            pass
        set_progress(100, "Report failed.")
    finally:
        try:
            _LOG_BASE_DIR = None
        except Exception:
            pass
        if engine is not None:
            try:
                engine.cleanup()
            except Exception:
                pass

# ---------------- GUI runner ----------------
def _start_report_thread(base_dir, config_file, palette, desc, report_mode):
    # Backwards-compatible helper for old call sites.
    threading.Thread(
        target=generate_report,
        args=(base_dir, config_file, palette, desc, report_mode),
        daemon=True
    ).start()

def _start_report_thread_selected(base_dir, config_file, palette, desc, *,
                                 include_assets: bool,
                                 include_other_maps: bool,
                                 include_index_statistics: bool,
                                 include_lines_and_segments: bool,
                                 include_atlas_maps: bool,
                                 include_analysis_presentation: bool,
                                 analysis_mode: str,
                                 analysis_area_left: str | None,
                                 analysis_area_right: str | None):
    threading.Thread(
        target=generate_report,
        kwargs={
            'base_dir': base_dir,
            'config_file': config_file,
            'palette_A2E': palette,
            'desc_A2E': desc,
            'report_mode': None,
            'include_assets': include_assets,
            'include_other_maps': include_other_maps,
            'include_index_statistics': include_index_statistics,
            'include_lines_and_segments': include_lines_and_segments,
            'include_atlas_maps': include_atlas_maps,
            'include_analysis_presentation': include_analysis_presentation,
            'analysis_mode': analysis_mode,
            'analysis_area_left': analysis_area_left,
            'analysis_area_right': analysis_area_right,
        },
        daemon=True
    ).start()

class ReportGeneratorWindow(QMainWindow):
    """PySide6 main window for the MESA report generator."""

    def __init__(self, base_dir: str, config_file: str, palette: dict, desc: dict, parent=None):
        super().__init__(parent)
        global log_widget, progress_var, link_var, _gui_window, _signals

        _gui_window = self
        _signals = _ReportSignals()

        self.setWindowTitle("MESA \u2013 Report generator")
        self.resize(900, 640)
        self.setMinimumSize(700, 500)

        try:
            ico = Path(base_dir) / "system_resources" / "mesa.ico"
            if ico.exists():
                self.setWindowIcon(QIcon(str(ico)))
        except Exception:
            pass

        central = QWidget(self)
        central.setObjectName("CentralHost")
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(10, 10, 10, 10)

        # --- Log area ---
        log_group = QGroupBox("Log")
        log_lay = QVBoxLayout(log_group)
        log_widget = QPlainTextEdit()
        log_widget.setReadOnly(True)
        log_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        log_lay.addWidget(log_widget)
        main_layout.addWidget(log_group)

        # --- Progress row ---
        pframe = QHBoxLayout()
        progress_var = QProgressBar()
        progress_var.setRange(0, 100)
        progress_var.setValue(0)
        progress_var.setTextVisible(True)
        progress_var.setFormat("%p%")
        progress_var.setAlignment(Qt.AlignCenter)
        progress_var.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        pframe.addWidget(progress_var, stretch=1)
        main_layout.addLayout(pframe)

        # --- Report options group ---
        action_group = QGroupBox("Report")
        action_layout = QVBoxLayout(action_group)

        # Include in report
        include_group = QGroupBox("Include in report")
        include_grid = QGridLayout(include_group)

        self._chk_assets = QCheckBox("Assets overview && statistics")
        self._chk_assets.setChecked(True)
        self._chk_analysis = QCheckBox("Analysis presentation (graphs)")
        self._chk_analysis.setChecked(False)
        self._chk_other_maps = QCheckBox("Other maps (basic_mosaic)")
        self._chk_other_maps.setChecked(True)
        self._chk_index_stats = QCheckBox("Index statistics")
        self._chk_index_stats.setChecked(True)
        self._chk_lines_segments = QCheckBox("Lines")
        self._chk_lines_segments.setChecked(True)
        self._chk_atlas = QCheckBox("Atlas maps (detailed)")
        self._chk_atlas.setChecked(False)

        _include_checks = [
            self._chk_assets, self._chk_analysis,
            self._chk_other_maps, self._chk_index_stats,
            self._chk_lines_segments, self._chk_atlas,
        ]
        for idx, chk in enumerate(_include_checks):
            include_grid.addWidget(chk, idx // 2, idx % 2)

        action_layout.addWidget(include_group)

        # Analysis presentation options
        analysis_group = QGroupBox("Analysis presentation options")
        analysis_grid = QGridLayout(analysis_group)

        analysis_grid.addWidget(QLabel("Mode:"), 0, 0)
        mode_widget = QWidget()
        mode_lay = QHBoxLayout(mode_widget)
        mode_lay.setContentsMargins(0, 0, 0, 0)
        self._radio_single = QRadioButton("Single area")
        self._radio_single.setChecked(True)
        self._radio_compare = QRadioButton("Compare two areas")
        self._radio_all = QRadioButton("All areas")
        self._mode_group = QButtonGroup(self)
        self._mode_group.addButton(self._radio_single)
        self._mode_group.addButton(self._radio_compare)
        self._mode_group.addButton(self._radio_all)
        mode_lay.addWidget(self._radio_single)
        mode_lay.addWidget(self._radio_compare)
        mode_lay.addWidget(self._radio_all)
        mode_lay.addStretch()
        analysis_grid.addWidget(mode_widget, 0, 1)

        try:
            _cfg_tmp = read_config(config_file)
            _gpq_tmp = parquet_dir_from_cfg(base_dir, _cfg_tmp)
        except Exception:
            _gpq_tmp = parquet_dir_from_cfg(base_dir, read_config(config_file))
        group_choices = _analysis_group_choices(_gpq_tmp)
        group_labels = [disp for _gid, disp in group_choices]
        self._label_to_id = {disp: gid for gid, disp in group_choices}

        analysis_grid.addWidget(QLabel("Area A:"), 1, 0)
        self._combo_left = QComboBox()
        self._combo_left.addItems(group_labels)
        self._combo_left.setMinimumWidth(260)
        analysis_grid.addWidget(self._combo_left, 1, 1)

        analysis_grid.addWidget(QLabel("Area B:"), 2, 0)
        self._combo_right = QComboBox()
        self._combo_right.addItems(group_labels)
        self._combo_right.setMinimumWidth(260)
        analysis_grid.addWidget(self._combo_right, 2, 1)

        if group_labels:
            self._combo_left.setCurrentIndex(0)
            self._combo_right.setCurrentIndex(1 if len(group_labels) > 1 else 0)

        analysis_grid.setColumnStretch(1, 1)
        action_layout.addWidget(analysis_group)

        # Create report button + hint label. Use the shared "primary" role
        # (gold-on-parchment) defined in ASSET_STYLESHEET, the same treatment
        # applied to Import / Build / Generate buttons across MESA, instead
        # of the bright Bootstrap green this dialog used to carry.
        btn_row = QHBoxLayout()
        create_btn = QPushButton("Create report")
        create_btn.setProperty("role", "primary")
        btn_row.addWidget(create_btn)
        hint_label = QLabel("The report will include only the selected sections.")
        btn_row.addWidget(hint_label)
        btn_row.addStretch()
        action_layout.addLayout(btn_row)

        main_layout.addWidget(action_group)

        # --- Bottom row: link + exit ---
        bottom_row = QHBoxLayout()
        link_var = QLabel("")
        link_var.setStyleSheet("QLabel { color: #4ea3ff; }")
        link_var.setFont(QFont("Segoe UI", 10))
        link_var.setCursor(Qt.PointingHandCursor)
        link_var.mousePressEvent = self._open_report_folder
        bottom_row.addWidget(link_var)
        bottom_row.addStretch()
        exit_btn = QPushButton("Exit")
        exit_btn.setObjectName("CornerExitButton")
        exit_btn.setStyleSheet("""
            QPushButton#CornerExitButton {
                background: #eadfc8; border: 1px solid #b79f73;
                border-radius: 4px; color: #453621;
                padding: 6px 18px;
            }
            QPushButton#CornerExitButton:hover { background: #e1d1ae; }
            QPushButton#CornerExitButton:pressed { background: #d4c094; }
        """)
        exit_btn.clicked.connect(self.close)
        bottom_row.addWidget(exit_btn)
        main_layout.addLayout(bottom_row)

        # --- Signals ---
        _signals.log_message.connect(self._on_log_message)
        _signals.progress_update.connect(self._on_progress_update)
        _signals.link_update.connect(self._on_link_update)

        # --- Control wiring ---
        self._chk_analysis.stateChanged.connect(self._sync_analysis_controls)
        self._mode_group.buttonClicked.connect(self._sync_analysis_controls)
        self._sync_analysis_controls()

        create_btn.clicked.connect(lambda: self._start_selected(base_dir, config_file, palette, desc))

        # Store references for use in callbacks
        self._base_dir = base_dir
        self._group_labels = group_labels

        write_to_log(f"Working directory: {base_dir}", base_dir)
        write_to_log("Ready. Select report contents, then press 'Create report'.", base_dir)

    # ---- Slots ----
    def _on_log_message(self, text: str):
        if log_widget is not None:
            log_widget.appendPlainText(text)

    def _on_progress_update(self, pct: float):
        if progress_var is not None:
            progress_var.setValue(int(pct))

    def _on_link_update(self, text: str):
        if link_var is not None:
            link_var.setText(text)

    def _sync_analysis_controls(self, *_args):
        enabled = self._chk_analysis.isChecked()
        if self._radio_single.isChecked():
            mode = "single"
        elif self._radio_compare.isChecked():
            mode = "compare"
        else:
            mode = "all"

        if not enabled or not self._group_labels:
            self._combo_left.setEnabled(False)
            self._combo_right.setEnabled(False)
            return
        if mode == "all":
            self._combo_left.setEnabled(False)
            self._combo_right.setEnabled(False)
        elif mode == "compare":
            self._combo_left.setEnabled(True)
            self._combo_right.setEnabled(True)
        else:
            self._combo_left.setEnabled(True)
            self._combo_right.setEnabled(False)

    def _start_selected(self, base_dir, config_file, palette, desc):
        if not any([
            self._chk_assets.isChecked(),
            self._chk_analysis.isChecked(),
            self._chk_other_maps.isChecked(),
            self._chk_index_stats.isChecked(),
            self._chk_lines_segments.isChecked(),
            self._chk_atlas.isChecked(),
        ]):
            write_to_log("No report sections selected. Tick at least one box.", base_dir)
            return

        analysis_left_id = None
        analysis_right_id = None
        if self._chk_analysis.isChecked():
            left_label = self._combo_left.currentText().strip()
            right_label = self._combo_right.currentText().strip()
            analysis_left_id = self._label_to_id.get(left_label) if left_label else None
            analysis_right_id = self._label_to_id.get(right_label) if right_label else None

        if self._radio_single.isChecked():
            mode = "single"
        elif self._radio_compare.isChecked():
            mode = "compare"
        else:
            mode = "all"

        _start_report_thread_selected(
            base_dir,
            config_file,
            palette,
            desc,
            include_assets=self._chk_assets.isChecked(),
            include_analysis_presentation=self._chk_analysis.isChecked(),
            analysis_mode=mode,
            analysis_area_left=analysis_left_id,
            analysis_area_right=analysis_right_id,
            include_other_maps=self._chk_other_maps.isChecked(),
            include_index_statistics=self._chk_index_stats.isChecked(),
            include_lines_and_segments=self._chk_lines_segments.isChecked(),
            include_atlas_maps=self._chk_atlas.isChecked(),
        )

    def _open_report_folder(self, event=None):
        if not last_report_path:
            return
        folder = os.path.dirname(last_report_path)
        if os.path.isdir(folder):
            try:
                os.startfile(folder)  # Windows
            except Exception:
                try:
                    subprocess.Popen(["explorer", folder])
                except Exception as ee:
                    write_to_log(f"Failed to open folder: {ee}", self._base_dir)
        else:
            write_to_log("Report folder not found.", self._base_dir)


def launch_gui(base_dir: str, config_file: str, palette: dict, desc: dict, theme: str = "", master=None):
    """Create and show the ReportGeneratorWindow. Returns the window instance."""
    app = QApplication.instance()
    own_app = app is None
    if own_app:
        app = QApplication([])
        apply_shared_stylesheet(app)

    window = ReportGeneratorWindow(base_dir, config_file, palette, desc)
    window.show()

    if own_app:
        app.exec()
    return window


# ---------------- In-process entry point (called by mesa.py via lazy import) ----------------
def run(base_dir: str, master=None) -> None:
    """Launch the report generator GUI in-process.

    mesa.py calls this instead of spawning a subprocess.
    """
    global PRIMARY_HEX, LIGHT_PRIMARY_HEX
    resolved = normalize_base_dir(base_dir)
    cfg_path = config_path(resolved)
    cfg = read_config(cfg_path)
    palette_A2E, desc_A2E = read_sensitivity_palette_and_desc(cfg_path)
    PRIMARY_HEX = cfg['DEFAULT'].get('ui_primary_color', PRIMARY_HEX).strip() or PRIMARY_HEX
    try:
        def _lighten(hexcol, amt=0.30):
            hexcol = hexcol.lstrip('#')
            r = int(hexcol[0:2],16); g = int(hexcol[2:4],16); b = int(hexcol[4:6],16)
            r = int(r + (255 - r)*amt); g = int(g + (255 - g)*amt); b = int(b + (255 - b)*amt)
            return f"#{r:02x}{g:02x}{b:02x}"
        LIGHT_PRIMARY_HEX = _lighten(PRIMARY_HEX, 0.30)
    except Exception:
        pass
    return launch_gui(resolved, cfg_path, palette_A2E, desc_A2E, master=master)


# ---------------- CLI ----------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Presentation report (GeoParquet per geocode, same-scale maps, line context + segments maps with buffers)')
    parser.add_argument('--original_working_directory', required=False, help='Path to running folder')
    parser.add_argument('--no-gui', action='store_true', help='Run directly without GUI')
    parser.add_argument('--report-mode', choices=['general', 'detailed'],
                        help='Select report detail level (general maps or detailed with atlas overviews).')
    parser.add_argument('--debug-atlas-sample',
                        help='Render a standalone sensitivity atlas map for the specified tile (name_gis).')
    parser.add_argument('--debug-atlas-size', type=int,
                        help='Optional max polygon count for --debug-atlas-sample (down-sample if exceeded).')
    args = parser.parse_args()

    base_dir = args.original_working_directory
    if not base_dir:
        base_dir = os.getcwd()
    base_dir = normalize_base_dir(base_dir)

    cfg_path = config_path(base_dir)
    cfg = read_config(cfg_path)

    # Sensitivity palette + descriptions from config (A-E)
    palette_A2E, desc_A2E = read_sensitivity_palette_and_desc(cfg_path)

    if args.debug_atlas_sample:
        out = debug_atlas_sample(
            base_dir,
            cfg,
            palette_A2E,
            desc_A2E,
            args.debug_atlas_sample,
            sample_size=args.debug_atlas_size
        )
        if out:
            print(f"Debug atlas sample saved to: {out}")
        else:
            print("Debug atlas sample failed. Check log for details.")
        sys.exit(0)

    # Optional override for brand color from config
    PRIMARY_HEX = cfg['DEFAULT'].get('ui_primary_color', PRIMARY_HEX).strip() or PRIMARY_HEX
    try:
        def _lighten(hexcol, amt=0.30):
            hexcol = hexcol.lstrip('#')
            r = int(hexcol[0:2],16); g = int(hexcol[2:4],16); b = int(hexcol[4:6],16)
            r = int(r + (255 - r)*amt); g = int(g + (255 - g)*amt); b = int(b + (255 - b)*amt)
            return f"#{r:02x}{g:02x}{b:02x}"
        LIGHT_PRIMARY_HEX = _lighten(PRIMARY_HEX, 0.30)
    except Exception:
        pass

    if args.no_gui:
        generate_report(base_dir, cfg_path, palette_A2E, desc_A2E,
                        report_mode=args.report_mode)
    else:
        launch_gui(base_dir, cfg_path, palette_A2E, desc_A2E)

