import locale
try:
    locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')
except Exception:
    # Fall back silently if locale isn't available on this system
    pass

import geopandas as gpd
import pandas as pd
import configparser
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import cm  # still used for ScalarMappable (not deprecated)
from matplotlib import colormaps as mpl_cmaps  # modern colormap access
import matplotlib.colors as mcolors
from matplotlib.path import Path as MplPath
try:
    from mpl_toolkits.axes_grid1.inset_locator import inset_axes
except Exception:
    inset_axes = None
import numpy as np
import argparse
import os
import sys
try:
    import contextily as ctx  # optional; heavy deps (rasterio) may be absent in EXE
except Exception:
    ctx = None
from matplotlib.patches import Rectangle, PathPatch
from matplotlib.ticker import MaxNLocator

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Image, Spacer,
    Table, TableStyle, PageBreak, HRFlowable
)
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import cm as RL_CM
from reportlab.lib.colors import HexColor
from PIL import Image as PILImage
import re
import tkinter as tk
from tkinter import scrolledtext
import ttkbootstrap as tb
from ttkbootstrap.constants import PRIMARY, WARNING
import threading
from pathlib import Path
import subprocess
import io, math, time, urllib.request
import sqlite3
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- UI / sizing constants ----------------
MAX_MAP_PX_HEIGHT = 2000           # hard cap for saved map PNG height (px)
MAX_MAP_CM_HEIGHT = 10.0           # map display cap inside PDF (cm)
RIBBON_CM_HEIGHT   = 0.6           # ribbon display height inside PDF (cm)
ATLAS_FIGURE_INCHES = (7.2, 7.2)   # atlas tiles render smaller to fit more per page
ATLAS_DOC_WIDTH_SCALE = 0.75       # reduce displayed width of atlas images inside documents

TILE_CACHE_MAX_AGE_DAYS = 30       # discard cached OSM tiles older than this (<=0 keeps forever)

OSM_ATTRIBUTION_TEXT = "© OpenStreetMap contributors"

_OSM_TILE_CACHE_LOGGED = False

# Basemap mode (set from config):
# - xyz: force built-in XYZ downloader with output/tile_cache
# - contextily: force contextily when available (fallback to xyz if not)
# - auto: prefer contextily, fallback to xyz
_REPORT_BASEMAP_MODE = "xyz"

# ---------------- GUI / globals ----------------
log_widget = None
progress_var = None
progress_label = None
last_report_path = None
link_var = None  # hyperlink label StringVar
atlas_geocode_var = None  # Combobox selection for atlas geocode group
_atlas_geocode_choices: list[str] = []  # cached list for GUI

SENSITIVITY_ORDER = ['A', 'B', 'C', 'D', 'E']
SENSITIVITY_UNKNOWN_COLOR = "#FF00F2"
_SENSITIVITY_NUMERIC_RANGES: list[tuple[str, float, float]] = []


# ---------------- Analysis presentation helpers (graphs per analysis area) ----------------
_ANALYSIS_GROUP_TABLE = "tbl_analysis_group.parquet"
_ANALYSIS_POLYGON_TABLE = "tbl_analysis_polygons.parquet"
_ANALYSIS_FLAT_TABLE = "tbl_analysis_flat.parquet"


def _analysis_read_table(gpq_dir: str, filename: str, *, geo: bool = False):
    path = os.path.join(gpq_dir, filename)
    if not os.path.exists(path):
        return None
    try:
        if geo:
            return gpd.read_parquet(path)
        return pd.read_parquet(path)
    except Exception:
        try:
            gdf = gpd.read_parquet(path)
            return pd.DataFrame(gdf.drop(columns=[c for c in ("geometry",) if c in gdf.columns]))
        except Exception:
            return None


def _analysis_group_choices(gpq_dir: str) -> list[tuple[str, str]]:
    df = _analysis_read_table(gpq_dir, _ANALYSIS_GROUP_TABLE, geo=False)
    if df is None or df.empty or "id" not in df.columns:
        return []
    choices: list[tuple[str, str]] = []
    for _, row in df.iterrows():
        gid = str(row.get("id", "") or "").strip()
        if not gid:
            continue
        name = str(row.get("name", "") or "").strip()
        display = name if name else gid
        choices.append((gid, f"{display} ({gid})"))
    return choices


def _analysis_group_title(gpq_dir: str, group_id: str) -> str:
    df = _analysis_read_table(gpq_dir, _ANALYSIS_GROUP_TABLE, geo=False)
    if df is None or df.empty:
        return str(group_id)
    try:
        match = df[df["id"].astype(str) == str(group_id)]
        if match.empty:
            return str(group_id)
        name = str(match.iloc[0].get("name", "") or "").strip()
        return name or str(group_id)
    except Exception:
        return str(group_id)


def _analysis_latest_run_filter(df: pd.DataFrame, group_id: str) -> pd.DataFrame:
    """Filter an analysis table to the latest run for the given group.

    tbl_analysis_flat.parquet can contain multiple analysis runs; without filtering,
    maps/statistics may mix runs.
    """
    if df is None or getattr(df, "empty", True):
        return df
    if "analysis_group_id" not in df.columns:
        return df
    try:
        mask = df["analysis_group_id"].astype(str) == str(group_id)
        sub = df.loc[mask].copy()
    except Exception:
        return df
    if sub.empty:
        return sub

    # Prefer run_id when present; otherwise use timestamp.
    if "analysis_run_id" in sub.columns:
        try:
            # Pick the most recent timestamp within each run_id, then choose the newest run.
            if "analysis_timestamp" in sub.columns:
                ts = pd.to_datetime(sub["analysis_timestamp"], errors="coerce")
                sub["__ts"] = ts
                run_ts = sub.groupby("analysis_run_id", dropna=False)["__ts"].max()
                chosen = run_ts.sort_values().index[-1]
                sub = sub[sub["analysis_run_id"] == chosen].copy()
                sub.drop(columns=["__ts"], inplace=True, errors="ignore")
                return sub
            # No timestamp: fall back to last run_id in sorted order
            chosen = sorted(sub["analysis_run_id"].astype(str).unique())[-1]
            return sub[sub["analysis_run_id"].astype(str) == chosen].copy()
        except Exception:
            sub.drop(columns=["__ts"], inplace=True, errors="ignore")
            return sub

    if "analysis_timestamp" in sub.columns:
        try:
            ts = pd.to_datetime(sub["analysis_timestamp"], errors="coerce")
            if ts.notna().any():
                latest = ts.max()
                return sub.loc[ts == latest].copy()
        except Exception:
            pass
    return sub


def _analysis_polygons_for_group(gpq_dir: str, group_id: str) -> gpd.GeoDataFrame:
    # tbl_analysis_flat.parquet contains *many* geometries per polygon (typically per geocode cell)
    # and is not suitable for rendering analysis area boundaries. We instead use it only to
    # determine which analysis_polygon_id are part of the selected group/run, then fetch the
    # actual polygon boundaries from tbl_analysis_polygons.parquet.
    polygon_ids: list[str] = []
    gdf_flat = _analysis_read_table(gpq_dir, _ANALYSIS_FLAT_TABLE, geo=True)
    if gdf_flat is not None and not getattr(gdf_flat, "empty", True):
        try:
            flat_df = pd.DataFrame(gdf_flat)
            flat_df = _analysis_latest_run_filter(flat_df, group_id)
            if "analysis_geocode" in flat_df.columns:
                # Match the rest of the report (basic_mosaic totals)
                mask = flat_df["analysis_geocode"].astype(str).str.lower() == "basic_mosaic"
                flat_df = flat_df.loc[mask].copy()
            if "analysis_polygon_id" in flat_df.columns:
                polygon_ids = [
                    str(v).strip()
                    for v in flat_df["analysis_polygon_id"].dropna().astype(str).unique().tolist()
                    if str(v).strip()
                ]
        except Exception:
            polygon_ids = []

    gdf = _analysis_read_table(gpq_dir, _ANALYSIS_POLYGON_TABLE, geo=True)
    if gdf is None or getattr(gdf, "empty", True):
        return gpd.GeoDataFrame()

    # First try: explicit polygon ids from the flat table.
    if polygon_ids and "id" in gdf.columns:
        try:
            sub = gdf[gdf["id"].astype(str).isin(set(polygon_ids))].copy()
            if not sub.empty:
                return gpd.GeoDataFrame(sub)
        except Exception:
            pass

    # Fallback: use group_id link in the polygons table.
    if "group_id" not in gdf.columns:
        return gpd.GeoDataFrame(gdf)
    try:
        subset = gdf[gdf["group_id"].astype(str) == str(group_id)].copy()
        return gpd.GeoDataFrame(subset)
    except Exception:
        return gpd.GeoDataFrame()


def _analysis_flat_for_group(gpq_dir: str, group_id: str) -> pd.DataFrame:
    gdf = _analysis_read_table(gpq_dir, _ANALYSIS_FLAT_TABLE, geo=True)
    if gdf is None or getattr(gdf, "empty", True):
        return pd.DataFrame()
    df = pd.DataFrame(gdf.drop(columns=[c for c in ("geometry",) if c in gdf.columns]))
    if "analysis_group_id" not in df.columns:
        return pd.DataFrame()
    try:
        subset = _analysis_latest_run_filter(df, group_id)
    except Exception:
        return pd.DataFrame()
    if "analysis_geocode" in subset.columns:
        mask = subset["analysis_geocode"].astype(str).str.lower() == "basic_mosaic"
        subset = subset.loc[mask]
    return subset


def _analysis_flat_geo_for_group(gpq_dir: str, group_id: str) -> gpd.GeoDataFrame:
    """GeoDataFrame subset of tbl_analysis_flat for mapping.

    Note: tbl_analysis_flat geometry is typically per geocode-cell polygon (many per analysis polygon).
    This is ideal for visualising sensitivity overlays within the analysis area boundary.
    """
    gdf = _analysis_read_table(gpq_dir, _ANALYSIS_FLAT_TABLE, geo=True)
    if gdf is None or getattr(gdf, "empty", True) or "geometry" not in gdf.columns:
        return gpd.GeoDataFrame()
    try:
        df = pd.DataFrame(gdf)
        df = _analysis_latest_run_filter(df, group_id)
        if "analysis_geocode" in df.columns:
            mask = df["analysis_geocode"].astype(str).str.lower() == "basic_mosaic"
            df = df.loc[mask].copy()
        out = gpd.GeoDataFrame(df, geometry="geometry", crs=getattr(gdf, "crs", None))
        out = out.dropna(subset=["geometry"]).copy()
        return out
    except Exception:
        return gpd.GeoDataFrame()


def _analysis_sensitivity_totals_km2(flat_df: pd.DataFrame) -> pd.DataFrame:
    if flat_df is None or flat_df.empty:
        return pd.DataFrame(columns=["Code", "Description", "Area (km²)"])
    code_col = "sensitivity_code_max" if "sensitivity_code_max" in flat_df.columns else "sensitivity_code"
    desc_col = "sensitivity_description_max" if "sensitivity_description_max" in flat_df.columns else (
        "sensitivity_description" if "sensitivity_description" in flat_df.columns else None
    )
    area_col = "analysis_area_m2" if "analysis_area_m2" in flat_df.columns else None
    if code_col not in flat_df.columns or area_col is None:
        return pd.DataFrame(columns=["Code", "Description", "Area (km²)"])

    tmp = flat_df.copy()
    tmp[area_col] = pd.to_numeric(tmp[area_col], errors="coerce").fillna(0.0)
    tmp[code_col] = tmp[code_col].astype(str).str.strip().str.upper().replace({"NAN": "", "NONE": ""})

    grouped = tmp.groupby(code_col, dropna=False)[area_col].sum()
    rows: list[dict[str, object]] = []
    for code, area_m2 in grouped.items():
        code_u = str(code or "").strip().upper()
        if not code_u:
            continue
        if float(area_m2 or 0.0) <= 0:
            continue
        desc = ""
        if desc_col and desc_col in tmp.columns:
            try:
                desc_vals = tmp.loc[tmp[code_col].astype(str).str.upper() == code_u, desc_col].dropna()
                if not desc_vals.empty:
                    desc = str(desc_vals.iloc[-1]).strip()
            except Exception:
                desc = ""
        rows.append({
            "Code": code_u,
            "Description": desc,
            "Area (km²)": float(area_m2) / 1_000_000.0,
        })
    out = pd.DataFrame(rows, columns=["Code", "Description", "Area (km²)"])
    if out.empty:
        return out
    # Sort codes consistently (A–E), then any other codes, then UNKNOWN/blank.
    order_map = {c: i for i, c in enumerate(SENSITIVITY_ORDER)}

    def _code_key(val: str) -> tuple[int, str]:
        s = str(val or "").strip().upper()
        if s in order_map:
            return (0, f"{order_map[s]:02d}")
        if s in ("", "NONE", "NAN"):
            return (3, "")
        if s == "UNKNOWN":
            return (2, s)
        return (1, s)

    out["__code_key"] = out["Code"].map(_code_key)
    out.sort_values(["__code_key"], inplace=True, ignore_index=True)
    out.drop(columns=["__code_key"], inplace=True)
    return out


def _analysis_code_color(code: str, palette_A2E: dict) -> str:
    fallback = None
    try:
        fallback = (palette_A2E or {}).get('UNKNOWN')
    except Exception:
        fallback = None
    if not fallback:
        fallback = SENSITIVITY_UNKNOWN_COLOR if 'SENSITIVITY_UNKNOWN_COLOR' in globals() else "#6c757d"

    try:
        entry = (palette_A2E or {}).get(str(code).strip().upper())
        if isinstance(entry, str) and entry.strip():
            return entry.strip()
        if isinstance(entry, dict) and entry.get("color"):
            return str(entry["color"]).strip() or fallback
    except Exception:
        pass
    return fallback


def _analysis_write_area_map_png(
    polygons_gdf: gpd.GeoDataFrame,
    out_path: str,
    *,
    base_dir: str | None,
    config_path: str | None,
    palette_A2E: dict,
    desc_A2E: dict,
    overlay_alpha: float = 0.65,
    flat_cells_gdf: gpd.GeoDataFrame | None = None,
) -> bool:
    """Render analysis area map with basemap + sensitivity overlay (no MBTiles).

    - Border is drawn from tbl_analysis_polygons.parquet (true analysis polygons).
    - Sensitivity overlay is derived from tbl_analysis_flat.parquet (geocode-cell polygons).
    """
    try:
        if polygons_gdf is None or polygons_gdf.empty or "geometry" not in polygons_gdf.columns:
            return False
        gdf = polygons_gdf.dropna(subset=["geometry"]).copy()
        if gdf.empty:
            return False

        g3857 = _safe_to_3857(gdf)
        if g3857.empty:
            return False

        bounds_3857 = _expand_bounds(g3857.total_bounds, pad_ratio=0.08)

        # Union geometry for clipping/masking
        clip_geom = None
        try:
            if hasattr(g3857, "union_all"):
                clip_geom = g3857.union_all()
            else:
                clip_geom = g3857.unary_union
        except Exception:
            clip_geom = None

        fig_h_in = 4.2
        fig_w_in = 5.8
        dpi = 180
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        minx, miny, maxx, maxy = bounds_3857
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)
        try:
            ax.set_aspect('equal', adjustable='box')
        except Exception:
            pass

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        # Sensitivity overlay from flat geocode-cell polygons (dissolved by sensitivity_code_max)
        try:
            if flat_cells_gdf is not None and not getattr(flat_cells_gdf, 'empty', True) and 'geometry' in flat_cells_gdf.columns:
                cells = flat_cells_gdf.dropna(subset=['geometry']).copy()
                if not cells.empty:
                    cells_3857 = _safe_to_3857(cells)
                    if not cells_3857.empty:
                        code_col = 'sensitivity_code_max' if 'sensitivity_code_max' in cells_3857.columns else (
                            'sensitivity_code' if 'sensitivity_code' in cells_3857.columns else None
                        )
                        if code_col:
                            cells_3857['__sens_code'] = (
                                cells_3857[code_col].astype(str).str.strip().str.upper().replace({'NAN': '', 'NONE': ''})
                            )
                            cells_3857 = cells_3857[cells_3857['__sens_code'] != ''].copy()
                            if not cells_3857.empty:
                                # Dissolve to reduce draw cost and make output stable.
                                dissolved = cells_3857.dissolve(by='__sens_code', as_index=False)
                                dissolved['__order'] = dissolved['__sens_code'].map({c: i for i, c in enumerate(SENSITIVITY_ORDER)}).fillna(99)
                                dissolved = dissolved.sort_values(['__order', '__sens_code'])
                                for _, row in dissolved.iterrows():
                                    code = str(row.get('__sens_code', '')).strip().upper()
                                    geom = row.get('geometry')
                                    if not code or geom is None:
                                        continue
                                    col = _analysis_code_color(code, palette_A2E)
                                    try:
                                        gpd.GeoSeries([geom], crs=dissolved.crs).plot(
                                            ax=ax,
                                            color=col,
                                            alpha=float(max(0.05, min(0.95, overlay_alpha))),
                                            edgecolor='none',
                                            linewidth=0.0,
                                            zorder=10,
                                        )
                                    except Exception:
                                        pass
        except Exception:
            pass

        # Border outline from analysis polygons (single union outline)
        try:
            gborder = g3857.copy()
            union_geom = None
            try:
                if hasattr(gborder, "union_all"):
                    union_geom = gborder.union_all()
                else:
                    union_geom = gborder.unary_union
            except Exception:
                union_geom = None
            if union_geom is not None and not getattr(union_geom, 'is_empty', True):
                gpd.GeoSeries([union_geom], crs=gborder.crs).boundary.plot(
                    ax=ax,
                    color="#2f3e46",
                    linewidth=1.3,
                    alpha=0.9,
                    zorder=30,
                )
        except Exception:
            pass

        _add_map_decorations(ax, bounds_3857, base_dir=base_dir, add_inset=False)
        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        return False


def _geom_to_mpl_path(geom) -> MplPath | None:
    """Convert a (Multi)Polygon-like geometry to a Matplotlib Path."""
    try:
        if geom is None:
            return None
        if getattr(geom, "is_empty", True):
            return None
        gtype = getattr(geom, "geom_type", "")
        if gtype == "Polygon":
            polygons = [geom]
        elif gtype == "MultiPolygon":
            polygons = list(getattr(geom, "geoms", []) or [])
        else:
            return None

        vertices: list[tuple[float, float]] = []
        codes: list[int] = []

        def _add_ring(coords):
            pts = list(coords)
            if len(pts) < 3:
                return
            # Ensure closed ring
            if pts[0] != pts[-1]:
                pts.append(pts[0])
            vertices.extend([(float(x), float(y)) for x, y in pts])
            codes.extend([MplPath.MOVETO] + [MplPath.LINETO] * (len(pts) - 2) + [MplPath.CLOSEPOLY])

        for poly in polygons:
            try:
                _add_ring(poly.exterior.coords)
                for interior in list(getattr(poly, "interiors", []) or []):
                    _add_ring(interior.coords)
            except Exception:
                continue

        if not vertices:
            return None
        return MplPath(vertices, codes)
    except Exception:
        return None


def _analysis_write_area_polygon_only_png(
    polygons_3857: gpd.GeoDataFrame,
    out_path: str,
    *,
    bounds_3857: tuple[float, float, float, float],
    base_dir: str | None,
    fill_alpha: float = 0.12,
) -> bool:
    """Fallback: basemap clipped to the study area polygon, with filled polygon."""
    try:
        if polygons_3857 is None or polygons_3857.empty:
            return False

        try:
            if hasattr(polygons_3857, "union_all"):
                union_geom = polygons_3857.union_all()
            else:
                union_geom = polygons_3857.unary_union
        except Exception:
            union_geom = None

        if union_geom is None or getattr(union_geom, "is_empty", True):
            return False

        fig_h_in = 4.2
        fig_w_in = 5.8
        dpi = 180
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        minx, miny, maxx, maxy = bounds_3857
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)
        try:
            ax.set_aspect('equal', adjustable='box')
        except Exception:
            pass

        # Basemap should cover the full extent (do NOT clip it)
        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        # Fill polygon (to make the study area stand out)
        try:
            polygons_3857.plot(
                ax=ax,
                color=LIGHT_PRIMARY_HEX,
                alpha=float(max(0.0, min(1.0, fill_alpha))),
                edgecolor='none',
                linewidth=0.0,
                zorder=9,
            )
        except Exception:
            pass

        # Intentionally no boundary outline here.
        # Outlines (often blue) make multi-polygons look like separate highlighted shapes.

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        try:
            plt.close('all')
        except Exception:
            pass
        return False


def _analysis_write_minimap_png(polygons_gdf: gpd.GeoDataFrame, out_path: str) -> bool:
    try:
        if polygons_gdf is None or polygons_gdf.empty or "geometry" not in polygons_gdf.columns:
            return False
        gdf = polygons_gdf.dropna(subset=["geometry"]).copy()
        if gdf.empty:
            return False
        fig, ax = plt.subplots(figsize=(3.2, 2.2), dpi=170)
        ax.set_axis_off()
        try:
            gdf.boundary.plot(ax=ax, linewidth=1.2, color="#2f3e46")
            gdf.plot(ax=ax, color="#adb5bd", alpha=0.35, edgecolor="#2f3e46", linewidth=1.0)
        except Exception:
            gdf.plot(ax=ax, color="#adb5bd", alpha=0.35)
        try:
            ax.set_aspect("equal", adjustable="box")
        except Exception:
            pass
        fig.tight_layout(pad=0.1)
        fig.savefig(out_path, bbox_inches="tight")
        plt.close(fig)
        return True
    except Exception:
        return False


def _analysis_write_sensitivity_bar_png(totals_df: pd.DataFrame, palette_A2E: dict, out_path: str, title: str) -> bool:
    try:
        if totals_df is None or totals_df.empty:
            return False
        raw = totals_df[["Code", "Area (km²)"]].copy()
        raw["Code"] = raw["Code"].astype(str).str.strip().str.upper()

        order_map = {c: i for i, c in enumerate(SENSITIVITY_ORDER)}

        def _code_key(val: str) -> tuple[int, str]:
            s = str(val or "").strip().upper()
            if s in order_map:
                return (0, f"{order_map[s]:02d}")
            if s in ("", "NONE", "NAN"):
                return (3, "")
            if s == "UNKNOWN":
                return (2, s)
            return (1, s)

        raw["__code_key"] = raw["Code"].map(_code_key)
        raw.sort_values(["__code_key"], inplace=True)

        codes = raw["Code"].astype(str).tolist()
        areas = raw["Area (km²)"] .astype(float).tolist()
        colors = [_analysis_code_color(c, palette_A2E) for c in codes]

        fig, ax = plt.subplots(figsize=(5.8, 2.8), dpi=170)
        ax.bar(codes, areas, color=colors, edgecolor="#333333", linewidth=0.4)
        ax.set_title(title, fontsize=11)
        ax.set_ylabel("Area (km²)")
        ax.yaxis.set_major_locator(MaxNLocator(nbins=5))
        ax.grid(axis="y", alpha=0.25)
        fig.tight_layout(pad=0.6)
        fig.savefig(out_path, bbox_inches="tight")
        plt.close(fig)
        return True
    except Exception:
        return False


def _analysis_write_relative_share_png(
    left_totals_df: pd.DataFrame,
    right_totals_df: pd.DataFrame,
    palette_A2E: dict,
    out_path: str,
    left_label: str,
    right_label: str,
) -> bool:
    """Render relative (percentage) composition charts for two areas."""
    try:
        def _to_map(df: pd.DataFrame) -> dict[str, float]:
            if df is None or df.empty:
                return {}
            m: dict[str, float] = {}
            for _, r in df.iterrows():
                code = str(r.get('Code', '')).strip().upper()
                if not code or code in {'NONE', 'NAN'}:
                    continue
                try:
                    km2 = float(r.get('Area (km²)', 0.0) or 0.0)
                except Exception:
                    km2 = 0.0
                if km2 > 0:
                    m[code] = m.get(code, 0.0) + km2
            return m

        left_map = _to_map(left_totals_df)
        right_map = _to_map(right_totals_df)
        codes = [c for c in SENSITIVITY_ORDER if left_map.get(c, 0) > 0 or right_map.get(c, 0) > 0]
        # Add any extra codes deterministically
        extras = sorted({*left_map.keys(), *right_map.keys()} - set(SENSITIVITY_ORDER))
        codes = codes + extras
        if not codes:
            return False

        left_total = sum(left_map.get(c, 0.0) for c in codes)
        right_total = sum(right_map.get(c, 0.0) for c in codes)
        if left_total <= 0 and right_total <= 0:
            return False

        def _shares(area_map: dict[str, float], total: float) -> list[float]:
            if total <= 0:
                return [0.0 for _ in codes]
            return [max(area_map.get(c, 0.0), 0.0) / total for c in codes]

        left_sh = _shares(left_map, left_total)
        right_sh = _shares(right_map, right_total)

        fig, ax = plt.subplots(figsize=(6.4, 2.2), dpi=170)
        ax.set_axis_off()

        y_positions = [1.0, 0.35]
        labels = [left_label, right_label]
        bar_h = 0.18
        for y, shares, label in zip(y_positions, [left_sh, right_sh], labels):
            x = 0.0
            for code, share in zip(codes, shares):
                if share <= 0:
                    continue
                col = _analysis_code_color(code, palette_A2E)
                ax.add_patch(Rectangle((x, y), share, bar_h, facecolor=col, edgecolor='white', lw=0.6, alpha=0.95))
                if share >= 0.08:
                    ax.text(x + share / 2, y + bar_h / 2, f"{code} {share*100:.0f}%", ha='center', va='center', fontsize=8, color='white')
                x += share
            ax.text(-0.02, y + bar_h / 2, label, ha='right', va='center', fontsize=9)

        ax.set_xlim(-0.12, 1.02)
        ax.set_ylim(0, 1.35)
        ax.text(0.0, 1.28, "Relative sensitivity composition (100%)", fontsize=10, fontweight='bold')
        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        try:
            plt.close('all')
        except Exception:
            pass
        return False


def _analysis_write_sankey_difference_png(
    left_totals_df: pd.DataFrame,
    right_totals_df: pd.DataFrame,
    palette_A2E: dict,
    out_path: str,
    left_label: str,
    right_label: str,
    ribbon_alpha: float = 0.45,
) -> bool:
    """Sankey-style comparison: same-code ribbons between left/right blocks (like data_analysis_presentation.py)."""
    try:
        def _to_map(df: pd.DataFrame) -> dict[str, float]:
            if df is None or df.empty:
                return {}
            m: dict[str, float] = {}
            for _, r in df.iterrows():
                code = str(r.get('Code', '')).strip().upper()
                if not code or code in {'NONE', 'NAN'}:
                    continue
                try:
                    km2 = float(r.get('Area (km²)', 0.0) or 0.0)
                except Exception:
                    km2 = 0.0
                if km2 > 0:
                    m[code] = m.get(code, 0.0) + km2
            return m

        left_map = _to_map(left_totals_df)
        right_map = _to_map(right_totals_df)
        codes_in_data = [c for c in SENSITIVITY_ORDER if left_map.get(c, 0) > 0 or right_map.get(c, 0) > 0]
        extras = sorted({*left_map.keys(), *right_map.keys()} - set(SENSITIVITY_ORDER))
        codes = codes_in_data + extras
        if not codes:
            return False

        left_total = sum(left_map.get(c, 0.0) for c in codes)
        right_total = sum(right_map.get(c, 0.0) for c in codes)
        scale_total = max(left_total, right_total)
        if scale_total <= 0:
            return False

        def _build_blocks(area_map: dict[str, float]) -> dict[str, dict[str, float]]:
            # Build stacked blocks top-down so sensitivity ordering reads A (top) -> E (bottom).
            values = [max(area_map.get(code, 0.0), 0.0) for code in codes]
            min_height = 0.010
            gap = 0.012
            heights = [max((v / scale_total), min_height) if v > 0 else 0.0 for v in values]

            non_zero = [h for h in heights if h > 0]
            span = sum(non_zero) + gap * (len(non_zero) - 1 if len(non_zero) > 1 else 0)

            y_bottom = 0.06
            y_top = 0.88
            available = max(0.05, y_top - y_bottom)
            scale = 1.0 if span <= available else available / span

            blocks: dict[str, dict[str, float]] = {}
            cursor_top = y_top
            for code, height, value in zip(codes, heights, values):
                if height <= 0:
                    continue
                h = height * scale
                blocks[code] = {'start': cursor_top - h, 'end': cursor_top, 'value': value, 'height': h}
                cursor_top -= h + gap * scale
            return blocks

        left_blocks = _build_blocks(left_map)
        right_blocks = _build_blocks(right_map)

        # NOTE: In the PDF renderer we never scale images *up* (only down).
        # So to make this diagram larger on the page we must render a reasonably
        # large PNG and with a taller aspect ratio (more height relative to width).
        fig, ax = plt.subplots(figsize=(10.0, 4.0), dpi=170)
        ax.axis('off')

        # Push the two stacks closer to the edges to avoid large left/right buffers in the exported PNG.
        x_left, x_right = 0.12, 0.88
        bar_w = 0.14

        def _draw_blocks(blocks: dict[str, dict[str, float]], x: float, align: str):
            for code, info in blocks.items():
                col = _analysis_code_color(code, palette_A2E)
                rect = PathPatch(
                    MplPath(
                        [
                            (x - bar_w / 2, info['start']),
                            (x + bar_w / 2, info['start']),
                            (x + bar_w / 2, info['end']),
                            (x - bar_w / 2, info['end']),
                            (x - bar_w / 2, info['start']),
                        ],
                        [MplPath.MOVETO, MplPath.LINETO, MplPath.LINETO, MplPath.LINETO, MplPath.CLOSEPOLY],
                    ),
                    facecolor=col,
                    edgecolor='#333333',
                    lw=0.5,
                    alpha=0.90,
                )
                ax.add_patch(rect)
                label_y = (info['start'] + info['end']) / 2
                # Keep labels inside the bars; labels outside the axes make bbox_inches='tight' add margins.
                if align == 'left':
                    label_x = x - bar_w / 2 + 0.012
                    ha = 'left'
                else:
                    label_x = x + bar_w / 2 - 0.012
                    ha = 'right'
                ax.text(label_x, label_y, f"{code}", fontsize=9, ha=ha, va='center', color='black')

        def _draw_ribbon(lb: dict[str, float], rb: dict[str, float], color: str, delta_km2: float):
            verts = [
                (x_left + bar_w / 2, lb['start']),
                (x_right - bar_w / 2, rb['start']),
                (x_right - bar_w / 2, rb['end']),
                (x_left + bar_w / 2, lb['end']),
                (x_left + bar_w / 2, lb['start']),
            ]
            patch = PathPatch(
                MplPath(verts, [MplPath.MOVETO, MplPath.LINETO, MplPath.LINETO, MplPath.LINETO, MplPath.CLOSEPOLY]),
                facecolor=color,
                alpha=float(max(0.05, min(0.90, ribbon_alpha))),
                edgecolor='none',
            )
            ax.add_patch(patch)
            mid_x = (x_left + x_right) / 2
            mid_y = (lb['start'] + lb['end'] + rb['start'] + rb['end']) / 4
            # Keep labels compact; show 0.00 for tiny deltas
            try:
                dv = float(delta_km2)
            except Exception:
                dv = 0.0
            if not math.isfinite(dv) or abs(dv) < 0.005:
                txt = '0.00 km²'
            else:
                txt = f"{dv:+.2f} km²"
            ax.text(mid_x, mid_y, txt, fontsize=8, ha='center', va='center', color='black')

        _draw_blocks(left_blocks, x_left, 'left')
        _draw_blocks(right_blocks, x_right, 'right')

        for code in codes:
            lb = left_blocks.get(code)
            rb = right_blocks.get(code)
            if not lb or not rb:
                continue
            col = _analysis_code_color(code, palette_A2E)
            delta = left_map.get(code, 0.0) - right_map.get(code, 0.0)
            _draw_ribbon(lb, rb, col, delta)

        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.text(x_left, 0.96, left_label, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(x_right, 0.96, right_label, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(0.5, 0.99, "Area difference Sankey", ha='center', va='top', fontsize=10)

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        return True
    except Exception:
        try:
            plt.close('all')
        except Exception:
            pass
        return False

class ReportEngine:
    """
    Helper responsible for rendering all cartographic artefacts and tracking temporary files.
    """

    def __init__(self,
                 base_dir: str,
                 tmp_dir: str,
                 palette: dict,
                 desc: dict,
                 config_path: str):
        self.base_dir = base_dir
        self.tmp_dir = Path(tmp_dir)
        self.tmp_dir.mkdir(parents=True, exist_ok=True)
        self.palette = palette
        self.desc = desc
        self.config_path = config_path
        self.generated_paths: list[Path] = []

    def register(self, path: str | Path):
        self.generated_paths.append(Path(path))

    def make_path(self, *parts: str, suffix: str = ".png") -> str:
        name = "_".join(parts) + suffix
        full = self.tmp_dir / name
        self.register(full)
        return str(full)

    def cleanup(self):
        for path in self.generated_paths:
            try:
                if path.exists():
                    path.unlink()
            except Exception:
                pass
        self.generated_paths.clear()

    def render_geocode_maps(self,
                            flat_df: gpd.GeoDataFrame,
                            set_progress_callback) -> tuple[list, list, list]:
        pages: list = []
        intro_table = None
        groups = []
        if flat_df.empty or 'name_gis_geocodegroup' not in flat_df.columns:
            return pages, intro_table, groups

        # Report maps: only basic_mosaic (or configured basic_group_name)
        try:
            cfg_local = read_config(self.config_path)
            basic_name = (cfg_local['DEFAULT'].get('basic_group_name', 'basic_mosaic') or 'basic_mosaic').strip()
        except Exception:
            basic_name = 'basic_mosaic'

        all_groups = (flat_df['name_gis_geocodegroup']
                      .astype('string')
                      .dropna().unique().tolist())
        all_groups = sorted(all_groups)

        # Prefer configured basic group if present, else fall back to a literal basic_mosaic, else first group.
        chosen = None
        for cand in (basic_name, 'basic_mosaic'):
            if any(str(g).strip().lower() == str(cand).strip().lower() for g in all_groups):
                chosen = cand
                break
        if chosen is None and all_groups:
            chosen = str(all_groups[0])
        groups = [chosen] if chosen else []

        counts = (flat_df.groupby('name_gis_geocodegroup')
                  .agg(total=('name_gis_geocodegroup', 'size'),
                       populated=('geometry', lambda s: int(s.notna().sum())))
                  .reset_index())
        # Only show the chosen (basic) group in the intro table.
        try:
            if chosen:
                mask = counts['name_gis_geocodegroup'].astype('string').str.strip().str.lower() == str(chosen).strip().lower()
                counts = counts[mask].copy()
        except Exception:
            pass

        intro_table = [["Geocode category", "Total objects", "With geometry"]]
        for _, row in counts.sort_values('name_gis_geocodegroup').iterrows():
            intro_table.append([
                str(row['name_gis_geocodegroup']),
                int(row['total']),
                int(row['populated']),
            ])

        fixed_bounds_3857 = compute_fixed_bounds_3857(flat_df, base_dir=self.base_dir)

        done = 0
        for gname in groups:
            # Use MBTiles rasters when available (preferred)
            safe = _safe_name(gname)
            mb_root = os.path.join(self.base_dir, "output", "mbtiles")

            # Other maps (non-index): these should come before the index sections in the report.
            layers = [
                ("sensitivity_max", "Sensitive areas (A–E)", "sensitivity"),
                ("importance_max", "Importance (max)", "importance_max"),
                ("groupstotal", "Asset groups total", "groupstotal"),
                ("assetstotal", "Assets total", "assetstotal"),
            ]

            pages.append(('heading(2)', f"Other maps: {gname}"))
            pages.append(('text', "Rendered from MBTiles with basemap (basic_mosaic only)."))
            pages.append(('rule', None))

            for suffix, title, slug in layers:
                mb_path = os.path.join(mb_root, f"{safe}_{suffix}.mbtiles")
                out_png = self.make_path("geocode", safe, slug)
                ok = False
                note = ""
                if os.path.exists(mb_path):
                    ok, note = render_mbtiles_to_png_best_fit(mb_path, out_png, base_dir=self.base_dir)
                else:
                    note = f"Missing MBTiles: {os.path.basename(mb_path)}"

                if ok and _file_ok(out_png):
                    pages.append(('heading(3)', title))
                    pages.append(('text', "Two-line max intro: MBTiles overlay on basemap."))
                    pages.append(('image_map', (title, out_png)))
                    pages.append(('new_page', None))
                else:
                    # Fallback (only for sensitivity): render polygons (no borders) if MBTiles missing
                    if suffix in ("sensitivity_max", "sensitivity"):
                        sub = flat_df[flat_df['name_gis_geocodegroup'].astype('string').str.strip().str.lower() == str(gname).strip().lower()].copy()
                        ok_poly = draw_group_map_sensitivity(sub, gname, self.palette, self.desc, out_png, fixed_bounds_3857, base_dir=self.base_dir)
                        if ok_poly and _file_ok(out_png):
                            pages.append(('heading(3)', title))
                            pages.append(('text', "Two-line max intro: polygons on basemap."))
                            pages.append(('image_map', (title, out_png)))
                            pages.append(('new_page', None))
                            continue
                    pages.append(('heading(3)', title))
                    pages.append(('text', f"Could not render map: {note or 'unknown error'}"))
                    pages.append(('new_page', None))

            done += 1
            if set_progress_callback:
                set_progress_callback(done, max(1, len(groups)))

        return pages, intro_table, groups

    def render_atlas_maps(self,
                          flat_df: gpd.GeoDataFrame,
                          atlas_df: gpd.GeoDataFrame,
                          atlas_geocode_pref: str | None,
                          include_atlas_maps: bool,
                          set_progress_callback) -> tuple[list, str | None]:
        pages = []
        atlas_geocode_selected = None
        if not include_atlas_maps or atlas_df is None or atlas_df.empty:
            return pages, atlas_geocode_selected

        flat_geos = flat_df if flat_df is not None else gpd.GeoDataFrame()
        flat_polys_3857 = gpd.GeoDataFrame()
        if not flat_geos.empty and 'geometry' in flat_geos.columns:
            flat_polys = flat_geos[flat_geos.geometry.type.isin(['Polygon','MultiPolygon'])].copy()
            if 'name_gis_geocodegroup' in flat_polys.columns:
                flat_polys['name_gis_geocodegroup'] = (
                    flat_polys['name_gis_geocodegroup']
                    .astype('string')
                    .str.strip()
                )
            flat_polys_3857 = _safe_to_3857(flat_polys)

        atlas_crs = atlas_df.crs
        atlas_df = atlas_df[atlas_df['geometry'].notna()].copy()

        def _pick_geocode_level():
            nonlocal atlas_geocode_selected
            levels = []
            if not flat_df.empty and 'name_gis_geocodegroup' in flat_df.columns:
                levels = sorted(flat_df['name_gis_geocodegroup'].astype('string').dropna().unique().tolist())
            if atlas_geocode_pref and atlas_geocode_pref in levels:
                atlas_geocode_selected = atlas_geocode_pref
            elif 'basic_mosaic' in levels:
                atlas_geocode_selected = 'basic_mosaic'
            elif levels:
                atlas_geocode_selected = levels[0]

        _pick_geocode_level()

        polys_for_atlas = flat_polys_3857
        if atlas_geocode_selected and not flat_polys_3857.empty:
            if 'name_gis_geocodegroup' in flat_polys_3857.columns:
                atlas_mask = flat_polys_3857['name_gis_geocodegroup'].astype('string').str.lower() == atlas_geocode_selected.lower()
                filtered = flat_polys_3857[atlas_mask].copy()
                if filtered.empty:
                    write_to_log(f"No atlas polygons found for geocode '{atlas_geocode_selected}'. Using all polygons.", self.base_dir)
                else:
                    polys_for_atlas = filtered
            else:
                write_to_log("Atlas polygons missing 'name_gis_geocodegroup'; using all polygons.", self.base_dir)

        bounds = compute_fixed_bounds_3857(flat_df, base_dir=self.base_dir)
        overview_png = self.make_path("atlas", "overview")
        ok_overview = draw_atlas_overview_map(atlas_df, atlas_crs, polys_for_atlas, overview_png, bounds, base_dir=self.base_dir)
        if ok_overview and _file_ok(overview_png):
            text = "Overview of all atlas tiles within the study area."
            if atlas_geocode_selected:
                text += f" Geocode level shown: <b>{atlas_geocode_selected}</b>."
            pages += [
                ('heading(3)', "Atlas overview"),
                ('text', text),
                ('image', ("Atlas tiles overview", overview_png)),
                ('new_page', None),
            ]

        atlas_total = len(atlas_df)
        for idx, tile_row in atlas_df.iterrows():
            safe_tile = _safe_name(tile_row.get('name_gis') or f"atlas_{idx+1}")
            sens_png = self.make_path("atlas", safe_tile, "sens")

            ok_sens = draw_atlas_map(tile_row, atlas_crs, polys_for_atlas, self.palette, self.desc, sens_png, bounds, base_dir=self.base_dir)

            has_entries = False
            title_raw = tile_row.get('title_user') or tile_row.get('name_gis') or safe_tile
            tile_id = tile_row.get('name_gis') or safe_tile
            heading = f"Atlas tile: {title_raw}" if str(title_raw) == str(tile_id) else f"Atlas tile: {title_raw} ({tile_id})"
            pages.append(('heading(3)', heading))
            info_parts = []
            if isinstance(tile_row.get('description'), str) and tile_row.get('description').strip():
                info_parts.append(tile_row.get('description').strip())
            if atlas_geocode_selected:
                info_parts.append(f"Geocode level: <b>{atlas_geocode_selected}</b>.")
            info_parts.append("Inset highlights tile within the study area.")
            pages.append(('text', " ".join(info_parts)))

            if ok_sens and _file_ok(sens_png):
                pages.append(('text', "Sensitivity (A–E palette)."))
                pages.append(('image', ("Sensitivity atlas map", sens_png)))
                has_entries = True
            if has_entries:
                pages.append(('spacer', 1))
            else:
                pages.pop()  # remove heading
                pages.pop()  # remove text

            if set_progress_callback:
                set_progress_callback(idx+1, max(1, atlas_total))

        return pages, atlas_geocode_selected

    def render_index_statistics(self,
                               flat_df: gpd.GeoDataFrame,
                               cfg: configparser.ConfigParser,
                               set_progress_callback=None) -> list:
        """Create one page per index showing basic_mosaic distribution."""
        pages: list = []
        if flat_df is None or flat_df.empty:
            return pages

        basic_name = (cfg["DEFAULT"].get("basic_group_name", "basic_mosaic") or "basic_mosaic").strip()

        candidates = [
            ("index_importance", "Importance index", "index_importance"),
            ("index_sensitivity", "Sensitivity index", "index_sensitivity"),
            ("index_owa", "OWA index", "index_owa"),
        ]

        available = [(col, title, mb_suffix) for (col, title, mb_suffix) in candidates if col in flat_df.columns]
        if not available:
            return pages

        mb_root = os.path.join(self.base_dir, "output", "mbtiles")
        safe_basic = _safe_name(basic_name)

        total = len(available)
        for i, (col, title, mb_suffix) in enumerate(available, start=1):
            out_png = self.make_path("index", _safe_name(col), "distribution")
            ok, note = create_index_area_distribution_chart(
                flat_df,
                index_col=col,
                output_path=out_png,
                basic_group_name=basic_name,
                base_dir=self.base_dir,
            )
            if ok and _file_ok(out_png):
                pages += [
                    ('heading(2)', f"{title} – statistics"),
                    ('text',
                     f"Computed from <b>tbl_flat</b> for geocode group <b>{basic_name}</b> only. "
                     "Bars show total polygon area (km²) per index value. "
                     "Line shows number of categories per index value (A–E if available; otherwise number of cells)."),
                    ('image', (f"{title} – area distribution", out_png)),
                    ('new_page', None),
                ]
            else:
                msg = note or "(chart not available)"
                pages += [
                    ('heading(2)', f"{title} – statistics"),
                    ('text', f"Could not generate chart for <b>{col}</b>: {msg}"),
                    ('new_page', None),
                ]

            # Index map page (placed immediately after its index statistics page)
            mb_path = os.path.join(mb_root, f"{safe_basic}_{mb_suffix}.mbtiles")
            map_png = self.make_path("index", safe_basic, _safe_name(mb_suffix), "map")
            okm = False
            note_m = ""
            if os.path.exists(mb_path):
                okm, note_m = render_mbtiles_to_png_best_fit(mb_path, map_png, base_dir=self.base_dir)
            else:
                note_m = f"Missing MBTiles: {os.path.basename(mb_path)}"

            pages.append(('heading(2)', f"{title} – map"))
            pages.append(('text', "Two-line max intro: MBTiles overlay on basemap."))
            if okm and _file_ok(map_png):
                pages.append(('image_map', (f"{title} – map", map_png)))
            else:
                pages.append(('text', f"Could not render map: {note_m or 'unknown error'}"))
            pages.append(('new_page', None))

            if set_progress_callback:
                try:
                    set_progress_callback(i, total)
                except Exception:
                    pass

        return pages

    def render_segments(self,
                        lines_df: gpd.GeoDataFrame,
                        segments_df: gpd.GeoDataFrame,
                        palette: dict,
                        base_dir: str,
                        set_progress_callback):
        pages_lines = []
        log_data = []
        if (lines_df.empty or segments_df.empty or
            {'name_gis','segment_id','sensitivity_code_max','sensitivity_code_min'}.issubset(segments_df.columns) is False or
            'length_m' not in lines_df.columns or 'geometry' not in lines_df.columns):
            return pages_lines, log_data

        total = len(lines_df)
        for idx, line in lines_df.iterrows():
            ln_visible = line['name_gis']
            ln_safe = _safe_name(ln_visible)
            length_m = float(line.get('length_m', 0) or 0)
            length_km = length_m / 1000.0
            segment_records = sort_segments_numerically(segments_df[segments_df['name_gis'] == ln_visible]).copy()

            if not segment_records.empty:
                segment_records['__sens_code_max'] = segment_records.apply(
                    lambda row: _normalize_sensitivity_code(
                        row.get('sensitivity_code_max'),
                        row.get('sensitivity_max')
                    ),
                    axis=1
                )
                segment_records['__sens_code_min'] = segment_records.apply(
                    lambda row: _normalize_sensitivity_code(
                        row.get('sensitivity_code_min'),
                        row.get('sensitivity_min')
                    ),
                    axis=1
                )

            context_img = self.make_path("line", ln_safe, "context")
            seg_map_max = self.make_path("line", ln_safe, "segments_max")
            seg_map_min = self.make_path("line", ln_safe, "segments_min")

            ok_context = draw_line_context_map(line, context_img, pad_ratio=1.0, rect_buffer_ratio=0.03, base_dir=base_dir)
            ok_max = draw_line_segments_map(segments_df, ln_visible, palette, seg_map_max, mode='max', pad_ratio=0.20, base_dir=base_dir)
            ok_min = draw_line_segments_map(segments_df, ln_visible, palette, seg_map_min, mode='min', pad_ratio=0.20, base_dir=base_dir)

            max_stats_img = self.make_path("line", ln_safe, "max_dist")
            min_stats_img = self.make_path("line", ln_safe, "min_dist")
            max_codes = segment_records.get('__sens_code_max', segment_records.get('sensitivity_code_max'))
            min_codes = segment_records.get('__sens_code_min', segment_records.get('sensitivity_code_min'))
            if max_codes is None:
                max_codes = pd.Series(dtype=object)
            if min_codes is None:
                min_codes = pd.Series(dtype=object)
            create_sensitivity_summary(max_codes, palette, max_stats_img)
            create_sensitivity_summary(min_codes, palette, min_stats_img)

            max_img = self.make_path("line", ln_safe, "max_ribbon")
            min_img = self.make_path("line", ln_safe, "min_ribbon")
            create_line_statistic_image(ln_visible, max_codes, palette, length_m, max_img)
            create_line_statistic_image(ln_visible, min_codes, palette, length_m, min_img)

            first_page = [
                ('heading(2)', f"Line: {ln_visible}"),
                ('text', f"This section summarizes sensitivity along the line <b>{ln_visible}</b> "
                         f"(total length <b>{length_km:.2f} km</b>, segments: <b>{len(segment_records)}</b>). "
                         "Below: geographical context and segments maps colored by sensitivity values, "
                         "followed by the distribution and a ribbon (maximum sensitivity).")
            ]

            if ok_context and _file_ok(context_img):
                first_page.append(('image', ("Geographical context", context_img)))
            if ok_max and _file_ok(seg_map_max):
                first_page.append(('image', ("Segments colored by maximum sensitivity", seg_map_max)))
            first_page.append(('image', ("Maximum sensitivity – distribution", max_stats_img)))
            first_page.append(('text', f"Distance (km): 0 – {length_km/2:.1f} – {length_km:.1f}"))
            first_page.append(('image_ribbon', ("Maximum sensitivity – along line", max_img)))
            first_page.append(('new_page', None))

            second_page = [
                ('heading(2)', f"Line: {ln_visible} (continued)"),
                ('text', "Minimum sensitivity map, distribution, and ribbon.")
            ]
            if ok_min and _file_ok(seg_map_min):
                second_page.append(('image', ("Segments colored by minimum sensitivity", seg_map_min)))
            second_page.append(('image', ("Minimum sensitivity – distribution", min_stats_img)))
            second_page.append(('text', f"Distance (km): 0 – {length_km/2:.1f} – {length_km:.1f}"))
            second_page.append(('image_ribbon', ("Minimum sensitivity – along line", min_img)))
            second_page.append(('new_page', None))

            pages_lines += first_page + second_page

            for _, seg in segment_records.iterrows():
                log_data.append({
                    'line_name': ln_visible,
                    'segment_id': seg['segment_id'],
                    'sensitivity_code_max': seg['sensitivity_code_max'],
                    'sensitivity_code_min': seg['sensitivity_code_min']
                })

            if set_progress_callback:
                set_progress_callback(idx+1, total)

        return pages_lines, log_data

# Primary UI color (Steel Blue by default)
PRIMARY_HEX = "#4682B4"        # Steel blue
LIGHT_PRIMARY_HEX = "#6fa6cf"  # lighter steel-blue

# ---------------- Config + path helpers ----------------
def read_config(file_name: str) -> configparser.ConfigParser:
    cfg = configparser.ConfigParser()
    try:
        cfg.read(file_name, encoding="utf-8")
    except Exception:
        pass
    if "DEFAULT" not in cfg:
        cfg["DEFAULT"] = {}
    return cfg

def config_path(base_dir: str) -> str:
    """Flat config: <base>/config.ini"""
    return os.path.join(base_dir, "config.ini")

def parquet_dir_from_cfg(base_dir: str, cfg: configparser.ConfigParser) -> str:
    sub = cfg["DEFAULT"].get("parquet_folder", "output/geoparquet")
    required = ("tbl_asset_object.parquet", "tbl_asset_group.parquet", "tbl_flat.parquet")

    base_path = Path(base_dir)
    sub_path = Path(sub)

    candidates: list[Path] = []
    seen: set[str] = set()

    def _register(path: Path):
        try:
            resolved = path.resolve()
        except Exception:
            resolved = path
        key = str(resolved)
        if key not in seen:
            seen.add(key)
            candidates.append(resolved)

    if sub_path.is_absolute():
        _register(sub_path)
    else:
        ancestors = [base_path]
        ancestors.extend(list(base_path.parents)[:4])
        for ancestor in ancestors:
            _register(ancestor / sub_path)
            _register(ancestor / "code" / sub_path)

    primary = candidates[0] if candidates else (base_path / sub_path)
    for cand in candidates:
        if all((cand / req).exists() for req in required):
            if cand != primary:
                write_to_log(f"Using parquet folder at {cand}", base_dir)
            return str(cand)

    try:
        primary.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    if all((primary / req).exists() for req in required):
        return str(primary)

    write_to_log(f"Expected parquet tables not found; using {primary}", base_dir)
    return str(primary)

# ---------------- Base dir normalization ----------------
def normalize_base_dir(path_str: str) -> str:
    """Return the logical app root for both .py and packaged .exe runs.
    - If started inside tools/ or system/ or code/, climb to parent.
    - Then climb up a few levels to find a folder containing 'output' & 'input'
      or containing 'tools' and 'config.ini'.
    """
    try:
        p = Path(path_str).resolve()
    except Exception:
        return path_str

    if p.name.lower() in ("tools", "system", "code"):
        p = p.parent

    q = p
    for _ in range(4):
        try:
            if (q / "output").exists() and (q / "input").exists():
                return str(q)
            if (q / "tools").exists() and (q / "config.ini").exists():
                return str(q)
        except Exception:
            pass
        q = q.parent
    return str(p)

# ---------------- Logging ----------------
def write_to_log(message: str, base_dir: str | None = None):
    timestamp = datetime.datetime.now().strftime("%Y.%m.%d %H:%M:%S")
    formatted = f"{timestamp} - {message}"
    try:
        if base_dir:
            with open(os.path.join(base_dir, "log.txt"), "a", encoding="utf-8") as f:
                f.write(formatted + "\n")
    except Exception:
        pass
    try:
        if log_widget and log_widget.winfo_exists():
            log_widget.insert(tk.END, formatted + "\n")
            log_widget.see(tk.END)
    except Exception:
        pass

_OUTPUT_ROOT: Path | None = None

def _output_candidates(base_dir: str | None) -> list[Path]:
    base = Path(base_dir) if base_dir else Path.cwd()
    try:
        base = base.resolve()
    except Exception:
        pass
    candidates: list[Path] = []
    # Prefer <base>/output (workspace-root output) whenever possible.
    # Only fall back to <base>/code/output if needed (e.g. when running with base_dir pointing at repo root
    # but output exists only under code/).
    candidates.append((base / "output").resolve())
    if base.name.lower() != "code":
        code_dir = (base / "code")
        if code_dir.exists():
            candidates.append((code_dir / "output").resolve())
    return candidates

def _resolve_output_root(base_dir: str | None) -> Path:
    global _OUTPUT_ROOT
    if _OUTPUT_ROOT is not None:
        return _OUTPUT_ROOT
    candidates = _output_candidates(base_dir)
    for idx, cand in enumerate(candidates):
        if cand.exists():
            _OUTPUT_ROOT = cand
            if idx > 0 and base_dir:
                write_to_log(f"Using fallback output folder: {cand}", base_dir)
            return _OUTPUT_ROOT
    target = candidates[0]
    try:
        target.mkdir(parents=True, exist_ok=True)
        if base_dir:
            write_to_log(f"Created output folder at {target}", base_dir)
    except Exception:
        pass
    _OUTPUT_ROOT = target
    return _OUTPUT_ROOT

def output_subpath(base_dir: str | None, *parts: str) -> Path:
    root = _resolve_output_root(base_dir)
    path = root.joinpath(*parts)
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    return path

# ---------------- Small utilities ----------------
def _file_ok(path: str) -> bool:
    try:
        return os.path.isfile(path) and os.path.getsize(path) > 0
    except Exception:
        return False

def _safe_name(name: str) -> str:
    # Windows-safe file name (keep visible name unchanged in titles)
    return re.sub(r'[^A-Za-z0-9_.-]+', '_', str(name))

def _cfg_getboolean(cfg: configparser.ConfigParser,
                    section: str,
                    option: str,
                    default: bool = False) -> bool:
    """
    Robust boolean reader for configparser with permissive parsing.
    Accepts truthy values like '1', 'true', 'yes', 'on'.
    """
    try:
        raw = cfg.get(section, option, fallback=None)
    except Exception:
        raw = None
    if raw is None:
        return default
    val = str(raw).strip().lower()
    if val in {"1", "true", "yes", "on"}:
        return True
    if val in {"0", "false", "no", "off"}:
        return False
    return default

def _available_geocode_levels(base_dir: str,
                              config_file: str) -> list[str]:
    """
    Return sorted list of geocode group names available in tbl_flat.
    Used for GUI combo box population. Falls back to an empty list on failure.
    """
    try:
        cfg = read_config(config_file)
        gpq_dir = parquet_dir_from_cfg(base_dir, cfg)
        flat_df = load_tbl_flat(gpq_dir, base_dir=base_dir)
        if flat_df.empty or 'name_gis_geocodegroup' not in flat_df.columns:
            return []
        levels = (flat_df['name_gis_geocodegroup']
                  .astype('string')
                  .dropna()
                  .unique()
                  .tolist())
        return sorted(levels)
    except Exception as exc:
        write_to_log(f"Failed to enumerate geocode levels for UI: {exc}", base_dir)
        return []

def _dpi_for_fig_height(fig_height_in: float, px_cap: int = MAX_MAP_PX_HEIGHT, min_dpi: int = 110, max_dpi: int = 300) -> int:
    """
    Choose a DPI so that fig_height_in * dpi <= px_cap, bounded by min/max dpi.
    """
    try:
        dpi_by_cap = int(px_cap / max(fig_height_in, 1e-6))
        return max(min_dpi, min(max_dpi, dpi_by_cap))
    except Exception:
        return 150

# ---------------- Palette / descriptions ----------------
def read_sensitivity_palette_and_desc(file_name: str):
    """
    Reads color palette and descriptions for A-E from config.ini.
    Returns: (colors: dict{A..E->hex, 'UNKNOWN'->hex}, desc: dict{A..E->str})
    """
    global SENSITIVITY_UNKNOWN_COLOR, _SENSITIVITY_NUMERIC_RANGES
    cfg = configparser.ConfigParser()
    try:
        cfg.read(file_name, encoding="utf-8")
    except Exception:
        pass

    unknown_col = '#BDBDBD'
    if cfg.has_section('VALID_VALUES'):
        try:
            unknown_col = cfg['VALID_VALUES'].get('category_colour_unknown', unknown_col).strip() or unknown_col
        except Exception:
            pass
    SENSITIVITY_UNKNOWN_COLOR = unknown_col

    ranges: list[tuple[str, float, float]] = []

    def _parse_range(range_str: str):
        try:
            parts = re.split(r'\s*[-]\s*', range_str.strip())
            if len(parts) == 2:
                lo = float(parts[0])
                hi = float(parts[1])
                if lo > hi:
                    lo, hi = hi, lo
                return lo, hi
        except Exception:
            return None
        return None

    colors_map, desc_map = {}, {}
    for code in ['A','B','C','D','E']:
        if cfg.has_section(code):
            col = cfg[code].get('category_colour', '').strip() or unknown_col
            colors_map[code] = col
            desc_map[code] = cfg[code].get('description', '').strip()
            range_str = cfg[code].get('range', '').strip()
            parsed = _parse_range(range_str) if range_str else None
            if parsed:
                ranges.append((code, parsed[0], parsed[1]))
        else:
            colors_map[code] = unknown_col
            desc_map[code] = ''
    colors_map['UNKNOWN'] = unknown_col
    _SENSITIVITY_NUMERIC_RANGES = ranges
    return colors_map, desc_map

def _normalize_sensitivity_code(code_val, numeric_val):
    if code_val is not None and not pd.isna(code_val):
        try:
            code_str = str(code_val).strip().upper()
        except Exception:
            code_str = str(code_val).upper()
        if code_str in SENSITIVITY_ORDER:
            return code_str
    numeric_code = _sensitivity_code_from_numeric(numeric_val)
    if numeric_code in SENSITIVITY_ORDER:
        return numeric_code
    return 'UNKNOWN'

def _prepare_sensitivity_annotations(df: gpd.GeoDataFrame,
                                     code_column: str,
                                     numeric_column: str) -> gpd.GeoDataFrame:
    if df.empty:
        return df.copy()
    codes = df.apply(
        lambda row: _normalize_sensitivity_code(
            row.get(code_column),
            row.get(numeric_column)
        ),
        axis=1
    )
    out = df.copy()
    out['__sens_code'] = codes.fillna('UNKNOWN')
    return out


def _colors_from_annotations(gdf: gpd.GeoDataFrame,
                             palette: dict[str, str]) -> pd.Series:
    """
    Map normalized sensitivity codes stored in '__sens_code' to palette colors.
    Falls back to the configured UNKNOWN color when missing.
    """
    fallback = palette.get('UNKNOWN', SENSITIVITY_UNKNOWN_COLOR)

    def _lookup(code):
        if pd.isna(code):
            return fallback
        try:
            key = str(code).strip().upper()
        except Exception:
            key = str(code).upper()
        return palette.get(key, fallback)

    return gdf['__sens_code'].apply(_lookup)

def _sensitivity_code_from_numeric(value: float | int | None) -> str | None:
    try:
        if value is None or (isinstance(value, float) and not math.isfinite(value)):
            return None
        val = float(value)
    except Exception:
        return None
    for code, lo, hi in _SENSITIVITY_NUMERIC_RANGES:
        if lo <= val <= hi:
            return code
    return None

def _resolve_sensitivity_color(code_value,
                               numeric_value,
                               palette: dict[str, str]) -> str:
    if code_value is not None:
        try:
            code_str = str(code_value).strip().upper()
        except Exception:
            code_str = str(code_value).upper()
        if code_str:
            color = palette.get(code_str)
            if color:
                return color
    numeric_code = _sensitivity_code_from_numeric(numeric_value)
    if numeric_code:
        color = palette.get(numeric_code)
        if color:
            return color
    return palette.get('UNKNOWN', SENSITIVITY_UNKNOWN_COLOR)



# ---------------- Generic plotting utils ----------------
def _safe_to_3857(gdf: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    if gdf.empty:
        return gdf
    g = gdf.copy()
    try:
        if g.crs is None:
            g = g.set_crs(4326)
        try:
            epsg = g.crs.to_epsg()
        except Exception:
            epsg = None
        if epsg != 3857:
            g = g.to_crs(3857)
    except Exception:
        try:
            g = g.set_crs(3857, allow_override=True)
        except Exception:
            pass
    return g

def _tile_cache_root(base_dir: str | None) -> Path:
    """
    Resolve the cache directory used for XYZ tiles. Defaults to <base>/output/tile_cache.
    """
    cache = output_subpath(base_dir, "tile_cache")
    try:
        cache.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass

    # Help users locate the cache on disk (log once per process).
    global _OSM_TILE_CACHE_LOGGED
    if not _OSM_TILE_CACHE_LOGGED and base_dir:
        _OSM_TILE_CACHE_LOGGED = True
        try:
            write_to_log(f"OSM tile cache folder: {cache}", base_dir)
        except Exception:
            pass
    return cache

def _load_cached_tile(cache_path: Path):
    if not cache_path.exists():
        return None
    if TILE_CACHE_MAX_AGE_DAYS > 0:
        try:
            age = time.time() - cache_path.stat().st_mtime
            if age > TILE_CACHE_MAX_AGE_DAYS * 86400:
                try:
                    cache_path.unlink()
                except Exception:
                    pass
                return None
        except Exception:
            return None
    try:
        with PILImage.open(cache_path) as pil_img:
            return pil_img.convert("RGB")
    except Exception:
        try:
            cache_path.unlink()
        except Exception:
            pass
        return None

def _save_tile_to_cache(cache_path: Path, data: bytes) -> None:
    try:
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        with open(cache_path, "wb") as f:
            f.write(data)
    except Exception:
        pass

def _plot_basemap(ax, crs_epsg=3857, base_dir: str | None = None):
    """Add a basemap under current axes.
    Prefers contextily; if unavailable, fetches OSM Web Mercator tiles and composites them.
    """
    mode = str(_REPORT_BASEMAP_MODE or "xyz").strip().lower()

    # Preferred path: contextily (unless forced xyz)
    if mode != "xyz" and ctx is not None:
        try:
            # We render attribution as plain text under the map in the DOCX output.
            # Try to suppress attribution embedded in the map image.
            try:
                ctx.add_basemap(
                    ax,
                    crs=f"EPSG:{crs_epsg}",
                    source=ctx.providers.OpenStreetMap.Mapnik,
                    attribution=None,
                )
            except TypeError:
                ctx.add_basemap(ax, crs=f"EPSG:{crs_epsg}", source=ctx.providers.OpenStreetMap.Mapnik)

            if base_dir:
                try:
                    write_to_log("Basemap source: contextily (no local XYZ tile cache used)", base_dir)
                except Exception:
                    pass

            # Defensive cleanup: remove any attribution text artists if present
            try:
                for t in list(getattr(ax, "texts", []) or []):
                    s = (getattr(t, "get_text", lambda: "")() or "")
                    if ("openstreetmap" in s.lower()) or ("©" in s) or ("copyright" in s.lower()):
                        try:
                            t.remove()
                        except Exception:
                            pass
            except Exception:
                pass
            return
        except Exception as e:
            write_to_log(f"Basemap via contextily failed, falling back to tiles: {e}", base_dir)

    # If contextily is forced, do not proceed with XYZ fallback.
    if mode == "contextily":
        return

    # Fallback: simple OSM tile fetch/composite in EPSG:3857 only
    if int(crs_epsg) != 3857:
        write_to_log("Basemap fallback only supports EPSG:3857; skipping.", base_dir)
        return

    try:
        # Current view in meters (WebMercator)
        minx, maxx = ax.get_xlim()
        miny, maxy = ax.get_ylim()

        def merc_to_lonlat(x, y):
            R = 6378137.0
            lon = (x / R) * 180.0 / math.pi
            lat = (2.0 * math.atan(math.exp(y / R)) - math.pi/2.0) * 180.0 / math.pi
            return lon, lat

        def lonlat_to_tile(lon, lat, z):
            n = 2 ** z
            x = int((lon + 180.0) / 360.0 * n)
            lat = max(-85.05112878, min(85.05112878, float(lat)))
            lat_rad = math.radians(lat)
            y = int((1.0 - math.log(math.tan(lat_rad) + 1.0 / math.cos(lat_rad)) / math.pi) / 2.0 * n)
            return x, y

        def tile_bounds_lonlat(z, x, y):
            n = 2.0 ** z
            minlon = x / n * 360.0 - 180.0
            maxlon = (x + 1) / n * 360.0 - 180.0
            def tiley_to_lat(t):
                Y = t / n
                lat_rad = math.atan(math.sinh(math.pi * (1 - 2 * Y)))
                return math.degrees(lat_rad)
            maxlat = tiley_to_lat(y)
            minlat = tiley_to_lat(y + 1)
            return (minlon, minlat, maxlon, maxlat)

        def lonlat_to_merc(lon, lat):
            R = 6378137.0
            x = lon * math.pi / 180.0 * R
            lat = max(-85.05112878, min(85.05112878, float(lat)))
            y = math.log(math.tan((90.0 + lat) * math.pi / 360.0)) * R
            return x, y

        # Decide a zoom level based on axis pixel size to keep labels crisp.
        # Aim for mosaic pixel dims >= 1.5x axes pixels (oversample to reduce blur).
        bbox = ax.get_window_extent()
        try:
            ax_pix_w = int(bbox.width)
            ax_pix_h = int(bbox.height)
        except Exception:
            # Fallback: estimate from figure size
            ax_pix_w = int(ax.figure.dpi * ax.figure.get_size_inches()[0])
            ax_pix_h = int(ax.figure.dpi * ax.figure.get_size_inches()[1])

        lon0, lat1 = merc_to_lonlat(minx, maxy)
        lon1, lat0 = merc_to_lonlat(maxx, miny)

        target_w = max(512, int(ax_pix_w * 1.5))
        target_h = max(512, int(ax_pix_h * 1.5))

        z = 3
        for test in range(3, 20):
            x0, y0 = lonlat_to_tile(lon0, lat0, test)
            x1, y1 = lonlat_to_tile(lon1, lat1, test)
            w = (abs(x1 - x0) + 1) * 256
            h = (abs(y1 - y0) + 1) * 256
            if w >= target_w and h >= target_h:
                z = test
                break

        tile_limit = 400 if getattr(sys, "frozen", False) else 900
        def _tile_range(z_level):
            tx0, ty0 = lonlat_to_tile(lon0, lat0, z_level)
            tx1, ty1 = lonlat_to_tile(lon1, lat1, z_level)
            return (min(tx0, tx1), max(tx0, tx1), min(ty0, ty1), max(ty0, ty1))

        xmin, xmax, ymin, ymax = _tile_range(z)
        tile_span = (xmax - xmin + 1) * (ymax - ymin + 1)
        while tile_span > tile_limit and z > 3:
            z -= 1
            xmin, xmax, ymin, ymax = _tile_range(z)
            tile_span = (xmax - xmin + 1) * (ymax - ymin + 1)

        if tile_span > tile_limit:
            try:
                ax.set_facecolor("#f5f5f5")
            except Exception:
                pass
            write_to_log(f"Basemap fallback skipped; tile span too large ({tile_span} tiles @ z{z}).", base_dir)
            return

        if tile_span > tile_limit * 0.6:
            write_to_log(f"Basemap fallback using reduced zoom z{z} ({tile_span} tiles).", base_dir)

        TILE = 256
        W, H = (xmax - xmin + 1) * TILE, (ymax - ymin + 1) * TILE

        mosaic = PILImage.new("RGB", (W, H), (240, 240, 240))

        ua = ("Mozilla/5.0 (compatible; MESA-Report/1.0)")
        opener = urllib.request.build_opener()
        opener.addheaders = [("User-Agent", ua)]
        urllib.request.install_opener(opener)

        def tile_url(z, x, y):
            return f"https://tile.openstreetmap.org/{z}/{x}/{y}.png"

        if base_dir:
            try:
                write_to_log("Basemap source: XYZ tiles (OpenStreetMap) with local cache", base_dir)
            except Exception:
                pass
        cache_root = _tile_cache_root(base_dir)
        cache_hits = 0
        fetched = 0

        for xi, x in enumerate(range(xmin, xmax+1)):
            for yi, y in enumerate(range(ymin, ymax+1)):
                cache_path = cache_root / str(z) / str(x) / f"{y}.png"
                img = _load_cached_tile(cache_path)
                if img is not None:
                    cache_hits += 1
                else:
                    try:
                        url = tile_url(z, x, y)
                        with urllib.request.urlopen(url, timeout=5) as resp:
                            data = resp.read()
                        img = PILImage.open(io.BytesIO(data)).convert("RGB")
                        _save_tile_to_cache(cache_path, data)
                        fetched += 1
                    except Exception:
                        img = None
                if img is not None:
                    mosaic.paste(img, (xi*TILE, yi*TILE))

        if (cache_hits or fetched) and base_dir:
            try:
                write_to_log(f"Basemap tiles z{z}: {cache_hits} cached, {fetched} fetched ({tile_span} total).", base_dir)
            except Exception:
                pass

        # Extent of mosaic in Mercator meters (XYZ scheme)
        west  = tile_bounds_lonlat(z, xmin, ymin)[0]  # minlon of top-left tile
        north = tile_bounds_lonlat(z, xmin, ymin)[3]  # maxlat of top-left tile
        east  = tile_bounds_lonlat(z, xmax, ymax)[2]  # maxlon of bottom-right tile
        south = tile_bounds_lonlat(z, xmax, ymax)[1]  # minlat of bottom-right tile

        lx, north_y = lonlat_to_merc(west,  north)
        rx, south_y = lonlat_to_merc(east,  south)

        # Show with correct extent; origin='upper' so row 0 is at 'north'
        ax.imshow(mosaic, extent=[lx, rx, south_y, north_y], origin="upper",
                  interpolation='bilinear', resample=True)
    except Exception as e:
        write_to_log(f"Basemap fallback failed: {e}", base_dir)

def _expand_bounds(bounds, pad_ratio=0.08):
    minx, miny, maxx, maxy = bounds
    dx, dy = maxx - minx, maxy - miny
    if dx <= 0 or dy <= 0:
        return bounds
    pad_x, pad_y = dx * pad_ratio, dy * pad_ratio
    return (minx - pad_x, miny - pad_y, maxx + pad_x, maxy + pad_y)

# ---------------- MBTiles rendering (raster) ----------------
WEBMERCATOR_MAX_LAT = 85.05112878
MAX_MAP_PX_WIDTH = 2000

def _clip_latlon_bounds(bounds: tuple[float, float, float, float]) -> tuple[float, float, float, float]:
    minlon, minlat, maxlon, maxlat = bounds
    minlon = max(-180.0, min(180.0, float(minlon)))
    maxlon = max(-180.0, min(180.0, float(maxlon)))
    minlat = max(-WEBMERCATOR_MAX_LAT, min(WEBMERCATOR_MAX_LAT, float(minlat)))
    maxlat = max(-WEBMERCATOR_MAX_LAT, min(WEBMERCATOR_MAX_LAT, float(maxlat)))
    if maxlon < minlon:
        minlon, maxlon = maxlon, minlon
    if maxlat < minlat:
        minlat, maxlat = maxlat, minlat
    # Avoid exact-180/85 edges causing off-by-one tile selection
    eps = 1e-9
    minlon = max(-180.0 + eps, minlon)
    maxlon = min(180.0 - eps, maxlon)
    minlat = max(-WEBMERCATOR_MAX_LAT + eps, minlat)
    maxlat = min(WEBMERCATOR_MAX_LAT - eps, maxlat)
    return minlon, minlat, maxlon, maxlat

def _lonlat_to_tile_fraction(lon: float, lat: float, z: int) -> tuple[float, float]:
    lat = max(-WEBMERCATOR_MAX_LAT, min(WEBMERCATOR_MAX_LAT, float(lat)))
    lon = max(-180.0, min(180.0, float(lon)))
    n = 2 ** int(z)
    x = (lon + 180.0) / 360.0 * n
    lat_rad = math.radians(lat)
    y = (1.0 - math.log(math.tan(lat_rad) + (1.0 / math.cos(lat_rad))) / math.pi) / 2.0 * n
    return x, y

def _mbtiles_metadata(conn: sqlite3.Connection) -> dict[str, str]:
    meta: dict[str, str] = {}
    try:
        cur = conn.cursor()
        cur.execute("SELECT name, value FROM metadata")
        for name, value in cur.fetchall():
            if name is None:
                continue
            meta[str(name)] = "" if value is None else str(value)
    except Exception:
        pass
    return meta

def _parse_mbtiles_bounds(meta: dict[str, str]) -> tuple[float, float, float, float] | None:
    raw = (meta.get("bounds") or "").strip()
    if not raw:
        return None
    try:
        parts = [float(x.strip()) for x in raw.split(",")]
        if len(parts) != 4:
            return None
        return _clip_latlon_bounds((parts[0], parts[1], parts[2], parts[3]))
    except Exception:
        return None

def _parse_mbtiles_zoom(meta: dict[str, str], key: str, default: int) -> int:
    try:
        return int(float((meta.get(key) or str(default)).strip()))
    except Exception:
        return int(default)

def _get_mbtiles_tile_bytes(conn: sqlite3.Connection, z: int, x: int, y_xyz: int) -> bytes | None:
    """Fetch tile bytes, trying TMS row convention first, then XYZ."""
    try:
        cur = conn.cursor()
        n = 2 ** int(z)
        y_tms = (n - 1) - int(y_xyz)
        cur.execute(
            "SELECT tile_data FROM tiles WHERE zoom_level=? AND tile_column=? AND tile_row=?",
            (int(z), int(x), int(y_tms)),
        )
        row = cur.fetchone()
        if row and row[0] is not None:
            return bytes(row[0])
        cur.execute(
            "SELECT tile_data FROM tiles WHERE zoom_level=? AND tile_column=? AND tile_row=?",
            (int(z), int(x), int(y_xyz)),
        )
        row = cur.fetchone()
        if row and row[0] is not None:
            return bytes(row[0])
    except Exception:
        return None
    return None

def lonlat_to_merc(lon: float, lat: float) -> tuple[float, float]:
    """Convert lon/lat (EPSG:4326) to WebMercator meters (EPSG:3857)."""
    R = 6378137.0
    lon_f = float(lon)
    lat_f = max(-85.05112878, min(85.05112878, float(lat)))
    x = lon_f * math.pi / 180.0 * R
    y = math.log(math.tan((90.0 + lat_f) * math.pi / 360.0)) * R
    return x, y


def merc_to_lonlat(x: float, y: float) -> tuple[float, float]:
    """Convert WebMercator meters (EPSG:3857) to lon/lat (EPSG:4326)."""
    R = 6378137.0
    lon = (float(x) / R) * 180.0 / math.pi
    lat = (2.0 * math.atan(math.exp(float(y) / R)) - math.pi / 2.0) * 180.0 / math.pi
    return lon, lat

def _nice_scale_length_m(target_m: float) -> float:
    """Pick a human-friendly scale bar length in meters (1/2/5 * 10^n) <= target."""
    try:
        t = float(target_m)
    except Exception:
        return 0.0
    if t <= 0:
        return 0.0
    exp = 10 ** math.floor(math.log10(t))
    for m in (5, 2, 1):
        val = m * exp
        if val <= t:
            return float(val)
    return float(exp)

def _add_map_decorations(ax,
                         extent_3857: tuple[float, float, float, float],
                         base_dir: str | None = None,
                         add_inset: bool = True):
    """Add north arrow (top-right), scale bar (bottom-left) and optional inset overview."""
    try:
        west_x, east_x, south_y, north_y = extent_3857
        width_m = float(east_x - west_x)
        height_m = float(north_y - south_y)
        if width_m <= 0 or height_m <= 0:
            return

        # North arrow (axes fraction coordinates)
        try:
            ax.annotate(
                "",
                xy=(0.95, 0.94),
                xytext=(0.95, 0.82),
                xycoords=ax.transAxes,
                textcoords=ax.transAxes,
                arrowprops=dict(arrowstyle='-|>', lw=1.6, color='black'),
                zorder=40,
            )
            ax.text(
                0.95, 0.95, "N",
                transform=ax.transAxes,
                ha='center', va='bottom',
                fontsize=12, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.15', fc='white', alpha=0.65, ec='none'),
                zorder=41,
            )
        except Exception:
            pass

        # Scale bar (data coordinates, bottom-left)
        try:
            pad_x = 0.04 * width_m
            pad_y = 0.04 * height_m
            target = 0.22 * width_m
            bar_len = _nice_scale_length_m(target)
            if bar_len > 0:
                x0 = west_x + pad_x
                y0 = south_y + pad_y
                x1 = x0 + bar_len
                tick_h = 0.012 * height_m
                ax.plot([x0, x1], [y0, y0], color='black', lw=3.0, zorder=40, solid_capstyle='butt')
                ax.plot([x0, x0], [y0, y0 + tick_h], color='black', lw=2.0, zorder=40)
                ax.plot([x1, x1], [y0, y0 + tick_h], color='black', lw=2.0, zorder=40)
                if bar_len >= 1000:
                    label = f"{bar_len/1000.0:g} km"
                else:
                    label = f"{bar_len:g} m"
                ax.text(
                    x0, y0 + tick_h * 1.4, label,
                    ha='left', va='bottom', fontsize=10,
                    bbox=dict(boxstyle='round,pad=0.15', fc='white', alpha=0.65, ec='none'),
                    zorder=41,
                )
        except Exception:
            pass

        # Inset overview (if available)
        if add_inset and inset_axes is not None:
            try:
                iax = inset_axes(ax, width="28%", height="28%", loc='upper left', borderpad=0.8)
                iax.set_axis_off()
                # Visible frame around the inset
                try:
                    iax.add_patch(Rectangle((0, 0), 1, 1,
                                            transform=iax.transAxes,
                                            fill=False,
                                            edgecolor='black',
                                            linewidth=1.0,
                                            zorder=100))
                except Exception:
                    pass
                # Expand view to provide context
                pad_ix = 2.5 * width_m
                pad_iy = 2.5 * height_m
                iax.set_xlim(west_x - pad_ix, east_x + pad_ix)
                iax.set_ylim(south_y - pad_iy, north_y + pad_iy)
                _plot_basemap(iax, crs_epsg=3857, base_dir=base_dir)
                rect = Rectangle((west_x, south_y), width_m, height_m,
                                 fill=False, edgecolor=PRIMARY_HEX, linewidth=1.6, alpha=0.95, zorder=50)
                iax.add_patch(rect)
            except Exception:
                pass
    except Exception:
        return

def render_mbtiles_to_png_best_fit(mbtiles_path: str,
                                  out_png: str,
                                  base_dir: str | None = None) -> tuple[bool, str]:
    """Stitch MBTiles into a single PNG for the MBTiles bounds.

    Chooses the highest zoom that fits within MAX_MAP_PX_WIDTH/MAX_MAP_PX_HEIGHT.
    """
    try:
        if not os.path.exists(mbtiles_path):
            return False, "mbtiles not found"
        conn = sqlite3.connect(mbtiles_path)
        try:
            meta = _mbtiles_metadata(conn)
            bounds = _parse_mbtiles_bounds(meta)
            if bounds is None:
                return False, "MBTiles missing bounds"
            minz = _parse_mbtiles_zoom(meta, "minzoom", 0)
            maxz = _parse_mbtiles_zoom(meta, "maxzoom", 19)
            tile_size = 256

            minlon, minlat, maxlon, maxlat = bounds

            # Pick best fit zoom
            chosen_z = maxz
            while chosen_z > minz:
                x_w, y_n = _lonlat_to_tile_fraction(minlon, maxlat, chosen_z)
                x_e, y_s = _lonlat_to_tile_fraction(maxlon, minlat, chosen_z)
                x0 = int(math.floor(min(x_w, x_e)))
                x1 = int(math.floor(max(x_w, x_e)))
                y0 = int(math.floor(min(y_n, y_s)))
                y1 = int(math.floor(max(y_n, y_s)))
                w_px = (x1 - x0 + 1) * tile_size
                h_px = (y1 - y0 + 1) * tile_size
                tiles = (x1 - x0 + 1) * (y1 - y0 + 1)
                if w_px <= MAX_MAP_PX_WIDTH and h_px <= MAX_MAP_PX_HEIGHT and tiles <= 400:
                    break
                chosen_z -= 1

            z = max(minz, chosen_z)

            x_w, y_n = _lonlat_to_tile_fraction(minlon, maxlat, z)
            x_e, y_s = _lonlat_to_tile_fraction(maxlon, minlat, z)
            x0 = int(math.floor(min(x_w, x_e)))
            x1 = int(math.floor(max(x_w, x_e)))
            y0 = int(math.floor(min(y_n, y_s)))
            y1 = int(math.floor(max(y_n, y_s)))

            width = (x1 - x0 + 1) * tile_size
            height = (y1 - y0 + 1) * tile_size
            if width <= 0 or height <= 0:
                return False, "Invalid tile bounds"

            canvas = PILImage.new("RGBA", (width, height), (0, 0, 0, 0))
            blank = PILImage.new("RGBA", (tile_size, tile_size), (0, 0, 0, 0))

            for xt in range(x0, x1 + 1):
                for yt in range(y0, y1 + 1):
                    tile_bytes = _get_mbtiles_tile_bytes(conn, z, xt, yt)
                    if tile_bytes:
                        try:
                            tile_img = PILImage.open(io.BytesIO(tile_bytes)).convert("RGBA")
                        except Exception:
                            tile_img = blank
                    else:
                        tile_img = blank
                    px = (xt - x0) * tile_size
                    py = (yt - y0) * tile_size
                    canvas.paste(tile_img, (px, py))

            # Crop to exact bounds
            left = int(round((min(x_w, x_e) - x0) * tile_size))
            right = int(round((max(x_w, x_e) - x0) * tile_size))
            top = int(round((min(y_n, y_s) - y0) * tile_size))
            bottom = int(round((max(y_n, y_s) - y0) * tile_size))
            left = max(0, min(width - 1, left))
            right = max(left + 1, min(width, right))
            top = max(0, min(height - 1, top))
            bottom = max(top + 1, min(height, bottom))
            canvas = canvas.crop((left, top, right, bottom))

            # Final safety resize
            if canvas.size[0] > MAX_MAP_PX_WIDTH or canvas.size[1] > MAX_MAP_PX_HEIGHT:
                ratio = min(MAX_MAP_PX_WIDTH / canvas.size[0], MAX_MAP_PX_HEIGHT / canvas.size[1])
                new_w = max(1, int(canvas.size[0] * ratio))
                new_h = max(1, int(canvas.size[1] * ratio))
                canvas = canvas.resize((new_w, new_h), PILImage.LANCZOS)

            # Composite over a basemap for context
            try:
                fig_h_in = 8.2
                fig_w_in = 6.2
                dpi = 180
                fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
                ax.set_axis_off()

                # Set extent first (basemap fetch depends on ax limits)
                west_x, north_y = lonlat_to_merc(minlon, maxlat)
                east_x, south_y = lonlat_to_merc(maxlon, minlat)
                ax.set_xlim(west_x, east_x)
                ax.set_ylim(south_y, north_y)
                try:
                    ax.set_aspect('equal', adjustable='box')
                except Exception:
                    pass

                # Background basemap in WebMercator
                _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

                # Overlay MBTiles image in correct extent
                ax.imshow(
                    canvas,
                    extent=[west_x, east_x, south_y, north_y],
                    origin="upper",
                    interpolation='nearest',
                    zorder=10,
                )
                _add_map_decorations(ax, (west_x, east_x, south_y, north_y), base_dir=base_dir, add_inset=True)
                # tight_layout frequently warns with inset axes; bbox_inches='tight' is sufficient here.
                plt.savefig(out_png, bbox_inches='tight')
                plt.close(fig)
                return True, ""
            except Exception:
                try:
                    plt.close('all')
                except Exception:
                    pass
                # Fallback: save the stitched overlay only
                canvas.save(out_png)
                return True, ""
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        write_to_log(f"MBTiles render failed ({mbtiles_path}): {e}", base_dir)
        return False, str(e)


def render_mbtiles_to_png_for_3857_bounds(
    mbtiles_path: str,
    out_png: str,
    bounds_3857: tuple[float, float, float, float],
    *,
    base_dir: str | None = None,
    overlay_alpha: float = 0.65,
    clip_geom_3857=None,
) -> tuple[bool, str]:
    """Stitch MBTiles into a PNG for an arbitrary EPSG:3857 bounds.

    The output is composited over a basemap; the MBTiles overlay is drawn with `overlay_alpha`.
    """
    try:
        if not os.path.exists(mbtiles_path):
            return False, "mbtiles not found"
        if bounds_3857 is None or len(bounds_3857) != 4:
            return False, "bounds missing"

        minx, miny, maxx, maxy = bounds_3857
        if not (math.isfinite(minx) and math.isfinite(miny) and math.isfinite(maxx) and math.isfinite(maxy)):
            return False, "invalid bounds"
        if maxx <= minx or maxy <= miny:
            return False, "invalid bounds"

        lon0, lat0 = merc_to_lonlat(minx, miny)
        lon1, lat1 = merc_to_lonlat(maxx, maxy)
        req_minlon = min(lon0, lon1)
        req_maxlon = max(lon0, lon1)
        req_minlat = min(lat0, lat1)
        req_maxlat = max(lat0, lat1)

        conn = sqlite3.connect(mbtiles_path)
        try:
            meta = _mbtiles_metadata(conn)
            mb_bounds = _parse_mbtiles_bounds(meta)
            if mb_bounds is None:
                return False, "MBTiles missing bounds"
            minz = _parse_mbtiles_zoom(meta, "minzoom", 0)
            maxz = _parse_mbtiles_zoom(meta, "maxzoom", 19)

            mb_minlon, mb_minlat, mb_maxlon, mb_maxlat = mb_bounds
            # Clamp requested area to what the MBTiles contains.
            minlon = max(req_minlon, mb_minlon)
            maxlon = min(req_maxlon, mb_maxlon)
            minlat = max(req_minlat, mb_minlat)
            maxlat = min(req_maxlat, mb_maxlat)
            if maxlon <= minlon or maxlat <= minlat:
                return False, "requested bounds outside MBTiles"

            tile_size = 256

            # Pick a zoom that fits within our pixel caps.
            chosen_z = maxz
            while chosen_z > minz:
                x_w, y_n = _lonlat_to_tile_fraction(minlon, maxlat, chosen_z)
                x_e, y_s = _lonlat_to_tile_fraction(maxlon, minlat, chosen_z)
                x0 = int(math.floor(min(x_w, x_e)))
                x1 = int(math.floor(max(x_w, x_e)))
                y0 = int(math.floor(min(y_n, y_s)))
                y1 = int(math.floor(max(y_n, y_s)))
                w_px = (x1 - x0 + 1) * tile_size
                h_px = (y1 - y0 + 1) * tile_size
                tiles = (x1 - x0 + 1) * (y1 - y0 + 1)
                if w_px <= MAX_MAP_PX_WIDTH and h_px <= MAX_MAP_PX_HEIGHT and tiles <= 400:
                    break
                chosen_z -= 1

            z = max(minz, chosen_z)

            x_w, y_n = _lonlat_to_tile_fraction(minlon, maxlat, z)
            x_e, y_s = _lonlat_to_tile_fraction(maxlon, minlat, z)
            x0 = int(math.floor(min(x_w, x_e)))
            x1 = int(math.floor(max(x_w, x_e)))
            y0 = int(math.floor(min(y_n, y_s)))
            y1 = int(math.floor(max(y_n, y_s)))

            width = (x1 - x0 + 1) * tile_size
            height = (y1 - y0 + 1) * tile_size
            if width <= 0 or height <= 0:
                return False, "Invalid tile bounds"

            canvas = PILImage.new("RGBA", (width, height), (0, 0, 0, 0))
            blank = PILImage.new("RGBA", (tile_size, tile_size), (0, 0, 0, 0))

            for xt in range(x0, x1 + 1):
                for yt in range(y0, y1 + 1):
                    tile_bytes = _get_mbtiles_tile_bytes(conn, z, xt, yt)
                    if tile_bytes:
                        try:
                            tile_img = PILImage.open(io.BytesIO(tile_bytes)).convert("RGBA")
                        except Exception:
                            tile_img = blank
                    else:
                        tile_img = blank
                    px = (xt - x0) * tile_size
                    py = (yt - y0) * tile_size
                    canvas.paste(tile_img, (px, py))

            # Crop to exact requested bounds
            left = int(round((min(x_w, x_e) - x0) * tile_size))
            right = int(round((max(x_w, x_e) - x0) * tile_size))
            top = int(round((min(y_n, y_s) - y0) * tile_size))
            bottom = int(round((max(y_n, y_s) - y0) * tile_size))
            left = max(0, min(width - 1, left))
            right = max(left + 1, min(width, right))
            top = max(0, min(height - 1, top))
            bottom = max(top + 1, min(height, bottom))
            canvas = canvas.crop((left, top, right, bottom))

            # Final safety resize
            if canvas.size[0] > MAX_MAP_PX_WIDTH or canvas.size[1] > MAX_MAP_PX_HEIGHT:
                ratio = min(MAX_MAP_PX_WIDTH / canvas.size[0], MAX_MAP_PX_HEIGHT / canvas.size[1])
                new_w = max(1, int(canvas.size[0] * ratio))
                new_h = max(1, int(canvas.size[1] * ratio))
                canvas = canvas.resize((new_w, new_h), PILImage.LANCZOS)

            # Composite over basemap
            try:
                fig_h_in = 4.2
                fig_w_in = 5.8
                dpi = 180
                fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
                ax.set_axis_off()

                west_x, north_y = lonlat_to_merc(minlon, maxlat)
                east_x, south_y = lonlat_to_merc(maxlon, minlat)
                ax.set_xlim(west_x, east_x)
                ax.set_ylim(south_y, north_y)
                try:
                    ax.set_aspect('equal', adjustable='box')
                except Exception:
                    pass

                _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)
                overlay_im = ax.imshow(
                    canvas,
                    extent=[west_x, east_x, south_y, north_y],
                    origin="upper",
                    interpolation='nearest',
                    alpha=float(max(0.0, min(1.0, overlay_alpha))),
                    zorder=10,
                )

                # Optional: clip everything to the study area polygon and fill it
                clip_patch = None
                try:
                    if clip_geom_3857 is not None and not getattr(clip_geom_3857, 'is_empty', True):
                        clip_path = _geom_to_mpl_path(clip_geom_3857)
                        if clip_path is not None:
                            clip_patch = PathPatch(clip_path, facecolor='none', edgecolor='none', transform=ax.transData)
                            ax.add_patch(clip_patch)
                except Exception:
                    clip_patch = None

                if clip_patch is not None:
                    # Clip only the overlay (basemap remains un-clipped)
                    try:
                        overlay_im.set_clip_path(clip_patch)
                    except Exception:
                        pass

                    # Subtle fill to make study area stand out
                    try:
                        ax.add_patch(
                            PathPatch(
                                clip_patch.get_path(),
                                transform=ax.transData,
                                facecolor=LIGHT_PRIMARY_HEX,
                                edgecolor='none',
                                alpha=0.10,
                                zorder=9,
                            )
                        )
                    except Exception:
                        pass

                    # Intentionally no outline: avoid blue borders and per-part outlines.

                _add_map_decorations(ax, (west_x, east_x, south_y, north_y), base_dir=base_dir, add_inset=False)
                plt.savefig(out_png, bbox_inches='tight')
                plt.close(fig)
                return True, ""
            except Exception:
                try:
                    plt.close('all')
                except Exception:
                    pass
                canvas.save(out_png)
                return True, ""
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        write_to_log(f"MBTiles bounds render failed ({mbtiles_path}): {e}", base_dir)
        return False, str(e)

# ---------------- tbl_flat (per-geocode maps) ----------------
def load_tbl_flat(parquet_dir: str, base_dir: str | None = None) -> gpd.GeoDataFrame:
    fp = os.path.join(parquet_dir, "tbl_flat.parquet")
    if not os.path.exists(fp):
        write_to_log("tbl_flat.parquet not found.", base_dir)
        return gpd.GeoDataFrame()
    try:
        gdf = gpd.read_parquet(fp)
        if 'geometry' not in gdf.columns or gdf.geometry.is_empty.all():
            write_to_log("tbl_flat has no valid geometry.", base_dir)
            return gpd.GeoDataFrame()
        return gdf
    except Exception as e:
        write_to_log(f"Failed reading tbl_flat: {e}", base_dir)
        return gpd.GeoDataFrame()

def compute_fixed_bounds_3857(flat_df: gpd.GeoDataFrame, base_dir: str | None = None):
    try:
        poly = flat_df[flat_df.geometry.type.isin(['Polygon','MultiPolygon'])].copy()
        if poly.empty:
            write_to_log("No polygons in tbl_flat to compute fixed bounds; using all geometry.", base_dir)
            poly = flat_df.copy()
        g3857 = _safe_to_3857(poly)
        if g3857.empty:
            return None
        b = g3857.total_bounds
        b = _expand_bounds(b, pad_ratio=0.08)
        return b
    except Exception as e:
        write_to_log(f"Failed computing fixed bounds: {e}", base_dir)
        return None

def _legend_for_sensitivity(ax, palette: dict, desc: dict, base_dir: str | None = None):
    try:
        y0 = 0.02
        dy = 0.06
        for i, code in enumerate(['A','B','C','D','E']):
            y = y0 + i*dy
            ax.add_patch(plt.Rectangle((0.02, y), 0.04, 0.035,
                                       transform=ax.transAxes,
                                       color=palette.get(code, '#BDBDBD'),
                                       ec='white', lw=0.8))
            label = f"{code}: {desc.get(code,'')}".strip() or code
            ax.text(0.07, y+0.005, label, transform=ax.transAxes,
                    fontsize=9, color='white',
                    bbox=dict(boxstyle="round,pad=0.2", fc="black", alpha=0.35, ec="none"))
    except Exception as e:
        write_to_log(f"Sensitivity legend failed: {e}", base_dir)

def draw_group_map_sensitivity(gdf_group: gpd.GeoDataFrame,
                               group_name: str,
                               palette: dict,
                               desc: dict,
                               out_path: str,
                               fixed_bounds_3857=None,
                               base_dir: str | None = None,
                               alpha: float = 0.95):
    try:
        g = gdf_group.copy()
        g = g[g.geometry.type.isin(['Polygon','MultiPolygon'])]
        if g.empty:
            write_to_log(f"[{group_name}] No polygon geometries for sensitivity map.", base_dir)
            return False

        g = _safe_to_3857(g)
        fig_h_in = 10.0  # square to preserve detail
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        if fixed_bounds_3857:
            minx, miny, maxx, maxy = fixed_bounds_3857
            ax.set_xlim(minx, maxx)
            ax.set_ylim(miny, maxy)

        # Draw basemap first, then polygons on top for clarity
        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        annotated = _prepare_sensitivity_annotations(g, 'sensitivity_code_max', 'sensitivity_max')
        if annotated.empty:
            write_to_log(f"[{group_name}] No sensitivity records to plot.", base_dir)
            plt.close(fig)
            return False

        colors = _colors_from_annotations(annotated, palette)
        annotated.plot(ax=ax,
                       color=colors,
                       edgecolor='none',
                       linewidth=0.0,
                       alpha=float(max(0.0, min(1.0, alpha))),
                       zorder=10)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=True)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Sensitivity map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Sensitivity map failed for {group_name}: {e}", base_dir)
        plt.close('all')
        return False

def draw_atlas_map(tile_row: pd.Series,
                   atlas_crs,
                   atlas_polys_3857: gpd.GeoDataFrame,
                   palette: dict,
                   desc: dict,
                   out_path: str,
                   global_bounds_3857=None,
                   base_dir: str | None = None) -> bool:
    """
    Render a single atlas tile sensitivity map using shared cartography.
    """
    try:
        geom = tile_row.get('geometry', None)
        tile_name = tile_row.get('name_gis', '?')
        if geom is None or getattr(geom, "is_empty", True):
            write_to_log(f"[Atlas {tile_name}] Missing geometry; skipping map.", base_dir)
            return False

        tile_gdf = gpd.GeoDataFrame([{'geometry': geom}], geometry='geometry', crs=atlas_crs)
        if tile_gdf.crs is None:
            tile_gdf.set_crs(4326, inplace=True)
        tile_3857 = _safe_to_3857(tile_gdf)
        if tile_3857.empty or tile_3857.geometry.is_empty.all():
            write_to_log(f"[Atlas {tile_name}] Reprojection failed; skipping map.", base_dir)
            return False

        tile_bounds = _expand_bounds(tile_3857.total_bounds, pad_ratio=0.05)

        subset = gpd.GeoDataFrame()
        if atlas_polys_3857 is not None and not atlas_polys_3857.empty:
            try:
                mask = atlas_polys_3857.geometry.intersects(tile_3857.iloc[0].geometry)
                subset = atlas_polys_3857[mask].copy()
            except Exception as e:
                write_to_log(f"[Atlas {tile_name}] Intersection failed: {e}", base_dir)
                subset = gpd.GeoDataFrame()

        fig_w_in, fig_h_in = ATLAS_FIGURE_INCHES
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()
        ax.set_xlim(tile_bounds[0], tile_bounds[2])
        ax.set_ylim(tile_bounds[1], tile_bounds[3])

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        drew_data = False
        if not subset.empty and "sensitivity_code_max" in subset.columns:
            subset_ann = _prepare_sensitivity_annotations(subset, "sensitivity_code_max", "sensitivity_max")
            if subset_ann.empty:
                write_to_log(f"[Atlas {tile_name}] Sensitivity annotations empty.", base_dir)
            else:
                colors = _colors_from_annotations(subset_ann, palette)
                subset_ann.plot(ax=ax,
                                color=colors,
                                edgecolor="none",
                                linewidth=0.0,
                                alpha=0.75,
                                zorder=10)
                drew_data = True
        elif not subset.empty:
            subset.plot(ax=ax,
                        color=SENSITIVITY_UNKNOWN_COLOR,
                        edgecolor="none",
                        linewidth=0.0,
                        alpha=0.75,
                        zorder=10)
            drew_data = True
        if not drew_data:
            plt.close(fig)
            write_to_log(f"[Atlas {tile_name}] No drawable data for sensitivity map.", base_dir)
            return False

        tile_3857.boundary.plot(ax=ax, edgecolor=PRIMARY_HEX, linewidth=1.6, zorder=12)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=True)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Atlas map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Atlas map failed for {tile_row.get('name_gis', '?')}: {e}", base_dir)
        plt.close('all')
        return False

def draw_atlas_overview_map(atlas_df: gpd.GeoDataFrame,
                            atlas_crs,
                            atlas_polys_3857: gpd.GeoDataFrame,
                            out_path: str,
                            global_bounds_3857=None,
                            base_dir: str | None = None) -> bool:
    """
    Draw an overview map showing every atlas tile in relation to the full dataset extent.
    """
    try:
        if atlas_df.empty or 'geometry' not in atlas_df.columns:
            write_to_log("Atlas overview skipped (no geometries).", base_dir)
            return False

        atlas_gdf = atlas_df.copy()
        atlas_gdf = atlas_gdf[atlas_gdf['geometry'].notna()]
        if atlas_gdf.empty:
            write_to_log("Atlas overview skipped (all geometries empty).", base_dir)
            return False

        if atlas_crs is not None:
            try:
                atlas_gdf = atlas_gdf.set_crs(atlas_crs, allow_override=True)
            except Exception:
                pass
        atlas_3857 = _safe_to_3857(atlas_gdf)
        if atlas_3857.empty:
            write_to_log("Atlas overview reprojection failed; skipping.", base_dir)
            return False

        if global_bounds_3857:
            bounds = global_bounds_3857
        else:
            bounds = _expand_bounds(atlas_3857.total_bounds, pad_ratio=0.08)

        fig_h_in = 10.0
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()
        ax.set_xlim(bounds[0], bounds[2])
        ax.set_ylim(bounds[1], bounds[3])

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        if atlas_polys_3857 is not None and not atlas_polys_3857.empty:
            atlas_polys_3857.plot(ax=ax, facecolor="#d9d9d9", edgecolor='white',
                                 linewidth=0.15, alpha=0.5, zorder=8)

        atlas_3857.boundary.plot(ax=ax, edgecolor=PRIMARY_HEX, linewidth=1.0, alpha=0.9, zorder=12)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=False)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Atlas overview map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Atlas overview map failed: {e}", base_dir)
        plt.close('all')
        return False

# ---------------- Lines: context + segments map ----------------
def draw_line_context_map(line_row: pd.Series, out_path: str,
                          pad_ratio: float = 1.0,
                          rect_buffer_ratio: float = 0.03,
                          base_dir: str | None = None):
    """
    Draw a geographical context map for the line.
    pad_ratio = 1.0 means a 100% buffer around the line's bounding box.
    """
    try:
        geom = line_row.get('geometry', None)
        if geom is None:
            write_to_log(f"[{line_row.get('name_gis','?')}] Missing geometry; cannot draw context map.", base_dir)
            return False

        g = gpd.GeoDataFrame([{'name_gis': line_row.get('name_gis','(line)'), 'geometry': geom}], crs=None)
        g = _safe_to_3857(g)
        if g.empty or g.geometry.is_empty.all():
            write_to_log(f"[{line_row.get('name_gis','?')}] Invalid geometry after reprojection.", base_dir)
            return False

        fig_h_in = 7.5
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        b = g.total_bounds
        minx, miny, maxx, maxy = _expand_bounds(b, pad_ratio)
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        # Red rectangle around line bbox (slightly expanded)
        bx0, by0, bx1, by1 = b
        bw, bh = (bx1 - bx0), (by1 - by0)
        rx0 = bx0 - bw * rect_buffer_ratio
        ry0 = by0 - bh * rect_buffer_ratio
        rw  = bw + 2 * bw * rect_buffer_ratio
        rh  = bh + 2 * bh * rect_buffer_ratio
        rect = Rectangle((rx0, ry0), rw, rh, fill=False, edgecolor='red', linewidth=2.0, alpha=0.95, zorder=20)
        ax.add_patch(rect)

        # Halo + line on top
        try:
            g.plot(ax=ax, color='none', edgecolor='white', linewidth=5.0, alpha=0.9, zorder=21)
        except Exception:
            pass
        g.plot(ax=ax, color='none', edgecolor=PRIMARY_HEX, linewidth=2.2, alpha=1.0, zorder=22)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=True)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Line context map saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Context map failed for line: {e}", base_dir)
        plt.close('all')
        return False

def draw_line_segments_map(segments_df: gpd.GeoDataFrame,
                           line_name: str,
                           palette: dict,
                           out_path: str,
                           mode: str = 'max',
                           pad_ratio: float = 0.20,
                           base_dir: str | None = None):
    """
    Draw segments for a given line, colored by sensitivity (max or min), with basemap.
    """
    try:
        if segments_df.empty or 'geometry' not in segments_df.columns:
            write_to_log(f"[{line_name}] No geometry in segments; skipping segments map.", base_dir)
            return False

        col = 'sensitivity_code_max' if mode == 'max' else 'sensitivity_code_min'
        segs = segments_df[(segments_df['name_gis'] == line_name) & (segments_df['geometry'].notna())].copy()
        if segs.empty:
            write_to_log(f"[{line_name}] No segments found for segments map.", base_dir)
            return False

        segs = _safe_to_3857(segs)
        if segs.empty:
            write_to_log(f"[{line_name}] Segments reprojection failed.", base_dir)
            return False

        fig_h_in = 7.5
        fig_w_in = 10.0
        dpi = _dpi_for_fig_height(fig_h_in)
        fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)
        ax.set_axis_off()

        b = segs.total_bounds
        minx, miny, maxx, maxy = _expand_bounds(b, pad_ratio)
        ax.set_xlim(minx, maxx)
        ax.set_ylim(miny, maxy)

        _plot_basemap(ax, crs_epsg=3857, base_dir=base_dir)

        polys = segs[segs.geometry.type.isin(['Polygon','MultiPolygon'])]
        lines = segs[segs.geometry.type.isin(['LineString','MultiLineString'])]
        value_col = 'sensitivity_max' if mode == 'max' else 'sensitivity_min'

        if not polys.empty:
            polys_ann = _prepare_sensitivity_annotations(polys, col, value_col)
            colors_polys = _colors_from_annotations(polys_ann, palette)
            polys_ann.plot(ax=ax,
                           color=colors_polys,
                           edgecolor='white',
                           linewidth=0.4,
                           alpha=0.95,
                           zorder=12)

        if not lines.empty:
            lines_ann = _prepare_sensitivity_annotations(lines, col, value_col)
            line_colors = _colors_from_annotations(lines_ann, palette)
            line_colors_rgba = [mcolors.to_rgba(c, alpha=1.0) for c in line_colors]
            try:
                lines.plot(ax=ax, color='white', linewidth=4.2, alpha=0.9, zorder=13)
            except Exception:
                pass
            lines.plot(ax=ax, color=line_colors_rgba, linewidth=2.4, alpha=1.0, zorder=14)

        try:
            minx, maxx = ax.get_xlim()
            miny, maxy = ax.get_ylim()
            _add_map_decorations(ax, (minx, maxx, miny, maxy), base_dir=base_dir, add_inset=True)
        except Exception:
            pass

        plt.savefig(out_path, bbox_inches='tight')
        plt.close(fig)
        write_to_log(f"Segments map ({mode}) saved: {out_path}", base_dir)
        return True
    except Exception as e:
        write_to_log(f"Segments map failed for {line_name}: {e}", base_dir)
        plt.close('all')
        return False

# ---------------- Existing (kept / refined) ----------------
def get_color_from_code(code, color_codes):
    return _resolve_sensitivity_color(code, None, color_codes)

def sort_segments_numerically(segments):
    def extract_number(segment_id):
        m = re.search(r'_(\d+)$', str(segment_id))
        return int(m.group(1)) if m else float('inf')
    s = segments.copy()
    s['sort_key'] = s['segment_id'].apply(extract_number)
    return s.sort_values(by='sort_key').drop(columns=['sort_key'])

def create_line_statistic_image(line_name, sensitivity_series, color_codes, length_m, output_path):
    """
    Generate ribbon PNG with NO axes/ticks/labels.
    Axes fill the full figure area; saved with bbox_inches=None and pad_inches=0
    so every PNG has an identical pixel canvas regardless of content.
    """
    segment_count = max(1, len(sensitivity_series))
    fig_h_in = 0.236  # ~6 mm
    dpi = 300
    fig_w_in = 12
    fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)

    # Fill figure with axes; no margins
    ax.set_position([0, 0, 1, 1])
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.axis('off')

    for i, code in enumerate(sensitivity_series):
        color = get_color_from_code(code, color_codes)
        ax.add_patch(plt.Rectangle((i/segment_count, 0), 1/segment_count, 1, color=color))

    # Save with constant canvas (no tight bbox, no padding)
    plt.savefig(output_path, dpi=dpi, bbox_inches=None, pad_inches=0)
    plt.close(fig)

def resize_image(image_path, max_width_px, max_height_px):
    """
    Resample the stored image file to be within (max_width_px, max_height_px) in *pixels*.
    """
    try:
        with PILImage.open(image_path) as img:
            width, height = img.size
            ratio = min(float(max_width_px) / width, float(max_height_px) / height)
            if ratio >= 1.0:
                return  # already within limits
            new_w, new_h = max(1, int(width * ratio)), max(1, int(height * ratio))
            img = img.resize((new_w, new_h), PILImage.LANCZOS)
            img.save(image_path)
    except Exception as e:
        write_to_log(f"Image resize failed for {image_path}: {e}")

def create_sensitivity_summary(sensitivity_series, color_codes, output_path):
    counts = sensitivity_series.value_counts().reindex(['A','B','C','D','E']).fillna(0)
    total = max(1, len(sensitivity_series))
    fallback_color = color_codes.get('UNKNOWN', SENSITIVITY_UNKNOWN_COLOR)

    fig, (ax_text, ax_bar) = plt.subplots(2, 1, figsize=(10, 1.7), height_ratios=[0.45, 0.55])
    plt.subplots_adjust(hspace=0.15)

    parts = [f"{c}: {int(counts[c])} ({counts[c]/total*100:.1f}%)" for c in ['A','B','C','D','E']]
    summary = "Distribution: " + " | ".join(parts)
    ax_text.text(0.5, 0.5, summary, ha='center', va='center', fontsize=10)
    ax_text.axis('off')

    left = 0
    for c in ['A','B','C','D','E']:
        w = counts[c] / total
        ax_bar.barh(0, w, left=left, color=color_codes.get(c, fallback_color), edgecolor='white')
        left += w
    ax_bar.set_xlim(0, 1); ax_bar.set_ylim(-0.5, 0.5); ax_bar.axis('off')

    plt.savefig(output_path, bbox_inches='tight', dpi=110)
    plt.close(fig)

def create_index_area_distribution_chart(flat_df: gpd.GeoDataFrame,
                                        index_col: str,
                                        output_path: str,
                                        basic_group_name: str = "basic_mosaic",
                                        base_dir: str | None = None) -> tuple[bool, str]:
    """Render a chart of total area per index value for basic_mosaic only.

    Bars: total polygon area (km²) per index value (0..100)
    Line (2nd axis): number of "categories" within each index value.
      - If sensitivity_code_max exists: distinct A–E count per index value.
      - Else: number of geocode cells per index value.
    Returns (ok, note).
    """
    try:
        if flat_df is None or flat_df.empty:
            return False, "tbl_flat is empty."
        if index_col not in flat_df.columns:
            return False, f"Missing column: {index_col}."
        if "area_m2" not in flat_df.columns:
            return False, "Missing column: area_m2."

        df = flat_df.copy()
        if "name_gis_geocodegroup" in df.columns and basic_group_name:
            mask = df["name_gis_geocodegroup"].astype("string").str.strip().str.lower() == str(basic_group_name).strip().lower()
            df = df[mask].copy()

        if df.empty:
            return False, f"No rows for geocode group '{basic_group_name}'."

        df["__idx__"] = pd.to_numeric(df[index_col], errors="coerce")
        df["__idx__"] = df["__idx__"].round().clip(0, 100)
        df["__area_m2__"] = pd.to_numeric(df["area_m2"], errors="coerce")
        df = df.dropna(subset=["__idx__", "__area_m2__"])
        if df.empty:
            return False, "No numeric index/area rows to chart."

        df["__idx__"] = df["__idx__"].astype(int)

        # Area (km²) per index value (0..100)
        area_by = df.groupby("__idx__")["__area_m2__"].sum()
        x = np.arange(0, 101)
        area_km2 = np.array([float(area_by.get(i, 0.0)) / 1e6 for i in x], dtype=float)

        # "Categories" per index value
        cat_label = "# categories"
        if "sensitivity_code_max" in df.columns:
            codes = df["sensitivity_code_max"].astype("string").str.strip().str.upper()
            tmp = pd.DataFrame({"idx": df["__idx__"], "code": codes})
            # count distinct A–E per idx (ignore blanks)
            tmp = tmp[tmp["code"].isin(["A", "B", "C", "D", "E"])].copy()
            cats_by = tmp.groupby("idx")["code"].nunique()
            cats = np.array([int(cats_by.get(i, 0)) for i in x], dtype=int)
        else:
            cat_label = "# cells"
            cnt_by = df.groupby("__idx__").size()
            cats = np.array([int(cnt_by.get(i, 0)) for i in x], dtype=int)

        total_area_km2 = float(area_km2.sum())
        nonzero_values = int(np.count_nonzero(area_km2 > 0))

        fig, ax1 = plt.subplots(figsize=(11.0, 6.2), dpi=140)
        ax2 = ax1.twinx()

        ax1.bar(x, area_km2, color=PRIMARY_HEX, alpha=0.85, width=1.0, linewidth=0)
        ax2.plot(x, cats, color="#222222", linewidth=1.5)

        ax1.set_xlim(0, 100)
        ax1.set_xlabel("Index value")
        ax1.set_ylabel("Total area (km²)")
        ax2.set_ylabel(cat_label)

        ax1.set_title(f"{index_col} distribution (basic_mosaic)\nTotal area: {total_area_km2:.1f} km² | Values present: {nonzero_values}")

        ax1.grid(True, axis="y", linestyle=":", alpha=0.35)
        ax1.set_xticks(np.arange(0, 101, 10))

        # Tighten y2 for category scale
        try:
            ax2.set_ylim(0, max(1, int(cats.max()) + 1))
        except Exception:
            pass

        fig.tight_layout()
        plt.savefig(output_path, bbox_inches="tight")
        plt.close(fig)
        return True, ""
    except Exception as e:
        try:
            plt.close('all')
        except Exception:
            pass
        write_to_log(f"Index distribution chart failed for {index_col}: {e}", base_dir)
        return False, str(e)

def debug_atlas_sample(base_dir: str,
                       cfg: configparser.ConfigParser,
                       palette: dict[str, str],
                       desc: dict[str, str],
                       tile_name: str,
                       geocode_level: str | None = None,
                       sample_size: int | None = None) -> str | None:
    """
    Render a quick atlas sensitivity map for a single tile using an optional
    subset of polygons. Helpful when debugging palette application.
    """
    gpq_dir = parquet_dir_from_cfg(base_dir, cfg)
    flat_df = load_tbl_flat(gpq_dir, base_dir=base_dir)
    if flat_df.empty:
        write_to_log("debug_atlas_sample: tbl_flat missing or empty.", base_dir)
        return None

    atlas_path = os.path.join(gpq_dir, "tbl_atlas.parquet")
    if not os.path.exists(atlas_path):
        write_to_log("debug_atlas_sample: tbl_atlas.parquet not found.", base_dir)
        return None

    try:
        atlas_df = gpd.read_parquet(atlas_path)
    except Exception as e:
        write_to_log(f"debug_atlas_sample: failed to read tbl_atlas ({e})", base_dir)
        return None

    if atlas_df.empty or 'geometry' not in atlas_df.columns:
        write_to_log("debug_atlas_sample: atlas table empty or missing geometry.", base_dir)
        return None

    tile_row = atlas_df[atlas_df['name_gis'] == tile_name]
    if tile_row.empty:
        write_to_log(f"debug_atlas_sample: tile '{tile_name}' not found in atlas.", base_dir)
        return None
    tile_row = tile_row.iloc[0]

    flat_polys = flat_df[flat_df.geometry.type.isin(['Polygon','MultiPolygon'])].copy()
    if flat_polys.empty:
        write_to_log("debug_atlas_sample: no polygon geometries in tbl_flat.", base_dir)
        return None

    if 'name_gis_geocodegroup' in flat_polys.columns:
        flat_polys['name_gis_geocodegroup'] = (
            flat_polys['name_gis_geocodegroup']
            .astype('string')
            .str.strip()
        )
    else:
        geocode_level = None  # cannot filter by level if column missing

    if geocode_level:
        mask_lvl = flat_polys['name_gis_geocodegroup'].astype('string').str.lower() == geocode_level.lower()
        filtered = flat_polys[mask_lvl].copy()
        if filtered.empty:
            write_to_log(f"debug_atlas_sample: geocode '{geocode_level}' produced no polygons; using full set.", base_dir)
        else:
            flat_polys = filtered

    flat_polys_3857 = _safe_to_3857(flat_polys)
    if flat_polys_3857.empty:
        write_to_log("debug_atlas_sample: reprojection yielded empty polygon set.", base_dir)
        return None

    tile_gdf = gpd.GeoDataFrame([tile_row], geometry='geometry', crs=atlas_df.crs)
    tile_3857 = _safe_to_3857(tile_gdf)
    if tile_3857.empty or tile_3857.geometry.is_empty.all():
        write_to_log(f"debug_atlas_sample: tile '{tile_name}' reprojection failed.", base_dir)
        return None

    try:
        intersects_mask = flat_polys_3857.geometry.intersects(tile_3857.iloc[0].geometry)
    except Exception as e:
        write_to_log(f"debug_atlas_sample: intersection failed ({e}).", base_dir)
        return None

    subset = flat_polys_3857[intersects_mask].copy()
    if subset.empty:
        write_to_log(f"debug_atlas_sample: no polygons intersect tile '{tile_name}'.", base_dir)
        return None

    if sample_size and sample_size > 0 and len(subset) > sample_size:
        subset = subset.sample(sample_size, random_state=0)
        write_to_log(f"debug_atlas_sample: down-sampled to {len(subset)} polygons.", base_dir)
    else:
        write_to_log(f"debug_atlas_sample: using {len(subset)} polygons.", base_dir)

    counts = _prepare_sensitivity_annotations(subset, 'sensitivity_code_max', 'sensitivity_max')['__sens_code'].value_counts(dropna=False)
    write_to_log(f"debug_atlas_sample sensitivity distribution: {counts.to_dict()}", base_dir)

    bounds = compute_fixed_bounds_3857(flat_df, base_dir=base_dir)
    out_path = output_subpath(base_dir, "tmp", f"debug_atlas_{_safe_name(tile_name)}.png")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    ok = draw_atlas_map(tile_row, atlas_df.crs, subset, palette, desc, out_path, bounds, base_dir=base_dir)
    if ok:
        write_to_log(f"debug_atlas_sample: saved {out_path}", base_dir)
        return out_path
    return None

def fetch_asset_group_statistics(asset_group_df: gpd.GeoDataFrame, asset_object_df: gpd.GeoDataFrame):
    """
    Returns per sensitivity code:
      - Number of Asset Objects
      - Active asset groups (distinct groups having ≥1 object)
    """
    if asset_group_df.empty:
        return pd.DataFrame(columns=['Sensitivity Code','Sensitivity Description','Active asset groups','Number of Asset Objects'])

    grp = asset_group_df[['id','sensitivity_code','sensitivity_description']].copy()

    if not asset_object_df.empty and 'ref_asset_group' in asset_object_df.columns:
        cnt = asset_object_df.groupby('ref_asset_group').size().rename('asset_objects_nr')
        merged = grp.merge(cnt, left_on='id', right_index=True, how='left')
    else:
        merged = grp.copy()
        merged['asset_objects_nr'] = 0

    merged['asset_objects_nr'] = merged['asset_objects_nr'].fillna(0).astype(int)

    active_groups = (merged[merged['asset_objects_nr'] > 0]
                     .groupby(['sensitivity_code','sensitivity_description'])
                     .agg(active_groups=('id','nunique'))
                     .reset_index())

    objects_per_code = (merged.groupby(['sensitivity_code','sensitivity_description'])
                        .agg(asset_objects=('asset_objects_nr','sum'))
                        .reset_index())

    out = objects_per_code.merge(active_groups,
                                 on=['sensitivity_code','sensitivity_description'],
                                 how='left').fillna({'active_groups':0})

    out = out.rename(columns={
        'sensitivity_code':'Sensitivity Code',
        'sensitivity_description':'Sensitivity Description',
        'active_groups':'Active asset groups',
        'asset_objects':'Number of Asset Objects'
    }).sort_values('Sensitivity Code')

    return out

def format_area(row):
    if row['geometry_type'] in ['Polygon', 'MultiPolygon']:
        try:
            area = float(row['total_area'])
        except Exception:
            return "-"
        return _format_area_m2(area)
    return "-"


def _format_area_m2(area_m2: float) -> str:
    """Format an area in m² as either m² (<0.1 km²) or km² (max 4 decimals)."""
    try:
        a = float(area_m2)
    except Exception:
        return "-"
    if not math.isfinite(a) or a < 0:
        return "-"
    # Use m² for small areas (<1 km² == 1,000,000 m²)
    if a < 1_000_000:
        return f"{a:.0f} m²"
    km2 = a / 1_000_000.0
    s = f"{km2:.4f}"  # max 4 decimals
    # Trim trailing zeros while keeping <= 4 decimals.
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return f"{s} km²"


def _format_area_km2(area_km2: float) -> str:
    """Format an area in km² as either m² (<1 km²) or km² (max 4 decimals)."""
    try:
        km2 = float(area_km2)
    except Exception:
        return "-"
    if not math.isfinite(km2) or km2 < 0:
        return "-"
    if km2 < 1.0:
        return _format_area_m2(km2 * 1_000_000.0)
    s = f"{km2:.4f}"  # max 4 decimals
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return f"{s} km²"


def _analysis_assets_within_area_table(
    asset_objects_df: gpd.GeoDataFrame,
    asset_groups_df: gpd.GeoDataFrame,
    area_polys_gdf: gpd.GeoDataFrame,
) -> list[list[str]] | None:
    """Return a table of polygon asset area inside the given analysis area polygon(s)."""
    try:
        if asset_objects_df is None or asset_objects_df.empty:
            return None
        if asset_groups_df is None or asset_groups_df.empty:
            return None
        if area_polys_gdf is None or area_polys_gdf.empty or 'geometry' not in area_polys_gdf.columns:
            return None
        if 'ref_asset_group' not in asset_objects_df.columns:
            return None

        base = asset_groups_df.drop(columns=[c for c in ['geometry'] if c in asset_groups_df.columns])
        ao = asset_objects_df.merge(base, left_on='ref_asset_group', right_on='id', how='left')
        if 'geometry' not in ao.columns:
            return None

        ao = ao[ao.geometry.notna()].copy()
        if ao.empty:
            return None
        ao['geometry_type'] = ao.geometry.type
        ao_poly = ao[ao['geometry_type'].isin(['Polygon', 'MultiPolygon'])].copy()
        if ao_poly.empty:
            return None

        # Ensure CRS, then compute intersection areas in an equal-area projection.
        if ao_poly.crs is None:
            ao_poly = ao_poly.set_crs(4326)
        ap = area_polys_gdf.dropna(subset=['geometry']).copy()
        if ap.empty:
            return None
        if ap.crs is None:
            ap = ap.set_crs(ao_poly.crs)

        try:
            ao_poly = ao_poly.to_crs('ESRI:54009')
            ap = ap.to_crs('ESRI:54009')
        except Exception:
            # If reprojection fails, proceed in current CRS (areas may be less accurate)
            pass

        try:
            if hasattr(ap, "union_all"):
                area_union = ap.union_all()
            else:
                area_union = ap.unary_union
        except Exception:
            area_union = None
        if area_union is None or getattr(area_union, 'is_empty', True):
            return None

        # Fast pre-filter by intersection
        try:
            ao_poly = ao_poly[ao_poly.intersects(area_union)].copy()
        except Exception:
            pass
        if ao_poly.empty:
            return None

        area_gdf = gpd.GeoDataFrame({'geometry': [area_union]}, geometry='geometry', crs=ao_poly.crs)
        try:
            inter = gpd.overlay(ao_poly, area_gdf, how='intersection')
        except Exception:
            # Fallback: compute per-geometry intersection without overlay
            try:
                ao_poly['__geom_i'] = ao_poly.geometry.apply(lambda g: g.intersection(area_union))
                inter = ao_poly.drop(columns=['geometry']).copy()
                inter = gpd.GeoDataFrame(inter, geometry='__geom_i', crs=ao_poly.crs)
                inter = inter.rename_geometry('geometry')
            except Exception:
                return None

        if inter is None or inter.empty:
            return None

        try:
            inter['area_m2'] = inter.geometry.area
        except Exception:
            return None

        title_col = 'title_fromuser' if 'title_fromuser' in inter.columns else None
        if title_col is None:
            return None

        agg = (inter.groupby([title_col], dropna=False)
               .agg(area_m2=('area_m2', 'sum'), objects=('geometry', 'size'))
               .reset_index())
        if agg.empty:
            return None

        agg = agg.sort_values('area_m2', ascending=False)

        table: list[list[str]] = [["Asset", "Area within area", "# objects (polygons)"]]
        for _, r in agg.iterrows():
            title = str(r[title_col]) if not pd.isna(r[title_col]) else "(Unknown)"
            area_str = _format_area_m2(float(r['area_m2']))
            try:
                obj_n = int(r['objects'])
            except Exception:
                obj_n = 0
            table.append([title, area_str, str(obj_n)])
        return table
    except Exception:
        return None

def calculate_group_statistics(asset_object_df: gpd.GeoDataFrame, asset_group_df: gpd.GeoDataFrame):
    if asset_object_df.empty or asset_group_df.empty:
        return pd.DataFrame(columns=['Title','Code','Description','Type','Total area','# objects'])
    base = asset_group_df.drop(columns=[c for c in ['geometry'] if c in asset_group_df.columns])
    gdf = asset_object_df.merge(base, left_on='ref_asset_group', right_on='id', how='left')
    gdf['geometry_type'] = gdf.geometry.type
    try:
        gdf = gdf.to_crs('ESRI:54009')
    except Exception:
        pass

    gdf = gdf.copy()
    gdf['area'] = np.nan
    poly_mask = gdf['geometry_type'].isin(['Polygon', 'MultiPolygon'])
    if poly_mask.any():
        try:
            gdf.loc[poly_mask, 'area'] = gdf.loc[poly_mask, 'geometry'].area
        except Exception:
            gdf.loc[poly_mask, 'area'] = gdf.loc[poly_mask, 'geometry'].area
    gdf2 = gdf

    stats = gdf2.groupby(
        ['title_fromuser','sensitivity_code','sensitivity_description','geometry_type'],
        dropna=False
    ).agg(total_area=('area','sum'), object_count=('geometry','size')).reset_index()

    stats['total_area'] = stats['total_area'].fillna(0)
    stats['total_area'] = stats.apply(format_area, axis=1)
    stats = stats.rename(columns={
        'title_fromuser':'Title',
        'sensitivity_code':'Code',
        'sensitivity_description':'Description',
        'geometry_type':'Type',
        'object_count':'# objects'
    }).sort_values('Code')
    return stats

def export_to_excel(df, fp):
    if not fp.lower().endswith('.xlsx'):
        fp += '.xlsx'
    df.to_excel(fp, index=False)

def fetch_lines_and_segments(parquet_dir: str):
    lines_pq    = os.path.join(parquet_dir, "tbl_lines.parquet")
    segments_pq = os.path.join(parquet_dir, "tbl_segment_flat.parquet")
    try:
        lines_df = gpd.read_parquet(lines_pq) if os.path.exists(lines_pq) else gpd.GeoDataFrame()
    except Exception:
        lines_df = gpd.GeoDataFrame()
    try:
        segments_df = gpd.read_parquet(segments_pq) if os.path.exists(segments_pq) else gpd.GeoDataFrame()
    except Exception:
        segments_df = gpd.GeoDataFrame()
    return lines_df, segments_df

# ---------------- PDF helpers ----------------
def line_up_to_pdf(order_list):
    elements = []
    styles = getSampleStyleSheet()

    primary = HexColor(PRIMARY_HEX)
    light_primary = HexColor(LIGHT_PRIMARY_HEX)
    heading_styles = {
        'H1': ParagraphStyle(
            name='H1', fontSize=20, leading=24, spaceBefore=6, spaceAfter=8,
            alignment=TA_CENTER, textColor=colors.black
        ),
        'H2Bar': ParagraphStyle(
            name='H2Bar', fontSize=16, leading=20, spaceBefore=10, spaceAfter=6,
            textColor=colors.white, backColor=light_primary, leftIndent=0, rightIndent=0
        ),
        'H3': ParagraphStyle(
            name='H3', fontSize=13, leading=16, spaceBefore=4, spaceAfter=4,
            textColor=primary
        ),
        'Body': styles['Normal']
    }

    # A4 text area constraints (points)
    max_image_width_pts = 16 * RL_CM                 # usable frame width
    default_max_image_height_pts = 24 * RL_CM
    map_max_height_pts = MAX_MAP_CM_HEIGHT * RL_CM

    ribbon_height_pts = RIBBON_CM_HEIGHT * RL_CM
    ribbon_width_pts  = max_image_width_pts          # <-- force ribbons to EXACT frame width

    def _add_heading(level, text):
        if level == 1:
            elements.append(Paragraph(text, heading_styles['H1']))
            elements.append(HRFlowable(width="100%", thickness=0.8, color=primary))
        elif level == 2:
            elements.append(Paragraph(text, heading_styles['H2Bar']))
        else:
            elements.append(Paragraph(text, heading_styles['H3']))
        elements.append(Spacer(1, 6))

    def _is_map_image(path: str) -> bool:
        name = os.path.basename(path).lower()
        return any(tok in name for tok in ['map_', '_segments_', '_context'])

    def _is_atlas_image(path: str) -> bool:
        return 'atlas' in os.path.basename(path).lower()

    for item in order_list:
        itype, ival = item

        if itype == 'text':
            if isinstance(ival, str) and os.path.isfile(ival):
                with open(ival, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                text = str(ival)
            text = text.replace("\n", "<br/>")
            elements.append(Paragraph(text, heading_styles['Body']))
            elements.append(Spacer(1, 8))

        elif itype == 'image':
            # ival: (heading_str, path)
            heading_str, path = ival
            if not _file_ok(path):
                elements.append(Paragraph(f"<i>(image missing: {os.path.basename(path)})</i>", heading_styles['Body']))
                elements.append(Spacer(1, 6))
                continue

            _add_heading(3, heading_str)

            # Cap stored pixel height for maps to keep file sizes reasonable.
            if _is_map_image(path):
                try:
                    resize_image(path, max_width_px=1_000_000, max_height_px=MAX_MAP_PX_HEIGHT)
                except Exception:
                    pass

            img = Image(path)
            img.hAlign = 'CENTER'

            # Dual constraint: width + height caps (atlas images get narrower bounds)
            max_h_pts = map_max_height_pts if _is_map_image(path) else default_max_image_height_pts
            width_cap_pts = max_image_width_pts
            if _is_atlas_image(path):
                width_cap_pts *= ATLAS_DOC_WIDTH_SCALE
                if _is_map_image(path):
                    max_h_pts = map_max_height_pts * ATLAS_DOC_WIDTH_SCALE
            w0, h0 = float(getattr(img, 'imageWidth', 0) or 0), float(getattr(img, 'imageHeight', 0) or 0)
            if w0 > 0 and h0 > 0:
                scale = min(width_cap_pts / w0, max_h_pts / h0, 1.0)
                img.drawWidth = w0 * scale
                img.drawHeight = h0 * scale

            elements.append(img)
            elements.append(Spacer(1, 6))

        elif itype == 'image_ribbon':
            # ival: (heading_str, path)
            heading_str, path = ival
            if not _file_ok(path):
                elements.append(Paragraph(f"<i>(image missing: {os.path.basename(path)})</i>", heading_styles['Body']))
                elements.append(Spacer(1, 6))
                continue
            _add_heading(3, heading_str)

            img = Image(path)
            # FORCE exact same display size for all ribbons (no aspect scaling surprises)
            img.drawWidth  = ribbon_width_pts
            img.drawHeight = ribbon_height_pts
            img.hAlign = 'CENTER'

            elements.append(img)
            elements.append(Spacer(1, 6))

        elif itype == 'table':
            df = pd.read_excel(ival)
            data = [df.columns.tolist()] + df.values.tolist()
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), primary),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('ALIGN',      (0, 0), (-1,-1), 'CENTER'),
                ('GRID',       (0, 0), (-1,-1), 0.5, colors.HexColor("#A9A9A9")),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.whitesmoke, colors.Color(0.97,0.97,0.97)])
            ]))
            elements.append(table)
            elements.append(Spacer(1, 8))

        elif itype == 'table_data':
            title, data_ll = ival
            _add_heading(3, title)
            table = Table(data_ll, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), primary),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('ALIGN',      (0, 0), (-1,-1), 'CENTER'),
                ('GRID',       (0, 0), (-1,-1), 0.5, colors.HexColor("#A9A9A9")),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.whitesmoke, colors.Color(0.97,0.97,0.97)])
            ]))
            elements.append(table)
            elements.append(Spacer(1, 8))

        elif itype.startswith('heading'):
            level = int(itype[-2])
            _add_heading(level, ival)

        elif itype == 'rule':
            elements.append(HRFlowable(width="100%", thickness=0.8, color=primary))
            elements.append(Spacer(1, 6))

        elif itype == 'spacer':
            lines = ival if ival else 1
            elements.append(Spacer(1, lines * 10))

        elif itype == 'new_page':
            elements.append(PageBreak())

    return elements

def compile_pdf(output_pdf, elements):
    doc = SimpleDocTemplate(output_pdf, pagesize=A4)
    primary = HexColor(PRIMARY_HEX)

    def add_header_footer(canvas, doc_obj):
        width, height = A4
        # Top steel-blue band
        canvas.saveState()
        canvas.setFillColor(primary)
        canvas.rect(0, height - 16, width, 16, fill=1, stroke=0)
        canvas.setFont('Helvetica-Bold', 10)
        canvas.setFillColor(colors.white)
        canvas.drawString(24, height - 12, "MESA – Report")
        canvas.restoreState()

        # Footer with page number
        canvas.saveState()
        canvas.setStrokeColor(primary)
        canvas.setLineWidth(0.5)
        canvas.line(24, 36, width - 24, 36)
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(colors.gray)
        page_num = canvas.getPageNumber()
        canvas.drawRightString(width - 24, 24, f"Page {page_num}")
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_header_footer, onLaterPages=add_header_footer)

def _clean_docx_text(txt: str) -> str:
    if txt is None:
        return ""
    s = str(txt)
    s = s.replace("<br/>", "\n").replace("<br>", "\n")
    s = re.sub(r"</?b>", "", s)
    s = re.sub(r"</?i>", "", s)
    s = re.sub(r"</?code>", "", s)
    return s

def compile_docx(output_docx: str, order_list: list):
    """Export the report to a Word document (.docx).

    Notes:
    - The table of contents (TOC) is a Word field and updates when opened/updated in Word.
    - Only real section headings (from 'heading(n)') use Word Heading styles.
      Image/table titles are bold paragraphs so they do NOT appear in the TOC.
    """
    doc = Document()
    doc.core_properties.title = "MESA Report"
    body_style = doc.styles["Normal"]
    body_style.font.name = "Calibri"
    body_style.font.size = Pt(10)
    try:
        body_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    except Exception:
        pass

    MAX_IMAGE_WIDTH_CM = 16.0
    MAX_ATLAS_HEIGHT_CM = 18.0

    def add_heading(level: int, text: str):
        clean = _clean_docx_text(text)
        lvl = max(1, min(4, int(level)))
        doc.add_heading(clean, level=lvl)

    def add_text(text: str):
        clean = _clean_docx_text(text)
        para = doc.add_paragraph(clean)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_bold_title(text: str):
        clean = _clean_docx_text(text)
        p = doc.add_paragraph()
        r = p.add_run(clean)
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_rule():
        run = doc.add_paragraph().add_run("―" * 40)
        run.font.size = Pt(9)

    def add_table(data):
        if not data:
            return
        rows = len(data)
        cols = len(data[0]) if rows else 0
        if rows <= 0 or cols <= 0:
            return
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for r in range(rows):
            for c in range(cols):
                try:
                    table.cell(r, c).text = _clean_docx_text(data[r][c])
                except Exception:
                    table.cell(r, c).text = ""
        # Bold header row
        try:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for rr in p.runs:
                        rr.font.bold = True
        except Exception:
            pass

    def _add_field(paragraph, instr: str, placeholder: str = ""):
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), instr)
        if placeholder:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = str(placeholder)
            r.append(t)
            fld.append(r)
        paragraph._p.append(fld)

    def _is_map_image(path: str) -> bool:
        name = os.path.basename(path).lower()
        return any(tok in name for tok in ['map_', '_segments_', '_context', 'mbtiles'])

    def _is_atlas_image(path: str) -> bool:
        return 'atlas' in os.path.basename(path).lower()

    def _needs_osm_attribution(path: str) -> bool:
        return _is_map_image(path) or _is_atlas_image(path)

    def _add_osm_attribution_paragraph():
        p = doc.add_paragraph()
        r = p.add_run(OSM_ATTRIBUTION_TEXT)
        r.italic = True
        r.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _heading_starts_map_block(idx: int, max_ahead: int = 10) -> bool:
        try:
            if idx < 0 or idx >= len(order_list):
                return False
            k0, _ = order_list[idx]
            if not str(k0).startswith("heading"):
                return False
            for j in range(idx + 1, min(len(order_list), idx + 1 + max_ahead)):
                kj, _pj = order_list[j]
                if kj == "new_page":
                    return False
                if kj == "image_map":
                    return True
                if str(kj).startswith("heading"):
                    return False
            return False
        except Exception:
            return False

    has_any_content = False
    last_was_page_break = False

    for idx, (kind, payload) in enumerate(order_list):
        if kind == 'toc':
            # TOC should start on a new page.
            if has_any_content and not last_was_page_break:
                doc.add_page_break()
                last_was_page_break = True
            add_bold_title("Contents")
            p = doc.add_paragraph()
            _add_field(p, 'TOC \\o "2-3" \\h \\z \\u', placeholder="(Table of contents will appear when fields are updated in Word)")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            has_any_content = True
            last_was_page_break = False

        elif isinstance(kind, str) and kind.startswith("heading"):
            # Keep map sections from starting at end of previous page.
            if _heading_starts_map_block(idx) and has_any_content and not last_was_page_break:
                doc.add_page_break()
                last_was_page_break = True
            try:
                level = int(kind[kind.find('(') + 1:kind.find(')')])
            except Exception:
                level = 2
            add_heading(level, payload)
            has_any_content = True
            last_was_page_break = False

        elif kind == "text":
            add_text(payload)
            has_any_content = True
            last_was_page_break = False

        elif kind in ("table", "table_data", "table_data_small"):
            title = None
            data = None
            if kind == "table":
                if isinstance(payload, str) and os.path.exists(payload):
                    try:
                        df = pd.read_excel(payload)
                        data = [df.columns.tolist()] + df.fillna("").astype(str).values.tolist()
                    except Exception:
                        data = None
                else:
                    data = payload
            else:
                try:
                    title, data = payload
                except Exception:
                    title, data = None, None
            if title:
                add_bold_title(title)
            if data:
                add_table(data)
            else:
                add_text("(table data unavailable)")
            has_any_content = True
            last_was_page_break = False

        elif kind in ("image", "image_ribbon", "image_map"):
            # payload: (title, path)
            try:
                title, path = payload
            except Exception:
                title, path = None, None
            if title:
                add_bold_title(title)
            if path and isinstance(path, str) and os.path.exists(path):
                try:
                    width_cm = float(MAX_IMAGE_WIDTH_CM)
                    base_name_lower = os.path.basename(path).lower()

                    # Atlas images slightly narrower
                    if _is_atlas_image(path):
                        width_cm *= float(ATLAS_DOC_WIDTH_SCALE)

                    # Ribbons: fixed height (keep consistent visual size)
                    if kind == "image_ribbon":
                        doc.add_picture(path, width=Cm(width_cm), height=Cm(float(RIBBON_CM_HEIGHT)))
                    # Atlas: clamp height to keep pages readable
                    elif _is_atlas_image(path):
                        with PILImage.open(path) as im:
                            w_px, h_px = im.size
                        if w_px > 0 and h_px > 0:
                            aspect = float(h_px) / float(w_px)
                            out_w_cm = float(width_cm)
                            out_h_cm = out_w_cm * aspect
                            if out_h_cm > float(MAX_ATLAS_HEIGHT_CM):
                                out_h_cm = float(MAX_ATLAS_HEIGHT_CM)
                                out_w_cm = out_h_cm / aspect
                            doc.add_picture(path, width=Cm(out_w_cm))
                        else:
                            doc.add_picture(path, width=Cm(width_cm))
                    else:
                        doc.add_picture(path, width=Cm(width_cm))

                    try:
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                    if _needs_osm_attribution(path):
                        _add_osm_attribution_paragraph()
                except Exception:
                    add_text(f"(image unavailable: {os.path.basename(path)})")
            else:
                if path:
                    add_text(f"(image missing: {os.path.basename(str(path))})")
                else:
                    add_text("(image missing)")
            has_any_content = True
            last_was_page_break = False

        elif kind == "rule":
            add_rule()
            has_any_content = True
            last_was_page_break = False

        elif kind == "spacer":
            try:
                n = int(payload) if payload is not None else 1
            except Exception:
                n = 1
            n = max(1, min(20, n))
            for _ in range(n):
                doc.add_paragraph("")
            has_any_content = True
            last_was_page_break = False

        elif kind == "new_page":
            # Avoid consecutive page breaks which can create empty pages.
            if last_was_page_break:
                continue
            doc.add_page_break()
            has_any_content = True
            last_was_page_break = True

        else:
            add_text(str(payload))
            has_any_content = True
            last_was_page_break = False

    doc.save(output_docx)


def _read_text_file(path: str) -> str:
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception:
        try:
            with open(path, 'r', encoding='cp1252') as f:
                return f.read()
        except Exception:
            return ""


_URL_RE = re.compile(r"\bhttps?://[^\s<>]+", re.IGNORECASE)


def _urls_to_footnotes(text: str) -> str:
    """Replace inline URLs with numeric footnotes appended at the end."""
    if not text:
        return ""
    urls: list[str] = []

    def _repl(m: re.Match) -> str:
        url = str(m.group(0))
        urls.append(url)
        return f"[{len(urls)}]"

    out = _URL_RE.sub(_repl, str(text))
    if not urls:
        return out
    foot = "\n\nFootnotes:\n" + "\n".join([f"[{i}] {u}" for i, u in enumerate(urls, start=1)])
    return out + foot

def set_progress(pct: float, message: str | None = None):
    try:
        pct = max(0.0, min(100.0, float(pct)))
        if progress_var is not None:
            progress_var.set(pct)
        if progress_label is not None:
            progress_label.config(text=f"{int(pct)}%")
        if message:
            write_to_log(message)
    except Exception:
        pass

# ---------------- Core: generate report ----------------
def generate_report(base_dir: str,
                    config_file: str,
                    palette_A2E: dict,
                    desc_A2E: dict,
                    report_mode: str | None = None,
                    atlas_geocode_level: str | None = None,
                    include_assets: bool = True,
                    include_other_maps: bool = True,
                    include_index_statistics: bool = True,
                    include_lines_and_segments: bool = True,
                    include_atlas_maps: bool | None = None,
                    include_analysis_presentation: bool = False,
                    analysis_mode: str = "single",
                    analysis_area_left: str | None = None,
                    analysis_area_right: str | None = None):
    engine: ReportEngine | None = None
    try:
        set_progress(3, "Initializing report generation …")
        cfg       = read_config(config_file)

        # Basemap mode: xyz (default), contextily, auto
        global _REPORT_BASEMAP_MODE
        try:
            _REPORT_BASEMAP_MODE = (cfg['DEFAULT'].get('report_basemap_mode', 'xyz') or 'xyz').strip().lower()
        except Exception:
            _REPORT_BASEMAP_MODE = 'xyz'
        if _REPORT_BASEMAP_MODE not in {'xyz', 'contextily', 'auto'}:
            _REPORT_BASEMAP_MODE = 'xyz'
        write_to_log(f"Report basemap mode: {_REPORT_BASEMAP_MODE}", base_dir)

        include_atlas_cfg = _cfg_getboolean(cfg, 'DEFAULT', 'report_include_atlas_maps', default=False)
        if include_atlas_maps is None:
            if report_mode is None:
                include_atlas_maps = include_atlas_cfg
            else:
                include_atlas_maps = (str(report_mode).lower() == "detailed")
        else:
            include_atlas_maps = bool(include_atlas_maps)
        atlas_geocode_pref = (atlas_geocode_level or
                              cfg['DEFAULT'].get('atlas_report_geocode_level', '') or '').strip()
        atlas_geocode_selected: str | None = None
        write_to_log(f"Report mode selected: {'Detailed (atlas included)' if include_atlas_maps else 'General maps only'}", base_dir)
        if include_atlas_maps:
            if atlas_geocode_pref:
                write_to_log(f"Atlas geocode preference: {atlas_geocode_pref}", base_dir)
        gpq_dir   = parquet_dir_from_cfg(base_dir, cfg)
        tmp_dir_path = output_subpath(base_dir, 'tmp')
        try:
            tmp_dir_path.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        tmp_dir   = str(tmp_dir_path)
        set_progress(6, "Paths prepared.")

        # Ensure the OSM tile cache folder exists in the expected place, even when basemap is served
        # via contextily (which doesn't use our XYZ-tile cache).
        try:
            _tile_cache_root(base_dir)
        except Exception:
            pass

        # Load key tables (conditionally)
        asset_objects_df = gpd.GeoDataFrame()
        asset_groups_df = gpd.GeoDataFrame()
        need_assets_tables = bool(include_assets or include_analysis_presentation)
        if need_assets_tables:
            asset_object_pq = os.path.join(gpq_dir, "tbl_asset_object.parquet")
            asset_group_pq  = os.path.join(gpq_dir, "tbl_asset_group.parquet")
            try:
                asset_objects_df = gpd.read_parquet(asset_object_pq) if os.path.exists(asset_object_pq) else gpd.GeoDataFrame()
            except Exception:
                asset_objects_df = gpd.GeoDataFrame()
            try:
                asset_groups_df = gpd.read_parquet(asset_group_pq) if os.path.exists(asset_group_pq) else gpd.GeoDataFrame()
            except Exception:
                asset_groups_df = gpd.GeoDataFrame()
            set_progress(10, "Loaded asset tables.")
        else:
            set_progress(8, "Skipping asset tables (not selected).")

        # ---- Assets statistics ----
        object_stats_xlsx = os.path.join(tmp_dir, 'asset_object_statistics.xlsx')
        ag_stats_xlsx = os.path.join(tmp_dir, 'asset_group_statistics.xlsx')
        if include_assets:
            write_to_log("Computing asset object statistics …", base_dir)
            group_stats_df = calculate_group_statistics(asset_objects_df, asset_groups_df)
            export_to_excel(group_stats_df, object_stats_xlsx)

            ag_stats_df = fetch_asset_group_statistics(asset_groups_df, asset_objects_df)
            export_to_excel(ag_stats_df, ag_stats_xlsx)
            set_progress(22, "Asset stats ready.")
        else:
            write_to_log("Skipping asset statistics (not selected).", base_dir)
            set_progress(12)

        need_flat = bool(include_other_maps or include_index_statistics or include_lines_and_segments or include_atlas_maps)
        flat_df = gpd.GeoDataFrame()
        if need_flat:
            flat_df = load_tbl_flat(gpq_dir, base_dir=base_dir)

        need_engine = bool(include_other_maps or include_index_statistics or include_lines_and_segments or include_atlas_maps)
        if need_engine:
            engine = ReportEngine(base_dir, tmp_dir, palette_A2E, desc_A2E, config_file)

        # ---- Lines & segments (grouped per line with context & segments maps) ----
        pages_lines = []
        if include_lines_and_segments and engine is not None:
            lines_df, segments_df = fetch_lines_and_segments(gpq_dir)

            def _progress_segments(done: int, total: int):
                set_progress(25 + int(10 * done / max(1, total)), f"Rendering line segment maps ({done}/{total})")

            pages_lines, log_data = engine.render_segments(lines_df, segments_df, palette_A2E, base_dir, _progress_segments)

            if log_data:
                log_df = pd.DataFrame(log_data)
                log_xlsx = os.path.join(tmp_dir, 'line_segment_log.xlsx')
                export_to_excel(log_df, log_xlsx)
                write_to_log(f"Segments log exported to {log_xlsx}", base_dir)
            else:
                write_to_log("Skipping line/segment pages (missing/empty).", base_dir)
            set_progress(35, "Lines/segments processed.")
        else:
            write_to_log("Skipping lines/segments section (not selected).", base_dir)
            set_progress(25)

        # ---- Atlas maps ----
        atlas_df = gpd.GeoDataFrame()
        atlas_geocode_selected = None
        if include_atlas_maps:
            atlas_pq = os.path.join(gpq_dir, "tbl_atlas.parquet")
            if os.path.exists(atlas_pq):
                try:
                    atlas_df = gpd.read_parquet(atlas_pq)
                except Exception as e:
                    write_to_log(f"Failed reading tbl_atlas: {e}", base_dir)
                    atlas_df = gpd.GeoDataFrame()
            else:
                write_to_log("Parquet table tbl_atlas not found; skipping atlas maps.", base_dir)

        atlas_geocode_pref = None
        if include_atlas_maps:
            atlas_geocode_pref = (atlas_geocode_level or cfg['DEFAULT'].get('atlas_report_geocode_level', '') or '').strip() or None
            if atlas_geocode_pref:
                write_to_log(f"Atlas geocode preference: {atlas_geocode_pref}", base_dir)

        def _progress_atlas(done: int, total: int):
            if include_atlas_maps:
                set_progress(37 + int(8 * done / max(1, total)), f"Rendering atlas tiles ({done}/{total})")

        atlas_pages = []
        if engine is not None:
            atlas_pages, atlas_geocode_selected = engine.render_atlas_maps(flat_df, atlas_df, atlas_geocode_pref, include_atlas_maps, _progress_atlas)
        if atlas_pages:
            write_to_log("Per-atlas maps created.", base_dir)
        elif include_atlas_maps:
            write_to_log("Atlas maps requested but none were rendered.", base_dir)
        if include_atlas_maps:
            set_progress(45, "Atlas maps processed.")
        else:
            # Atlas is intentionally excluded in Basic/General report mode.
            set_progress(38)

        # ---- Other maps (basic_mosaic only) ----
        def _progress_geocode(done: int, total: int):
            start = 45 if include_atlas_maps else 38
            set_progress(start + int(30 * done / max(1, total)), f"Rendered maps for group {done}/{total}")

        geocode_pages = []
        geocode_intro_table = None
        geocode_groups = []
        if include_other_maps and engine is not None:
            geocode_pages, geocode_intro_table, geocode_groups = engine.render_geocode_maps(flat_df, _progress_geocode)
            if geocode_pages:
                write_to_log("Per-geocode maps created.", base_dir)
            else:
                write_to_log("tbl_flat missing or no 'name_gis_geocodegroup'; skipping per-geocode maps.", base_dir)
            set_progress(75 if include_atlas_maps else 68, "Per-geocode maps completed.")
        else:
            write_to_log("Skipping other maps section (not selected).", base_dir)
            set_progress(55 if include_atlas_maps else 45)

        # ---- Index statistics (basic_mosaic only; each index includes its map page) ----
        def _progress_indexes(done: int, total: int):
            start = 45 if include_atlas_maps else 38
            # keep this light: just a small bump in the progress bar
            set_progress(start + int(5 * done / max(1, total)), f"Rendering index charts ({done}/{total})")

        index_pages = []
        if include_index_statistics and engine is not None:
            index_pages = engine.render_index_statistics(flat_df, cfg, _progress_indexes)
        else:
            write_to_log("Skipping index statistics section (not selected).", base_dir)
        # ---- Compose PDF ----
        timestamp = datetime.datetime.now().strftime("%Y.%m.%d %H:%M:%S")

        assets_area_note = (
            "Note on areas: areas are computed only for polygon geometries. "
            "Points and lines have no surface area, therefore their 'Total area' is shown as “–”. "
            "Line and point assets are still counted in '# objects' and can be visualized on maps."
        )

        contents_lines = []
        if include_assets:
            contents_lines.append("- Assets overview & statistics")
        if include_analysis_presentation:
            mode = str(analysis_mode or "single").strip().lower()
            if mode == "compare":
                contents_lines.append("- Analysis presentation (compare two areas)")
            elif mode == "all":
                contents_lines.append("- Analysis presentation (all areas)")
            else:
                contents_lines.append("- Analysis presentation (single area)")
        if include_other_maps:
            contents_lines.append("- Other maps (basic_mosaic)")
        if include_index_statistics:
            contents_lines.append("- Index statistics + maps (basic_mosaic)")
        if include_atlas_maps and atlas_pages:
            contents_lines.append("- Atlas tile maps (per atlas object with inset)")
        if include_lines_and_segments:
            contents_lines.append("- Lines & segments (grouped by line; context & segments maps, distributions, ribbons)")
        contents_text = "<br/>".join(contents_lines)

        about_path = os.path.join(base_dir, "docs", "report_about.md")
        about_text = _urls_to_footnotes(_read_text_file(about_path)).strip() if os.path.exists(about_path) else ""

        # Build initial order_list up to About section
        order_list = [
            ('heading(1)', "MESA report"),
            ('text', f"Timestamp: {timestamp}"),
            ('spacer', 2),
            ('heading(2)', "About this report"),
            ('text', about_text or "(About text missing: docs/report_about.md)"),
            ('rule', None),
            # Insert a Word TOC field after About section
            ('toc', None),
            ('new_page', None),
        ]

        # After all sections are added to order_list, insert the dynamic Contents page
        # (This requires a second pass after order_list is fully built, so patch the PDF builder to do this)

        if include_assets:
            order_list.extend([
                ('heading(2)', "Assets – overview"),
                ('text', "Asset objects by group (count and area, where applicable)."),
                ('text', assets_area_note),
                ('table', object_stats_xlsx),
                ('spacer', 1),
                ('text', "Asset groups – sensitivity distribution (count of objects) and number of active groups (distinct groups with ≥1 object)."),
                ('table', ag_stats_xlsx),
                ('new_page', None),
            ])

        if include_analysis_presentation:
            write_to_log("Preparing analysis presentation section …", base_dir)
            try:
                gpq_dir_analysis = gpq_dir
                groups = _analysis_group_choices(gpq_dir_analysis)
                id_list = [gid for gid, _disp in groups]

                def _add_single(group_id: str):
                    group_title = _analysis_group_title(gpq_dir_analysis, group_id)
                    order_list.append(('heading(2)', f"Analysis – {group_title}"))

                    polys = _analysis_polygons_for_group(gpq_dir_analysis, group_id)
                    flat_cells = _analysis_flat_geo_for_group(gpq_dir_analysis, group_id)
                    minimap_path = os.path.join(tmp_dir, f"analysis_minimap_{group_id}.png")
                    if _analysis_write_area_map_png(
                        polys,
                        minimap_path,
                        base_dir=base_dir,
                        config_path=config_file,
                        palette_A2E=palette_A2E,
                        desc_A2E=desc_A2E,
                        overlay_alpha=0.65,
                        flat_cells_gdf=flat_cells,
                    ):
                        order_list.append(('image', ("Area overview (basemap + sensitivity)", minimap_path)))
                    else:
                        order_list.append(('text', "(Area map unavailable: missing analysis polygons and/or basemap/mbtiles.)"))

                    assets_tbl = _analysis_assets_within_area_table(asset_objects_df, asset_groups_df, polys)
                    if assets_tbl:
                        order_list.append(('table_data', ("Assets within area (polygon area)", assets_tbl)))

                    flat_sub = _analysis_flat_for_group(gpq_dir_analysis, group_id)
                    totals = _analysis_sensitivity_totals_km2(flat_sub)
                    chart_path = os.path.join(tmp_dir, f"analysis_sensitivity_max_{group_id}.png")
                    if _analysis_write_sensitivity_bar_png(totals, palette_A2E, chart_path, "Sensitivity (max) totals"):
                        order_list.append(('image', ("Sensitivity (max) totals", chart_path)))
                    else:
                        order_list.append(('text', "(Sensitivity chart unavailable: no analysis results.)"))

                    if totals is not None and not totals.empty:
                        cols = ["Code", "Description", "Area"]
                        rows = []
                        for _, r in totals.iterrows():
                            code = str(r.get("Code", ""))
                            desc = str(r.get("Description", ""))
                            area = _format_area_km2(r.get("Area (km²)", 0.0))
                            rows.append([code, desc, area])
                        table = [cols] + rows
                        order_list.append(('table_data', ("Sensitivity totals (max)", table)))
                    order_list.append(('new_page', None))

                mode = str(analysis_mode or "single").strip().lower()
                if mode == "all":
                    if not id_list:
                        order_list.append(('heading(2)', "Analysis"))
                        order_list.append(('text', "(No analysis groups found. Run data_analysis_setup.py / analysis_process.py first.)"))
                        order_list.append(('new_page', None))
                    else:
                        for gid in id_list:
                            _add_single(gid)
                elif mode == "compare":
                    left = (analysis_area_left or "").strip()
                    right = (analysis_area_right or "").strip()
                    if not left or not right:
                        order_list.append(('heading(2)', "Analysis"))
                        order_list.append(('text', "(Compare mode selected, but one or both areas were not chosen.)"))
                        order_list.append(('new_page', None))
                    else:
                        left_title = _analysis_group_title(gpq_dir_analysis, left)
                        right_title = _analysis_group_title(gpq_dir_analysis, right)
                        order_list.append(('heading(2)', f"Analysis – Compare: {left_title} vs {right_title}"))

                        left_polys = _analysis_polygons_for_group(gpq_dir_analysis, left)
                        right_polys = _analysis_polygons_for_group(gpq_dir_analysis, right)
                        left_cells = _analysis_flat_geo_for_group(gpq_dir_analysis, left)
                        right_cells = _analysis_flat_geo_for_group(gpq_dir_analysis, right)
                        left_map = os.path.join(tmp_dir, f"analysis_minimap_{left}.png")
                        right_map = os.path.join(tmp_dir, f"analysis_minimap_{right}.png")
                        if _analysis_write_area_map_png(
                            left_polys,
                            left_map,
                            base_dir=base_dir,
                            config_path=config_file,
                            palette_A2E=palette_A2E,
                            desc_A2E=desc_A2E,
                            overlay_alpha=0.65,
                            flat_cells_gdf=left_cells,
                        ):
                            order_list.append(('image', (f"Area map – {left_title}", left_map)))

                        left_assets_tbl = _analysis_assets_within_area_table(asset_objects_df, asset_groups_df, left_polys)
                        if left_assets_tbl:
                            order_list.append(('table_data', (f"Assets within area – {left_title}", left_assets_tbl)))

                        if _analysis_write_area_map_png(
                            right_polys,
                            right_map,
                            base_dir=base_dir,
                            config_path=config_file,
                            palette_A2E=palette_A2E,
                            desc_A2E=desc_A2E,
                            overlay_alpha=0.65,
                            flat_cells_gdf=right_cells,
                        ):
                            order_list.append(('image', (f"Area map – {right_title}", right_map)))

                        right_assets_tbl = _analysis_assets_within_area_table(asset_objects_df, asset_groups_df, right_polys)
                        if right_assets_tbl:
                            order_list.append(('table_data', (f"Assets within area – {right_title}", right_assets_tbl)))

                        left_totals = _analysis_sensitivity_totals_km2(_analysis_flat_for_group(gpq_dir_analysis, left))
                        right_totals = _analysis_sensitivity_totals_km2(_analysis_flat_for_group(gpq_dir_analysis, right))
                        left_chart = os.path.join(tmp_dir, f"analysis_sensitivity_max_{left}.png")
                        right_chart = os.path.join(tmp_dir, f"analysis_sensitivity_max_{right}.png")
                        _analysis_write_sensitivity_bar_png(left_totals, palette_A2E, left_chart, f"Sensitivity (max) – {left_title}")
                        _analysis_write_sensitivity_bar_png(right_totals, palette_A2E, right_chart, f"Sensitivity (max) – {right_title}")
                        if os.path.exists(left_chart):
                            order_list.append(('image', (f"Sensitivity (max) – {left_title}", left_chart)))
                        if os.path.exists(right_chart):
                            order_list.append(('image', (f"Sensitivity (max) – {right_title}", right_chart)))

                        # Relative-size (100%) composition chart
                        rel_chart = os.path.join(tmp_dir, f"analysis_sensitivity_relative_{left}_vs_{right}.png")
                        if _analysis_write_relative_share_png(left_totals, right_totals, palette_A2E, rel_chart, left_title, right_title):
                            order_list.append(('image', ("Sensitivity composition (relative)", rel_chart)))

                        # Sankey-style difference diagram (ribbons between left/right blocks)
                        sankey_chart = os.path.join(tmp_dir, f"analysis_sensitivity_sankey_{left}_vs_{right}.png")
                        if _analysis_write_sankey_difference_png(left_totals, right_totals, palette_A2E, sankey_chart, left_title, right_title):
                            order_list.append(('image', ("Area difference Sankey", sankey_chart)))

                        lm = {r["Code"]: float(r["Area (km²)"]) for _, r in left_totals.iterrows()} if left_totals is not None and not left_totals.empty else {}
                        rm = {r["Code"]: float(r["Area (km²)"]) for _, r in right_totals.iterrows()} if right_totals is not None and not right_totals.empty else {}
                        order_map = {c: i for i, c in enumerate(SENSITIVITY_ORDER)}

                        def _code_key(val: str) -> tuple[int, str]:
                            s = str(val or "").strip().upper()
                            if s in order_map:
                                return (0, f"{order_map[s]:02d}")
                            if s in ("", "NONE", "NAN"):
                                return (3, "")
                            if s == "UNKNOWN":
                                return (2, s)
                            return (1, s)

                        all_codes = sorted(set(lm.keys()) | set(rm.keys()), key=_code_key)
                        comp_rows = [["Code", f"{left_title} (km²)", f"{right_title} (km²)", "Difference (L - R)"]]
                        for code in all_codes:
                            lv = lm.get(code, 0.0)
                            rv = rm.get(code, 0.0)
                            comp_rows.append([code, f"{lv:,.2f}", f"{rv:,.2f}", f"{(lv-rv):+,.2f}"])
                        order_list.append(('table_data', ("Sensitivity comparison (max)", comp_rows)))
                        order_list.append(('new_page', None))
                else:
                    gid = (analysis_area_left or "").strip() or (id_list[0] if id_list else "")
                    if not gid:
                        order_list.append(('heading(2)', "Analysis"))
                        order_list.append(('text', "(No analysis groups found. Run data_analysis_setup.py / analysis_process.py first.)"))
                        order_list.append(('new_page', None))
                    else:
                        _add_single(gid)
            except Exception as exc:
                write_to_log(f"Analysis presentation section failed: {exc}", base_dir)
                order_list.append(('heading(2)', "Analysis"))
                order_list.append(('text', "(Analysis presentation could not be generated; check log for details.)"))
                order_list.append(('new_page', None))

        if include_other_maps:
            order_list.extend([
                ('heading(2)', "Other maps"),
                ('text',
                 "Maps in this section are rendered for <b>basic_mosaic</b> only (other geocodes are omitted). "
                 "Rendered from MBTiles (raster) with a basemap background."),
            ])

            if geocode_intro_table is not None:
                order_list.append(('table_data', ("Available geocode categories and object counts", geocode_intro_table)))
                order_list.append(('rule', None))

            order_list.extend(geocode_pages)

        if include_index_statistics:
            order_list.extend([
                ('heading(2)', "Index statistics"),
                ('text', "One page per index based on <b>basic_mosaic</b> values from <b>tbl_flat</b>. Each index includes its map page."),
            ])

        if index_pages:
            order_list.extend(index_pages)
        else:
            order_list.append(('text', "No index statistics could be generated (missing data/columns)."))
            order_list.append(('new_page', None))
        if atlas_pages:
            order_list.append(('heading(2)', "Atlas maps"))
            atlas_intro = ("Each atlas tile focuses on a subset of the study area. "
                           "An inset map indicates where the tile sits relative to the full extent.")
            if atlas_geocode_selected:
                atlas_intro += f" Geocode level shown: <b>{atlas_geocode_selected}</b>."
            order_list.append(('text', atlas_intro))
            order_list.extend(atlas_pages)

        if include_lines_and_segments and pages_lines:
            order_list.append(('heading(2)', "Lines and segments"))
            order_list.append(('text', "Images contain only cartography—no embedded titles. "
                                       "Ribbons are fixed to 0.6 cm high and the full text width. "
                                       "Distance markers are written as text before each ribbon."))
            order_list.append(('rule', None))
            order_list.append(('new_page', None))
            order_list.extend(pages_lines)

        set_progress(86, "Composing Word report …")
        elements = line_up_to_pdf(order_list)

        ts_docx = datetime.datetime.now().strftime("%Y-%m-%d_%H%M")
        output_docx_path = output_subpath(base_dir, "reports", f"MESA-report_{ts_docx}.docx")
        output_docx = str(output_docx_path)
        compile_docx(output_docx, order_list)
        engine.cleanup()
        engine = None
        set_progress(100, "Report completed.")

        global last_report_path
        last_report_path = output_docx
        write_to_log(f"Word report created: {output_docx}", base_dir)

        try:
            if link_var:
                link_var.set("Open report folder")
        except Exception:
            pass

    except Exception as e:
        write_to_log(f"ERROR during report generation: {e}", base_dir)
        set_progress(100, "Report failed.")
    finally:
        if engine is not None:
            try:
                engine.cleanup()
            except Exception:
                pass

# ---------------- GUI runner ----------------
def _start_report_thread(base_dir, config_file, palette, desc, report_mode, atlas_geocode):
    # Backwards-compatible helper for old call sites.
    threading.Thread(
        target=generate_report,
        args=(base_dir, config_file, palette, desc, report_mode, atlas_geocode),
        daemon=True
    ).start()

def _start_report_thread_selected(base_dir, config_file, palette, desc, *, atlas_geocode,
                                 include_assets: bool,
                                 include_other_maps: bool,
                                 include_index_statistics: bool,
                                 include_lines_and_segments: bool,
                                 include_atlas_maps: bool,
                                 include_analysis_presentation: bool,
                                 analysis_mode: str,
                                 analysis_area_left: str | None,
                                 analysis_area_right: str | None):
    threading.Thread(
        target=generate_report,
        kwargs={
            'base_dir': base_dir,
            'config_file': config_file,
            'palette_A2E': palette,
            'desc_A2E': desc,
            'report_mode': None,
            'atlas_geocode_level': atlas_geocode,
            'include_assets': include_assets,
            'include_other_maps': include_other_maps,
            'include_index_statistics': include_index_statistics,
            'include_lines_and_segments': include_lines_and_segments,
            'include_atlas_maps': include_atlas_maps,
            'include_analysis_presentation': include_analysis_presentation,
            'analysis_mode': analysis_mode,
            'analysis_area_left': analysis_area_left,
            'analysis_area_right': analysis_area_right,
        },
        daemon=True
    ).start()

def launch_gui(base_dir: str, config_file: str, palette: dict, desc: dict, theme: str):
    global log_widget, progress_var, progress_label, link_var, atlas_geocode_var, _atlas_geocode_choices
    root = tb.Window(themename=theme)
    root.title("MESA – Report generator")
    try:
        ico = Path(base_dir) / "system_resources" / "mesa.ico"
        if ico.exists():
            root.iconbitmap(str(ico))
    except Exception:
        pass

    frame = tb.LabelFrame(root, text="Log", bootstyle="info")
    frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
    log_widget = scrolledtext.ScrolledText(frame, height=14)
    log_widget.pack(fill=tk.BOTH, expand=True)

    pframe = tk.Frame(root); pframe.pack(pady=6)
    progress_var = tk.DoubleVar(value=0.0)
    pbar = tb.Progressbar(pframe, orient="horizontal", length=300, mode="determinate",
                          variable=progress_var, bootstyle="info")
    pbar.pack(side=tk.LEFT, padx=6)
    progress_label = tk.Label(pframe, text="0%", width=5, anchor="w")
    progress_label.pack(side=tk.LEFT)

    atlas_frame = tb.LabelFrame(root, text="Atlas geocode level", bootstyle="secondary")
    atlas_frame.pack(padx=10, pady=(0, 10), fill=tk.X)
    _atlas_geocode_choices = _available_geocode_levels(base_dir, config_file)
    default_level = None
    if _atlas_geocode_choices:
        default_level = 'basic_mosaic' if 'basic_mosaic' in _atlas_geocode_choices else _atlas_geocode_choices[0]
    atlas_geocode_var = tk.StringVar(value=default_level or '')
    atlas_combo = tb.Combobox(atlas_frame, textvariable=atlas_geocode_var,
                              values=tuple(_atlas_geocode_choices),
                              state="readonly" if _atlas_geocode_choices else "disabled",
                              width=30)
    atlas_combo.grid(row=0, column=0, padx=6, pady=4, sticky="w")
    if _atlas_geocode_choices:
        tk.Label(atlas_frame, text="Used for atlas sensitivity and environment maps.",
                 anchor="w").grid(row=0, column=1, padx=6, pady=4, sticky="w")
    else:
        tk.Label(atlas_frame, text="(No geocode levels detected yet)", anchor="w")\
          .grid(row=0, column=1, padx=6, pady=4, sticky="w")

    action_frame = tb.LabelFrame(root, text="Report", bootstyle="primary")
    action_frame.pack(padx=10, pady=(0, 10), fill=tk.X)
    action_frame.columnconfigure(1, weight=1)

    contents_frame = tk.Frame(action_frame)
    contents_frame.grid(row=0, column=0, columnspan=2, padx=6, pady=(6, 2), sticky="w")

    contents_frame.columnconfigure(0, weight=1)
    contents_frame.columnconfigure(1, weight=1)

    tk.Label(contents_frame, text="Include in report:").grid(row=0, column=0, columnspan=2, sticky="w", padx=(0, 10))

    include_assets_var = tk.BooleanVar(value=True)
    include_analysis_var = tk.BooleanVar(value=False)
    include_other_maps_var = tk.BooleanVar(value=True)
    include_index_statistics_var = tk.BooleanVar(value=True)
    include_lines_and_segments_var = tk.BooleanVar(value=True)
    include_atlas_maps_var = tk.BooleanVar(value=False)

    _include_options = [
        ("Assets overview & statistics", include_assets_var),
        ("Analysis presentation (graphs)", include_analysis_var),
        ("Other maps (basic_mosaic)", include_other_maps_var),
        ("Index statistics", include_index_statistics_var),
        ("Lines & segments", include_lines_and_segments_var),
        ("Atlas maps (detailed)", include_atlas_maps_var),
    ]
    for idx, (label, var) in enumerate(_include_options):
        row = 1 + (idx // 2)
        col = idx % 2
        tb.Checkbutton(
            contents_frame,
            text=label,
            variable=var,
            bootstyle="round-toggle",
        ).grid(row=row, column=col, sticky="w", pady=2, padx=(0, 16) if col == 0 else (0, 0))

    analysis_frame = tb.LabelFrame(action_frame, text="Analysis presentation options", bootstyle="secondary")
    analysis_frame.grid(row=1, column=0, columnspan=2, padx=6, pady=(4, 2), sticky="ew")
    analysis_frame.columnconfigure(1, weight=1)

    analysis_mode_var = tk.StringVar(value="single")
    analysis_left_var = tk.StringVar(value="")
    analysis_right_var = tk.StringVar(value="")

    tk.Label(analysis_frame, text="Mode:").grid(row=0, column=0, padx=(6, 2), pady=4, sticky="w")
    mode_holder = tk.Frame(analysis_frame)
    mode_holder.grid(row=0, column=1, padx=2, pady=4, sticky="w")
    tb.Radiobutton(mode_holder, text="Single area", value="single", variable=analysis_mode_var,
                   bootstyle="secondary-toolbutton").pack(side=tk.LEFT, padx=(0, 6))
    tb.Radiobutton(mode_holder, text="Compare two areas", value="compare", variable=analysis_mode_var,
                   bootstyle="secondary-toolbutton").pack(side=tk.LEFT, padx=(0, 6))
    tb.Radiobutton(mode_holder, text="All areas", value="all", variable=analysis_mode_var,
                   bootstyle="secondary-toolbutton").pack(side=tk.LEFT)

    try:
        _cfg_tmp = read_config(config_file)
        _gpq_tmp = parquet_dir_from_cfg(base_dir, _cfg_tmp)
    except Exception:
        _gpq_tmp = parquet_dir_from_cfg(base_dir, read_config(config_file))
    group_choices = _analysis_group_choices(_gpq_tmp)
    group_labels = [disp for _gid, disp in group_choices]
    label_to_id = {disp: gid for gid, disp in group_choices}

    tk.Label(analysis_frame, text="Area A:").grid(row=1, column=0, padx=(6, 2), pady=4, sticky="w")
    analysis_left_combo = tb.Combobox(analysis_frame, textvariable=analysis_left_var, values=tuple(group_labels), width=44,
                                      state="readonly" if group_labels else "disabled")
    analysis_left_combo.grid(row=1, column=1, padx=2, pady=4, sticky="w")

    tk.Label(analysis_frame, text="Area B:").grid(row=2, column=0, padx=(6, 2), pady=4, sticky="w")
    analysis_right_combo = tb.Combobox(analysis_frame, textvariable=analysis_right_var, values=tuple(group_labels), width=44,
                                       state="readonly" if group_labels else "disabled")
    analysis_right_combo.grid(row=2, column=1, padx=2, pady=4, sticky="w")

    if group_labels:
        analysis_left_var.set(group_labels[0])
        analysis_right_var.set(group_labels[1] if len(group_labels) > 1 else group_labels[0])

    def _sync_analysis_controls(*_args):
        enabled = bool(include_analysis_var.get())
        mode = str(analysis_mode_var.get() or "single").lower()
        if not enabled or not group_labels:
            analysis_left_combo.configure(state="disabled")
            analysis_right_combo.configure(state="disabled")
            return
        if mode == "all":
            analysis_left_combo.configure(state="disabled")
            analysis_right_combo.configure(state="disabled")
        elif mode == "compare":
            analysis_left_combo.configure(state="readonly")
            analysis_right_combo.configure(state="readonly")
        else:
            analysis_left_combo.configure(state="readonly")
            analysis_right_combo.configure(state="disabled")

    include_analysis_var.trace_add("write", _sync_analysis_controls)
    analysis_mode_var.trace_add("write", _sync_analysis_controls)
    _sync_analysis_controls()

    def _start_selected():
        if not any([
            include_assets_var.get(),
            include_analysis_var.get(),
            include_other_maps_var.get(),
            include_index_statistics_var.get(),
            include_lines_and_segments_var.get(),
            include_atlas_maps_var.get(),
        ]):
            write_to_log("No report sections selected. Tick at least one box.", base_dir)
            return

        analysis_left_id = None
        analysis_right_id = None
        if include_analysis_var.get():
            left_label = (analysis_left_var.get() or "").strip()
            right_label = (analysis_right_var.get() or "").strip()
            analysis_left_id = label_to_id.get(left_label) if left_label else None
            analysis_right_id = label_to_id.get(right_label) if right_label else None

        _start_report_thread_selected(
            base_dir,
            config_file,
            palette,
            desc,
            atlas_geocode=atlas_geocode_var.get(),
            include_assets=include_assets_var.get(),
            include_analysis_presentation=include_analysis_var.get(),
            analysis_mode=analysis_mode_var.get(),
            analysis_area_left=analysis_left_id,
            analysis_area_right=analysis_right_id,
            include_other_maps=include_other_maps_var.get(),
            include_index_statistics=include_index_statistics_var.get(),
            include_lines_and_segments=include_lines_and_segments_var.get(),
            include_atlas_maps=include_atlas_maps_var.get(),
        )

    tb.Button(
        action_frame,
        text="Create report",
        bootstyle="success",
        width=18,
        command=_start_selected,
    ).grid(row=2, column=0, padx=6, pady=(6, 6), sticky="w")
    tk.Label(
        action_frame,
        text="The report will include only the selected sections.",
        anchor="w",
        justify="left",
        wraplength=420,
    ).grid(row=2, column=1, padx=(4, 8), pady=(6, 6), sticky="w")

    btn_frame = tk.Frame(root); btn_frame.pack(pady=(0, 6))
    tb.Button(btn_frame, text="Exit", bootstyle=WARNING, command=root.destroy).pack(side=tk.RIGHT, padx=6)

    # Live link to report folder
    def open_report_folder(event=None):
        if not last_report_path:
            return
        folder = os.path.dirname(last_report_path)
        if os.path.isdir(folder):
            try:
                os.startfile(folder)  # Windows
            except Exception:
                try:
                    subprocess.Popen(["explorer", folder])
                except Exception as ee:
                    write_to_log(f"Failed to open folder: {ee}", base_dir)
        else:
            write_to_log("Report folder not found.", base_dir)

    link_var = tk.StringVar(value="")
    link_label = tk.Label(root, textvariable=link_var, fg="#4ea3ff", cursor="hand2",
                          font=("Segoe UI", 10))
    link_label.pack(pady=(2, 8))
    link_label.bind("<Button-1>", open_report_folder)

    write_to_log(f"Working directory: {base_dir}", base_dir)
    write_to_log("Ready. Select report contents, then press 'Create report'.", base_dir)
    root.mainloop()

# ---------------- CLI ----------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Presentation report (GeoParquet per geocode, same-scale maps, line context + segments maps with buffers)')
    parser.add_argument('--original_working_directory', required=False, help='Path to running folder')
    parser.add_argument('--no-gui', action='store_true', help='Run directly without GUI')
    parser.add_argument('--report-mode', choices=['general', 'detailed'],
                        help='Select report detail level (general maps or detailed with atlas overviews).')
    parser.add_argument('--atlas-geocode',
                        help='Override geocode level to use for atlas maps (e.g. basic_mosaic, H3_R7).')
    parser.add_argument('--debug-atlas-sample',
                        help='Render a standalone sensitivity atlas map for the specified tile (name_gis).')
    parser.add_argument('--debug-atlas-geocode',
                        help='Optional geocode level to filter polygons when using --debug-atlas-sample.')
    parser.add_argument('--debug-atlas-size', type=int,
                        help='Optional max polygon count for --debug-atlas-sample (down-sample if exceeded).')
    args = parser.parse_args()

    base_dir = args.original_working_directory
    if not base_dir:
        base_dir = os.getcwd()
    base_dir = normalize_base_dir(base_dir)

    cfg_path = config_path(base_dir)
    cfg = read_config(cfg_path)
    theme = cfg['DEFAULT'].get('ttk_bootstrap_theme', 'flatly')

    # Sensitivity palette + descriptions from config (A–E)
    palette_A2E, desc_A2E = read_sensitivity_palette_and_desc(cfg_path)

    if args.debug_atlas_sample:
        out = debug_atlas_sample(
            base_dir,
            cfg,
            palette_A2E,
            desc_A2E,
            args.debug_atlas_sample,
            geocode_level=args.debug_atlas_geocode,
            sample_size=args.debug_atlas_size
        )
        if out:
            print(f"Debug atlas sample saved to: {out}")
        else:
            print("Debug atlas sample failed. Check log for details.")
        sys.exit(0)

    # Optional override for brand color from config
    PRIMARY_HEX = cfg['DEFAULT'].get('ui_primary_color', PRIMARY_HEX).strip() or PRIMARY_HEX
    try:
        def _lighten(hexcol, amt=0.30):
            hexcol = hexcol.lstrip('#')
            r = int(hexcol[0:2],16); g = int(hexcol[2:4],16); b = int(hexcol[4:6],16)
            r = int(r + (255 - r)*amt); g = int(g + (255 - g)*amt); b = int(b + (255 - b)*amt)
            return f"#{r:02x}{g:02x}{b:02x}"
        LIGHT_PRIMARY_HEX = _lighten(PRIMARY_HEX, 0.30)
    except Exception:
        pass

    if args.no_gui:
        generate_report(base_dir, cfg_path, palette_A2E, desc_A2E,
                        report_mode=args.report_mode,
                        atlas_geocode_level=args.atlas_geocode)
    else:
        launch_gui(base_dir, cfg_path, palette_A2E, desc_A2E, theme)

